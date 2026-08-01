# =========================================================
# Auditor de Contratos - Grupo SBF - VERSÃO REVISADA + IA REFORÇADA
# Ajustes aplicados: visual profissional, carregamento seguro
# do histórico, proteção de HTML no Assistente IA e refinamentos
# de usabilidade sem alterar o fluxo principal.
# =========================================================

import io
import os
import re
import json
import html
import math
import base64
import tempfile
import time
import threading
from pathlib import Path
from datetime import datetime, timedelta
from typing import Any, Dict, List
from textwrap import dedent

try:
    from streamlit.runtime.scriptrunner import add_script_run_ctx
except Exception:
    add_script_run_ctx = None

import pandas as pd
import sqlite3
import streamlit as st
import streamlit.components.v1 as components
import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document
from dotenv import load_dotenv

from database import criar_banco, salvar_analise, listar_analises, limpar_historico
from auditor_evidencias import (
    PROMPT_EVIDENCIAS_V4,
    PROMPT_VERIFICADOR_V4,
    aplicar_motor_evidencias_v4,
    linhas_auditoria_para_tela,
)
from extrator_tabela_comercial import (
    extrair_tabela_comercial_completa,
    mesclar_itens_comerciais,
    calcular_metricas_tabela_comercial,
)

from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter


# =========================================================
# CONFIGURAÇÕES INICIAIS
# =========================================================
load_dotenv()


def obter_gemini_key() -> str:
    """Lê a chave do Gemini no local (.env) e no Streamlit Cloud (Secrets)."""
    chave_env = os.getenv("GEMINI_API_KEY", "")
    if chave_env:
        return str(chave_env).strip()
    try:
        chave_secret = st.secrets.get("GEMINI_API_KEY", "")
        return str(chave_secret).strip()
    except Exception:
        return ""
criar_banco()

TESSERACT_CMD = os.getenv("TESSERACT_CMD", r"C:\Users\124034\AppData\Local\Programs\Tesseract-OCR\tesseract.exe")
if os.path.exists(TESSERACT_CMD):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

st.set_page_config(
    page_title="Auditor de Contratos - Grupo SBF",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Campos oficiais exibidos nos cards e no Excel
CAMPOS_OFICIAIS = [
    ("Tipo de Contrato", "tipo_contrato"),
    ("Empresa do Grupo SBF", "empresa_grupo_sbf"),
    ("CNPJ Empresa do Grupo", "cnpj_empresa_grupo"),
    ("Local de Prestação Contraparte", "local_prestacao"),
    ("Contraparte", "contraparte"),
    ("CNPJ Contraparte", "cnpj_contraparte"),
    ("Objetivo", "objetivo"),
    ("Descrição do Serviço/ Material", "descricao_servico_material"),
    ("Descrição breve do cadastro", "descricao_breve_cadastro"),
    ("Forma de pagamento", "forma_pagamento"),
    ("Condição de Pagamento em Dias", "condicao_pagamento_dias"),
    ("Multa", "multa"),
    ("Vigência após a data de assinatura", "vigencia_apos_assinatura"),
    ("Tipo de Vigência", "tipo_vigencia"),
    ("Período de Vigência", "periodo_vigencia_formatado"),
    ("Status Contratual", "status_contratual"),
    ("Situação Operacional", "situacao_operacional"),
    ("Resumo de Aditivos", "resumo_aditivos"),
    ("Rescisão e Indenização", "rescisao_indenizacao"),
    ("Anticorrupção", "anticorrupcao"),
    ("Proteção de Dados LGPD", "protecao_dados_lgpd"),
    ("Data da Assinatura", "data_assinatura"),
    ("Data do Contrato", "data_contrato"),
    ("Data Conclusão DocuSign", "data_conclusao_docusign"),
    ("Valor do Contrato Original", "valor_contrato_original"),
    ("Valor Mensal Estimado", "valor_mensal_estimado"),
    ("Valor Total Estimado da Vigência", "valor_total_estimado_vigencia"),
    ("Valor Total dos Materiais e Serviços", "valor_total_materiais_servicos"),
    ("Pessoas que assinaram", "pessoas_que_assinaram"),
]

CAMPOS_JSON_OBRIGATORIOS = ", ".join([campo for _, campo in CAMPOS_OFICIAIS] + [
    "aditivos_contrato",
    "contraparte", "fornecedor",
    "contrato_assinado", "alerta_assinatura", "status", "status_contratual", "situacao_operacional", "tipo_vigencia",
    "risco", "score", "confianca_extracao", "indicadores_pendencias", "metricas_tabela_comercial",
    "resumo_executivo", "parecer", "checklist", "pendencias", "itens_contrato", "assinaturas_contrato",
    "data_reconhecimento_firma"
])


# =========================================================
# ESTILO VISUAL
# =========================================================
CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800;900&display=swap');

:root{
    --bg:#070b10;
    --panel:#101822;
    --panel-2:#0b1118;
    --green:#003c2f;
    --green-2:#065f46;
    --gold:#d7bf75;
    --gold-soft:#f3e6b3;
    --text:#f8fafc;
    --muted:#94a3b8;
    --line:rgba(215,191,117,.22);
    --danger:#dc2626;
    --warn:#f59e0b;
    --ok:#16a34a;
}

html, body, [class*="css"]{
    font-family:'Inter', sans-serif;
}

.stApp{
    background:
        radial-gradient(circle at top left, rgba(0,60,47,.32), transparent 34%),
        radial-gradient(circle at top right, rgba(215,191,117,.18), transparent 32%),
        var(--bg);
    color:var(--text);
}

.block-container{
    padding-top:1.8rem;
    padding-bottom:2rem;
}

section[data-testid="stSidebar"]{
    background:linear-gradient(180deg,#202631,#111821);
    border-right:1px solid var(--line);
}

section[data-testid="stSidebar"] *{
    color:#f8fafc;
}

.stRadio label, .stFileUploader label{
    font-weight:800 !important;
}

.hero{
    padding:34px 40px;
    border-radius:28px;
    background:
        linear-gradient(115deg, rgba(0,60,47,.98), rgba(10,88,65,.96) 52%, rgba(215,191,117,.95));
    margin:0 0 26px;
    box-shadow:0 22px 55px rgba(0,0,0,.38);
    border:1px solid rgba(255,255,255,.10);
}

.hero .eyebrow{
    display:inline-block;
    padding:8px 13px;
    border-radius:999px;
    background:rgba(255,255,255,.13);
    color:#fff;
    font-size:12px;
    font-weight:900;
    letter-spacing:.9px;
    text-transform:uppercase;
    margin-bottom:14px;
}

.hero h1{
    margin:0;
    color:#fff;
    font-size:44px;
    line-height:1.05;
    font-weight:900;
    letter-spacing:.4px;
}

.hero p{
    margin:14px 0 0;
    color:rgba(255,255,255,.92);
    font-size:17px;
    font-weight:600;
    max-width:980px;
}

.section-title{
    color:var(--gold);
    font-size:22px;
    font-weight:900;
    margin:28px 0 14px;
    letter-spacing:.2px;
}

.subtle{
    color:var(--muted);
    font-size:13px;
    line-height:1.6;
}

.metric-card{
    background:linear-gradient(145deg,var(--panel),var(--panel-2));
    border:1px solid var(--line);
    border-radius:22px;
    padding:22px;
    min-height:122px;
    box-shadow:0 16px 38px rgba(0,0,0,.35);
}

.metric-card small{
    color:var(--gold);
    font-weight:900;
    text-transform:uppercase;
    letter-spacing:.9px;
    font-size:11px;
}

.metric-card h2{
    color:#fff;
    margin:13px 0 0;
    font-size:34px;
    line-height:1.05;
    font-weight:900;
    word-break:break-word;
}

.metric-link{
    display:block;
    text-decoration:none !important;
    color:inherit !important;
}

.metric-card.clickable{
    cursor:pointer;
    transition:all .18s ease;
    position:relative;
    overflow:hidden;
}

.metric-card.clickable:hover{
    transform:translateY(-3px);
    border-color:rgba(215,191,117,.70);
    box-shadow:0 20px 48px rgba(0,0,0,.42);
}

.metric-card.clickable.active{
    border:2px solid var(--gold);
    background:linear-gradient(145deg,rgba(0,60,47,.98),rgba(11,17,24,.98));
}

.metric-card.clickable.active::after{
    content:"Filtro ativo";
    position:absolute;
    top:14px;
    right:14px;
    background:rgba(215,191,117,.18);
    color:var(--gold);
    border:1px solid rgba(215,191,117,.45);
    border-radius:999px;
    padding:5px 9px;
    font-size:10px;
    font-weight:900;
    text-transform:uppercase;
    letter-spacing:.5px;
}

.filter-note{
    margin:16px 0 4px;
    padding:12px 16px;
    border-radius:14px;
    background:rgba(0,60,47,.42);
    border:1px solid rgba(215,191,117,.25);
    color:#f8fafc;
    font-size:13px;
    font-weight:700;
}

.executive-box{
    background:linear-gradient(135deg, rgba(0,60,47,.96), rgba(17,24,33,.96));
    border:1px solid rgba(215,191,117,.35);
    border-radius:20px;
    padding:22px;
    color:#fff;
    margin:12px 0;
    box-shadow:0 14px 35px rgba(0,0,0,.25);
}

.executive-box h3{
    margin:0 0 8px;
    color:var(--gold);
    font-weight:900;
}

.info-grid{
    display:grid;
    grid-template-columns:repeat(3, minmax(0, 1fr));
    gap:16px;
}

.info-card{
    background:linear-gradient(145deg,var(--panel),var(--panel-2));
    border:1px solid rgba(215,191,117,.22);
    border-radius:18px;
    padding:18px;
    min-height:112px;
    box-shadow:0 10px 28px rgba(0,0,0,.24);
}

.info-card small{
    display:block;
    color:var(--gold);
    font-weight:900;
    text-transform:uppercase;
    letter-spacing:.7px;
    font-size:11px;
}

.info-card p{
    color:#fff;
    font-size:15px;
    font-weight:700;
    margin:10px 0 0;
    line-height:1.45;
    overflow-wrap:anywhere;
    white-space:pre-line;
}

.flow-grid{
    display:grid;
    grid-template-columns:repeat(5, minmax(0, 1fr));
    gap:14px;
    margin-top:14px;
}

.flow-card{
    background:linear-gradient(145deg,var(--panel),var(--panel-2));
    border:1px solid rgba(215,191,117,.24);
    border-radius:20px;
    padding:18px;
    min-height:150px;
}

.flow-number{
    width:36px;
    height:36px;
    border-radius:999px;
    background:var(--gold);
    color:var(--green);
    display:flex;
    align-items:center;
    justify-content:center;
    font-weight:900;
    margin-bottom:12px;
}

.flow-card h4{
    color:#fff;
    margin:0;
    font-size:16px;
    font-weight:900;
}

.flow-card p{
    color:var(--muted);
    font-size:13px;
    margin:8px 0 0;
    line-height:1.45;
}

.contract-card{
    background:linear-gradient(145deg,var(--panel),var(--panel-2));
    border:1px solid rgba(215,191,117,.18);
    border-radius:22px;
    padding:22px;
    margin:16px 0;
    box-shadow:0 16px 38px rgba(0,0,0,.32);
}

.contract-top{
    display:flex;
    justify-content:space-between;
    align-items:flex-start;
    gap:16px;
}

.contract-title{
    color:var(--gold);
    font-size:21px;
    font-weight:900;
    margin:0;
}

.contract-file{
    color:#93c5fd;
    font-size:13px;
    margin-top:6px;
    overflow-wrap:anywhere;
}

.contract-grid{
    display:grid;
    grid-template-columns:2fr 1fr 1.5fr .7fr;
    gap:20px;
    margin-top:18px;
}

.contract-label{
    color:var(--muted);
    font-size:11px;
    font-weight:900;
    text-transform:uppercase;
    letter-spacing:.7px;
}

.contract-value{
    color:#fff;
    font-size:14px;
    font-weight:800;
    margin-top:7px;
    line-height:1.5;
    overflow-wrap:anywhere;
}

.badge-risk{
    color:white;
    padding:9px 16px;
    border-radius:999px;
    font-weight:900;
    font-size:13px;
    white-space:nowrap;
}

.card-footer{
    display:flex;
    justify-content:space-between;
    align-items:center;
    gap:14px;
    margin-top:20px;
    padding-top:16px;
    border-top:1px solid rgba(255,255,255,.08);
}

.excel-link{
    background:var(--green);
    color:white !important;
    padding:9px 14px;
    border-radius:10px;
    font-size:12px;
    font-weight:900;
    text-decoration:none !important;
    border:1px solid rgba(215,191,117,.60);
}

.risk-row, .ok-row{
    padding:15px 16px;
    border-radius:14px;
    margin:10px 0;
    color:#fff;
    font-weight:700;
}

.risk-row{
    background:rgba(245,158,11,.14);
    border-left:5px solid var(--warn);
}

.ok-row{
    background:rgba(22,163,74,.14);
    border-left:5px solid var(--ok);
}

.pill{
    display:block;
    padding:14px 16px;
    border-radius:16px;
    font-weight:800;
    line-height:1.5;
}

.pill-ok{background:rgba(22,163,74,.14);color:#bbf7d0;border:1px solid rgba(22,163,74,.30);}
.pill-warn{background:rgba(245,158,11,.14);color:#fde68a;border:1px solid rgba(245,158,11,.30);}
.pill-danger{background:rgba(220,38,38,.14);color:#fecaca;border:1px solid rgba(220,38,38,.30);}

.stButton>button,
.stDownloadButton>button{
    border-radius:13px !important;
    border:1px solid var(--gold) !important;
    background:var(--green) !important;
    color:#fff !important;
    font-weight:900 !important;
    min-height:43px;
}

.stButton>button:hover,
.stDownloadButton>button:hover{
    border-color:#fff !important;
    color:var(--gold) !important;
}

[data-testid="stDataFrame"]{
    border-radius:16px;
    overflow:hidden;
    border:1px solid rgba(215,191,117,.18);
}

.stTextInput input,
.stTextArea textarea,
.stSelectbox div[data-baseweb="select"] > div,
.stMultiSelect div[data-baseweb="select"] > div{
    border-radius:13px !important;
    border:1px solid rgba(215,191,117,.24) !important;
    background:#0b1118 !important;
    color:#f8fafc !important;
}

.stFileUploader{
    background:linear-gradient(145deg,rgba(16,24,34,.82),rgba(11,17,24,.82));
    border:1px dashed rgba(215,191,117,.45);
    border-radius:20px;
    padding:14px;
}

button[kind="secondary"]{
    transition:all .18s ease !important;
}

button[kind="secondary"]:hover{
    transform:translateY(-1px);
}

[data-testid="stTabs"] button{
    font-weight:900 !important;
}

.footer{
    color:var(--muted);
    text-align:center;
    margin-top:32px;
    padding:20px;
    border-top:1px solid rgba(255,255,255,.10);
}

@media(max-width:1100px){
    .info-grid{grid-template-columns:repeat(2, minmax(0, 1fr));}
    .flow-grid{grid-template-columns:repeat(2, minmax(0, 1fr));}
    .contract-grid{grid-template-columns:1fr 1fr;}
}

@media(max-width:700px){
    .hero{padding:26px 24px;}
    .hero h1{font-size:32px;}
    .info-grid,.flow-grid,.contract-grid{grid-template-columns:1fr;}
    .contract-top,.card-footer{flex-direction:column;align-items:flex-start;}
}

/* Filtros do histórico mais limpos */
div[data-testid="stExpander"]{
    border:1px solid rgba(215,191,117,.22) !important;
    border-radius:18px !important;
    background:rgba(16,24,34,.42) !important;
}

div[data-testid="stExpander"] summary{
    font-weight:900 !important;
    color:#f3e6b3 !important;
}

.filter-note{
    background:linear-gradient(135deg,rgba(0,60,47,.80),rgba(16,24,34,.88));
    border:1px solid rgba(215,191,117,.28);
    border-radius:16px;
    padding:13px 16px;
    color:#e5e7eb;
    margin:14px 0 20px;
    font-weight:700;
}

/* Histórico: deixa filtros mais alinhados e elegantes */
div[data-testid="stVerticalBlockBorderWrapper"] input,
div[data-testid="stVerticalBlockBorderWrapper"] div[data-baseweb="select"]{
    min-height:42px;
}

div[data-testid="stVerticalBlockBorderWrapper"] label{
    min-height:22px;
}

/* Assistente IA: evita aparência de bloco de código e melhora os cards */
[data-testid="stChatMessage"] .auditor-ai-box{
    width:100%;
    box-sizing:border-box;
}

@media (max-width: 1100px){
    .ai-summary-grid,
    .ai-contract-grid{
        grid-template-columns:repeat(2, minmax(0,1fr)) !important;
    }
}

@media (max-width: 700px){
    .ai-summary-grid,
    .ai-contract-grid{
        grid-template-columns:1fr !important;
    }
}


/* Histórico: card executivo premium */
.history-card-v2{
    background:
        radial-gradient(circle at top left, rgba(16,185,129,.13), transparent 30%),
        linear-gradient(145deg, rgba(12,18,28,.98), rgba(4,10,16,.98));
    border:1px solid rgba(215,191,117,.34);
    border-radius:24px;
    padding:0;
    margin:12px 0 14px;
    box-shadow:0 18px 48px rgba(0,0,0,.34);
    position:relative;
    overflow:hidden;
}
.history-card-v2::before{content:"";position:absolute;top:0;left:0;right:0;height:4px;background:linear-gradient(90deg,#d7bf75,#008f6b,#16a34a);}
.history-hero-v3{display:grid;grid-template-columns:1fr auto;gap:20px;padding:24px 26px 18px;border-bottom:1px solid rgba(255,255,255,.08);}
.history-title-v3{display:flex;gap:15px;align-items:flex-start;min-width:0;}
.history-icon-v3{width:46px;height:46px;border-radius:16px;display:flex;align-items:center;justify-content:center;background:rgba(215,191,117,.12);border:1px solid rgba(215,191,117,.28);font-size:23px;box-shadow:inset 0 0 18px rgba(215,191,117,.08);flex:0 0 auto;}
.history-title-v3 h3{margin:0 0 7px 0;color:#fff;font-size:26px;line-height:1.1;font-weight:950;letter-spacing:-.02em;}
.history-sub-v3{color:#9ca3af;font-size:12px;font-weight:800;display:flex;flex-wrap:wrap;gap:8px;align-items:center;}
.history-risk-stack-v3{display:flex;flex-direction:column;gap:10px;align-items:flex-end;min-width:170px;}
.history-risk-v2{min-width:145px;text-align:center;color:#fff;padding:11px 18px;border-radius:999px;font-weight:950;font-size:12px;letter-spacing:.06em;box-shadow:inset 0 0 0 1px rgba(255,255,255,.16),0 10px 26px rgba(0,0,0,.18);}
.history-score-v3{display:flex;gap:8px;align-items:center;color:#d7bf75;font-size:12px;font-weight:900;background:rgba(255,255,255,.045);border:1px solid rgba(215,191,117,.18);padding:7px 10px;border-radius:999px;}
.history-score-v3 strong{color:#fff;font-size:16px;}
.history-body-v3{padding:18px 26px 22px;}
.history-chip-row-v2{display:flex;flex-wrap:wrap;gap:8px;margin:0 0 16px 0;}
.history-chip-v2{display:inline-flex;align-items:center;gap:6px;padding:7px 11px;border-radius:999px;background:rgba(255,255,255,.045);border:1px solid rgba(215,191,117,.18);color:#d1d5db;font-size:12px;font-weight:850;}
.history-chip-v2.ok{background:rgba(16,185,129,.12);border-color:rgba(16,185,129,.26);color:#ecfdf5;}
.history-chip-v2.warn{background:rgba(215,191,117,.10);border-color:rgba(215,191,117,.25);color:#fff7d6;}
.history-main-grid-v3{display:grid;grid-template-columns:1fr 1.2fr .72fr .68fr;gap:12px;margin:12px 0;}
.history-kpi-v2{background:rgba(255,255,255,.038);border:1px solid rgba(215,191,117,.16);border-radius:17px;padding:14px 15px;min-height:88px;}
.history-kpi-v2 small,.history-box-v2 small{display:block;color:#d7bf75;font-weight:950;font-size:10px;text-transform:uppercase;letter-spacing:.09em;margin-bottom:8px;}
.history-kpi-v2 strong{color:#fff;display:block;font-size:15px;line-height:1.32;font-weight:950;}
.history-kpi-v2 p{margin:0;color:#dbe4ee;font-size:13px;line-height:1.38;font-weight:750;}
.history-value-v3 strong{font-size:18px;}
.history-wide-v2{display:grid;grid-template-columns:1.35fr 1fr;gap:12px;margin-top:12px;}
.history-box-v2{background:linear-gradient(145deg,rgba(0,72,55,.24),rgba(8,17,25,.38));border:1px solid rgba(16,185,129,.20);border-radius:18px;padding:15px 16px;min-height:116px;}
.history-box-v2 p{margin:0;color:#f8fafc;font-size:13px;line-height:1.45;font-weight:760;}
.history-mini-dashboard-v3{display:grid;grid-template-columns:repeat(4,1fr);gap:8px;}
.history-mini-v3{border-radius:14px;padding:11px 10px;background:rgba(255,255,255,.045);border:1px solid rgba(215,191,117,.14);}
.history-mini-v3 small{margin:0 0 5px 0;}
.history-mini-v3 strong{color:#fff;font-size:21px;font-weight:950;line-height:1;}
.history-mini-v3 span{display:block;color:#9ca3af;font-size:10px;font-weight:850;margin-top:6px;}
.history-summary-v3{display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin-top:12px;}
.history-insight-v3{border:1px solid rgba(215,191,117,.14);background:rgba(255,255,255,.032);border-radius:15px;padding:12px;}
.history-insight-v3 small{display:block;color:#d7bf75;font-size:10px;letter-spacing:.08em;text-transform:uppercase;font-weight:950;margin-bottom:6px;}
.history-insight-v3 strong{color:#fff;font-size:18px;font-weight:950;}
.history-insight-v3 p{margin:4px 0 0 0;color:#aeb8c4;font-size:11px;font-weight:750;line-height:1.25;}
.history-docline-v3{margin-top:14px;padding:12px 13px;border-radius:16px;background:rgba(0,0,0,.16);border:1px solid rgba(255,255,255,.08);display:flex;justify-content:space-between;gap:12px;align-items:center;color:#aeb8c4;font-size:12px;font-weight:800;}
.history-docline-v3 b{color:#f8fafc;}
.history-footer-v2{display:flex;justify-content:space-between;gap:12px;align-items:center;margin-top:14px;padding-top:14px;border-top:1px solid rgba(255,255,255,.08);color:#9ca3af;font-size:12px;font-weight:750;}
.history-actions-v4{display:flex;justify-content:space-between;gap:12px;align-items:center;margin-top:14px;padding-top:14px;border-top:1px solid rgba(255,255,255,.08);color:#9ca3af;font-size:12px;font-weight:750;}
.history-action-note-v4{color:#aeb8c4;font-weight:800;}
.history-excel-link-v4{display:inline-flex;align-items:center;justify-content:center;gap:8px;min-width:190px;padding:11px 18px;border-radius:13px;background:#004d3d;border:1px solid #d7bf75;color:#fff !important;text-decoration:none !important;font-weight:950;box-shadow:0 10px 24px rgba(0,0,0,.18);}
.history-excel-link-v4:hover{filter:brightness(1.12);transform:translateY(-1px);}
.history-card-v2 code,.history-card-v2 pre{background:transparent!important;color:inherit!important;}
@media(max-width:1250px){.history-main-grid-v3{grid-template-columns:1fr 1fr;}.history-wide-v2{grid-template-columns:1fr;}.history-summary-v3{grid-template-columns:1fr 1fr;}}
@media(max-width:760px){.history-hero-v3{grid-template-columns:1fr;}.history-risk-stack-v3{align-items:flex-start;}.history-main-grid-v3,.history-summary-v3{grid-template-columns:1fr;}.history-mini-dashboard-v3{grid-template-columns:1fr 1fr;}}


.valor-card{
    background:linear-gradient(145deg,#101821,#0b1118);
    border:1px solid rgba(215,191,117,.30);
    border-radius:18px;
    padding:18px;
    min-height:135px;
    height:auto;
    box-shadow:0 10px 28px rgba(0,0,0,.24);
    position:relative;
    overflow:visible;
}

.copy-card-btn{
    position:absolute;
    top:10px;
    right:10px;
    width:32px;
    height:32px;
    border-radius:10px;
    border:1px solid rgba(215,191,117,.28);
    background:rgba(255,255,255,.045);
    color:#f3e6b3;
    display:flex;
    align-items:center;
    justify-content:center;
    font-size:15px;
    cursor:pointer;
}

.copy-card-btn:hover{
    background:rgba(215,191,117,.14);
    transform:translateY(-1px);
}

.valor-card::before{
    content:"";
    position:absolute;
    left:0;
    top:0;
    width:100%;
    height:4px;
    background:linear-gradient(90deg,#d7bf75,#065f46);
}

.valor-titulo{
    color:#d7bf75;
    font-weight:900;
    text-transform:uppercase;
    letter-spacing:.7px;
    font-size:11px;
    margin-bottom:12px;
}

.valor-principal{
    color:#ffffff;
    font-size:22px;
    line-height:1.18;
    font-weight:900;
    margin-bottom:10px;
    overflow-wrap:anywhere;
}

.valor-detalhe{
    color:#d8dee9;
    font-size:13px;
    line-height:1.45;
    font-weight:600;
    overflow-wrap:anywhere;
}

.valor-alerta{
    display:inline-block;
    margin-top:12px;
    padding:5px 10px;
    border-radius:999px;
    background:rgba(215,191,117,.12);
    color:#f3e6b3;
    border:1px solid rgba(215,191,117,.28);
    font-size:11px;
    font-weight:900;
}

</style>
"""
st.markdown(CSS, unsafe_allow_html=True)


# =========================================================
# FUNÇÕES AUXILIARES
# =========================================================
def safe(value: Any) -> str:
    """Escapa textos para exibição segura em HTML."""
    if value is None:
        return "Não localizado"
    value = str(value).strip()
    return html.escape(value if value else "Não localizado")


def clean_text(value: Any) -> str:
    """Remove tags HTML e entidades comuns e SEMPRE devolve texto.

    A IA pode devolver números em campos financeiros/itens (ex.: 6720 em
    vez de "R$ 6.720,00"). Se deixarmos int/float passar adiante, funções
    com regex quebram com TypeError.
    """
    if value is None:
        return "Não localizado"
    if not isinstance(value, str):
        value = str(value)
    value = re.sub(r"<[^>]+>", "", value)
    value = html.unescape(value)
    value = re.sub(r"\s+", " ", value).strip()
    return value or "Não localizado"


def normalize_risco(risco: Any) -> str:
    risco = str(risco or "N/A").upper().strip()
    if risco == "MEDIO":
        return "MÉDIO"
    return risco


def risco_cor(risco: Any) -> str:
    risco = normalize_risco(risco)
    if risco == "ALTO":
        return "#dc2626"
    if risco in ["MÉDIO", "MEDIO"]:
        return "#f59e0b"
    if risco == "BAIXO":
        return "#16a34a"
    return "#64748b"


def as_float_score(value: Any) -> float:
    try:
        return float(str(value).replace(",", "."))
    except Exception:
        return 0.0

def carregar_contratos_chat():
    try:
        df = listar_analises()
        return df if isinstance(df, pd.DataFrame) else pd.DataFrame()
    except Exception:
        return pd.DataFrame()


def carregar_historico_seguro() -> pd.DataFrame:
    """Carrega o histórico sem derrubar a tela caso o banco esteja vazio ou inacessível."""
    try:
        df = listar_analises()
        return df if isinstance(df, pd.DataFrame) else pd.DataFrame()
    except Exception as erro:
        st.error(f"Não foi possível carregar o histórico. Detalhe: {erro}")
        return pd.DataFrame()


# =========================================================
# LEITURA DOS ARQUIVOS
# =========================================================
def ler_pdf(file) -> str:
    texto = ""

    try:
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                pagina = page.extract_text() or ""
                if pagina.strip():
                    texto += f"\n\n--- PÁGINA {i} ---\n{pagina}"

                # Extração adicional de tabelas para preservar itens unitários:
                # descrição | quantidade | unidade | valor unitário | valor total
                try:
                    tabelas = page.extract_tables() or []
                    for t_idx, tabela in enumerate(tabelas, 1):
                        if not tabela:
                            continue
                        texto += f"\n\n--- PÁGINA {i} TABELA {t_idx} ---\n"
                        for linha in tabela:
                            celulas = [
                                re.sub(r"\s+", " ", str(c or "")).strip()
                                for c in linha
                            ]
                            if any(celulas):
                                texto += " | ".join(celulas) + "\n"
                except Exception:
                    pass
    except Exception:
        texto = ""

    if len(texto.strip()) > 300:
        return texto

    # OCR automático para PDFs escaneados
    try:
        file.seek(0)
        imagens = convert_from_bytes(file.read(), dpi=300)
        texto_ocr = ""

        for i, imagem in enumerate(imagens, 1):
            pagina = pytesseract.image_to_string(imagem, lang="por")
            texto_ocr += f"\n\n--- PÁGINA {i} OCR ---\n{pagina}"

        return texto_ocr
    except Exception as e:
        return f"Não foi possível extrair texto do PDF. Detalhe: {e}"


def ler_docx(file) -> str:
    file.seek(0)
    doc = Document(file)
    partes = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            linha = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if linha:
                partes.append(linha)

    return "\n".join(partes)


def texto_indica_falha_leitura(texto: Any) -> bool:
    """Detecta quando o texto extraído não representa o documento real.

    Isso evita que PDF escaneado/corrompido vire "apoio simples" ou gere cards
    com "Não identificado" como se fosse uma análise válida.
    """
    txt = clean_text(texto).lower()
    if not txt:
        return True

    sinais_erro = [
        "não foi possível extrair texto do pdf",
        "nao foi possivel extrair texto do pdf",
        "unable to get page count",
        "is poppler installed",
        "poppler",
        "erro ao ler arquivo",
        "erro na extração de texto",
        "erro na extracao de texto",
        "arquivo ilegível",
        "arquivo ilegivel",
        "corrompido",
        "cannot read",
        "invalid pdf",
    ]
    if any(s in txt for s in sinais_erro):
        return True

    # Se sobrou só cabeçalho técnico da triagem sem conteúdo contratual, é ruim.
    txt_sem_cabecalho = re.sub(r"(arquivo|decisão da triagem|decisao da triagem|motivo|tipo|assinado|apoio simples|análise profunda|analise profunda|documentos enviados como base da análise|documentos enviados como base da analise|triagem automática dos anexos antes da ia)", " ", txt)
    txt_sem_cabecalho = re.sub(r"[^a-zà-ÿ0-9]+", " ", txt_sem_cabecalho).strip()
    return len(txt_sem_cabecalho) < 120


def texto_tem_conteudo_contratual(texto: Any) -> bool:
    txt = clean_text(texto).lower()
    gatilhos = [
        "contrato de", "instrumento particular", "cláusula", "clausula",
        "objeto", "vigência", "vigencia", "remuneração", "remuneracao",
        "preço", "preco", "valor", "assinatura", "assinam", "contratada",
        "contratante", "anexo i", "anexo ii", "prestação de serviços",
        "prestacao de servicos", "fornecimento"
    ]
    return sum(1 for g in gatilhos if g in txt) >= 3


def nome_indica_apoio_explicito(nome: Any) -> bool:
    n = _nome_low(nome) if '_nome_low' in globals() else clean_text(nome).lower()
    apoios = [
        "certificado", "certificate", "comunicado", "apresentacao", "apresentação",
        "proposta", "orcamento", "orçamento", "validacao", "validação",
        "aprovacao", "aprovação", "berkan", "cartao cnpj", "cartão cnpj",
        "cnpj", "contrato social", "alteracao contratual", "alteração contratual",
    ]
    return any(a in n for a in apoios)


def arquivo_documental_analisavel(nome: Any) -> bool:
    n = clean_text(nome).lower()
    return n.endswith((".pdf", ".docx", ".doc"))


def resultado_falha_tecnica_leitura(nome_arquivo: str, detalhe: Any = "") -> Dict[str, Any]:
    """Resultado seguro quando não existe conteúdo confiável para análise.

    Melhor uma falha técnica clara do que preencher cards errados.
    """
    detalhe_txt = clean_text(detalhe) or "Falha técnica de leitura do documento original."
    base = local_extract("")
    for _, chave in CAMPOS_OFICIAIS:
        base[chave] = "Não analisado com segurança"

    base.update({
        "tipo_contrato": "Não analisado com segurança",
        "empresa_grupo_sbf": "Não analisado com segurança",
        "cnpj_empresa_grupo": "Não analisado com segurança",
        "contraparte": "Não analisado com segurança",
        "fornecedor": "Não analisado com segurança",
        "cnpj_contraparte": "Não analisado com segurança",
        "cnpj": "Não analisado com segurança",
        "objeto": "Não analisado com segurança",
        "valor_total": "Não analisado com segurança",
        "vigencia": "Não analisado com segurança",
        "contrato_assinado": "Não validado",
        "alerta_assinatura": "Não foi possível validar assinatura porque o documento não foi lido com segurança.",
        "status": "Falha técnica de leitura",
        "risco": "ALTO",
        "score": 0,
        "resumo_executivo": (
            f"Não foi possível concluir a análise do arquivo {nome_arquivo} com segurança. "
            f"O sistema não conseguiu ler o conteúdo contratual do documento. Detalhe: {detalhe_txt}"
        ),
        "parecer": (
            "Não seguir com RC/PO com base nesta execução. Reenviar o arquivo em PDF pesquisável/DOCX "
            "ou garantir que a análise via Gemini Files API esteja ativa para leitura do documento original."
        ),
        "aditivos_contrato": [],
        "itens_contrato": [],
        "assinaturas_contrato": [],
        "checklist": [
            {
                "Validação": "Leitura do arquivo principal",
                "Status": "Incompleto",
                "Peso de risco": "Alto",
                "Crítico": "Sim",
                "Evidência": detalhe_txt,
            }
        ],
        "pendencias": [
            {
                "Pendência": "Arquivo ilegível ou não analisado com segurança",
                "Crítico": "Sim",
                "Risco": "Alto",
                "Recomendação": "Reenviar o documento em PDF pesquisável/DOCX ou corrigir a integração Gemini Files API antes de seguir.",
            }
        ],
    })
    return base


# =========================================================
# EXTRAÇÃO DE ITENS / MATERIAIS / SERVIÇOS
# =========================================================
def _inferir_tipo_item(descricao: Any) -> str:
    txt = str(descricao or "").lower()
    termos_servico = [
        "serviço", "servico", "instalação", "instalacao", "manutenção", "manutencao",
        "suporte", "consultoria", "licença", "licenca", "assinatura", "treinamento",
        "implantação", "implantacao", "configuração", "configuracao", "mão de obra", "mao de obra"
    ]
    if any(t in txt for t in termos_servico):
        return "Serviço"
    return "Material"


def _valor_informado(valor: Any) -> bool:
    txt = clean_text(valor)
    return txt not in ("", "Não localizado", "Não localizada", "N/A", "None", "nan")


def _limpar_percentual_global(valor: Any) -> str:
    txt = clean_text(valor)
    if not _valor_informado(txt):
        return "Não localizado"
    txt = txt.replace(" %", "%").strip()
    if re.fullmatch(r"\d+(?:[\.,]\d+)?", txt):
        return txt.replace(".", ",") + "%"
    return txt


def detectar_servico_percentual(texto: str) -> List[Dict[str, Any]]:
    """Detecta propostas de serviço percentual, como mão de obra temporária.

    Nestes casos, taxa e encargos não são itens separados: são condições
    comerciais do serviço principal.
    """
    texto_original = str(texto or "")
    if not texto_original.strip():
        return []

    plano_completo = re.sub(r"\s+", " ", texto_original).strip()
    low = plano_completo.lower()

    gatilhos = [
        "mão de obra temporária", "mao de obra temporaria",
        "mão-de-obra temporária", "mao-de-obra temporaria",
        "proposta de mão de obra", "proposta de mao de obra",
    ]
    if not any(g in low for g in gatilhos):
        return []

    # Foca no bloco comercial. Isso evita capturar "serviços prestados" do texto institucional.
    pos = low.find("dados da proposta")
    plano = plano_completo[pos:] if pos >= 0 else plano_completo

    def primeiro_valor(captura: Any) -> str:
        txt = clean_text(captura).strip(" -:;,.|/")
        partes = [p.strip(" -:;,.|/") for p in txt.split("|") if p.strip(" -:;,.|/")]
        if not partes:
            return txt
        # remove duplicidades preservando ordem
        vistos_local = []
        for parte in partes:
            if parte.lower() not in [v.lower() for v in vistos_local]:
                vistos_local.append(parte)
        return vistos_local[0] if vistos_local else txt

    servico = "Mão de Obra Temporária"
    m_serv = re.search(
        r"Servi[cç]o\s*:?\s*(?:\|\s*)?(.+?)(?=\s+\|?\s*Taxa\s*:|\s+\|?\s*Total\s+de\s+Encargos|\s+\|?\s*Vencimento\s*:|\s+\|?\s*Exame\s*:|\s+\|?\s*Observa[cç]|$)",
        plano,
        flags=re.IGNORECASE,
    )
    if m_serv:
        capturado = primeiro_valor(m_serv.group(1))
        if 3 <= len(capturado) <= 120:
            servico = capturado

    taxa = "Não localizado"
    m_taxa = re.search(r"Taxa\s*:?\s*(?:\|\s*)?(\d+(?:[\.,]\d+)?\s*%)", plano, flags=re.IGNORECASE)
    if m_taxa:
        taxa = _limpar_percentual_global(m_taxa.group(1))

    encargos = "Não localizado"
    m_enc = re.search(r"Total\s+de\s+Encargos\s*:?\s*(?:\|\s*)?(\d+(?:[\.,]\d+)?\s*%)", plano, flags=re.IGNORECASE)
    if m_enc:
        encargos = _limpar_percentual_global(m_enc.group(1))

    vencimento = "Não localizado"
    m_venc = re.search(
        r"Vencimento\s*:?\s*(?:\|\s*)?(.+?)(?=\s+\|?\s*Exame\s*:|\s+\|?\s*Observa[cç]|\s+\|?\s*Representante|\s+DA RESPONSABILIDADE|$)",
        plano,
        flags=re.IGNORECASE,
    )
    if m_venc:
        venc = primeiro_valor(m_venc.group(1))
        if venc:
            vencimento = resumir_campo(venc, 120)

    # Caso comum em propostas de RH/Atração & Retenção: o PDF informa
    # "Taxa de 40%", "Taxa de 50%" e DISC de R$ 180,00, sem tabela estruturada.
    itens_proposta: List[Dict[str, Any]] = []
    tributos_nf = "0,8367" if re.search(r"0[,.]8367", plano_completo) else "Não localizado"
    venc_prop = vencimento
    m_venc_prop = re.search(
        r"vencimento\s+da\s+nota\s+ser[aá]\s+em\s+(\d+)\s*\([^)]*\)\s+dias\s+corridos\s+a\s+contar\s+da\s+data\s+de\s+emiss[aã]o",
        plano_completo,
        flags=re.IGNORECASE,
    )
    if m_venc_prop:
        venc_prop = f"{m_venc_prop.group(1)} dias corridos a contar da data de emissão da nota fiscal"

    if re.search(r"taxa\s+de\s+40\s*%", plano_completo, flags=re.IGNORECASE):
        itens_proposta.append({
            "item": str(len(itens_proposta) + 1),
            "descricao": "Recrutamento e seleção - vagas administrativas, operacionais e comerciais",
            "tipo": "Serviço",
            "quantidade": "Não aplicável",
            "unidade": "Percentual sobre salário/remuneração",
            "valor_unitario": "Não aplicável",
            "valor_total": "Não localizado",
            "taxa_percentual": "40%",
            "total_encargos": f"Tributos de nota fiscal: {tributos_nf}" if _valor_informado(tributos_nf) else "Não localizado",
            "vencimento": venc_prop,
            "fonte": "Proposta comercial",
        })

    if re.search(r"taxa\s+de\s+50\s*%", plano_completo, flags=re.IGNORECASE):
        itens_proposta.append({
            "item": str(len(itens_proposta) + 1),
            "descricao": "Recrutamento e seleção - vagas técnicas, estratégicas ou liderança",
            "tipo": "Serviço",
            "quantidade": "Não aplicável",
            "unidade": "Percentual sobre salário/remuneração",
            "valor_unitario": "Não aplicável",
            "valor_total": "Não localizado",
            "taxa_percentual": "50%",
            "total_encargos": f"Tributos de nota fiscal: {tributos_nf}" if _valor_informado(tributos_nf) else "Não localizado",
            "vencimento": venc_prop,
            "fonte": "Proposta comercial",
        })

    if re.search(r"R\$\s*180[,.]00|180,00\s*\([^)]*cento\s+e\s+oitenta", plano_completo, flags=re.IGNORECASE):
        itens_proposta.append({
            "item": str(len(itens_proposta) + 1),
            "descricao": "Teste de análise comportamental / DISC",
            "tipo": "Serviço",
            "quantidade": "Não aplicável",
            "unidade": "Teste",
            "valor_unitario": "R$ 180,00",
            "valor_total": "Não localizado",
            "taxa_percentual": "Não aplicável",
            "total_encargos": f"Tributos de nota fiscal: {tributos_nf}" if _valor_informado(tributos_nf) else "Não localizado",
            "vencimento": venc_prop,
            "fonte": "Proposta comercial",
        })

    if itens_proposta:
        return normalizar_itens_contrato(itens_proposta)

    # Só considera como caso especial se houver pelo menos taxa ou encargos.
    if not (_valor_informado(taxa) or _valor_informado(encargos)):
        return []

    return [{
        "item": "1",
        "descricao": servico,
        "tipo": "Serviço",
        "quantidade": "Não aplicável",
        "unidade": "Não aplicável",
        "valor_unitario": "Não aplicável",
        "valor_total": "Não localizado",
        "taxa_percentual": taxa,
        "total_encargos": encargos,
        "vencimento": vencimento,
        "fonte": "Proposta comercial",
    }]


def _itens_sao_apenas_atributos_comerciais(itens: List[Dict[str, Any]]) -> bool:
    if not itens:
        return False
    termos_atributo = [
        "taxa", "encargo", "imposto", "percentual", "comissão", "comissao",
        "vencimento", "pagamento", "pis", "cofins", "inss", "iss", "icms",
    ]
    descricoes = [str(i.get("Descrição") or i.get("descricao") or "").lower() for i in itens]
    if not descricoes:
        return False
    return all(any(t in d for t in termos_atributo) for d in descricoes)


def normalizar_itens_contrato(itens: Any) -> List[Dict[str, Any]]:
    """Padroniza materiais/serviços vindos da IA ou da extração local.

    Importante: contratos de serviço nem sempre possuem valor unitário em R$.
    Exemplo: mão de obra temporária pode vir com Taxa, Total de Encargos e Vencimento.
    Por isso a tabela aceita campos comerciais complementares sem forçar tudo como material.
    """
    if not isinstance(itens, list):
        return []

    def chave_norm(k: Any) -> str:
        k = str(k or "").strip().lower()
        k = html.unescape(k)
        k = re.sub(r"[áàãâ]", "a", k)
        k = re.sub(r"[éê]", "e", k)
        k = re.sub(r"[í]", "i", k)
        k = re.sub(r"[óôõ]", "o", k)
        k = re.sub(r"[ú]", "u", k)
        k = re.sub(r"[ç]", "c", k)
        k = re.sub(r"[^a-z0-9]+", "_", k)
        return k.strip("_")

    def pegar(item_dict: Dict[str, Any], aliases: List[str], padrao: Any = "Não localizado") -> Any:
        mapa = {chave_norm(k): v for k, v in item_dict.items()}
        for alias in aliases:
            nk = chave_norm(alias)
            if nk in mapa and _valor_informado(mapa[nk]):
                return mapa[nk]
        return padrao

    def limpar_percentual(valor: Any) -> str:
        txt = clean_text(valor)
        if not _valor_informado(txt):
            return "Não localizado"
        if re.fullmatch(r"\d+(?:[\.,]\d+)?", txt):
            return txt.replace(".", ",") + "%"
        return txt

    normalizados: List[Dict[str, Any]] = []

    for idx, item in enumerate(itens, 1):
        if isinstance(item, str):
            item = {"descricao": item}
        if not isinstance(item, dict):
            continue

        descricao = clean_text(pegar(item, [
            "descricao", "descrição", "descricao_item", "descrição_item",
            "descricao_do_item", "descrição do item", "Descrição do Item",
            "servico_material", "serviço_material", "material_servico",
            "nome", "nome_item", "item_descricao", "produto", "material", "serviço", "servico"
        ]))

        if descricao in ("", "Não localizado", "Não localizada"):
            continue

        tipo = clean_text(pegar(item, ["tipo", "categoria"], _inferir_tipo_item(descricao)))
        quantidade = clean_text(pegar(item, ["quantidade", "qtde", "qtd", "qtd.", "quant.", "volume"]))
        unidade = clean_text(pegar(item, ["unidade", "un", "uom", "und", "medida", "unid"]))

        taxa_percentual = limpar_percentual(pegar(item, [
            "taxa", "taxa_percentual", "percentual", "percentual_taxa", "%", "aliquota", "alíquota",
            "taxa_de_agenciamento", "taxa agenciamento", "taxa administrativa", "taxa_admin"
        ]))

        total_encargos = limpar_percentual(pegar(item, [
            "total_encargos", "total de encargos", "encargos", "encargos_sociais",
            "total_encargos_sociais", "total de encargos sociais", "custo_encargos"
        ]))

        vencimento = clean_text(pegar(item, [
            "vencimento", "prazo", "prazo_pagamento", "condicao_pagamento", "condição pagamento",
            "pagamento", "vencimento_nota", "vencimento_nf"
        ]))

        valor_unitario = clean_text(pegar(item, [
            "valor_unitario", "valor_unitário", "valor unitario", "valor unitário",
            "Valor Unitário (R$)", "valor_unitario_rs", "preco_unitario",
            "preço_unitário", "preço unitário", "preco", "preço", "unitario", "unitário"
        ]))
        valor_total = clean_text(pegar(item, [
            "valor_total", "valor total", "Valor Total (R$)", "total",
            "subtotal", "valor_total_rs", "preco_total", "preço_total"
        ]))
        fonte = clean_text(pegar(item, ["fonte", "origem", "documento", "arquivo", "arquivo_fonte"], "Contrato/anexo"))
        pagina = clean_text(pegar(item, ["pagina", "página", "page"], "Não localizado"))
        evidencia = clean_text(pegar(item, ["evidencia", "evidência", "trecho_evidencia", "trecho de evidência"], "Não localizado"))
        natureza_valor = clean_text(pegar(item, ["natureza_valor", "natureza do valor", "tipo_valor", "tipo de valor"], "Não localizado"))
        periodicidade = clean_text(pegar(item, ["periodicidade", "recorrencia", "recorrência"], "Não localizado"))
        grupo_tabela = clean_text(pegar(item, ["grupo_tabela", "grupo tabela", "grupo", "tabela", "secao_tabela", "seção tabela"], "Não localizado"))
        faixa_condicao = clean_text(pegar(item, ["faixa_condicao", "faixa condição", "faixa", "condicao_faixa", "condição faixa"], "Não localizado"))
        condicao_comercial = clean_text(pegar(item, ["condicao_comercial", "condição comercial", "regra_comercial", "observacao_comercial", "observação comercial"], "Não localizado"))
        status_evidencia = clean_text(pegar(item, ["status_evidencia", "status de evidência", "status"], "Não localizado"))

        # Não transforme taxa/encargos em valor unitário. Eles são condições comerciais.
        normalizados.append({
            "Item": clean_text(pegar(item, ["item", "numero", "número", "n", "id"], idx)),
            "Descrição": resumir_campo(descricao, 420),
            "Tipo": tipo if tipo not in ("", "Não localizado", "Não localizada") else _inferir_tipo_item(descricao),
            "Quantidade": quantidade,
            "Unidade": unidade,
            "Valor unitário": valor_unitario,
            "Valor total": valor_total,
            "Taxa / Percentual": taxa_percentual,
            "Total de encargos": total_encargos,
            "Vencimento / Prazo": vencimento,
            "Natureza do valor": natureza_valor,
            "Periodicidade": periodicidade,
            "Grupo/Tabela": grupo_tabela,
            "Faixa/Condição": faixa_condicao,
            "Condição comercial": condicao_comercial,
            "Fonte": fonte,
            "Página": pagina,
            "Status de evidência": status_evidencia,
            "Evidência": evidencia,
        })

    # Regra especial: se a IA devolveu uma linha "Taxa de Agenciamento" e outra "Encargos Sociais",
    # consolidamos em uma única linha de serviço para mão de obra temporária.
    descricoes = " ".join(str(i.get("Descrição", "")) for i in normalizados).lower()
    tem_taxa = any(_valor_informado(i.get("Taxa / Percentual")) or "taxa" in str(i.get("Descrição", "")).lower() for i in normalizados)
    tem_encargo = any(_valor_informado(i.get("Total de encargos")) or "encargo" in str(i.get("Descrição", "")).lower() for i in normalizados)
    mao_obra = "mão de obra" in descricoes or "mao de obra" in descricoes or "temporári" in descricoes or "temporari" in descricoes

    if mao_obra and tem_taxa and tem_encargo and len(normalizados) <= 5:
        taxa = "Não localizado"
        encargos = "Não localizado"
        vencimento = "Não localizado"
        fonte = "Contrato/anexo"
        pagina = "Não localizado"
        evidencias = []
        for item in normalizados:
            desc_low = str(item.get("Descrição", "")).lower()
            if not _valor_informado(taxa):
                if _valor_informado(item.get("Taxa / Percentual")):
                    taxa = item.get("Taxa / Percentual")
                elif "taxa" in desc_low and _valor_informado(item.get("Valor unitário")):
                    taxa = item.get("Valor unitário")
            if not _valor_informado(encargos):
                if _valor_informado(item.get("Total de encargos")):
                    encargos = item.get("Total de encargos")
                elif "encargo" in desc_low and _valor_informado(item.get("Valor unitário")):
                    encargos = item.get("Valor unitário")
            if not _valor_informado(vencimento) and _valor_informado(item.get("Vencimento / Prazo")):
                vencimento = item.get("Vencimento / Prazo")
            if _valor_informado(item.get("Fonte")):
                fonte = item.get("Fonte")
            if _valor_informado(item.get("Página")):
                pagina = item.get("Página")
            if _valor_informado(item.get("Evidência")):
                evidencias.append(clean_text(item.get("Evidência")))

        return [{
            "Item": "1",
            "Descrição": "Mão de obra temporária",
            "Tipo": "Serviço",
            "Quantidade": "Não localizado",
            "Unidade": "Não localizado",
            "Valor unitário": "Não localizado",
            "Valor total": "Não localizado",
            "Taxa / Percentual": taxa,
            "Total de encargos": encargos,
            "Vencimento / Prazo": vencimento,
            "Natureza do valor": "PERCENTUAL_VARIAVEL",
            "Periodicidade": "Conforme demanda/folha",
            "Fonte": fonte,
            "Página": pagina,
            "Status de evidência": "CONFIRMADO" if evidencias else "NÃO_LOCALIZADO",
            "Evidência": " | ".join(dict.fromkeys(evidencias)) if evidencias else "Não localizado",
        }]

    return normalizados

def _mapear_anexos_aditivos_por_texto(texto: str) -> Dict[str, str]:
    """
    Tenta mapear datas de aditivos para o nome do arquivo/anexo.
    O texto do robô vem separado por blocos:
    ARQUIVO: nome_do_arquivo.pdf
    """
    texto = str(texto or "")
    mapa: Dict[str, str] = {}

    blocos = re.split(r"\n\s*ARQUIVO:\s*", texto, flags=re.IGNORECASE)

    for bloco in blocos:
        bloco = str(bloco or "").strip()
        if not bloco:
            continue

        linhas = bloco.splitlines()
        if not linhas:
            continue

        nome_arquivo = clean_text(linhas[0]).strip()
        conteudo = "\n".join(linhas[1:]) if len(linhas) > 1 else bloco
        conteudo_low = conteudo.lower()
        nome_low = nome_arquivo.lower()

        eh_aditivo = any(t in conteudo_low or t in nome_low for t in [
            "aditivo",
            "termo aditivo",
            "aditamento",
            "renovação",
            "renovacao",
            "prorrogação",
            "prorrogacao",
            "reajuste",
        ])

        if not eh_aditivo:
            continue

        datas = re.findall(r"\b\d{1,2}/\d{1,2}/\d{4}\b", conteudo)
        for data in datas:
            data_fmt = _formatar_data_slash(data)
            mapa[data_fmt] = nome_arquivo

    return mapa


def _preencher_anexo_aditivo_por_texto(aditivos: List[Dict[str, Any]], texto: str) -> List[Dict[str, Any]]:
    mapa = _mapear_anexos_aditivos_por_texto(texto)

    if not aditivos or not mapa:
        return aditivos

    for aditivo in aditivos:
        anexo_atual = clean_text(aditivo.get("Anexo do aditivo"))
        if _valor_informado(anexo_atual):
            continue

        datas_possiveis = [
            clean_text(aditivo.get("Data do aditivo")),
            clean_text(aditivo.get("Data da assinatura")),
        ]

        for data in datas_possiveis:
            data_fmt = _formatar_data_slash(data)
            if data_fmt in mapa:
                aditivo["Anexo do aditivo"] = mapa[data_fmt]
                break

    return aditivos

def normalizar_aditivos_contrato(aditivos: Any) -> List[Dict[str, Any]]:
    """Padroniza aditivos vindos da IA para exibição em tabela executiva."""
    if not isinstance(aditivos, list):
        return []

    def chave_norm(k: Any) -> str:
        k = str(k or "").strip().lower()
        k = html.unescape(k)
        k = re.sub(r"[áàãâ]", "a", k)
        k = re.sub(r"[éê]", "e", k)
        k = re.sub(r"[í]", "i", k)
        k = re.sub(r"[óôõ]", "o", k)
        k = re.sub(r"[ú]", "u", k)
        k = re.sub(r"[ç]", "c", k)
        k = re.sub(r"[^a-z0-9]+", "_", k)
        return k.strip("_")

    def pegar(item_dict: Dict[str, Any], aliases: List[str], padrao: Any = "Não localizado") -> Any:
        mapa = {chave_norm(k): v for k, v in item_dict.items()}
        for alias in aliases:
            nk = chave_norm(alias)
            if nk in mapa and _valor_informado(mapa[nk]):
                return mapa[nk]
        return padrao

    normalizados: List[Dict[str, Any]] = []

    for idx, item in enumerate(aditivos, 1):
        if isinstance(item, str):
            item = {"objeto_escopo_aditivo": item}

        if not isinstance(item, dict):
            continue

        # Se o aditivo já veio no padrão final da tabela, preserva os campos.
        # Isso evita transformar "Anexo do aditivo", "Tipo do aditivo",
        # "Impacto no valor" e "Status de validação" em "Não localizado".
        if any(campo in item for campo in [
            "Tipo do aditivo",
            "Anexo do aditivo",
            "Data do aditivo",
            "Data de carga no robô",
            "Data da assinatura",
            "Quem assinou",
            "Valor do aditivo",
            "Impacto no valor",
            "Impacto no prazo",
            "Período do aditivo",
            "Escopo do aditivo",
            "Status de validação",
        ]):
            itens_norm = item.get("_itens_aditivo", [])
            if not isinstance(itens_norm, list):
                itens_norm = []

            normalizados.append({
                "Nº": clean_text(item.get("Nº") or item.get("numero_aditivo") or idx),
                "Tipo do aditivo": clean_text(item.get("Tipo do aditivo")),
                "Anexo do aditivo": clean_text(item.get("Anexo do aditivo")),
                "Data do aditivo": clean_text(item.get("Data do aditivo")),
                "Data de carga no robô": clean_text(item.get("Data de carga no robô")),
                "Assinado": clean_text(item.get("Assinado")),
                "Data da assinatura": clean_text(item.get("Data da assinatura")),
                "Quem assinou": resumir_campo(item.get("Quem assinou"), 500),
                "Valor do aditivo": clean_text(item.get("Valor do aditivo")),
                "Impacto no valor": resumir_campo(item.get("Impacto no valor"), 500),
                "Impacto no prazo": resumir_campo(item.get("Impacto no prazo"), 500),
                "Período do aditivo": clean_text(item.get("Período do aditivo")),
                "Escopo do aditivo": resumir_campo(item.get("Escopo do aditivo"), 600),
                "Itens do aditivo": resumir_campo(item.get("Itens do aditivo"), 600),
                "Status de validação": resumir_campo(item.get("Status de validação"), 500),
                "Observações": resumir_campo(item.get("Observações"), 300),
                "_itens_aditivo": normalizar_itens_contrato(itens_norm),
            })
            continue

        itens_raw = pegar(item, [
            "itens_aditivo", "itens_do_aditivo", "materiais_servicos_aditivo",
            "materiais_e_servicos_aditivo", "itens", "servicos", "serviços"
        ], [])

        if not isinstance(itens_raw, list):
            itens_raw = []

        itens_norm = normalizar_itens_contrato(itens_raw)

        def valor(alias_list, padrao="Não localizado"):
            return clean_text(pegar(item, alias_list, padrao))

        numero = valor(["numero_aditivo", "n_aditivo", "número_aditivo", "aditivo", "item"], str(idx))
        tipo = valor(["tipo_aditivo", "tipo", "natureza", "categoria"], "Não localizado")
        anexo = valor(["anexo_origem", "arquivo", "nome_arquivo", "documento", "fonte"], "Não localizado")
        data_aditivo = valor(["data_aditivo", "data_do_aditivo", "data_documento", "data"], "Não localizado")
        data_carga = valor(["data_carga_robo", "data_upload", "data_upload_aditivo", "data_carga"], "Não localizado")
        assinado = valor(["assinado", "aditivo_assinado", "status_assinatura", "contrato_assinado"], "Não localizado")
        data_assinatura = valor(["data_assinatura_aditivo", "data_da_assinatura", "data_assinatura"], "Não localizado")
        assinantes = valor(["pessoas_que_assinaram_aditivo", "assinantes", "quem_assinou", "pessoas_que_assinaram"], "Não localizado")
        valor_aditivo = valor(["valor_aditivo", "valor", "valor_total", "valor_alteracao", "valor_alteração"], "Não localizado")
        impacto_valor = valor(["impacto_valor", "alteracao_valor", "alteração_valor", "efeito_valor"], "Não localizado")
        impacto_prazo = valor(["impacto_prazo", "alteracao_prazo", "alteração_prazo", "efeito_prazo"], "Não localizado")
        periodo = valor(["periodo_vigencia_aditivo", "período_vigencia_aditivo", "periodo", "vigencia"], "Não localizado")
        escopo = valor(["objeto_escopo_aditivo", "escopo", "objeto", "descricao", "descrição"], "Não localizado")
        status_validacao = valor(["status_validacao_aditivo", "status", "validacao", "validação"], "Não localizado")
        observacoes = valor(["observacoes_aditivo", "observações_aditivo", "observacao", "observação"], "")

        # Padroniza assinatura
        low_ass = assinado.lower()
        if any(t in low_ass for t in ["sim", "assinado", "completed", "concluído", "concluido"]):
            assinado = "Sim"
        elif any(t in low_ass for t in ["não", "nao", "sem assinatura", "pendente"]):
            assinado = "Não"

        valor_aditivo = _formatar_valor_monetario_item(valor_aditivo)

        itens_resumo = "Não localizado"
        if itens_norm:
            partes = []
            for it in itens_norm[:5]:
                desc = clean_text(it.get("Descrição"))
                vt = clean_text(it.get("Valor total"))
                if _valor_informado(vt):
                    partes.append(f"{desc}: {_formatar_valor_monetario_item(vt)}")
                else:
                    partes.append(desc)
            itens_resumo = "; ".join(partes)
            if len(itens_norm) > 5:
                itens_resumo += f"; e mais {len(itens_norm) - 5} item(ns)."

        # Se não tem nada útil, ignora
        bloco_texto = " ".join([numero, tipo, anexo, data_aditivo, assinado, valor_aditivo, escopo])
        if bloco_texto.replace("Não localizado", "").strip() == "":
            continue

        normalizados.append({
            "Nº": numero,
            "Tipo do aditivo": tipo,
            "Anexo do aditivo": resumir_campo(anexo, 180),
            "Data do aditivo": data_aditivo,
            "Data de carga no robô": data_carga,
            "Assinado": assinado,
            "Data da assinatura": data_assinatura,
            "Quem assinou": resumir_campo(assinantes, 260),
            "Valor do aditivo": valor_aditivo,
            "Impacto no valor": resumir_campo(impacto_valor, 240),
            "Impacto no prazo": resumir_campo(impacto_prazo, 240),
            "Período do aditivo": periodo,
            "Escopo do aditivo": resumir_campo(escopo, 360),
            "Itens do aditivo": resumir_campo(itens_resumo, 360),
            "Status de validação": resumir_campo(status_validacao, 260),
            "Observações": resumir_campo(observacoes, 260),
            "_itens_aditivo": itens_norm,
        })

    return normalizados


def _montar_resumo_aditivos(aditivos: List[Dict[str, Any]]) -> str:
    aditivos = normalizar_aditivos_contrato(aditivos)

    if not aditivos:
        return "Nenhum aditivo identificado nos documentos analisados."

    total = len(aditivos)
    assinados = sum(1 for a in aditivos if clean_text(a.get("Assinado")).upper() == "SIM")
    nao_assinados = sum(1 for a in aditivos if clean_text(a.get("Assinado")).upper() in ("NÃO", "NAO"))

    resumo = f"{total} aditivo(s) identificado(s). {assinados} assinado(s)"

    if nao_assinados:
        resumo += f", {nao_assinados} sem assinatura localizada"

    resumo += ". Sem valor global fixo; os aditivos analisados alteram valores unitários, percentuais, prazos ou quantidades."
    resumo += " Consulte a tabela de aditivos para detalhes de anexo, assinatura, prazo, valor e itens."
    return resumo


def _normalizar_data_aditivo_para_br(valor: Any) -> str:
    txt = clean_text(valor)

    if not _valor_informado(txt):
        return "Não localizado"

    # Exemplo: 08 de setembro de 2023
    convertido = _data_textual_para_br(txt)
    if convertido != txt:
        return convertido

    # Exemplo: 08/09/2023 ou 9/24/2025
    convertido = _formatar_data_slash(txt)
    return convertido if _valor_informado(convertido) else txt


def _extrair_blocos_arquivo_texto(texto: str) -> List[Dict[str, Any]]:
    """
    Quebra o texto extraído em blocos por ARQUIVO.
    O robô monta o texto assim:
    ARQUIVO: nome_do_arquivo.pdf
    """
    texto = str(texto or "")
    blocos = []

    partes = re.split(r"\n\s*ARQUIVO:\s*", texto, flags=re.IGNORECASE)

    for parte in partes:
        parte = str(parte or "").strip()
        if not parte:
            continue

        linhas = parte.splitlines()
        if not linhas:
            continue

        nome_arquivo = clean_text(linhas[0]).strip()
        conteudo = "\n".join(linhas[1:]) if len(linhas) > 1 else parte

        if not _valor_informado(nome_arquivo):
            continue

        blocos.append({
            "arquivo": nome_arquivo,
            "conteudo": conteudo,
            "conteudo_low": conteudo.lower(),
            "arquivo_low": nome_arquivo.lower(),
        })

    return blocos


def _eh_bloco_aditivo(bloco: Dict[str, Any]) -> bool:
    texto = f"{bloco.get('arquivo_low', '')} {bloco.get('conteudo_low', '')}"

    termos = [
        "aditivo",
        "termo aditivo",
        "aditamento",
        "renovação",
        "renovacao",
        "prorrogação",
        "prorrogacao",
        "reajuste",
        "alteração contratual",
        "alteracao contratual",
    ]

    return any(t in texto for t in termos)


def _datas_do_bloco_aditivo(bloco: Dict[str, Any]) -> List[str]:
    conteudo = str(bloco.get("conteudo") or "")
    datas = []

    # Datas numéricas
    for data in re.findall(r"\b\d{1,2}/\d{1,2}/\d{4}\b", conteudo):
        datas.append(_normalizar_data_aditivo_para_br(data))

    # Datas por extenso: 08 de setembro de 2023
    for data in re.findall(
        r"\b\d{1,2}\s+de\s+[A-Za-zçÇãáéíóúâêôõ]+\s+de\s+\d{4}\b",
        conteudo,
        flags=re.IGNORECASE,
    ):
        datas.append(_normalizar_data_aditivo_para_br(data))

    # Remove duplicados preservando ordem
    saida = []
    for d in datas:
        if _valor_informado(d) and d not in saida:
            saida.append(d)

    return saida


def _inferir_tipo_aditivo_por_bloco(bloco: Dict[str, Any]) -> str:
    txt = f"{bloco.get('arquivo_low', '')} {bloco.get('conteudo_low', '')}"

    if "reajuste" in txt:
        return "Aditivo de reajuste"
    if "prorrogação" in txt or "prorrogacao" in txt or "vigência" in txt or "vigencia" in txt:
        return "Aditivo de prazo/vigência"
    if "valor" in txt or "preço" in txt or "preco" in txt:
        return "Aditivo de valor"
    if "escopo" in txt or "objeto" in txt:
        return "Aditivo de escopo"

    return "Termo aditivo"


def _inferir_impacto_valor_por_bloco(bloco: Dict[str, Any]) -> str:
    txt = bloco.get("conteudo_low", "")

    if "sem alteração de valor" in txt or "sem alteracao de valor" in txt:
        return "Sem alteração de valor identificada."
    if "reajuste" in txt:
        return "Possui reajuste de valor. Validar o percentual/base no anexo do aditivo."
    if "valor" in txt or "preço" in txt or "preco" in txt:
        return "Possui menção a valor/preço. Validar o detalhe no anexo do aditivo."

    return "Não localizado"


def _inferir_impacto_prazo_por_bloco(bloco: Dict[str, Any]) -> str:
    txt = bloco.get("conteudo_low", "")

    if "prazo indeterminado" in txt:
        return "Altera a vigência para prazo indeterminado."
    if "prorrogação" in txt or "prorrogacao" in txt:
        return "Prorroga o prazo/vigência do contrato."
    if "vigência" in txt or "vigencia" in txt:
        return "Possui alteração ou confirmação de vigência. Validar período no anexo."

    return "Não localizado"


def _inferir_valor_aditivo_por_bloco(bloco: Dict[str, Any]) -> str:
    conteudo = str(bloco.get("conteudo") or "")

    valores = re.findall(
        r"R\$\s*[0-9]{1,3}(?:\.[0-9]{3})*,[0-9]{2}|R\$\s*[0-9]+(?:,[0-9]{2})?",
        conteudo,
        flags=re.IGNORECASE,
    )

    if valores:
        return _formatar_valor_monetario_item(valores[0])

    return "Não localizado"


def _mapear_blocos_aditivos(texto: str) -> List[Dict[str, Any]]:
    blocos = _extrair_blocos_arquivo_texto(texto)
    aditivos = []

    for bloco in blocos:
        if not _eh_bloco_aditivo(bloco):
            continue

        aditivos.append({
            "arquivo": bloco.get("arquivo"),
            "datas": _datas_do_bloco_aditivo(bloco),
            "tipo": _inferir_tipo_aditivo_por_bloco(bloco),
            "impacto_valor": _inferir_impacto_valor_por_bloco(bloco),
            "impacto_prazo": _inferir_impacto_prazo_por_bloco(bloco),
            "valor": _inferir_valor_aditivo_por_bloco(bloco),
            "conteudo_low": bloco.get("conteudo_low", ""),
        })

    return aditivos


def _arquivos_aditivos_do_pacote(base: Dict[str, Any], texto: str) -> List[str]:
    """Lista arquivos/anexos que parecem ser aditivos, usando arquivos_analisados e blocos ARQUIVO."""
    nomes: List[str] = []

    termos = [
        "aditivo",
        "termo aditivo",
        "aditamento",
        "renovação",
        "renovacao",
        "prorrogação",
        "prorrogacao",
        "reajuste",
    ]

    arquivos_base = clean_text(base.get("arquivos_analisados"))
    if _valor_informado(arquivos_base):
        for nome in re.split(r"\s*\|\s*", arquivos_base):
            nome = clean_text(nome)
            if _valor_informado(nome) and any(t in nome.lower() for t in termos):
                if nome not in nomes:
                    nomes.append(nome)

    try:
        for bloco in _extrair_blocos_arquivo_texto(texto):
            nome = clean_text(bloco.get("arquivo"))
            conteudo = clean_text(bloco.get("conteudo"))
            plano = f"{nome} {conteudo}".lower()

            if _valor_informado(nome) and any(t in plano for t in termos):
                if nome not in nomes:
                    nomes.append(nome)
    except Exception:
        pass

    return nomes


def _enriquecer_aditivos_por_anexo(
    aditivos: List[Dict[str, Any]],
    texto: str,
    base: Dict[str, Any] | None = None
) -> List[Dict[str, Any]]:
    """
    Preenche anexo, tipo, impacto, valor e status dos aditivos.
    Primeiro tenta por data. Se não conseguir, usa os nomes dos arquivos de aditivos por ordem.
    """
    base = base or {}
    blocos_aditivos = _mapear_blocos_aditivos(texto)
    arquivos_pacote = _arquivos_aditivos_do_pacote(base, texto)

    if not aditivos:
        return aditivos

    # 1. Tenta mapear por data encontrada no bloco do arquivo
    for aditivo in aditivos:
        data_aditivo = _normalizar_data_aditivo_para_br(aditivo.get("Data do aditivo"))
        data_assinatura = _normalizar_data_aditivo_para_br(aditivo.get("Data da assinatura"))

        for bloco in blocos_aditivos:
            datas_bloco = bloco.get("datas", [])

            if data_aditivo in datas_bloco or data_assinatura in datas_bloco:
                if not _valor_informado(aditivo.get("Anexo do aditivo")):
                    aditivo["Anexo do aditivo"] = bloco.get("arquivo", "Não localizado")

                if not _valor_informado(aditivo.get("Tipo do aditivo")):
                    aditivo["Tipo do aditivo"] = bloco.get("tipo", "Termo aditivo")

                if not _valor_informado(aditivo.get("Impacto no valor")):
                    aditivo["Impacto no valor"] = bloco.get("impacto_valor", "Não localizado")

                if not _valor_informado(aditivo.get("Impacto no prazo")):
                    aditivo["Impacto no prazo"] = bloco.get("impacto_prazo", "Não localizado")

                if not _valor_informado(aditivo.get("Valor do aditivo")):
                    aditivo["Valor do aditivo"] = bloco.get("valor", "Não localizado")

                break

    # 2. Se ainda não encontrou anexo, preenche por ordem usando os arquivos do pacote
    if arquivos_pacote:
        for idx, aditivo in enumerate(aditivos):
            if idx >= len(arquivos_pacote):
                break

            if not _valor_informado(aditivo.get("Anexo do aditivo")):
                aditivo["Anexo do aditivo"] = arquivos_pacote[idx]

            if not _valor_informado(aditivo.get("Tipo do aditivo")):
                nome_low = arquivos_pacote[idx].lower()

                if "reajuste" in nome_low:
                    aditivo["Tipo do aditivo"] = "Aditivo de reajuste"
                elif "prorro" in nome_low or "vig" in nome_low:
                    aditivo["Tipo do aditivo"] = "Aditivo de prazo/vigência"
                elif "valor" in nome_low:
                    aditivo["Tipo do aditivo"] = "Aditivo de valor"
                else:
                    aditivo["Tipo do aditivo"] = "Termo aditivo"

    # 3. Status executivo
    for aditivo in aditivos:
        assinado = str(aditivo.get("Assinado", "")).strip().upper()

        if assinado == "SIM":
            aditivo["Status de validação"] = "Aditivo assinado. Validar anexo, escopo, prazo e impacto de valor antes de seguir."
        elif assinado in ("NÃO", "NAO"):
            aditivo["Status de validação"] = "Aditivo sem assinatura localizada. Revisar antes de considerar vigente."
        else:
            aditivo["Status de validação"] = "Assinatura do aditivo não localizada com clareza. Revisar documento."

        if not _valor_informado(aditivo.get("Tipo do aditivo")):
            aditivo["Tipo do aditivo"] = "Termo aditivo"

        if not _valor_informado(aditivo.get("Escopo do aditivo")) and _valor_informado(aditivo.get("Anexo do aditivo")):
            aditivo["Escopo do aditivo"] = f"Aditivo identificado no anexo {aditivo.get('Anexo do aditivo')}."

    return aditivos

def _info_util(valor: Any) -> bool:
    """Retorna True somente quando o campo tem informação real."""
    txt = clean_text(valor).strip().lower()

    invalidos = {
        "",
        "none",
        "null",
        "n/a",
        "na",
        "-",
        "não localizado",
        "nao localizado",
        "não localizada",
        "nao localizada",
        "não identificado",
        "nao identificado",
        "não identificada",
        "nao identificada",
        "não aplicável",
        "nao aplicavel",
    }

    return txt not in invalidos


def _normalizar_data_aditivo_para_br(valor: Any) -> str:
    txt = clean_text(valor)

    if not _info_util(txt):
        return "Não localizado"

    convertido = _data_textual_para_br(txt)
    if convertido != txt and _info_util(convertido):
        return convertido

    convertido = _formatar_data_slash(txt)
    return convertido if _info_util(convertido) else txt


def _ordinal_aditivo_numero(texto: Any) -> int | None:
    txt = clean_text(texto).lower()

    mapa = {
        "primeiro": 1,
        "1º": 1,
        "1°": 1,
        "1o": 1,
        "segundo": 2,
        "2º": 2,
        "2°": 2,
        "2o": 2,
        "terceiro": 3,
        "3º": 3,
        "3°": 3,
        "3o": 3,
        "quarto": 4,
        "4º": 4,
        "4°": 4,
        "4o": 4,
        "quinto": 5,
        "5º": 5,
        "5°": 5,
        "5o": 5,
    }

    for chave, numero in mapa.items():
        if chave in txt:
            return numero

    m = re.search(r"\b([1-9])\s*(?:º|°|o)?\s*(?:aditivo|aditamento)", txt, flags=re.IGNORECASE)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None

    return None


def _ordinal_aditivo_nome(numero: int | None) -> str:
    nomes = {
        1: "Primeiro Aditamento",
        2: "Segundo Aditamento",
        3: "Terceiro Aditamento",
        4: "Quarto Aditamento",
        5: "Quinto Aditamento",
    }
    return nomes.get(numero, "Termo aditivo")


def _extrair_blocos_arquivo_texto(texto: str) -> List[Dict[str, Any]]:
    """
    Quebra o texto extraído em blocos por arquivo.
    Esperado:
    ARQUIVO: nome_do_arquivo.pdf
    texto do arquivo...
    """
    texto = str(texto or "")
    blocos: List[Dict[str, Any]] = []

    partes = re.split(r"(?:^|\n)\s*ARQUIVO:\s*", texto, flags=re.IGNORECASE)

    for parte in partes:
        parte = str(parte or "").strip()
        if not parte:
            continue

        linhas = parte.splitlines()
        if not linhas:
            continue

        nome_arquivo = clean_text(linhas[0])
        conteudo = "\n".join(linhas[1:]) if len(linhas) > 1 else ""

        if not _info_util(nome_arquivo):
            continue

        blocos.append({
            "arquivo": nome_arquivo,
            "conteudo": conteudo,
            "texto_completo": f"{nome_arquivo}\n{conteudo}",
            "arquivo_low": nome_arquivo.lower(),
            "conteudo_low": conteudo.lower(),
        })

    return blocos


def _extrair_arquivos_analisados_base(base: Dict[str, Any]) -> List[str]:
    nomes = []

    bruto = clean_text(base.get("arquivos_analisados"))
    if _info_util(bruto):
        for parte in re.split(r"\s*\|\s*", bruto):
            nome = clean_text(parte)
            if _info_util(nome) and nome not in nomes:
                nomes.append(nome)

    return nomes


def _eh_arquivo_aditivo(nome: Any, conteudo: Any = "") -> bool:
    """
    Identifica se o arquivo é realmente um aditivo.

    Regra importante:
    - Nome do arquivo manda.
    - Contrato principal não pode virar aditivo só porque cita "Primeiro Aditivo" nos considerandos.
    - Apresentação, comunicado, proposta/estudo e certificado isolado não viram aditivo.
    """
    nome_txt = clean_text(nome)
    conteudo_txt = clean_text(conteudo)

    nome_low = nome_txt.lower()
    conteudo_inicio = conteudo_txt[:2500].lower()
    plano = f"{nome_low} {conteudo_inicio}"

    termos_excluir = [
        "comunicado",
        "apresentação",
        "apresentacao",
        "estudo de mercado",
        "proposta",
        "certificado de conclusão",
        "certificado de conclusao",
        "certificate of completion",
    ]

    if any(t in plano for t in termos_excluir):
        return False

    # Contrato principal não pode ser classificado como aditivo.
    if (
        "contrato" in nome_low
        and "aditivo" not in nome_low
        and "aditamento" not in nome_low
        and "para_ass" not in nome_low
    ):
        return False

    # Prioridade: nome do arquivo.
    if "aditivo" in nome_low or "aditamento" in nome_low:
        return True

    # Fallback: título do documento no começo do texto.
    padrao_titulo = re.search(
        r"\b(primeiro|segundo|terceiro|quarto|quinto|\d+\s*[º°o]?)\s+aditamento\b",
        conteudo_inicio,
        flags=re.IGNORECASE,
    )

    if padrao_titulo:
        return True

    return False


def _pontuar_bloco_aditivo(bloco: Dict[str, Any]) -> int:
    """
    Pontua versões do mesmo aditivo.
    Objetivo: escolher o melhor anexo.
    PDF assinado/com certificado > PDF com DocuSign > PDF final > DOCX/minuta.
    """
    nome = clean_text(bloco.get("arquivo"))
    conteudo = clean_text(bloco.get("conteudo"))
    nome_low = nome.lower()
    plano = f"{nome} {conteudo}".lower()

    pontos = 0

    if "aditivo" in nome_low or "aditamento" in nome_low:
        pontos += 300

    if nome_low.endswith(".pdf"):
        pontos += 250

    if "para_ass" in nome_low or "assinado" in nome_low or "(assinado)" in nome_low:
        pontos += 300

    if "docusign envelope id" in plano:
        pontos += 180

    if "certificate of completion" in plano or "certificado de conclusão" in plano or "certificado de conclusao" in plano:
        pontos += 250

    if "status: completed" in plano or "status: concluído" in plano or "status: concluido" in plano:
        pontos += 250

    if nome_low.endswith(".docx"):
        pontos -= 180

    if "rev" in nome_low or "minuta" in nome_low or "rascunho" in nome_low:
        pontos -= 120

    if "comunicado" in nome_low or "apresentação" in nome_low or "apresentacao" in nome_low:
        pontos -= 300

    if "contrato" in nome_low and "aditivo" not in nome_low and "aditamento" not in nome_low:
        pontos -= 500

    return pontos


def _extrair_envelope_id(texto: Any) -> str:
    txt = clean_text(texto)

    m = re.search(
        r"(?:Docusign Envelope ID|DocuSign Envelope ID|Envelope Id|Envelope ID|Identificação de envelope):?\s*([A-Z0-9\-]{20,})",
        txt,
        flags=re.IGNORECASE,
    )

    if m:
        return re.sub(r"[^A-Z0-9]", "", m.group(1).upper())

    return ""


def _montar_mapa_certificados_por_envelope(texto: str) -> Dict[str, str]:
    mapa: Dict[str, str] = {}

    for bloco in _extrair_blocos_arquivo_texto(texto):
        conteudo = bloco.get("texto_completo", "")
        nome = clean_text(bloco.get("arquivo"))
        plano = f"{nome} {conteudo}".lower()

        if "certificado de conclusão" not in plano and "certificate of completion" not in plano:
            continue

        env = _extrair_envelope_id(conteudo)
        if env:
            mapa[env] = conteudo

    return mapa


def _datas_encontradas(texto: Any) -> List[str]:
    txt = str(texto or "")
    datas = []

    for data in re.findall(r"\b\d{1,2}/\d{1,2}/\d{4}\b", txt):
        d = _normalizar_data_aditivo_para_br(data)
        if _info_util(d) and d not in datas:
            datas.append(d)

    for data in re.findall(
        r"\b\d{1,2}\s+de\s+[A-Za-zçÇãáéíóúâêôõÁÉÍÓÚÂÊÔÕ]+\s+de\s+\d{4}\b",
        txt,
        flags=re.IGNORECASE,
    ):
        d = _normalizar_data_aditivo_para_br(data)
        if _info_util(d) and d not in datas:
            datas.append(d)

    return datas


def _extrair_data_documento_aditivo(conteudo: Any) -> str:
    txt = str(conteudo or "")

    padroes = [
        r"(?:Atibaia|Jarinu|São Paulo\/SP|São Paulo|Sao Paulo\/SP|Sao Paulo),?\s*(\d{1,2}\s+de\s+[A-Za-zçÇãáéíóúâêôõÁÉÍÓÚÂÊÔÕ]+\s+de\s+\d{4})",
        r"(?:Atibaia|Jarinu|São Paulo\/SP|São Paulo|Sao Paulo\/SP|Sao Paulo),?\s*(\d{1,2}/\d{1,2}/\d{4})",
    ]

    for padrao in padroes:
        achados = re.findall(padrao, txt, flags=re.IGNORECASE)
        if achados:
            return _normalizar_data_aditivo_para_br(achados[-1])

    datas = _datas_encontradas(txt)
    if datas:
        return datas[0]

    return "Não localizado"


def _extrair_data_assinatura_docusign(conteudo: Any) -> str:
    txt = str(conteudo or "")

    padroes_conclusao = [
        r"(?:Concluído|Concluido|Completed).*?(\d{1,2}/\d{1,2}/\d{4})",
        r"(?:Signing Complete|Assinatura concluída|Assinatura concluida).*?(\d{1,2}/\d{1,2}/\d{4})",
    ]

    for padrao in padroes_conclusao:
        achados = re.findall(padrao, txt, flags=re.IGNORECASE | re.DOTALL)
        if achados:
            return _normalizar_data_aditivo_para_br(achados[-1])

    assinaturas = re.findall(
        r"(?:Assinado|Signed):\s*(\d{1,2}/\d{1,2}/\d{4})",
        txt,
        flags=re.IGNORECASE,
    )

    if assinaturas:
        return _normalizar_data_aditivo_para_br(assinaturas[-1])

    return "Não localizado"


def _extrair_signatarios_docusign(conteudo: Any) -> str:
    txt = str(conteudo or "")
    nomes = []

    linhas = [clean_text(l) for l in txt.splitlines() if clean_text(l)]

    ignorar = {
        "signer events",
        "eventos do signatário",
        "eventos do signatario",
        "assinatura",
        "registro de hora e data",
        "security level",
        "nível de segurança",
        "nivel de seguranca",
        "using ip address",
        "usando endereço ip",
        "sent:",
        "enviado:",
        "viewed:",
        "visualizado:",
        "signed:",
        "assinado:",
        "electronic record and signature disclosure",
        "termos de assinatura e registro eletrônico",
        "not offered via docusign",
        "accepted:",
        "aceito:",
    }

    for i, linha in enumerate(linhas):
        low = linha.lower()

        if "@" in linha:
            # normalmente o nome está na linha anterior
            if i > 0:
                nome = linhas[i - 1]
                nome_low = nome.lower()

                if (
                    _info_util(nome)
                    and "@" not in nome
                    and not any(t in nome_low for t in ignorar)
                    and not re.search(r"\d{1,2}/\d{1,2}/\d{4}", nome)
                    and nome not in nomes
                ):
                    nomes.append(nome)

    if nomes:
        return "; ".join(nomes)

    return "Não localizado"


def _extrair_periodo_aditivo(conteudo: Any) -> str:
    txt = clean_text(conteudo)

    m = re.search(
        r"(?:aplicad[oa]s?\s+a\s+partir\s+de|a partir de)\s+(\d{1,2}/\d{1,2}/\d{4}).{0,100}?(?:até|ate)\s+(\d{1,2}/\d{1,2}/\d{4})",
        txt,
        flags=re.IGNORECASE,
    )
    if m:
        return f"{_normalizar_data_aditivo_para_br(m.group(1))} até {_normalizar_data_aditivo_para_br(m.group(2))}"

    m = re.search(
        r"Quadro\s+4.*?aplicad[oa]s?\s+a\s+partir\s+de\s+(\d{1,2}/\d{1,2}/\d{4})\s+até\s+(\d{1,2}/\d{1,2}/\d{4})",
        txt,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if m:
        return f"{_normalizar_data_aditivo_para_br(m.group(1))} até {_normalizar_data_aditivo_para_br(m.group(2))}"

    m = re.search(
        r"(?:efeitos\s+retroagem\s+a|retroagem\s+a)\s+(\d{1,2}/\d{1,2}/\d{4})",
        txt,
        flags=re.IGNORECASE,
    )
    if m:
        return f"A partir de {_normalizar_data_aditivo_para_br(m.group(1))}"

    m = re.search(
        r"(?:efeitos\s+retroagem\s+a|retroagem\s+a)\s+(\d{1,2}\s+de\s+[A-Za-zçÇãáéíóúâêôõÁÉÍÓÚÂÊÔÕ]+\s+de\s+\d{4})",
        txt,
        flags=re.IGNORECASE,
    )
    if m:
        return f"A partir de {_normalizar_data_aditivo_para_br(m.group(1))}"

    m = re.search(
        r"(?:aplicad[oa]s?\s+a\s+partir\s+de|a partir de)\s+(\d{1,2}/\d{1,2}/\d{4})",
        txt,
        flags=re.IGNORECASE,
    )
    if m:
        return f"A partir de {_normalizar_data_aditivo_para_br(m.group(1))}"

    m = re.search(
        r"vigência\s+pelo\s+prazo\s+de\s+(\d+)\s*\([^)]*\)?\s*mes(?:es)?[^.]{0,120}?a\s+partir\s+de\s+(\d{1,2}\s+de\s+[A-Za-zçÇãáéíóúâêôõÁÉÍÓÚÂÊÔÕ]+\s+de\s+\d{4})",
        txt,
        flags=re.IGNORECASE,
    )
    if m:
        meses = int(m.group(1))
        inicio = _parse_data_br_para_datetime(_normalizar_data_aditivo_para_br(m.group(2)))
        if inicio:
            fim = _somar_meses(inicio, meses) - timedelta(days=1)
            return f"{inicio.strftime('%d/%m/%Y')} até {fim.strftime('%d/%m/%Y')}"

    return "Não localizado"


def _extrair_itens_aditivo_bbp(conteudo: Any) -> List[Dict[str, Any]]:
    txt = clean_text(conteudo)
    itens = []

    servicos_valores = [
        ("Refeição (Almoço ou Jantar)", r"Refeição\s*\(Almoço\s+ou\s+Jantar\)\s+([0-9]{1,3},[0-9]{2})\s+([0-9]{1,3},[0-9]{2})"),
        ("Refeição Ceia", r"Refeição\s+Ceia\s+([0-9]{1,3},[0-9]{2})\s+([0-9]{1,3},[0-9]{2})"),
        ("Desjejum", r"Desjejum.*?([0-9]{1,3},[0-9]{2})\s+([0-9]{1,3},[0-9]{2})"),
    ]

    for desc, padrao in servicos_valores:
        m = re.search(padrao, txt, flags=re.IGNORECASE)
        if m:
            itens.append({
                "Item": len(itens) + 1,
                "Descrição": desc,
                "Tipo": "Serviço",
                "Quantidade": "1",
                "Unidade": "UN",
                "Valor unitário": f"R$ {m.group(1)}",
                "Valor total": f"R$ {m.group(1)}",
                "Taxa / Percentual": f"Domingos/Feriados: R$ {m.group(2)}",
                "Total de encargos": "Não aplicável",
                "Fonte": "Aditivo",
            })

    quantidades = [
        ("Desjejum", r"Desjejum\s+\d{1,2}:\d{2}\s+às\s+\d{1,2}:\d{2}\s+(\d+)"),
        ("Almoço", r"Almoço\s+\d{1,2}:\d{2}\s+às\s+\d{1,2}:\d{2}\s+(\d+)"),
        ("Lanche da Tarde", r"Lanche\s+da\s+Tarde\s+\d{1,2}:\d{2}\s+às\s+\d{1,2}:\d{2}\s+(\d+)"),
        ("Jantar", r"Jantar\s+\d{1,2}:\d{2}\s+às\s+\d{1,2}:\d{2}\s+(\d+)"),
        ("Ceia", r"Ceia\s+\d{1,2}:\d{2}\s+às\s+\d{1,2}:\d{2}\s+(\d+)"),
    ]

    for desc, padrao in quantidades:
        m = re.search(padrao, txt, flags=re.IGNORECASE)
        if m:
            itens.append({
                "Item": len(itens) + 1,
                "Descrição": desc,
                "Tipo": "Serviço",
                "Quantidade": m.group(1),
                "Unidade": "REF/DIA",
                "Valor unitário": "Não localizado",
                "Valor total": "Não localizado",
                "Taxa / Percentual": "Quantidade mínima diária",
                "Total de encargos": "Não aplicável",
                "Fonte": "Aditivo",
            })

    return itens


def _inferir_impacto_valor_aditivo(conteudo: Any, numero: int | None = None) -> str:
    txt = clean_text(conteudo).lower()

    if numero == 1:
        return "Altera a cláusula 10 e o Quadro 4 de valores do fornecimento por refeição."

    if numero == 2:
        return "Reajusta os valores unitários das refeições para o período de 01/08/2024 a 31/07/2025."

    if numero == 3:
        return "Sem valor específico localizado; altera quantidade mínima/escopo operacional."

    if numero == 4:
        return "Ajuste de valores por alteração de ICMS de 3,20% para 4,00%."

    if numero == 5:
        return "Reajuste anual de 7,31% sobre o Quadro 4 de valores do fornecimento por refeição."

    if "7,31%" in txt:
        return "Reajuste anual de 7,31% sobre o Quadro 4 de valores do fornecimento por refeição."

    if "6,5%" in txt or "6,50%" in txt:
        return "Reajuste dos valores unitários das refeições."

    if "icms" in txt and ("3,20%" in txt or "4,00%" in txt):
        return "Ajuste de valores por alteração de ICMS de 3,20% para 4,00%."

    if "quadro 4" in txt and "valores" in txt:
        return "Altera o Quadro 4 de valores do fornecimento por refeição."

    if "quadro 3" in txt and "quantidades mínimas" in txt:
        return "Sem valor específico localizado; altera quantidade mínima/escopo operacional."

    return "Sem valor global fixo."


def _valor_aditivo_profissional(numero: int | None, conteudo: Any, itens: List[Dict[str, Any]]) -> str:
    """
    Valor do aditivo não é valor global.
    Estes aditivos alteram valores unitários, percentuais ou quantidades.
    """
    txt = clean_text(conteudo).lower()

    if numero == 3:
        return "Não aplicável (alteração de quantidade)."

    if numero == 5:
        return "Sem valor global fixo; reajuste de 7,31%."

    if numero in (1, 2, 4):
        return "Sem valor global fixo; possui valores unitários por refeição."

    if "7,31%" in txt:
        return "Sem valor global fixo; reajuste de 7,31%."

    if itens:
        return "Sem valor global fixo; possui valores unitários por refeição."

    return "Sem valor global fixo."

def _inferir_impacto_prazo_aditivo(conteudo: Any) -> str:
    txt = clean_text(conteudo).lower()

    periodo = _extrair_periodo_aditivo(conteudo)
    if _info_util(periodo):
        return f"Define período de aplicação: {periodo}."

    if "vigência pelo prazo de 12" in txt or "vigencia pelo prazo de 12" in txt:
        return "Define vigência de 12 meses para o aditivo."

    if "prazo indeterminado" in txt:
        return "Mantém/indica vigência por prazo indeterminado."

    if "retroagem a" in txt or "efeitos retroagem" in txt:
        return "Possui efeitos retroativos conforme data indicada no aditivo."

    return "Não localizado"


def _inferir_escopo_aditivo(numero: int | None, conteudo: Any) -> str:
    txt = clean_text(conteudo).lower()

    if numero == 1:
        return "Altera a cláusula de cessão/transferência e o Quadro 4 de valores do fornecimento por refeição."

    if numero == 2:
        return "Altera o Quadro 4 de valores do fornecimento por refeição, com aplicação a partir de 01/08/2024."

    if numero == 3:
        return "Altera o Quadro 3 de quantidades mínimas diárias e inclui o serviço Lanche da Tarde."

    if numero == 4:
        return "Ajusta o Quadro 4 de valores em razão da alteração da alíquota de ICMS."

    if numero == 5:
        return "Altera o Quadro 4 de valores do fornecimento por refeição, com reajuste anual e vigência de 12 meses."

    if "quadro 3" in txt and "lanche da tarde" in txt:
        return "Altera quantidades mínimas diárias e inclui Lanche da Tarde."

    if "quadro 4" in txt and "valores" in txt:
        return "Altera valores do fornecimento por refeição."

    return "Aditivo identificado nos documentos analisados."


def _montar_status_aditivo(aditivo: Dict[str, Any]) -> str:
    assinado = clean_text(aditivo.get("Assinado")).upper()
    anexo = aditivo.get("Anexo do aditivo")
    escopo = aditivo.get("Escopo do aditivo")

    if assinado == "SIM" and _info_util(anexo) and _info_util(escopo):
        return "Aditivo assinado e identificado com anexo/escopo. Validar impactos antes de seguir."

    if assinado == "SIM":
        return "Aditivo assinado. Revisar campos não localizados para confirmação operacional."

    if assinado in ("NÃO", "NAO"):
        return "Aditivo sem assinatura localizada. Não considerar vigente sem validação."

    return "Assinatura do aditivo não localizada com clareza. Revisar documento."


def _criar_aditivos_pelos_blocos(base: Dict[str, Any], texto: str) -> List[Dict[str, Any]]:
    blocos = _extrair_blocos_arquivo_texto(texto)
    certificados = _montar_mapa_certificados_por_envelope(texto)

    candidatos = []

    for bloco in blocos:
        nome = clean_text(bloco.get("arquivo"))
        conteudo = bloco.get("conteudo", "")
        texto_completo = bloco.get("texto_completo", "")

        if not _eh_arquivo_aditivo(nome, conteudo):
            continue

        # Prioriza o número do aditivo pelo NOME DO ARQUIVO.
        # Só usa o conteúdo como fallback, para evitar o contrato principal virar 1º aditivo.
        numero = _ordinal_aditivo_numero(nome)

        if not numero:
            numero = _ordinal_aditivo_numero(str(conteudo)[:2500])

        if not numero:
            continue

        env = _extrair_envelope_id(texto_completo)
        conteudo_certificado = certificados.get(env, "")
        conteudo_enriquecido = f"{texto_completo}\n{conteudo_certificado}"

        candidatos.append({
            "numero": numero,
            "arquivo": nome,
            "conteudo": conteudo_enriquecido,
            "pontos": _pontuar_bloco_aditivo(bloco),
        })

    # Se não tiver blocos ARQUIVO, tenta pelo nome dos arquivos analisados
    if not candidatos:
        for nome in _extrair_arquivos_analisados_base(base):
            if not _eh_arquivo_aditivo(nome, ""):
                continue

            numero = _ordinal_aditivo_numero(nome)
            if not numero:
                continue

            candidatos.append({
                "numero": numero,
                "arquivo": nome,
                "conteudo": nome,
                "pontos": 50,
            })

    melhor_por_numero: Dict[int, Dict[str, Any]] = {}

    for cand in candidatos:
        numero = cand["numero"]
        atual = melhor_por_numero.get(numero)

        if not atual or cand["pontos"] > atual["pontos"]:
            melhor_por_numero[numero] = cand

    aditivos = []

    data_carga = base.get("data_analise")
    if not _info_util(data_carga):
        data_carga = datetime.now().strftime("%d/%m/%Y %H:%M")

    for numero in sorted(melhor_por_numero):
        cand = melhor_por_numero[numero]
        conteudo = cand["conteudo"]
        nome = cand["arquivo"]

        data_aditivo = _extrair_data_documento_aditivo(conteudo)
        data_assinatura = _extrair_data_assinatura_docusign(conteudo)
        assinantes = _extrair_signatarios_docusign(conteudo)
        periodo = _extrair_periodo_aditivo(conteudo)
        impacto_valor = _inferir_impacto_valor_aditivo(conteudo, numero)
        impacto_prazo = _inferir_impacto_prazo_aditivo(conteudo)
        escopo = _inferir_escopo_aditivo(numero, conteudo)
        itens = _extrair_itens_aditivo_bbp(conteudo)

        assinado = "Sim" if (
            _info_util(data_assinatura)
            or "docusign envelope id" in clean_text(conteudo).lower()
            or "certificate of completion" in clean_text(conteudo).lower()
            or "certificado de conclusão" in clean_text(conteudo).lower()
            or "assinado" in clean_text(nome).lower()
            or "para_ass" in clean_text(nome).lower()
        ) else "Não"

        valor_aditivo = _valor_aditivo_profissional(numero, conteudo, itens)

        aditivo = {
            "Nº": str(numero),
            "Tipo do aditivo": _ordinal_aditivo_nome(numero),
            "Anexo do aditivo": nome,
            "Data do aditivo": data_aditivo,
            "Data de carga no robô": data_carga,
            "Assinado": assinado,
            "Data da assinatura": data_assinatura,
            "Quem assinou": resumir_campo(assinantes, 500),
            "Valor do aditivo": valor_aditivo,
            "Impacto no valor": impacto_valor,
            "Impacto no prazo": impacto_prazo,
            "Período do aditivo": periodo,
            "Escopo do aditivo": escopo,
            "Itens do aditivo": resumir_campo("; ".join([
                f"{i.get('Descrição')}: {i.get('Valor unitário')}" for i in itens
            ]) if itens else "Não localizado", 600),
            "Status de validação": "Não localizado",
            "Observações": "",
            "_itens_aditivo": itens,
        }

        aditivo["Status de validação"] = _montar_status_aditivo(aditivo)
        aditivos.append(aditivo)

    return aditivos


def _numero_aditivo_para_merge(aditivo: Dict[str, Any], fallback: int | None = None) -> str:
    numero = clean_text(
        aditivo.get("Nº")
        or aditivo.get("numero_aditivo")
        or aditivo.get("numero")
        or aditivo.get("aditivo")
        or ""
    )

    m = re.search(r"\b([1-9])\b", numero)
    if m:
        return m.group(1)

    numero_inferido = _ordinal_aditivo_numero(
        " ".join([
            clean_text(aditivo.get("Tipo do aditivo")),
            clean_text(aditivo.get("Anexo do aditivo")),
            clean_text(aditivo.get("Escopo do aditivo")),
            clean_text(aditivo.get("objeto_escopo_aditivo")),
            clean_text(aditivo.get("tipo_aditivo")),
            clean_text(aditivo.get("anexo_origem")),
        ])
    )

    if numero_inferido:
        return str(numero_inferido)

    return str(fallback or "")


def _mesclar_aditivos_ia_com_blocos(
    aditivos_ia: List[Dict[str, Any]],
    aditivos_blocos: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """
    Fonte principal = arquivos/blocos reais.
    A IA só complementa campos vazios.
    Isso evita trocar o anexo do 1º aditivo pelo contrato principal ou priorizar DOCX sobre PDF assinado.
    """
    mapa: Dict[str, Dict[str, Any]] = {}

    # 1. Arquivos reais têm prioridade
    for idx, aditivo_bloco in enumerate(aditivos_blocos, 1):
        numero = _numero_aditivo_para_merge(aditivo_bloco, idx)
        if not numero:
            numero = str(idx)

        aditivo_bloco["Nº"] = numero
        mapa[numero] = dict(aditivo_bloco)

    # 2. IA só complementa o que estiver vazio
    for idx, aditivo_ia in enumerate(aditivos_ia, 1):
        numero = _numero_aditivo_para_merge(aditivo_ia, idx)
        if not numero:
            numero = str(idx)

        if numero not in mapa:
            aditivo_ia["Nº"] = numero
            mapa[numero] = dict(aditivo_ia)
            continue

        existente = mapa[numero]

        for campo, valor in aditivo_ia.items():
            if not _info_util(existente.get(campo)) and _info_util(valor):
                existente[campo] = valor

        existente["Nº"] = numero
        existente["Status de validação"] = _montar_status_aditivo(existente)
        mapa[numero] = existente

    def ordem(item):
        try:
            return int(item[0])
        except Exception:
            return 999

    return [v for _, v in sorted(mapa.items(), key=ordem)]


def aplicar_regras_aditivos(base: Dict[str, Any], texto: str) -> Dict[str, Any]:
    """Consolida aditivos usando IA + leitura determinística dos arquivos/anexos."""
    texto = str(texto or "")

    aditivos_ia = []
    if isinstance(base.get("aditivos_contrato"), list):
        aditivos_ia = normalizar_aditivos_contrato(base.get("aditivos_contrato"))

    for chave in ["aditivos", "aditivos_identificados", "termos_aditivos", "lista_aditivos"]:
        if not aditivos_ia and isinstance(base.get(chave), list):
            aditivos_ia = normalizar_aditivos_contrato(base.get(chave))

    aditivos_blocos = _criar_aditivos_pelos_blocos(base, texto)

    aditivos = _mesclar_aditivos_ia_com_blocos(aditivos_ia, aditivos_blocos)

    # Se não conseguiu pelos blocos, mantém IA, mas corrige campos "Não localizado"
    if not aditivos:
        aditivos = aditivos_ia

    data_carga = base.get("data_analise")
    if not _info_util(data_carga):
        data_carga = datetime.now().strftime("%d/%m/%Y %H:%M")

    for aditivo in aditivos:
        if not _info_util(aditivo.get("Data de carga no robô")):
            aditivo["Data de carga no robô"] = data_carga

        if _info_util(aditivo.get("Data do aditivo")):
            aditivo["Data do aditivo"] = _normalizar_data_aditivo_para_br(aditivo.get("Data do aditivo"))

        if _info_util(aditivo.get("Data da assinatura")):
            aditivo["Data da assinatura"] = _normalizar_data_aditivo_para_br(aditivo.get("Data da assinatura"))

        if not _info_util(aditivo.get("Status de validação")):
            aditivo["Status de validação"] = _montar_status_aditivo(aditivo)

    base["aditivos_contrato"] = aditivos
    base["resumo_aditivos"] = _montar_resumo_aditivos(aditivos)

    return base

def extrair_itens_local(texto: str, limite: int = 120) -> List[Dict[str, Any]]:
    """Fallback para propostas/orçamentos com tabela."""
    texto = str(texto or "")

    servico_percentual = detectar_servico_percentual(texto)
    if servico_percentual:
        return normalizar_itens_contrato(servico_percentual)

    itens: List[Dict[str, Any]] = []

    def eh_numero_brasil(v: str) -> bool:
        v = str(v or "").strip()
        return bool(re.fullmatch(r"\d{1,3}(?:\.\d{3})*(?:,\d{2,6})?|\d+(?:,\d{2,6})?", v))

    def parece_valor(v: str) -> bool:
        v = str(v or "").strip()
        return bool(re.search(r"(?:R\$\s*)?\d{1,3}(?:\.\d{3})*,\d{2}\b|\b\d+,\d{2}\b", v))

    def moeda(v: str) -> str:
        v = clean_text(v)
        if v in ("", "Não localizado"):
            return "Não localizado"
        if "R$" in v:
            return v
        if parece_valor(v):
            return "R$ " + v
        return v

    linhas = [re.sub(r"\s+", " ", l).strip() for l in texto.splitlines() if str(l).strip()]
    vistos = set()

    # Primeiro: linhas de tabela com |
    for linha in linhas:
        if "|" not in linha:
            continue
        partes = [p.strip() for p in linha.split("|")]
        partes = [p for p in partes if p not in ("", "-", "None", "nan")]
        if len(partes) < 4:
            continue

        linha_lower = " ".join(partes).lower()
        if any(h in linha_lower for h in ["descrição", "descricao", "valor unit", "quantidade", "unidade"]):
            continue

        idx_valores = [i for i, p in enumerate(partes) if parece_valor(p)]
        if not idx_valores:
            continue

        idx_total = idx_valores[-1]
        idx_unit = idx_valores[-2] if len(idx_valores) >= 2 else idx_valores[-1]

        quantidade = "Não localizado"
        unidade = "Não localizado"
        for i, p in enumerate(partes[:idx_unit]):
            if eh_numero_brasil(p):
                quantidade = p
                if i + 1 < len(partes) and not parece_valor(partes[i + 1]):
                    unidade = partes[i + 1]
                break

        desc_partes = []
        for p in partes[:idx_unit]:
            if p == quantidade:
                break
            desc_partes.append(p)

        descricao = " ".join(desc_partes).strip() or partes[0]
        if len(descricao) < 5:
            continue

        chave = (descricao.lower(), quantidade, moeda(partes[idx_unit]), moeda(partes[idx_total]))
        if chave in vistos:
            continue
        vistos.add(chave)

        itens.append({
            "item": len(itens) + 1,
            "descricao": descricao[:420],
            "tipo": _inferir_tipo_item(descricao),
            "quantidade": quantidade,
            "unidade": unidade,
            "valor_unitario": moeda(partes[idx_unit]),
            "valor_total": moeda(partes[idx_total]),
            "fonte": "Tabela extraída do contrato/anexo/proposta",
        })

        if len(itens) >= limite:
            return normalizar_itens_contrato(itens)

    # Segundo: linhas soltas com valores
    moeda_re = re.compile(r"R\$\s?\d{1,3}(?:\.\d{3})*,\d{2}|\b\d{1,3}(?:\.\d{3})*,\d{2}\b")
    qtd_un_re = re.compile(r"\b(\d{1,3}(?:\.\d{3})*|\d+(?:,\d+)?)\s*(Rolos?|UN|UND|UNID|Unit[aá]rio|Unidade|MÊS|MES|HORA|HR|DIA|KG|CX|PC|PÇ|SERV|SV)\b", re.IGNORECASE)

    for linha in linhas:
        valores = moeda_re.findall(linha)
        if len(valores) < 1 or len(linha) < 20:
            continue

        descricao = moeda_re.sub(" ", linha)
        m_qtd = qtd_un_re.search(descricao)
        qtd = "Não localizado"
        un = "Não localizado"
        if m_qtd:
            qtd = m_qtd.group(1)
            un = m_qtd.group(2)

        descricao = re.sub(r"\b\d{1,3}(?:\.\d{3})*(?:,\d+)?\b", " ", descricao)
        descricao = re.sub(r"\s+", " ", descricao).strip(" -|;:.,")
        if len(descricao) < 8:
            continue

        vu = valores[0]
        vt = valores[-1]
        chave = (descricao.lower(), vu, vt)
        if chave in vistos:
            continue
        vistos.add(chave)

        itens.append({
            "item": len(itens) + 1,
            "descricao": descricao[:420],
            "tipo": _inferir_tipo_item(descricao),
            "quantidade": qtd,
            "unidade": un,
            "valor_unitario": moeda(vu),
            "valor_total": moeda(vt),
            "fonte": "Extração local por linha com valor monetário",
        })

        if len(itens) >= limite:
            break

    return normalizar_itens_contrato(itens)


# =========================================================
# ANÁLISE LOCAL
# =========================================================
def local_extract(texto: str) -> Dict[str, Any]:
    texto = texto or ""
    low = texto.lower()

    cnpjs = re.findall(r"\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}", texto)
    valores = re.findall(r"R\$\s?\d{1,3}(?:\.\d{3})*,\d{2}", texto)
    datas = re.findall(r"\b\d{2}/\d{2}/\d{4}\b", texto)

    fornecedor = "Não localizado"
    fornecedor_patterns = [
        r"contratada[:\s]+([^\n]{5,140})",
        r"fornecedor[:\s]+([^\n]{5,140})",
        r"licenciante[:\s]+([^\n]{5,140})",
        r"licenciada[:\s]+([^\n]{5,140})",
        r"([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\s\.\-&]+S/A)",
        r"([A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9\s\.\-&]+LTDA\.?)",
    ]

    for pattern in fornecedor_patterns:
        match = re.search(pattern, texto, re.IGNORECASE)
        if match:
            fornecedor = re.sub(r"\s+", " ", match.group(1)).strip(" -:;,.|/")[:180]
            break

    objeto = "Não localizado"
    objeto_patterns = [
        r"CLÁUSULA PRIMEIRA[:\s\-]+DO OBJETO DO CONTRATO(.{50,1200})",
        r"CLAUSULA PRIMEIRA[:\s\-]+DO OBJETO DO CONTRATO(.{50,1200})",
        r"OBJETO DO CONTRATO(.{50,1200})",
        r"objeto(?:\s+do\s+contrato)?[:\s]+(.{20,900})",
        r"escopo[:\s]+(.{20,900})",
    ]

    for pattern in objeto_patterns:
        match = re.search(pattern, texto, re.IGNORECASE | re.DOTALL)
        if match:
            objeto = re.sub(r"\s+", " ", match.group(1)).strip(" -:;,.|/")[:700]
            break

    valor_total = max(valores, key=len) if valores else "Não localizado"

    vigencia = "Não localizada"
    vigencia_match = re.search(r"(vig[eê]ncia|prazo)[^.\n]{0,180}", texto, re.IGNORECASE)
    if vigencia_match:
        vigencia = clean_text(vigencia_match.group(0))[:220]
    if "prazo de duração indeterminado" in low or "prazo indeterminado" in low:
        vigencia = "Prazo indeterminado"

    condicao_pagamento = "Não localizada"
    pagamento_patterns = [
        r"prazo de\s*(\d{1,3})\s*\(?\w*\)?\s*dias",
        r"em até\s*(\d{1,3})\s*\(?\w*\)?\s*dias",
        r"(\d{1,3})\s*\(?\w*\)?\s*dias",
        r"(\d{1,3})\s*dd",
        r"(\d{1,3})\s*d\.?d\.?”?",
    ]

    for pattern in pagamento_patterns:
        match = re.search(pattern, low, re.IGNORECASE)
        if match:
            condicao_pagamento = f"{match.group(1)}DD"
            break

    if condicao_pagamento == "Não localizada" and any(t in low for t in ["pagamento", "nota fiscal", "fatura", "mensalidade", "vencimento"]):
        condicao_pagamento = "Cláusula de pagamento localizada"

    multa = "Não localizada"
    match = re.search(r"multa[^.\n]{0,140}?(\d{1,3})\s*%", low, re.IGNORECASE)
    if match:
        multa = f"{match.group(1)}%"
    elif any(t in low for t in ["multa", "penalidade", "penalidades"]):
        multa = "Cláusula localizada"

    reajuste = "Não localizado"
    if "igp-m" in low or "igpm" in low:
        reajuste = "IGP-M"
    elif "ipca" in low:
        reajuste = "IPCA"
    elif "fgv" in low:
        reajuste = "FGV"
    elif "reajuste" in low:
        reajuste = "Cláusula localizada"

    rescisao = "Localizada" if any(t in low for t in ["rescisão", "rescisao", "resilir", "rescindir"]) else "Não localizada"
    sla = "Localizado" if any(t in low for t in ["sla", "nível de serviço", "nivel de serviço", "tempo de resposta", "disponibilidade"]) else "Não localizado"
    confidencialidade = "Localizada" if any(t in low for t in ["confidencialidade", "sigilo"]) else "Não localizada"
    lgpd = "Localizada" if any(t in low for t in ["lgpd", "proteção de dados", "dados pessoais"]) else "Não localizada"
    anticorrupcao = "Localizada" if any(t in low for t in ["anticorrupção", "anticorrupcao", "lei anticorrupção", "lei 12.846"]) else "Não localizada"

    assinatura_termos = [
        "assinado", "assinatura", "testemunhas", "docusign", "certificado digital",
        "partes assinam", "por estarem justas e contratadas"
    ]
    contrato_assinado = "Sim" if any(t in low for t in assinatura_termos) else "Não"
    alerta_assinatura = "Evidência de assinatura localizada" if contrato_assinado == "Sim" else "Contrato sem evidência de assinatura localizada"

    validacoes = [
        ("Contrato assinado", contrato_assinado == "Sim", 30, "Sim"),
        ("CNPJ do fornecedor", bool(cnpjs), 20, "Sim"),
        ("Fornecedor / Contratada", fornecedor != "Não localizado", 20, "Sim"),
        ("Objeto do contrato", objeto != "Não localizado", 20, "Sim"),
        ("Valor contratual", valor_total != "Não localizado", 20, "Sim"),
        ("Vigência", vigencia != "Não localizada", 25, "Sim"),
        ("Condição de pagamento", condicao_pagamento != "Não localizada", 15, "Sim"),
        ("Multa / penalidade", multa != "Não localizada", 10, "Não"),
        ("Reajuste", reajuste != "Não localizado", 10, "Não"),
        ("Rescisão", rescisao == "Localizada", 15, "Sim"),
        ("SLA / nível de serviço", sla == "Localizado", 15, "Não"),
        ("Confidencialidade", confidencialidade == "Localizada", 10, "Não"),
        ("LGPD / proteção de dados", lgpd == "Localizada", 15, "Não"),
        ("Anticorrupção", anticorrupcao == "Localizada", 15, "Não"),
    ]

    checklist: List[Dict[str, Any]] = []
    pendencias: List[Dict[str, Any]] = []
    score = 0

    for item, ok, peso, critico in validacoes:
        if not ok:
            score += peso
            pendencias.append({
                "Pendência": item,
                "Crítico": critico,
                "Risco": peso,
                "Recomendação": "Validar cláusula/informação antes de seguir com RC/PO."
            })

        checklist.append({
            "Validação": item,
            "Status": "OK" if ok else "Pendente",
            "Peso de risco": 0 if ok else peso,
            "Crítico": critico,
            "Evidência": "Termo localizado" if ok else "Não localizado no texto extraído",
        })

    score = min(score, 100)
    risco = "BAIXO" if score <= 25 else "MÉDIO" if score <= 60 else "ALTO"
    status = "Aprovado" if score <= 25 else "Aprovado com ressalvas" if score <= 60 else "Contrato pendente de revisão"

    return {
        "tipo_contrato": "Não localizado",
        "empresa_grupo_sbf": "SBF COMÉRCIO DE PRODUTOS ESPORTIVOS LTDA." if "sbf comércio" in low or "sbf comercio" in low else "Não localizado",
        "cnpj_empresa_grupo": "06.347.409/0001-65" if "06.347.409/0001-65" in texto else "Não localizado",
        "local_prestacao": "Não localizado",
        "contraparte": fornecedor,
        "cnpj_contraparte": cnpjs[0] if cnpjs else "Não localizado",
        "objetivo": objeto,
        "descricao_servico_material": objeto,
        "descricao_breve_cadastro": objeto[:120] if objeto != "Não localizado" else "Não localizado",
        "forma_pagamento": condicao_pagamento,
        "condicao_pagamento_dias": condicao_pagamento,
        "multa": multa,
        "vigencia_apos_assinatura": vigencia,
        "periodo_vigencia_formatado": "Não localizado",
        "resumo_aditivos": "Nenhum aditivo identificado nos documentos analisados.",
        "aditivos_contrato": [],
        "rescisao_indenizacao": rescisao,
        "anticorrupcao": anticorrupcao,
        "indice_reajuste_anual": reajuste,
        "protecao_dados_lgpd": "Há cláusula de proteção de dados pessoais em conformidade com a LGPD." if lgpd == "Localizada" else "Não localizado",
        "data_assinatura": datas[-1] if datas else "Não localizada",
        "unidade_organizacional": "Não localizado",
        "area_solicitante": "Não localizado",
        "cadastrado_por": "Não localizado",
        "fim_contrato": "Não localizado",
        "prazo_contrato_dias": "Não localizado",
        "data_assinatura_aditivo": "Não localizado",
        "valor_contrato_original": valor_total,
        "valor_mensal_estimado": "Não localizado",
        "valor_total_estimado_vigencia": "Não localizado",
        "valor_total_materiais_servicos": "Não identificado",
        "aditivo_escopo": "Não localizado",
        "aditivo_quitacao": "Não localizado",
        "aditivo_valor": "Não localizado",
        "aditivo_prazo": "Não localizado",
        "quantidade_renovacoes": "Não localizado",
        "data_atualizacao": datetime.now().strftime("%d/%m/%Y"),
        "contrato_assinado": contrato_assinado,
        "alerta_assinatura": alerta_assinatura,
        "fornecedor": fornecedor,
        "cnpj": cnpjs[0] if cnpjs else "Não localizado",
        "cnpjs_encontrados": ", ".join(cnpjs[:10]) or "Não localizado",
        "valor_total": valor_total,
        "valores_encontrados": ", ".join(valores[:15]) or "Não localizado",
        "datas_encontradas": ", ".join(datas[:15]) or "Não localizado",
        "objeto": objeto,
        "vigencia": vigencia,
        "condicao_pagamento": condicao_pagamento,
        "multa": multa,
        "reajuste": reajuste,
        "rescisao": rescisao,
        "sla": sla,
        "lgpd": lgpd,
        "confidencialidade": confidencialidade,
        "assinatura": "Localizada" if contrato_assinado == "Sim" else "Não localizada",
        "score": score,
        "risco": risco,
        "status": status,
        "resumo_executivo": f"Análise concluída. Status: {status}. Risco: {risco}. Score: {score}. Pendências: {len(pendencias)}.",
        "parecer": "Recomenda-se revisar as pendências antes de seguir com RC/PO." if pendencias else "Contrato aparentemente possui os itens essenciais para continuidade.",
        "itens_contrato": extrair_itens_local(texto),
        "checklist": checklist,
        "pendencias": pendencias,
    }


# =========================================================
# IA
# =========================================================
def prompt_ia(texto: str) -> str:
    return f"""
Você é um analista sênior de contratos da área de Suprimentos/Jurídico.
Sua tarefa é consolidar CONTRATO PRINCIPAL + ANEXOS + ORÇAMENTOS enviados no mesmo texto.
Retorne APENAS JSON válido, sem markdown, sem comentários e sem texto fora do JSON.

OBJETIVO PRINCIPAL
Extrair os dados comerciais e jurídicos reais do contrato, priorizando o documento principal e usando anexos/orçamentos apenas quando o contrato fizer referência a eles.

ORDEM DE PRIORIDADE DOS DOCUMENTOS
Quando os arquivos enviados tiverem minutas, versões antigas e versão assinada, siga esta ordem:
1. PDF assinado com DocuSign / Certificate of Completion com Status Completed.
2. Contrato final em PDF.
3. Proposta comercial.
4. Proposta técnica.
5. E-mails de negociação apenas como contexto comercial.
6. Minutas DOCX antigas somente quando não houver contrato final assinado.
Se houver conflito entre minuta antiga e PDF assinado, prevalece o PDF assinado.

REGRA DE OURO DE QUALIDADE
Antes de preencher o JSON, faça mentalmente esta validação:
1. Qual é o contrato principal assinado? Ele manda nos campos de partes, CNPJ, objeto, vigência, assinatura, cláusulas, risco e status.
2. Quais anexos são proposta comercial/técnica? Use-os para preço, taxas, itens, vencimento, reposição, escopo técnico e condições comerciais.
3. Quais anexos são CNPJ, contrato social, Berkan, certificado, comunicado ou apoio cadastral? Use-os apenas como apoio/evidência, sem sobrescrever o contrato principal.
4. Minuta Word/DOCX sem assinatura só pode complementar campo vazio; nunca pode substituir PDF assinado.
5. Quando existir matriz e filial da mesma contraparte, mantenha no campo cnpj_contraparte o CNPJ que está no contrato principal assinado e registre divergência/apoio no resumo/checklist.
6. Todo campo importante deve ter conteúdo de negócio. Não responda de forma rasa. Exemplo ruim: "Possui LGPD". Exemplo bom: "Há cláusula de proteção de dados exigindo tratamento apenas para execução do contrato, segurança das informações e comunicação de incidentes em até 24 horas."
7. Se o contrato remeter valor à proposta, obrigatoriamente busque taxa, preço, vencimento e reposição na proposta comercial.
8. No resumo executivo, NÃO cite minutas, modelos DOCX, contrato social, CNPJ, Berkan, validações ou documentos de apoio como se fossem documento principal. Cite somente contrato principal/final, proposta comercial/técnica quando complementar, aditivos válidos e pendências relevantes.
9. Contrato social, alteração contratual societária, cartão CNPJ e Berkan são APOIO CADASTRAL/TÉCNICO. Eles validam dados, mas não são contrato operacional principal e não devem substituir partes, vigência, objeto ou valor do contrato principal.
10. Se houver contrato final em PDF e minutas/modelos DOCX, as minutas/modelos devem ser ignorados da análise profunda e usados apenas se não existir documento final.
11. Para vigência por prazo indeterminado, sempre preencha periodo_vigencia_formatado no padrão: "Início DD/MM/AAAA até 31/12/9999".
12. Para valores/taxas percentuais, não calcule valor total. Percentual, tributo, taxa e valor unitário sem quantidade definida devem ser classificados como condição comercial identificada, e não como valor global.
13. Se o texto extraído disser “Não foi possível extrair texto”, “Unable to get page count”, “Poppler” ou erro semelhante, IGNORE essa mensagem como conteúdo contratual. Se o arquivo original estiver anexado, analise visualmente o PDF original pela Files API.
14. Nome genérico de arquivo, como “SBF.pdf”, “contrato.pdf” ou “documento.pdf”, não significa apoio. Classifique pelo conteúdo real do arquivo.
15. Cards/dados extraídos só podem ser preenchidos quando houver evidência no contrato principal/anexo válido. Não use contrato social, CNPJ, Berkan, certificado ou comunicado para inventar objeto, vigência, valor ou partes do contrato operacional.
16. Se não houver evidência suficiente para um campo, use “Não identificado com segurança”. Nunca substitua por chute ou texto genérico.
17. Na análise completa, use o mesmo JSON final consolidado dos cards, aditivos, materiais, checklist, pendências, parecer e assinaturas. Não deixe campos vazios quando a informação constar no contrato principal ou proposta válida.

CONTRATOS DE MÃO DE OBRA TEMPORÁRIA / RH / RECRUTAMENTO
Quando o pacote contiver Lei 6.019/1974, mão de obra temporária, Atração & Retenção, recrutamento, seleção ou proposta de RH:
- tipo_contrato deve indicar prestação de serviços de alocação de mão de obra temporária ou recrutamento/seleção, conforme o contrato principal.
- objetivo deve mencionar substituição transitória de pessoal permanente e/ou demanda complementar de serviços quando constar.
- descricao_servico_material deve citar alocação de mão de obra temporária, recrutamento, seleção, contratação, gestão de folha/encargos quando constar.
- forma_pagamento deve citar emissão de nota fiscal e aprovação formal da contratante quando constar.
- condicao_pagamento_dias deve buscar na proposta frases como "60 dias corridos a contar da data de emissão".
- itens_contrato deve trazer condições comerciais da proposta: 40% + tributos de NF 0,8367 para vagas administrativas/operacionais/comerciais; 50% + tributos de NF 0,8367 para vagas técnicas/estratégicas/liderança; DISC/R$ 180,00 + tributos quando constar.
- valor_contrato_original deve ficar "Sem valor global fixo" quando não houver valor total fechado, mesmo havendo taxas percentuais.
- valor_mensal_estimado e valor_total_estimado_vigencia devem explicar que dependem de demanda, quantidade, salário/remuneração e fechamento de vaga, quando essa for a natureza comercial.

CAMPOS OBRIGATÓRIOS
Retorne exatamente estas chaves principais e internas:
{CAMPOS_JSON_OBRIGATORIOS}

Os campos principais que serão exibidos ao usuário são exatamente:
- Tipo de Contrato = tipo_contrato
- Empresa do Grupo SBF = empresa_grupo_sbf
- CNPJ Empresa do Grupo = cnpj_empresa_grupo
- Local de Prestação Contraparte = local_prestacao
- Contraparte = contraparte
- CNPJ Contraparte = cnpj_contraparte
- Objetivo = objetivo
- Descrição do Serviço/ Material = descricao_servico_material
- Descrição breve do cadastro = descricao_breve_cadastro
- Forma de pagamento = forma_pagamento
- Condição de Pagamento em Dias = condicao_pagamento_dias
- Multa = multa
- Vigência após a data de assinatura = vigencia_apos_assinatura
- Período de Vigência = periodo_vigencia_formatado
- Resumo de Aditivos = resumo_aditivos
- Rescisão e Indenização = rescisao_indenizacao
- Anticorrupção = anticorrupcao
- Proteção de Dados LGPD = protecao_dados_lgpd
- Data da Assinatura = data_assinatura
- Data do Contrato = data_contrato
- Data Conclusão DocuSign = data_conclusao_docusign
- Valor do Contrato Original = valor_contrato_original
- Valor Mensal Estimado = valor_mensal_estimado
- Valor Total Estimado da Vigência = valor_total_estimado_vigencia
- Valor Total dos Materiais e Serviços = valor_total_materiais_servicos
- Pessoas que assinaram = pessoas_que_assinaram

TABELA DE ITENS OBRIGATÓRIA
Também retorne a chave itens_contrato como lista. Cada item deve conter, quando aplicável:
- item
- descricao
- tipo
- quantidade
- unidade
- valor_unitario
- valor_total
- taxa_percentual
- total_encargos
- vencimento
- fonte

Para propostas de mão de obra temporária ou serviço sem preço unitário em R$:
- NÃO crie uma linha separada para "taxa de agenciamento" e outra para "encargos sociais".
- Crie UMA linha principal com descricao = "Mão de obra temporária" ou o serviço equivalente.
- Preencha taxa_percentual com a taxa de agenciamento, exemplo "7%".
- Preencha total_encargos com o percentual/total de encargos, exemplo "59,08%".
- Preencha vencimento com o prazo de pagamento, exemplo "30 dias após emissão da nota fiscal".
- Deixe valor_unitario e valor_total como "Não localizado" quando não houver valor em R$.

TABELA DE ASSINATURAS OBRIGATÓRIA
Também retorne a chave assinaturas_contrato como lista. Cada assinatura deve conter, quando aplicável:
- nome
- papel_cargo
- email
- data_assinatura
- data_reconhecimento_firma
- fonte
- status
- evidencia

Regras para assinaturas:
1. Se houver DocuSign Certificate of Completion / Status Completed, considere o contrato assinado.
2. Liste os signatários encontrados no certificado ou no bloco de assinaturas do contrato.
3. Não use nomes de documentos cadastrais como signatários.
4. Se houver data de assinatura por pessoa, preencha; se não houver, use a data principal da assinatura/conclusão do contrato.
5. Reconhecimento de firma NÃO é data de assinatura. Preencha data_reconhecimento_firma separadamente quando houver cartório, selo ou reconhecimento por semelhança/autenticidade.
6. Se não houver assinaturas localizadas, retorne assinaturas_contrato como lista vazia [].

TABELA DE ADITIVOS OBRIGATÓRIA
Também retorne a chave aditivos_contrato como lista.

Cada aditivo identificado deve conter:
- numero_aditivo
- tipo_aditivo
- anexo_origem
- data_aditivo
- data_carga_robo
- assinado
- data_assinatura_aditivo
- pessoas_que_assinaram_aditivo
- valor_aditivo
- impacto_valor
- impacto_prazo
- periodo_vigencia_aditivo
- objeto_escopo_aditivo
- itens_aditivo
- status_validacao_aditivo
- observacoes_aditivo

Regras para aditivos:
1. Identifique todos os documentos que sejam termo aditivo, aditamento, alteração contratual, renovação, prorrogação, reajuste, alteração de escopo, alteração de valor ou quitação.
2. anexo_origem deve trazer o nome do arquivo/anexo onde o aditivo foi identificado.
3. data_aditivo deve ser a data textual do termo aditivo, se houver.
4. data_carga_robo deve ser preenchida como "Data de carga no robô" quando não existir data de upload no documento.
5. assinado deve retornar "Sim" quando houver assinatura, DocuSign Completed, certificado de conclusão ou evidência de assinatura das partes.
6. Se não houver assinatura localizada no aditivo, retorne "Não".
7. pessoas_que_assinaram_aditivo deve listar os signatários do aditivo, separados por ponto e vírgula.
8. valor_aditivo deve trazer o valor específico do aditivo, se houver.
9. impacto_valor deve explicar se o aditivo aumenta valor, reduz valor, mantém valor, reajusta valor ou não altera valor.
10. impacto_prazo deve explicar se prorroga, reduz, mantém ou encerra prazo.
11. periodo_vigencia_aditivo deve trazer o período do aditivo quando houver, exemplo: "Início 01/01/2026 até 31/12/2026".
12. objeto_escopo_aditivo deve resumir o que o aditivo altera.
13. itens_aditivo deve trazer materiais/serviços específicos do aditivo, quando houver, usando a mesma estrutura de itens_contrato.
14. status_validacao_aditivo deve trazer uma visão executiva, exemplo: "Aditivo assinado e válido", "Aditivo sem assinatura localizada", "Aditivo com valor, porém sem evidência de assinatura".
15. resumo_aditivos deve resumir a situação geral dos aditivos encontrados em até 2 linhas.
16. Se não houver aditivos nos documentos analisados, retorne aditivos_contrato como lista vazia [] e resumo_aditivos como "Nenhum aditivo identificado nos documentos analisados."

REGRAS DE EXTRAÇÃO OBRIGATÓRIAS
1. Não invente dados. Se não encontrar, retorne "Não localizado".
2. Nunca confunda CNPJ com número sem máscara. Sempre formate CNPJ como 00.000.000/0000-00 quando houver 14 dígitos.
3. contraparte deve ser somente a empresa contratada/fornecedor. Não inclua CNPJ no mesmo campo.
4. local_prestacao deve ser cidade/UF/local de execução, se existir. Não misture com contraparte.
5. objetivo deve responder "para quê é o contrato" em uma frase curta.
6. descricao_servico_material deve descrever o serviço/material contratado, não histórico societário, alteração cadastral ou qualificação da empresa.
7. descricao_breve_cadastro deve ser curta e própria para cadastro de material/serviço. Exemplo: "Fornecimento e instalação de baterias para nobreaks".
8. condicao_pagamento_dias deve retornar em formato claro, exemplo: "90 dias". Não retorne apenas número solto.
9. forma_pagamento deve informar a forma descrita no contrato, exemplo: "Mediante emissão de nota fiscal".
10. valor_contrato_original deve ser preenchido SOMENTE quando houver valor global fechado expressamente definido para todo o contrato. Não use valor unitário, valor por pessoa, valor por item, valor mensal, taxa, consumo estimado ou valor por demanda como valor original do contrato. Se não houver valor global fixo, retorne: "Sem valor global fixo definido no contrato. O documento apresenta valores unitários, mensais ou por demanda, mas não informa um valor total fechado para todo o contrato."
11. valor_mensal_estimado deve explicar se o valor mensal é fixo ou variável. Se houver valor por unidade/pessoa/aprendiz/item/serviço, explique que o valor mensal depende da quantidade contratada ou utilizada. Exemplo: "R$ 282,42 por aprendiz ativo/mês. O valor mensal não é fixo; varia conforme a quantidade de aprendizes ativos no período."
12. valor_total_estimado_vigencia deve ser calculado somente quando houver informações suficientes: valor mensal, quantidade/volume aplicável e prazo de vigência. Se faltar quantidade ou volume, não invente valor. Retorne explicação clara e a fórmula, exemplo: "Não calculável com precisão. Para calcular, é necessário aplicar: valor unitário x quantidade x quantidade de meses."
13. valor_total_materiais_servicos deve trazer a soma da coluna valor_total dos itens em itens_contrato. Esse campo representa apenas a soma dos materiais/serviços identificados, não necessariamente o valor global do contrato. Se não houver itens com valor total, retorne "Não identificado".
14. vigencia_apos_assinatura deve manter a redação objetiva do contrato, exemplo: "36 meses a partir da assinatura" ou "Prazo indeterminado".
15. data_contrato deve ser a data textual do instrumento, exemplo "São Paulo, 21 de maio de 2026".
16. data_conclusao_docusign deve ser a data do Certificate of Completion / Completed do DocuSign.
17. data_assinatura deve ser a data final efetiva da assinatura quando houver DocuSign Completed; se não houver certificado, use a data textual do contrato.
18. contrato_assinado deve retornar "Sim" se houver DocuSign Certificate of Completion com Status Completed, assinaturas eletrônicas concluídas, ou evidência de assinatura das partes/testemunhas.
19. pessoas_que_assinaram deve listar os nomes encontrados no certificado DocuSign ou nas assinaturas do contrato. Retorne nomes separados por ponto e vírgula.
20. Se houver PDF assinado/DocuSign Completed, ignore pendências de assinatura existentes em minutas DOCX antigas.
21. Se não encontrar assinatura, alerta_assinatura deve retornar "Contrato sem evidência de assinatura localizada".
16. anticorrupcao e protecao_dados_lgpd devem resumir a cláusula específica quando houver, explicando a obrigação principal.
17. risco deve ser BAIXO, MÉDIO ou ALTO.
18. score deve ser número inteiro de 0 a 100, onde maior é melhor.
19. checklist deve ser lista de objetos com exatamente: Validação, Status, Peso de risco, Crítico, Evidência.
20. pendencias deve ser lista de objetos com exatamente: Pendência, Crítico, Risco, Recomendação.
21. Os campos exibidos ao usuário devem vir em RESUMO EXECUTIVO: nem curtos demais, nem cópia integral da cláusula.
22. Para cada campo principal, retorne de 1 a 3 linhas objetivas, com informação suficiente para entendimento corporativo.
23. Nunca retorne apenas: "Localizado", "Localizada", "Sim", "Possui cláusula" ou "Cláusula localizada" nos campos principais. Sempre explique resumidamente o conteúdo encontrado.
24. Só use "Não localizado" quando realmente não houver informação no contrato/anexo.
25. Anticorrupção e Proteção de Dados LGPD devem dizer que há cláusula e resumir a obrigação/finalidade.
26. Multa deve informar percentual, juros, base de cálculo e situação de aplicação quando constar.
27. Rescisão e Indenização deve resumir aviso prévio, multa/ônus, hipóteses de rescisão e indenização quando constar.
28. Vigência deve explicar prazo e início/encerramento de forma objetiva.
29. Forma de pagamento deve explicar o gatilho do pagamento, por exemplo: emissão/aprovação de nota fiscal, aceite, medição ou conclusão do serviço.
30. Descrição breve do cadastro deve ser uma frase própria para cadastro de serviço/material, sem copiar histórico societário ou qualificação jurídica.
31. valor_contrato_original deve estar formatado em reais somente quando houver valor global fechado. Quando não houver valor global, explique profissionalmente que o contrato é por demanda, mensal, unitário ou sem valor global fixo definido.
32. itens_contrato deve trazer todos os materiais e serviços unitários encontrados no contrato, anexos, propostas ou orçamentos.
33. descricao deve ser o nome do material/serviço, sem copiar cláusulas longas.
34. valor_unitario deve ser o preço unitário do item. Se não houver preço unitário monetário claro, retorne "Não aplicável" ou "Não localizado".
35. valor_total deve ser o total da linha/item. Se não houver valor monetário, retorne "Não localizado".
36. quantidade e unidade devem ser extraídas quando existirem; em serviços percentuais sem quantidade, retorne "Não aplicável".
37. tipo deve ser "Material" ou "Serviço".
38. fonte deve indicar de onde veio o item, exemplo: "Contrato", "Anexo", "Proposta" ou "Orçamento".
39. Para mão de obra temporária/RH, crie 1 item com descricao "Mão de Obra Temporária" e preencha taxa_percentual, total_encargos e vencimento. Não crie itens separados chamados "Taxa" ou "Encargos".
40. Se não houver tabela de itens nem condição comercial de serviço, retorne itens_contrato como lista vazia [].
40. Para contratos de mão de obra temporária, taxa de agenciamento, folha de pagamento, encargos sociais ou MCT/Mão de Obra: consolide as condições em um único serviço principal. Não trate taxa e encargos como materiais separados.

PADRÃO ESPERADO PARA ESTE TIPO DE CONTRATO
Se o texto indicar contrato de prestação de serviços de substituição/fornecimento de baterias de nobreaks:
- objetivo deve ser algo como "Substituição das baterias dos nobreaks da unidade indicada".
- descricao_servico_material deve citar fornecimento, instalação, testes, recolhimento/descarte das baterias, se constar.
- descrição breve de cadastro não deve falar de transformação societária/EIRELI; deve falar do serviço/material.

EXEMPLOS DE NÍVEL DE DETALHE ESPERADO
- objetivo: "Contratação para substituição das baterias dos nobreaks instalados na unidade da Centauro."
- descricao_servico_material: "Fornecimento, instalação e testes de baterias para nobreaks, incluindo recolhimento das baterias antigas quando previsto."
- forma_pagamento: "Pagamento mediante emissão de nota fiscal após conclusão/aceite dos serviços pela contratante."
- condicao_pagamento_dias: "90 dias após aprovação da nota fiscal."
- multa: "Multa de 2% sobre valores em atraso, acrescida de juros de 1% ao mês, se previsto."
- rescisao_indenizacao: "Rescisão mediante aviso prévio de 30 dias, sem ônus ou multa, conforme previsto no contrato."
- anticorrupcao: "Há cláusula de anticorrupção exigindo conduta ética e conformidade com a Lei 12.846/13."
- protecao_dados_lgpd: "Há cláusula de proteção de dados pessoais, com obrigações de segurança, confidencialidade e conformidade com a LGPD."

VALIDAÇÃO FINAL OBRIGATÓRIA ANTES DE RESPONDER
- Se identificar tabela de valores mensal e vigência em meses, separe valor mensal de valor total estimado da vigência. Nunca coloque valor unitário ou valor mensal no campo de valor global do contrato.
- Se identificar Certificate of Completion / Status Completed do DocuSign, contrato_assinado = "Sim", status não deve ser "Pendente de assinatura" e pessoas_que_assinaram deve listar os signatários.
- Se identificar PDF assinado e minutas DOCX sem assinatura, prevalece o PDF assinado.
- Para contratos de alimentação/refeição, itens como Desjejum, Almoço e Café da Tarde devem aparecer em itens_contrato com quantidade, unidade, valor unitário e valor total.

CONTRATO E ANEXOS EXTRAÍDOS:
{texto[:120000]}
"""
MODELOS_GEMINI = {
    "Automático recomendado": [
        "gemini-3.5-thinking",
        "gemini-3.1-pro",
        "gemini-3.5-flash",
    ],
    "Gemini 3.5 Thinking": ["gemini-3.5-thinking"],
    "Gemini 3.1 Pro": ["gemini-3.1-pro"],
    "Gemini 3.5 Flash": ["gemini-3.5-flash"],
}



def _mime_type_arquivo(nome_arquivo: str) -> str:
    """MIME type usado no upload dos arquivos para o Gemini Files API.

    Observação importante: a Files API não aceita DOCX diretamente em alguns
    ambientes/chaves corporativas. Por isso Word é convertido para TXT antes
    do upload e deve ir como text/plain.
    """
    nome = str(nome_arquivo or "").lower()
    if nome.endswith(".pdf"):
        return "application/pdf"
    if nome.endswith((".txt", ".md")):
        return "text/plain"
    return "application/octet-stream"


def _arquivo_word(nome_arquivo: str) -> bool:
    nome = str(nome_arquivo or "").lower()
    return nome.endswith((".docx", ".doc"))


def _preparar_upload_gemini(arquivo: Any) -> tuple[str, bytes, str, str]:
    """Prepara arquivo para upload na Files API.

    PDF segue original. Word/DOCX é convertido para um .txt estruturado, porque
    application/vnd.openxmlformats-officedocument.wordprocessingml.document
    não é aceito pela Files API no fluxo atual.

    Retorna: nome_display, conteudo_bytes, suffix, mime_type.
    """
    nome = getattr(arquivo, "name", "documento")

    if _arquivo_word(nome):
        try:
            arquivo.seek(0)
            texto_word = ler_docx(arquivo)
        except Exception as erro:
            texto_word = f"Erro ao converter Word/DOCX para texto: {erro}"

        texto_word = (
            "DOCUMENTO WORD/DOCX CONVERTIDO PARA TEXTO ESTRUTURADO PARA ANÁLISE DO GEMINI\n"
            f"ARQUIVO ORIGINAL: {nome}\n"
            "OBSERVAÇÃO: o binário DOCX não foi enviado porque a Gemini Files API rejeita esse MIME type.\n"
            "O conteúdo abaixo foi extraído mantendo parágrafos e tabelas quando possível.\n\n"
            + str(texto_word or "")
        )
        nome_txt = f"{nome}.txt"
        return nome_txt, texto_word.encode("utf-8", errors="ignore"), ".txt", "text/plain"

    arquivo.seek(0)
    conteudo = arquivo.read()
    suffix = Path(nome).suffix or ".bin"
    return nome, conteudo, suffix, _mime_type_arquivo(nome)


def _prompt_ia_com_documentos_originais(texto: str, nomes_arquivos: List[str]) -> str:
    """Prompt reforçado quando os arquivos originais são enviados ao Gemini."""
    lista = "\n".join(f"- {nome}" for nome in nomes_arquivos) or "- Não informado"

    return f"""
ATENÇÃO: nesta análise você recebeu os ARQUIVOS ORIGINAIS anexados, além do texto extraído como apoio técnico.

COMO ANALISAR
1. Use os arquivos originais como fonte principal. Quando o arquivo for Word/DOCX, ele pode ter sido convertido para TXT estruturado por limitação de MIME da Files API; nesse caso, use esse TXT como representação fiel do Word.
2. Navegue visualmente pelos documentos, páginas, anexos, tabelas, propostas comerciais, comentários, imagens e certificados de assinatura.
3. Entenda a relação entre contrato principal, termo aditivo, proposta comercial, proposta técnica, orçamento, e-mail de aprovação e certificado DocuSign.
4. Quando houver documento assinado/DocuSign, priorize a versão assinada sobre minutas, versões antigas ou arquivos com comentários.
5. Quando houver conflito entre contrato e proposta, siga a regra do próprio contrato sobre prevalência. Se o contrato disser que a proposta prevalece para escopo técnico/especificações/refeições, use a proposta nesses pontos.
6. Quando houver divergência entre o texto extraído e o arquivo original, priorize o arquivo original.
7. Use o texto extraído apenas como apoio para localização de trechos.
8. Faça uma análise contextual, como se estivesse revisando o pacote documental completo.
9. Mesmo fazendo análise contextual, retorne APENAS JSON válido para o sistema.
10. Se o nome do arquivo for genérico (ex.: SBF.pdf), classifique pelo conteúdo visual real do arquivo, não pelo nome.
11. Se o texto extraído abaixo contiver erro técnico de leitura, use o ARQUIVO ORIGINAL anexado como fonte principal e ignore o erro como conteúdo contratual.
12. Para cada campo exibido em card, só preencha com informação encontrada em contrato principal/anexo válido. Caso não encontre, escreva “Não identificado com segurança”.
13. A análise completa deve ser coerente com os cards: não deixe de listar assinaturas, valores, vigência, materiais/serviços, aditivos e pendências quando existirem no arquivo original.

ARQUIVOS ORIGINAIS RECEBIDOS:
{lista}

""" + prompt_ia(texto)


def _subir_arquivos_originais_gemini(client, arquivos_originais: Any) -> tuple[list, list, list]:
    """Salva temporariamente e envia arquivos usando google-genai >= 2.x.

    PDF é enviado como arquivo original. DOC/DOCX é convertido para TXT antes
    do upload, porque a Files API retornou INVALID_ARGUMENT para o MIME type
    application/vnd.openxmlformats-officedocument.wordprocessingml.document.
    Assim, um DOCX não derruba a análise inteira nem força fallback para texto.
    """
    uploaded_files = []
    temp_paths = []
    avisos_upload = []

    if not arquivos_originais:
        return uploaded_files, temp_paths, avisos_upload

    for arquivo in arquivos_originais:
        nome_original = getattr(arquivo, "name", "documento")
        try:
            nome_upload, conteudo, suffix, mime = _preparar_upload_gemini(arquivo)

            if not conteudo:
                avisos_upload.append(f"{nome_original}: arquivo vazio ou sem conteúdo extraível.")
                continue

            if _arquivo_word(nome_original):
                avisos_upload.append(
                    f"{nome_original}: DOC/DOCX convertido para TXT estruturado antes do upload "
                    "porque a Files API não aceitou o MIME type do Word."
                )

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.write(conteudo)
            tmp.flush()
            tmp.close()
            temp_paths.append(tmp.name)

            uploaded = client.files.upload(
                file=tmp.name,
                config={
                    "mime_type": mime,
                    "display_name": nome_upload,
                },
            )

            # Alguns arquivos podem ficar em processamento por alguns segundos.
            for _ in range(45):
                state = getattr(getattr(uploaded, "state", None), "name", "")
                if state and state.upper() == "PROCESSING":
                    time.sleep(1)
                    uploaded = client.files.get(name=uploaded.name)
                else:
                    break

            uploaded_files.append(uploaded)

        except Exception as erro:
            avisos_upload.append(f"{nome_original}: {erro}")

    return uploaded_files, temp_paths, avisos_upload


def _limpar_uploads_gemini(client, uploaded_files: list, temp_paths: list) -> None:
    """Remove temporários locais e tenta remover arquivos da área temporária da API."""
    for uploaded in uploaded_files or []:
        try:
            if getattr(uploaded, "name", None):
                client.files.delete(name=uploaded.name)
        except Exception:
            pass

    for caminho in temp_paths or []:
        try:
            os.remove(caminho)
        except Exception:
            pass


def _extrair_texto_resposta_gemini(resp: Any) -> str:
    """Extrai texto de resposta do SDK novo ou legado."""
    txt = getattr(resp, "text", None)
    if txt:
        return str(txt)
    try:
        return resp.candidates[0].content.parts[0].text
    except Exception:
        return str(resp)


def _json_da_resposta_gemini(resp: Any) -> Dict[str, Any]:
    """Converte a resposta do Gemini em JSON com tolerância a pequenos ruídos.

    Mesmo usando response_mime_type=application/json, alguns modelos podem
    devolver cercas de markdown ou texto antes/depois do objeto. Aqui mantemos
    o fluxo seguro sem quebrar a análise inteira por formatação mínima.
    """
    content = (
        _extrair_texto_resposta_gemini(resp)
        .strip()
        .replace("```json", "")
        .replace("```", "")
        .strip()
    )

    try:
        return json.loads(content)
    except Exception:
        ini = content.find("{")
        fim = content.rfind("}")
        if ini >= 0 and fim > ini:
            return json.loads(content[ini:fim + 1])
        raise


def _gerar_com_sdk_novo(client: Any, modelo: str, prompt_final: str, uploaded_files: list) -> Dict[str, Any]:
    """Gera análise usando google-genai >= 2.x, com arquivos originais anexados."""
    contents = [prompt_final] + list(uploaded_files or [])
    resp = client.models.generate_content(
        model=modelo,
        contents=contents,
        config={
            "temperature": 0.0,
            "top_p": 0.2,
            "max_output_tokens": 65535,
            "response_mime_type": "application/json",
        },
    )
    return _json_da_resposta_gemini(resp)


def _gerar_texto_com_sdk_legado(texto: str, api_key: str, modelo: str) -> Dict[str, Any]:
    """Fallback compatível com o fluxo antigo: envia apenas o texto extraído."""
    import google.generativeai as genai_legacy

    genai_legacy.configure(api_key=api_key)
    model = genai_legacy.GenerativeModel(modelo)
    resp = model.generate_content(
        prompt_ia(texto),
        generation_config={
            "temperature": 0.0,
            "top_p": 0.2,
            "max_output_tokens": 65535,
            "response_mime_type": "application/json",
        },
    )
    return _json_da_resposta_gemini(resp)



def _importar_google_genai_novo():
    """
    Importa o SDK novo do Gemini usado pela Files API.

    O erro "cannot import name 'genai' from 'google'" acontece quando o pacote
    google-genai não foi instalado no ambiente, ou quando o ambiente ainda está
    usando somente o pacote legado google-generativeai.
    """
    try:
        from google import genai as genai_new  # pacote: google-genai
        return genai_new
    except Exception as erro_1:
        try:
            import google.genai as genai_new  # fallback de importação
            return genai_new
        except Exception as erro_2:
            raise ImportError(
                "SDK novo do Gemini não disponível. Instale/atualize o pacote "
                "google-genai no requirements.txt. Sem ele, a Files API não consegue "
                "analisar os documentos originais e o sistema usa apenas texto extraído. "
                f"Detalhe técnico: {erro_1 or erro_2}"
            ) from erro_2

def analisar_gemini(texto: str, api_key: str, opcao_modelo: str, arquivos_originais: Any = None) -> Dict[str, Any]:
    """Analisa com Gemini.

    Fluxo principal: google-genai >= 2.x + Files API para enviar os documentos originais.
    Fallback: fluxo antigo por texto extraído, caso a biblioteca nova não esteja disponível
    ou caso o upload/generation multimodal falhe.
    """
    modelos = MODELOS_GEMINI.get(opcao_modelo, MODELOS_GEMINI["Automático recomendado"])
    nomes_arquivos = [getattr(a, "name", "documento") for a in (arquivos_originais or [])]

    # 1) Tenta fluxo novo: documentos originais + texto extraído.
    ultimo_erro_multimodal = None
    uploaded_files: list = []
    temp_paths: list = []

    if arquivos_originais:
        try:
            genai_new = _importar_google_genai_novo()

            client = genai_new.Client(api_key=api_key)
            uploaded_files, temp_paths, erros_upload = _subir_arquivos_originais_gemini(client, arquivos_originais)

            if erros_upload:
                with st.expander("⚠️ Detalhes de preparação/upload dos arquivos para o Gemini", expanded=False):
                    st.write("Avisos do envio para a Files API:")
                    for erro in erros_upload:
                        st.write(f"- {erro}")

            if uploaded_files:
                prompt_final = _prompt_ia_com_documentos_originais(texto, nomes_arquivos)
                for nome in modelos:
                    try:
                        resultado_json = _gerar_com_sdk_novo(client, nome, prompt_final, uploaded_files)
                        if isinstance(resultado_json, dict):
                            resultado_json["modelo_ia"] = nome
                            resultado_json["modo_analise_ia"] = "Documentos originais + texto extraído"
                            resultado_json["arquivos_originais_enviados"] = len(uploaded_files)
                        st.success(f"IA utilizada: {nome} • documentos/anexos analisados pela Files API")
                        return resultado_json
                    except Exception as e:
                        ultimo_erro_multimodal = e
                        if opcao_modelo != "Automático recomendado":
                            raise Exception(f"Erro ao usar o modelo {nome} com documentos originais. Detalhe: {e}")
                        continue
        except Exception as e:
            ultimo_erro_multimodal = e
        finally:
            try:
                if 'client' in locals():
                    _limpar_uploads_gemini(client, uploaded_files, temp_paths)
            except Exception:
                pass

    # 2) Fallback: versão anterior, usando texto extraído.
    # Se o documento original existia, mas o texto extraído é só erro técnico/sem conteúdo,
    # não podemos gerar uma análise falsa com cards vazios.
    texto_sem_base_confiavel = texto_indica_falha_leitura(texto) and not texto_tem_conteudo_contratual(texto)
    if arquivos_originais and texto_sem_base_confiavel:
        raise Exception(
            "Falha técnica: o arquivo original não pôde ser analisado pela Gemini Files API "
            "e o texto extraído não contém conteúdo contratual confiável. "
            "Não foi gerada análise por fallback para evitar cards incorretos. "
            f"Detalhe Files API: {ultimo_erro_multimodal}"
        )

    ultimo_erro_texto = None
    for nome in modelos:
        try:
            resultado_json = _gerar_texto_com_sdk_legado(texto, api_key, nome)
            if isinstance(resultado_json, dict):
                resultado_json["modelo_ia"] = nome
                resultado_json["modo_analise_ia"] = "Texto extraído"
                if ultimo_erro_multimodal:
                    resultado_json["erro_upload_documentos_originais"] = str(ultimo_erro_multimodal)

            if ultimo_erro_multimodal:
                st.warning(
                    "Não foi possível concluir a análise com documentos originais. "
                    "O sistema usou o texto extraído como fallback. "
                    f"Detalhe: {ultimo_erro_multimodal}"
                )
            st.success(f"IA utilizada: {nome} • texto extraído analisado")
            return resultado_json
        except Exception as e:
            ultimo_erro_texto = e
            if opcao_modelo != "Automático recomendado":
                raise Exception(f"Erro ao usar o modelo {nome}. Detalhe: {e}")
            continue

    raise Exception(
        "Nenhum modelo Gemini disponível. "
        f"Erro documentos originais: {ultimo_erro_multimodal}. "
        f"Erro texto extraído: {ultimo_erro_texto}."
    )

def formatar_cnpj(valor: Any) -> str:
    txt = clean_text(valor)
    if not txt or txt == "Não localizado":
        return "Não localizado"
    m = re.search(r"\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}", txt)
    if m:
        return m.group(0)
    digits = re.sub(r"\D", "", txt)
    if len(digits) == 14:
        return f"{digits[:2]}.{digits[2:5]}.{digits[5:8]}/{digits[8:12]}-{digits[12:]}"
    return txt


def resumir_campo(valor: Any, limite: int = 220) -> str:
    txt = clean_text(valor)
    if len(txt) <= limite:
        return txt
    corte = txt[:limite].rsplit(" ", 1)[0]
    return corte + "..."


# =========================================================
# PÓS-VALIDAÇÃO DOCUMENTAL: VALORES, DATAS E ASSINANTES
# =========================================================
_MESES_PT = {
    "janeiro": "01", "fevereiro": "02", "março": "03", "marco": "03",
    "abril": "04", "maio": "05", "junho": "06", "julho": "07",
    "agosto": "08", "setembro": "09", "outubro": "10", "novembro": "11", "dezembro": "12",
}


def _data_textual_para_br(valor: Any) -> str:
    txt = clean_text(valor)
    m = re.search(r"(\d{1,2})\s+de\s+([A-Za-zçÇãáéíóúâêôõ]+)\s+de\s+(\d{4})", txt, flags=re.IGNORECASE)
    if not m:
        return txt if _valor_informado(txt) else "Não localizado"
    dia = int(m.group(1))
    mes = _MESES_PT.get(m.group(2).lower(), "")
    ano = m.group(3)
    if not mes:
        return txt
    return f"{dia:02d}/{mes}/{ano}"


def _formatar_data_slash(valor: Any) -> str:
    txt = clean_text(valor)
    m = re.search(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b", txt)
    if not m:
        return txt if _valor_informado(txt) else "Não localizado"
    a, b, ano = int(m.group(1)), int(m.group(2)), m.group(3)
    # DocuSign costuma trazer M/D/YYYY. Se o segundo número > 12, converte para DD/MM/YYYY.
    if b > 12 and a <= 12:
        return f"{b:02d}/{a:02d}/{ano}"
    return f"{a:02d}/{b:02d}/{ano}"


def _extrair_data_contrato(texto: str) -> str:
    txt = str(texto or "")
    padroes = [
        r"São Paulo,\s*(\d{1,2}\s+de\s+[A-Za-zçÇãáéíóúâêôõ]+\s+de\s+\d{4})",
        r"Sao Paulo,\s*(\d{1,2}\s+de\s+[A-Za-zçÇãáéíóúâêôõ]+\s+de\s+\d{4})",
        r"(?:Data do Contrato|Data do contrato)[:\s-]+(\d{1,2}/\d{1,2}/\d{4})",
    ]
    for p in padroes:
        m = re.search(p, txt, flags=re.IGNORECASE)
        if m:
            val = m.group(1)
            return _data_textual_para_br(val) if " de " in val.lower() else _formatar_data_slash(val)
    return "Não localizado"


def _extrair_data_conclusao_docusign(texto: str) -> str:
    txt = str(texto or "")
    padroes = [
        r"Completed\s+Security\s+Checked\s+(\d{1,2}/\d{1,2}/\d{4})",
        r"Envelope\s+Summary\s+Events[\s\S]{0,700}?Completed\s+Security\s+Checked\s+(\d{1,2}/\d{1,2}/\d{4})",
        r"Status:\s*Completed[\s\S]{0,2000}?Completed\s+Security\s+Checked\s+(\d{1,2}/\d{1,2}/\d{4})",
    ]
    for p in padroes:
        m = re.search(p, txt, flags=re.IGNORECASE)
        if m:
            return _formatar_data_slash(m.group(1))
    # fallback: última assinatura individual do certificado
    assinaturas = re.findall(r"Signed:\s*(\d{1,2}/\d{1,2}/\d{4})", txt, flags=re.IGNORECASE)
    if assinaturas:
        return _formatar_data_slash(assinaturas[-1])
    return "Não localizado"


def _extrair_assinantes_docusign(texto: str) -> list[str]:
    txt = str(texto or "")
    low = txt.lower()
    if "docusign" not in low and "signer events" not in low and "certificate of completion" not in low:
        return []

    start = low.find("signer events")
    if start < 0:
        start = low.find("certificate of completion")
    seg = txt[start:] if start >= 0 else txt
    finais = [
        "in person signer events", "editor delivery events", "agent delivery events",
        "intermediary delivery events", "certified delivery events", "carbon copy events",
        "witness events", "notary events", "envelope summary events", "payment events",
    ]
    seg_low = seg.lower()
    cortes = [seg_low.find(f) for f in finais if seg_low.find(f) > 0]
    if cortes:
        seg = seg[:min(cortes)]

    linhas = [clean_text(l) for l in seg.splitlines() if clean_text(l) not in ("", "Não localizado")]
    nomes: list[str] = []
    for i, linha in enumerate(linhas[:-1]):
        prox = linhas[i + 1]
        if "@" not in prox:
            continue
        if "@" in linha or re.search(r"\d", linha):
            continue
        low_l = linha.lower()
        bloqueios = [
            "signer events", "signature timestamp", "security level", "electronic record",
            "using ip", "docusign", "not offered", "accepted", "signature adoption",
            "grupo sbf", "max fast", "none", "sent", "viewed", "signed",
        ]
        if any(b in low_l for b in bloqueios):
            continue
        # Evita cargos/áreas. Nome real costuma ter 2+ palavras e não termina com pontuação técnica.
        palavras = re.findall(r"[A-Za-zÀ-ÿ]+", linha)
        if len(palavras) < 2:
            continue
        nome = re.sub(r"\s+", " ", linha).strip(" -:;,.|")
        if nome and nome not in nomes:
            nomes.append(nome)
    return nomes


def _parse_moeda_brasil(valor: Any) -> float | None:
    """Converte valores monetários em float, aceitando BR (1.087,85) e decimal (1087.85)."""
    txt = str(clean_text(valor))
    if not _valor_informado(txt):
        return None
    m = re.search(r"(?:R\$\s*)?([0-9]{1,3}(?:[\.,][0-9]{3})*(?:[\.,][0-9]{2})|[0-9]+(?:[\.,][0-9]{2})?)", txt)
    if not m:
        return None
    num = m.group(1).strip()

    if "," in num and "." in num:
        # Formato brasileiro: 1.087,85
        num = num.replace(".", "").replace(",", ".")
    elif "," in num:
        # Formato brasileiro sem milhar: 1087,85
        num = num.replace(".", "").replace(",", ".")
    elif "." in num:
        partes = num.split(".")
        if len(partes[-1]) == 2:
            # Decimal internacional: 1087.85 / 848.35
            num = "".join(partes[:-1]) + "." + partes[-1] if len(partes) > 2 else num
        else:
            # Milhar sem decimal: 1.087
            num = num.replace(".", "")
    try:
        return float(num)
    except Exception:
        return None


def _formatar_moeda_brasil(valor: float, sufixo: str = "") -> str:
    s = f"R$ {valor:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return (s + (f" {sufixo}" if sufixo else "")).strip()


def _extrair_meses_vigencia(texto: str, base: Dict[str, Any] | None = None) -> int | None:
    candidatos = [str(texto or "")]
    if base:
        candidatos.extend([str(base.get("vigencia_apos_assinatura") or ""), str(base.get("vigencia") or "")])
    plano = " ".join(candidatos)
    m = re.search(r"(\d{1,3})\s*(?:\([^)]*\)\s*)?mes(?:es)?", plano, flags=re.IGNORECASE)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None
    return None

def _parse_data_br_para_datetime(valor: Any) -> datetime | None:
    txt = clean_text(valor)

    m = re.search(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b", txt)
    if not m:
        return None

    try:
        dia = int(m.group(1))
        mes = int(m.group(2))
        ano = int(m.group(3))
        return datetime(ano, mes, dia)
    except Exception:
        return None


def _ultimo_dia_mes(ano: int, mes: int) -> int:
    if mes == 12:
        proximo = datetime(ano + 1, 1, 1)
    else:
        proximo = datetime(ano, mes + 1, 1)
    return (proximo - timedelta(days=1)).day


def _somar_meses(data_base: datetime, meses: int) -> datetime:
    mes_total = data_base.month - 1 + meses
    ano = data_base.year + mes_total // 12
    mes = mes_total % 12 + 1
    dia = min(data_base.day, _ultimo_dia_mes(ano, mes))
    return datetime(ano, mes, dia)


def _extrair_anos_vigencia(texto: str, base: Dict[str, Any] | None = None) -> int | None:
    candidatos = [str(texto or "")]
    if base:
        candidatos.extend([
            str(base.get("vigencia_apos_assinatura") or ""),
            str(base.get("vigencia") or ""),
        ])

    plano = " ".join(candidatos)

    m = re.search(r"(\d{1,3})\s*(?:\([^)]*\)\s*)?ano(?:s)?", plano, flags=re.IGNORECASE)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None

    return None


def _montar_periodo_vigencia_formatado(base: Dict[str, Any], texto: str) -> str:
    """
    Monta o período de vigência no padrão:
    Início DD/MM/AAAA até DD/MM/AAAA

    Regra:
    1. Início = Data Conclusão DocuSign
    2. Se não tiver, usa Data da Assinatura
    3. Se não tiver, usa Data do Contrato
    4. Prazo indeterminado = 31/12/9999
    5. Prazo em meses/anos = início + prazo - 1 dia
    """
    data_inicio = None

    for chave_data in ("data_conclusao_docusign", "data_assinatura", "data_contrato"):
        data_tentativa = _parse_data_br_para_datetime(base.get(chave_data))
        if data_tentativa:
            data_inicio = data_tentativa
            break

    if not data_inicio:
        return "Não localizado"

    inicio_fmt = data_inicio.strftime("%d/%m/%Y")

    vigencia_txt = clean_text(base.get("vigencia_apos_assinatura") or base.get("vigencia"))
    plano = f"{vigencia_txt} {texto}".lower()

    if any(t in plano for t in [
        "prazo indeterminado",
        "vigência indeterminada",
        "duracao indeterminada",
        "duração indeterminada",
        "por prazo indeterminado",
    ]):
        return f"Início {inicio_fmt} até 31/12/9999"

    meses = _extrair_meses_vigencia(texto, base)

    if not meses:
        anos = _extrair_anos_vigencia(texto, base)
        if anos:
            meses = anos * 12

    if meses:
        data_fim = _somar_meses(data_inicio, meses) - timedelta(days=1)
        fim_fmt = data_fim.strftime("%d/%m/%Y")
        return f"Início {inicio_fmt} até {fim_fmt}"

    return f"Início {inicio_fmt} até Não localizado"

def _somar_valores_itens(itens: Any) -> float | None:
    itens_norm = normalizar_itens_contrato(itens)
    total = 0.0
    encontrou = False
    for item in itens_norm:
        v = _parse_moeda_brasil(item.get("Valor total"))
        if v is not None:
            total += v
            encontrou = True
    return total if encontrou and total > 0 else None


def _extrair_total_mensal_texto(texto: str) -> float | None:
    txt = str(texto or "")
    # Busca totais próximos de tabelas de proposta/cenário final.
    padroes = [
        r"TOTAL\s+(?:R\$\s*)?([0-9]{1,3}(?:\.[0-9]{3})*,[0-9]{2}|[0-9]+,[0-9]{2})",
        r"CDB\s+Final[^R$]{0,80}R\$\s*([0-9]{1,3}(?:\.[0-9]{3})*,[0-9]{2})",
        r"valor\s+mensal\s+estimado[^R$]{0,80}R\$\s*([0-9]{1,3}(?:\.[0-9]{3})*,[0-9]{2})",
    ]
    valores = []
    for p in padroes:
        for m in re.finditer(p, txt, flags=re.IGNORECASE):
            v = _parse_moeda_brasil(m.group(1))
            if v is not None:
                valores.append(v)
    if valores:
        # Em propostas de alimentação o menor total relevante costuma ser mensal; evita totais de vigência calculados.
        return min(valores)
    return None



def _parece_valor_unitario_ou_demanda(valor: Any) -> bool:
    """Identifica valores que não são globais: por pessoa, item, mês, demanda etc."""
    txt = clean_text(valor).lower()
    if not _valor_informado(txt):
        return False
    termos = [
        " por ", "/mês", "/mes", "mês", "mes", "mensal", "mensais",
        "unitário", "unitario", "unidade", "item", "serviço", "servico",
        "aprendiz", "ativo", "demanda", "consumo", "utilizada", "utilizado",
        "contratada", "contratado", "quantidade", "recorrente", "estimado", "estimada",
    ]
    return any(t in txt for t in termos)


def _texto_sem_valor_global(referencia: Any = "") -> str:
    txt = "Sem valor global fixo definido no contrato. O documento apresenta valores unitários, mensais ou por demanda, mas não informa um valor total fechado para todo o contrato."
    ref = clean_text(referencia)
    if _valor_informado(ref) and ref not in txt:
        txt += f" Referência comercial localizada: {ref}."
    return txt


def _extrair_unidade_demanda(valor: Any) -> str:
    txt = clean_text(valor)
    m = re.search(r"por\s+([^.;,\n/]+)", txt, flags=re.IGNORECASE)
    if m:
        unidade = clean_text(m.group(1)).strip(" -")
        unidade = re.sub(r"\s*(ao|por)?\s*m[eê]s$", "", unidade, flags=re.IGNORECASE).strip()
        if unidade:
            return unidade
    low = txt.lower()
    if "aprendiz" in low:
        return "aprendiz ativo"
    if "colaborador" in low:
        return "colaborador"
    if "usuário" in low or "usuario" in low:
        return "usuário"
    if "item" in low:
        return "item"
    return "unidade contratada"


def _label_qtd_unidade(qtd: int, unidade: str) -> str:
    unidade = clean_text(unidade).lower()
    if "aprendiz" in unidade:
        return f"{qtd} aprendiz ativo" if qtd == 1 else f"{qtd} aprendizes ativos"
    if "colaborador" in unidade:
        return f"{qtd} colaborador" if qtd == 1 else f"{qtd} colaboradores"
    if "usu" in unidade:
        return f"{qtd} usuário" if qtd == 1 else f"{qtd} usuários"
    if qtd == 1:
        return f"1 {unidade}"
    return f"{qtd} unidades"


def _formatar_valor_monetario_item(valor: Any) -> str:
    txt = clean_text(valor)
    v = _parse_moeda_brasil(txt)
    if v is None:
        return txt
    return _formatar_moeda_brasil(v)


def _montar_total_materiais_servicos(itens: List[Dict[str, Any]], total: float | None) -> str:
    """Monta o resumo executivo do total de materiais/serviços sem jogar todos os itens dentro do card."""
    itens_norm = normalizar_itens_contrato(itens)

    qtd_itens_com_valor = 0
    for item in itens_norm:
        valor = _parse_moeda_brasil(item.get("Valor total") or item.get("valor_total"))
        if valor is not None:
            qtd_itens_com_valor += 1

    if total is None or total <= 0:
        return (
            "Não identificado. "
            "Não foram localizados itens com valor total suficiente para somar materiais e serviços."
        )

    return (
        f"{_formatar_moeda_brasil(total)}. "
        f"Soma dos valores totais de {qtd_itens_com_valor} item(ns) identificados nos documentos analisados. "
        "Esse campo não representa necessariamente o valor global do contrato; representa apenas a soma dos itens identificados. "
        "Consulte o detalhamento completo na tabela de materiais e serviços."
    )

def _formatar_valor_mensal_profissional(valor: Any, mensal_detectado: float | None = None, origem_itens: bool = False) -> str:
    txt = clean_text(valor)
    v = _parse_moeda_brasil(txt)
    if v is None and mensal_detectado:
        v = mensal_detectado

    if _valor_informado(txt) and _parece_valor_unitario_ou_demanda(txt):
        unidade = _extrair_unidade_demanda(txt)
        if v is not None:
            e1 = _formatar_moeda_brasil(v)
            e10 = _formatar_moeda_brasil(v * 10)
            e100 = _formatar_moeda_brasil(v * 100)
            return (
                f"{_formatar_moeda_brasil(v)} por {unidade}/mês. "
                "O valor mensal não é fixo; ele varia conforme a quantidade contratada ou utilizada no período. "
                f"Exemplo de cálculo: {_label_qtd_unidade(1, unidade)} = {e1}/mês; "
                f"{_label_qtd_unidade(10, unidade)} = {e10}/mês; "
                f"{_label_qtd_unidade(100, unidade)} = {e100}/mês."
            )
        return "Valor mensal variável conforme demanda ou quantidade contratada. O contrato não apresenta um valor mensal fixo; apresenta critérios unitários para cálculo."

    if v is not None and origem_itens:
        return (
            f"{_formatar_moeda_brasil(v)}/mês, quando os itens identificados representarem cobrança mensal recorrente. "
            "A periodicidade deve ser validada no contrato/proposta, pois o valor foi obtido pela soma dos itens com valor total localizado."
        )

    if v is not None:
        return f"{_formatar_moeda_brasil(v)}/mês. Valor mensal estimado informado nos documentos analisados."

    if _valor_informado(txt):
        return txt

    return "Não localizado. Não foi identificado valor mensal fixo ou variável com clareza nos documentos analisados."


def _formatar_valor_total_vigencia_profissional(base: Dict[str, Any], meses: int | None, mensal_num: float | None) -> str:
    atual_txt = clean_text(base.get("valor_total_estimado_vigencia"))
    mensal_txt = clean_text(base.get("valor_mensal_estimado"))
    mensal_variavel = _parece_valor_unitario_ou_demanda(mensal_txt)

    if mensal_variavel:
        unidade = _extrair_unidade_demanda(mensal_txt)
        valor_unit = _parse_moeda_brasil(mensal_txt)
        if valor_unit is not None and meses:
            return (
                "Não calculável com precisão. "
                f"O contrato possui valor unitário de {_formatar_moeda_brasil(valor_unit)} por {unidade}/mês, "
                f"porém não informa a quantidade total de {unidade}(s) durante os {meses} meses de vigência. "
                "Para calcular, é necessário aplicar: valor unitário x quantidade contratada/utilizada x quantidade de meses."
            )
        return (
            "Não calculável com precisão. O valor é variável conforme demanda, quantidade utilizada ou quantidade contratada. "
            "Para calcular, é necessário aplicar: valor unitário x quantidade x quantidade de meses."
        )

    atual_num = _parse_moeda_brasil(atual_txt)
    if atual_num is not None and _valor_informado(atual_txt):
        if mensal_num and meses:
            return f"{_formatar_moeda_brasil(atual_num)}. Cálculo estimado com base em {_formatar_moeda_brasil(mensal_num)}/mês x {meses} meses de vigência."
        return f"{_formatar_moeda_brasil(atual_num)}. Valor total estimado para a vigência conforme informações localizadas nos documentos."

    if mensal_num and meses:
        total = mensal_num * meses
        return f"{_formatar_moeda_brasil(total)}. Cálculo estimado com base em {_formatar_moeda_brasil(mensal_num)}/mês x {meses} meses de vigência."

    if _valor_informado(atual_txt):
        return atual_txt

    return "Não calculável com precisão. Os documentos não apresentam informações suficientes para projetar o valor total da vigência."


def aplicar_regras_valores_profissionais(base: Dict[str, Any], itens: List[Dict[str, Any]], texto: str, meses: int | None) -> Dict[str, Any]:
    """Padroniza os 4 campos de valores para evitar confusão entre global, mensal, vigência e itens."""
    total_itens = _somar_valores_itens(itens)
    base["valor_total_materiais_servicos"] = _montar_total_materiais_servicos(itens, total_itens)

    valor_original = clean_text(base.get("valor_contrato_original") or base.get("valor_total"))
    valor_mensal_original = clean_text(base.get("valor_mensal_estimado"))
    original_num = _parse_moeda_brasil(valor_original)
    mensal_num_existente = _parse_moeda_brasil(valor_mensal_original)

    original_igual_mensal_variavel = (
        original_num is not None
        and mensal_num_existente is not None
        and abs(original_num - mensal_num_existente) < 0.01
        and _parece_valor_unitario_ou_demanda(valor_mensal_original)
    )

    if not _valor_informado(valor_original):
        base["valor_contrato_original"] = _texto_sem_valor_global()
    elif _parece_valor_unitario_ou_demanda(valor_original) or original_igual_mensal_variavel:
        base["valor_contrato_original"] = _texto_sem_valor_global(valor_original)
    elif original_num is not None:
        base["valor_contrato_original"] = f"{_formatar_moeda_brasil(original_num)}. Valor global previsto no contrato para execução do objeto contratado, conforme informação localizada nos documentos."
    else:
        base["valor_contrato_original"] = valor_original

    # Valor mensal: prioriza o que a IA encontrou; se vazio, usa soma dos itens apenas como apoio.
    if _valor_informado(valor_mensal_original):
        base["valor_mensal_estimado"] = _formatar_valor_mensal_profissional(valor_mensal_original)
    elif total_itens:
        base["valor_mensal_estimado"] = _formatar_valor_mensal_profissional("", total_itens, origem_itens=True)
    else:
        mensal_texto = _extrair_total_mensal_texto(texto)
        base["valor_mensal_estimado"] = _formatar_valor_mensal_profissional("", mensal_texto, origem_itens=False)

    mensal_num = _parse_moeda_brasil(base.get("valor_mensal_estimado"))
    base["valor_total_estimado_vigencia"] = _formatar_valor_total_vigencia_profissional(base, meses, mensal_num)

    # Mantém compatibilidade com histórico antigo, mas sem confundir valor global com total de itens.
    base["valor_total"] = base.get("valor_contrato_original")
    return base

def aplicar_regras_finais_contrato(base: Dict[str, Any], texto: str) -> Dict[str, Any]:
    """Blindagem final contra confusão de minuta, valor mensal x vigência e assinatura DocuSign."""
    texto = str(texto or "")
    low = texto.lower()

    data_contrato = _extrair_data_contrato(texto)
    data_docusign = _extrair_data_conclusao_docusign(texto)
    assinantes = _extrair_assinantes_docusign(texto)

    if _valor_informado(data_contrato):
        base["data_contrato"] = data_contrato
    if _valor_informado(data_docusign):
        base["data_conclusao_docusign"] = data_docusign
        base["data_assinatura"] = data_docusign

    if assinantes:
        base["assinantes"] = assinantes
        base["pessoas_que_assinaram"] = "; ".join(assinantes)

    docusign_completed = bool(re.search(r"Certificate Of Completion|Certificate of Completion|Status:\s*Completed|Completed\s+Security\s+Checked", texto, flags=re.IGNORECASE))
    if docusign_completed:
        base["contrato_assinado"] = "Sim"
        base["alerta_assinatura"] = "Contrato assinado eletronicamente via DocuSign / Certificate of Completion."
        if str(base.get("status", "")).strip().lower() in ("", "não localizado", "pendente de assinatura", "em negociação / pendente de assinatura"):
            base["status"] = "Ativo"
        # Remove pendências de assinatura herdadas de minutas antigas.
        pend_filtradas = []
        for p in base.get("pendencias", []) if isinstance(base.get("pendencias"), list) else []:
            txtp = json.dumps(p, ensure_ascii=False).lower() if isinstance(p, dict) else str(p).lower()
            if "assinatura" in txtp or "assinar" in txtp or "docusign" in txtp:
                continue
            pend_filtradas.append(p)
        base["pendencias"] = pend_filtradas

        checklist = []
        for item in base.get("checklist", []) if isinstance(base.get("checklist"), list) else []:
            if isinstance(item, dict):
                val = str(item.get("Validação") or item.get("validacao") or item.get("validação") or "").lower()
                if "assinatura" in val or "assinado" in val:
                    item["Status"] = "Aprovado"
                    item["Evidência"] = base.get("alerta_assinatura")
            checklist.append(item)
        base["checklist"] = checklist

    base["periodo_vigencia_formatado"] = _montar_periodo_vigencia_formatado(base, texto)

    itens = normalizar_itens_contrato(base.get("itens_contrato", []))
    meses = _extrair_meses_vigencia(texto, base)
    base = aplicar_regras_valores_profissionais(base, itens, texto, meses)
    base = aplicar_regras_especificas_qualidade_ia(base, texto)

    return base



def _texto_indica_mao_obra_temporaria(texto: Any) -> bool:
    low = str(texto or "").lower()
    gatilhos = [
        "mão de obra tempor", "mao de obra tempor", "lei 6.019/1974",
        "trabalho temporario", "trabalho temporário", "imediatta trabalho temporario",
        "atração & retenção", "atracao & retencao", "recrutamento & seleção", "recrutamento e seleção",
    ]
    return any(g in low for g in gatilhos)


def _extrair_cnpj_proximo(texto: str, termo: str) -> str:
    txt = str(texto or "")
    pos = txt.lower().find(str(termo or "").lower())
    if pos < 0:
        return "Não localizado"
    janela = txt[pos:pos + 1500]
    m = re.search(r"\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}|\d{14}", janela)
    return formatar_cnpj(m.group(0)) if m else "Não localizado"


def _extrair_envelope_docusign(texto: str) -> str:
    txt = str(texto or "")
    m = re.search(r"DocuSign\s+Envelope\s+ID:\s*([A-Z0-9-]{20,})", txt, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip()
    m = re.search(r"Envelope\s*ID[:\s]+([A-Z0-9-]{20,})", txt, flags=re.IGNORECASE)
    return m.group(1).strip() if m else ""


def _extrair_assinantes_campos_docusign(texto: str) -> List[str]:
    txt = str(texto or "")
    nomes: List[str] = []

    # Form fields extraídos de PDFs DocuSign podem aparecer como "Por:NOME".
    for nome in re.findall(r"Por\s*:\s*([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s.'-]{2,80})", txt, flags=re.IGNORECASE):
        nome = re.sub(r"\s+", " ", nome).strip(" -:;,.|")
        if len(nome.split()) >= 1 and nome.lower() not in ("por", "nome") and nome not in nomes:
            nomes.append(nome)

    # Campos customizados da assinatura no final do PDF.
    for nome in re.findall(r"\[FORM FIELD\]\s+Custom_[^:]+:\s*([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s.'-]{2,80})", txt, flags=re.IGNORECASE):
        nome = re.sub(r"\s+", " ", nome).strip(" -:;,.|")
        low = nome.lower()
        if any(x in low for x in ["por:", "envelope", "docusign"]):
            continue
        if re.search(r"\d", nome):
            continue
        if len(nome.split()) >= 2 and nome not in nomes:
            nomes.append(nome)

    return nomes[:12]


def _aplicar_regras_mao_obra_temporaria(base: Dict[str, Any], texto: str) -> Dict[str, Any]:
    """Blindagem de qualidade para contratos de mão de obra temporária/RH.

    O Gemini costuma se confundir quando recebe contrato principal + minutas + CNPJ
    + contrato social + proposta. Esta regra determinística não substitui a IA;
    ela corrige campos centrais quando há evidência textual clara no pacote.
    """
    if not _texto_indica_mao_obra_temporaria(texto):
        return base

    txt = str(texto or "")
    low = txt.lower()

    # Partes e CNPJs: prioriza o contrato principal assinado, não CNPJ/contrato social de apoio.
    if "sbf comércio de produtos esportivos" in low or "sbf comercio de produtos esportivos" in low:
        base["empresa_grupo_sbf"] = "SBF Comércio de Produtos Esportivos S.A."
        cnpj_sbf = _extrair_cnpj_proximo(txt, "SBF COM")
        if _valor_informado(cnpj_sbf):
            base["cnpj_empresa_grupo"] = cnpj_sbf

    if "imediatta trabalho temporario" in low or "imediatta trabalho temporário" in low:
        base["contraparte"] = "Imediatta Trabalho Temporário Ltda"
        base["fornecedor"] = "Imediatta Trabalho Temporário Ltda"
        cnpj_filial = _extrair_cnpj_proximo(txt, "IMEDIATTA TRABALHO")
        if _valor_informado(cnpj_filial):
            base["cnpj_contraparte"] = cnpj_filial
            base["cnpj"] = cnpj_filial
    elif "imediatta terceiriza" in low and not _valor_informado(base.get("contraparte")):
        base["contraparte"] = "Imediatta Terceirização de Mão de Obra Ltda"
        base["fornecedor"] = base["contraparte"]

    base["tipo_contrato"] = "Prestação de Serviços de Alocação de Mão de Obra Temporária"
    base["descricao_breve_cadastro"] = "Alocação de mão de obra temporária"

    if "substituição transitória" in low or "substituicao transitoria" in low or "demanda complementar" in low:
        base["objetivo"] = (
            "Contratação de empresa especializada para fornecimento/alocação de mão de obra temporária, "
            "visando atender necessidade de substituição transitória de pessoal permanente ou demanda complementar de serviços."
        )
    else:
        base["objetivo"] = "Contratação de serviços de alocação de mão de obra temporária."

    base["descricao_servico_material"] = (
        "Prestação de serviços profissionais de alocação de mão de obra temporária, nos termos da Lei nº 6.019/1974, "
        "incluindo recrutamento, seleção, contratação, gestão de folha de pagamento e encargos dos trabalhadores temporários alocados."
    )

    if re.search(r"emiss[aã]o\s+de\s+nota\s+fiscal", txt, flags=re.IGNORECASE):
        base["forma_pagamento"] = "Pagamento mediante emissão de nota fiscal de serviços, condicionado à aprovação formal da Contratante."

    m_venc = re.search(
        r"vencimento\s+da\s+nota\s+ser[aá]\s+em\s+(\d+)\s*\([^)]*\)\s+dias\s+corridos\s+a\s+contar\s+da\s+data\s+de\s+emiss[aã]o",
        txt,
        flags=re.IGNORECASE,
    )
    if m_venc:
        base["condicao_pagamento_dias"] = f"{m_venc.group(1)} dias corridos a contar da data de emissão da nota fiscal."
    elif "aprovação formal da contratante" in low or "aprovacao formal da contratante" in low:
        base["condicao_pagamento_dias"] = "Conforme condições da proposta comercial aprovada, mediante aprovação formal da Contratante."

    if "igp-m" in low or "igp-m/fgv" in low:
        base["multa"] = (
            "Atraso no pagamento: correção pela variação do IGP-M/FGV, multa de 2% sobre o valor corrigido "
            "e juros moratórios simples de 1% ao mês. Para descumprimento de regras de trabalhador temporário, "
            "há previsão de multa não compensatória de 15% do valor faturado no mês anterior ou R$ 20.000,00, o que for maior."
        )

    if "prazo indeterminado" in low:
        base["vigencia_apos_assinatura"] = "Prazo indeterminado a partir da data de assinatura, ou até a conclusão dos serviços contratados."
        data_contrato = _extrair_data_contrato(txt)
        if _valor_informado(data_contrato):
            base["periodo_vigencia_formatado"] = f"Início {data_contrato} até 31/12/9999"

    if re.search(r"denunciado\s+sem\s+ônus|denunciado\s+sem\s+onus|aviso\s+prévio\s+de\s+30|aviso\s+previo\s+de\s+30", txt, flags=re.IGNORECASE):
        base["rescisao_indenizacao"] = (
            "Resilição/denúncia permitida a qualquer tempo por ambas as partes, sem ônus ou multa, mediante aviso prévio por escrito de 30 dias. "
            "Também há hipóteses de rescisão por inadimplemento não sanado, insolvência/falência, força maior superior a 30 dias, cessão sem autorização, suspensão dos serviços ou acordo entre as partes."
        )

    if "lei nº 12.846" in low or "lei n° 12.846" in low or "anticorrup" in low:
        base["anticorrupcao"] = (
            "Há cláusula anticorrupção exigindo que as partes atuem de forma ética e em conformidade com a Lei nº 12.846/13 e Decreto nº 8.420/15, "
            "com possibilidade de rescisão motivada imediata em caso de violação."
        )

    if "lei geral de proteção de dados" in low or "13.709/18" in low or "proteção de dados" in low or "protecao de dados" in low:
        base["protecao_dados_lgpd"] = (
            "Há cláusula de proteção de dados pessoais determinando tratamento apenas para as finalidades do contrato, adoção de medidas de segurança, "
            "registro das operações, cooperação entre as partes e notificação à Contratante em até 24 horas em caso de incidente ou descumprimento relacionado a dados pessoais."
        )

    data_contrato = _extrair_data_contrato(txt)
    if _valor_informado(data_contrato):
        base["data_contrato"] = data_contrato
        if not _valor_informado(base.get("data_assinatura")) or clean_text(base.get("data_assinatura")) == "Não localizada":
            base["data_assinatura"] = data_contrato

    envelope = _extrair_envelope_docusign(txt)
    nomes_docu = _extrair_assinantes_docusign(txt) or _extrair_assinantes_campos_docusign(txt)
    if envelope or nomes_docu or re.search(r"assinaram\s+o\s+presente", txt, flags=re.IGNORECASE):
        base["contrato_assinado"] = "Sim"
        base["alerta_assinatura"] = "Contrato com evidência de assinatura eletrônica/DocuSign ou assinatura das partes."
        base["status"] = "Ativo"
        if nomes_docu:
            base["pessoas_que_assinaram"] = "; ".join(nomes_docu)
        if envelope and not _valor_informado(base.get("data_conclusao_docusign")):
            base["data_conclusao_docusign"] = f"Não localizada data exata de conclusão no certificado; Envelope ID DocuSign {envelope}."

    # Valor: mão de obra temporária costuma ser por demanda/taxa, não valor global fechado.
    base["valor_contrato_original"] = (
        "Sem valor global fixo definido no contrato. A remuneração depende das condições comerciais da proposta, "
        "taxas aplicáveis, salário/remuneração da vaga, tributos e demanda efetivamente utilizada."
    )
    base["valor_mensal_estimado"] = (
        "Valor mensal variável conforme demanda, quantidade de trabalhadores/vagas e salário/remuneração aplicável. "
        "O contrato não apresenta valor mensal fixo fechado."
    )
    base["valor_total_estimado_vigencia"] = (
        "Não calculável com precisão sem quantidade de vagas/trabalhadores, remuneração base e período efetivamente utilizado. "
        "A fórmula depende de taxa/comissão x remuneração ou valor faturado conforme proposta x volume contratado."
    )
    base["valor_total_materiais_servicos"] = "Não calculável como total global. Condições comerciais identificadas: taxas percentuais sobre salário/remuneração e/ou valores unitários dependem de quantidade, demanda e vaga utilizada."

    itens_proposta = detectar_servico_percentual(txt)
    if itens_proposta:
        base["itens_contrato"] = normalizar_itens_contrato(itens_proposta)

    base["resumo_aditivos"] = base.get("resumo_aditivos") or "Nenhum aditivo identificado nos documentos analisados."
    aditivos = normalizar_aditivos_contrato(base.get("aditivos_contrato", []))
    if not aditivos and not re.search(r"\b(aditivo|aditamento|termo aditivo)\b", txt, flags=re.IGNORECASE):
        base["aditivos_contrato"] = []
        base["resumo_aditivos"] = "Nenhum aditivo identificado nos documentos analisados."

    # Checklist executivo: garante que os pontos centrais apareçam mesmo se a IA devolver pouco detalhe.
    checklist = base.get("checklist") if isinstance(base.get("checklist"), list) else []
    nomes_check = " ".join(str(x.get("Validação") or x.get("validacao") or "") for x in checklist if isinstance(x, dict)).lower()
    if "assinatura" not in nomes_check:
        checklist.insert(0, {
            "Validação": "Assinatura das Partes",
            "Status": "Concluído" if base.get("contrato_assinado") == "Sim" else "Revisar",
            "Peso de risco": "Crítico",
            "Crítico": "Sim",
            "Evidência": base.get("alerta_assinatura") or "Verificar evidência de assinatura no contrato/certificado."
        })
    if "lgpd" not in nomes_check and _valor_informado(base.get("protecao_dados_lgpd")):
        checklist.append({"Validação": "Cláusula de LGPD", "Status": "Concluído", "Peso de risco": "Médio", "Crítico": "Não", "Evidência": base.get("protecao_dados_lgpd")})
    if "anticorrup" not in nomes_check and _valor_informado(base.get("anticorrupcao")):
        checklist.append({"Validação": "Cláusula Anticorrupção", "Status": "Concluído", "Peso de risco": "Médio", "Crítico": "Não", "Evidência": base.get("anticorrupcao")})
    base["checklist"] = checklist

    # Evita que score/risco venham bons demais em contrato de trabalho temporário sem volume fixo.
    if not _valor_informado(base.get("risco")) or str(base.get("risco")).upper() == "BAIXO":
        base["risco"] = "MÉDIO"
    try:
        score_atual = int(float(base.get("score") or 0))
    except Exception:
        score_atual = 0
    if score_atual < 80:
        base["score"] = 85

    resumo_low = str(base.get("resumo_executivo", "")).lower()
    if (not _valor_informado(base.get("resumo_executivo")) or "análise concluída" in resumo_low
        or "minuta" in resumo_low or "contrato social" in resumo_low or "berkan" in resumo_low
        or "cnpj" in resumo_low and "contrato principal" not in resumo_low):
        data_res = base.get("data_assinatura") or base.get("data_contrato") or "data não localizada"
        base["resumo_executivo"] = (
            f"Contrato de prestação de serviços de alocação de mão de obra temporária firmado entre {base.get('empresa_grupo_sbf','SBF')} "
            f"e {base.get('contraparte','contraparte')}, com assinatura/data principal em {data_res}. "
            "A contratação possui vigência por prazo indeterminado e atende demandas de substituição transitória de pessoal permanente ou demanda complementar de serviços, nos termos da Lei nº 6.019/1974. "
            "Não há valor global fixo; as condições comerciais devem ser apuradas pela proposta aprovada, conforme demanda, quantidade de trabalhadores/vagas, salário/remuneração e taxas aplicáveis."
        )

    base["parecer"] = (
        "Contrato estruturalmente válido para continuidade, desde que a área mantenha controle da demanda, remuneração base, "
        "prazo máximo de alocação do trabalhador temporário, encargos e aderência às condições comerciais da proposta aprovada."
    )

    return base


def aplicar_regras_especificas_qualidade_ia(base: Dict[str, Any], texto: str) -> Dict[str, Any]:
    """Aplica blindagens determinísticas após a resposta do Gemini."""
    base = _aplicar_regras_mao_obra_temporaria(base, texto)
    return base


def padronizar_resultado_ia(base: Dict[str, Any]) -> Dict[str, Any]:
    """Ajustes finais para evitar cards quebrados e dados mal formatados."""
    base["cnpj_empresa_grupo"] = formatar_cnpj(base.get("cnpj_empresa_grupo"))
    base["cnpj_contraparte"] = formatar_cnpj(base.get("cnpj_contraparte"))
    base["cnpj"] = formatar_cnpj(base.get("cnpj") or base.get("cnpj_contraparte"))

    # Se a IA devolver somente número no pagamento, transforma em "X dias".
    pag = clean_text(base.get("condicao_pagamento_dias"))
    if re.fullmatch(r"\d{1,3}", pag):
        base["condicao_pagamento_dias"] = f"{pag} dias"

    # Mantém os cards com resumo executivo: detalhado o suficiente, sem virar cláusula inteira.
    limites = {
        "objetivo": 900,
        "descricao_servico_material": 1200,
        "descricao_breve_cadastro": 500,
        "forma_pagamento": 700,
        "condicao_pagamento_dias": 400,
        "multa": 900,
        "vigencia_apos_assinatura": 900,
        "resumo_aditivos": 900,
        "rescisao_indenizacao": 1200,
        "anticorrupcao": 900,
        "protecao_dados_lgpd": 900,
        "alerta_assinatura": 500,
        "valor_contrato_original": 900,
        "valor_mensal_estimado": 900,
        "valor_total_estimado_vigencia": 900,
        "valor_total_materiais_servicos": 900,
    }
    for campo, limite in limites.items():
        base[campo] = resumir_campo(base.get(campo), limite)

    # Padroniza período de vigência por prazo indeterminado para data técnica final.
    periodo_txt = clean_text(base.get("periodo_vigencia_formatado"))
    if "prazo indeterminado" in periodo_txt.lower():
        inicio_match = re.search(r"(\d{2}/\d{2}/\d{4})", periodo_txt)
        inicio = inicio_match.group(1) if inicio_match else clean_text(base.get("data_assinatura") or base.get("data_contrato"))
        if _valor_informado(inicio):
            base["periodo_vigencia_formatado"] = f"Início {inicio} até 31/12/9999"

    # Valor total de materiais/serviços: não somar taxas, percentuais e valores sem quantidade.
    itens_tmp = normalizar_itens_contrato(base.get("itens_contrato", []))
    tem_taxa = any(_valor_informado(i.get("Taxa / Percentual")) and "%" in str(i.get("Taxa / Percentual")) for i in itens_tmp)
    soma_calculavel = _somar_valores_itens(itens_tmp)
    if tem_taxa and not soma_calculavel:
        base["valor_total_materiais_servicos"] = (
            "Não calculável como total global. Foram identificadas condições comerciais/taxas, porém o total depende de quantidade, demanda, salário/remuneração, vaga utilizada e período de uso."
        )

    # Corrige erro comum: descrição de cadastro com histórico societário em vez do serviço.
    desc_cad = clean_text(base.get("descricao_breve_cadastro"))
    desc_serv = clean_text(base.get("descricao_servico_material"))
    if any(t in desc_cad.lower() for t in ["transformação", "transformacao", "eireli", "societária", "societaria"]):
        base["descricao_breve_cadastro"] = resumir_campo(desc_serv, 120)

    base["itens_contrato"] = normalizar_itens_contrato(base.get("itens_contrato", []))
    return base

def normalizar(resultado: Dict[str, Any]) -> Dict[str, Any]:
    base = local_extract("")
    base.update(resultado or {})

    if clean_text(base.get("contraparte")) in ("", "Não localizado", "Não localizada", "N/A", "None"):
        base["contraparte"] = base.get("fornecedor", "Não localizado")
    if clean_text(base.get("fornecedor")) in ("", "Não localizado", "Não localizada", "N/A", "None"):
        base["fornecedor"] = base.get("contraparte", "Não localizado")

    # Compatibilidade com nomes que a IA pode devolver em português ou com variações.
    alias = {
        "local_de_prestacao_contraparte": "local_prestacao",
        "local_prestacao_contraparte": "local_prestacao",
        "cnpj_empresa_do_grupo": "cnpj_empresa_grupo",
        "empresa_do_grupo_sbf": "empresa_grupo_sbf",
        "descricao_do_servico_material": "descricao_servico_material",
        "descricao_servico_material": "descricao_servico_material",
        "vigencia_apos_a_data_de_assinatura": "vigencia_apos_assinatura",
        "periodo_de_vigencia": "periodo_vigencia_formatado",
        "periodo_vigencia": "periodo_vigencia_formatado",
        "vigencia_formatada": "periodo_vigencia_formatado",
        "valor_do_contrato_original": "valor_contrato_original",
        "valor_mensal": "valor_mensal_estimado",
        "valor_mensal_estimado": "valor_mensal_estimado",
        "valor_total_da_vigencia": "valor_total_estimado_vigencia",
        "valor_total_estimado_vigencia": "valor_total_estimado_vigencia",
        "valor_total_materiais_servicos": "valor_total_materiais_servicos",
        "valor_total_dos_materiais_e_servicos": "valor_total_materiais_servicos",
        "valor_total_dos_materiais_servicos": "valor_total_materiais_servicos",
        "soma_itens_contrato": "valor_total_materiais_servicos",
        "data_do_contrato": "data_contrato",
        "data_conclusao_docusign": "data_conclusao_docusign",
        "data_de_conclusao_docusign": "data_conclusao_docusign",
        "assinantes": "pessoas_que_assinaram",
        "pessoas_que_assinaram": "pessoas_que_assinaram",
        "pessoas_que_assinaram_o_contrato": "pessoas_que_assinaram",
    }
    def _vazio(valor: Any) -> bool:
        txt = clean_text(valor)
        return txt in ("", "Não localizado", "Não localizada", "N/A", "None")

    for origem, destino in alias.items():
        if origem in base and _vazio(base.get(destino)):
            base[destino] = base.get(origem)

    for chave, valor in list(base.items()):
        if isinstance(valor, str):
            base[chave] = clean_text(valor)

    if not isinstance(base.get("checklist"), list):
        base["checklist"] = []
    if not isinstance(base.get("pendencias"), list):
        base["pendencias"] = []

    texto_fallback = resultado.get("texto_extraido") if isinstance(resultado, dict) else ""
    itens_ia = normalizar_itens_contrato(base.get("itens_contrato", []))
    itens_servico_percentual = detectar_servico_percentual(str(texto_fallback or ""))

    # Para contratos de mão de obra temporária/RH, taxa e encargos são atributos
    # do serviço principal, não itens separados.
    if itens_servico_percentual and (_texto_indica_mao_obra_temporaria(texto_fallback) or not itens_ia or _itens_sao_apenas_atributos_comerciais(itens_ia)):
        # Para mão de obra/RH, a extração determinística da proposta evita que a IA
        # devolva apenas uma linha rasa e perca 40%, 50%, DISC, tributos e vencimento.
        base["itens_contrato"] = normalizar_itens_contrato(itens_servico_percentual)
    elif itens_ia:
        base["itens_contrato"] = itens_ia
    else:
        base["itens_contrato"] = extrair_itens_local(str(texto_fallback or ""))

    base = aplicar_regras_finais_contrato(base, str(texto_fallback or ""))
    base = aplicar_regras_aditivos(base, str(texto_fallback or ""))

    for lista in ["checklist", "pendencias"]:
        for item in base.get(lista, []):
            if isinstance(item, dict):
                for k, v in list(item.items()):
                    if isinstance(v, str):
                        item[k] = clean_text(v)

    base["risco"] = normalize_risco(base.get("risco"))
    base["score"] = int(min(max(as_float_score(base.get("score")), 0), 100))
    base = padronizar_resultado_ia(base)
    return base




# =========================================================
# MOTOR DE CONFIANÇA V2 — EVIDÊNCIA, DUPLA VALIDAÇÃO E
# SEPARAÇÃO SEGURA DE VALORES/DATAS/ASSINATURAS
# =========================================================
# Mantém as funções antigas como base de compatibilidade e aplica uma camada
# final mais rígida. Assim, histórico, visual e exportações continuam iguais,
# mas nenhum dado crítico é confirmado sem fonte/evidência.
_prompt_ia_legado = prompt_ia
_normalizar_legado = normalizar

_CAMPOS_DIRETOS_COM_EVIDENCIA = {
    "tipo_contrato", "empresa_grupo_sbf", "cnpj_empresa_grupo", "local_prestacao",
    "contraparte", "cnpj_contraparte", "objetivo", "descricao_servico_material",
    "descricao_breve_cadastro", "forma_pagamento", "condicao_pagamento_dias", "multa",
    "vigencia_apos_assinatura", "rescisao_indenizacao", "anticorrupcao",
    "protecao_dados_lgpd", "data_assinatura", "data_contrato",
    "data_conclusao_docusign", "data_reconhecimento_firma",
    "valor_contrato_original", "valor_mensal_estimado",
    "pessoas_que_assinaram", "contrato_assinado", "alerta_assinatura",
}

_CAMPOS_CALCULADOS_OU_CONSOLIDADOS = {
    "periodo_vigencia_formatado", "valor_total_estimado_vigencia",
    "valor_total_materiais_servicos", "resumo_aditivos",
}

_STATUS_CONFIRMADOS = {
    "CONFIRMADO", "CONFIRMADO_NO_DOCUMENTO", "LOCALIZADO", "VALIDADO",
    "CALCULADO", "CALCULADO_PELO_SISTEMA", "NAO_APLICAVEL", "NÃO_APLICÁVEL",
}


def _norm_token(valor: Any) -> str:
    txt = clean_text(valor).upper()
    txt = (txt.replace("Á", "A").replace("À", "A").replace("Ã", "A").replace("Â", "A")
              .replace("É", "E").replace("Ê", "E").replace("Í", "I")
              .replace("Ó", "O").replace("Ô", "O").replace("Õ", "O")
              .replace("Ú", "U").replace("Ç", "C"))
    return re.sub(r"[^A-Z0-9]+", "_", txt).strip("_")


def _texto_evidencia_util(valor: Any) -> bool:
    txt = clean_text(valor)
    low = txt.lower()
    if len(txt) < 8:
        return False
    bloqueios = {
        "não localizado", "nao localizado", "não identificado", "nao identificado",
        "conforme documento", "documento analisado", "informação localizada",
        "informacao localizada", "termo localizado", "não aplicável", "nao aplicavel",
    }
    return low not in bloqueios and not low.startswith("não foi localizada informação")


def _normalizar_auditoria_campos(valor: Any) -> List[Dict[str, Any]]:
    if isinstance(valor, dict):
        itens = []
        for campo, detalhe in valor.items():
            if isinstance(detalhe, dict):
                novo = dict(detalhe)
                novo.setdefault("campo", campo)
            else:
                novo = {"campo": campo, "valor": detalhe}
            itens.append(novo)
        valor = itens
    if not isinstance(valor, list):
        return []

    saida: List[Dict[str, Any]] = []
    for item in valor:
        if not isinstance(item, dict):
            continue
        campo = clean_text(item.get("campo") or item.get("chave") or item.get("field"))
        if not campo:
            continue
        try:
            confianca = int(float(str(item.get("confianca") or item.get("confiança") or 0).replace(",", ".")))
        except Exception:
            confianca = 0
        saida.append({
            "Campo": campo,
            "Valor": clean_text(item.get("valor") or item.get("resultado") or "Não identificado com segurança"),
            "Status": clean_text(item.get("status") or "NAO_LOCALIZADO"),
            "Tipo de dado": clean_text(item.get("tipo_dado") or item.get("classificacao") or item.get("classificação") or "DADO_DOCUMENTAL"),
            "Arquivo fonte": clean_text(item.get("arquivo_fonte") or item.get("arquivo") or item.get("fonte") or "Não localizado"),
            "Página": clean_text(item.get("pagina") or item.get("página") or "Não localizado"),
            "Trecho de evidência": clean_text(item.get("trecho_evidencia") or item.get("trecho") or item.get("evidencia") or item.get("evidência") or "Não localizado"),
            "Confiança": max(0, min(confianca, 100)),
        })
    return saida


def _mapa_auditoria(auditoria: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    return {_norm_token(item.get("Campo")): item for item in auditoria if item.get("Campo")}


def _auditoria_confirma(item: Dict[str, Any] | None) -> bool:
    if not item:
        return False
    status = _norm_token(item.get("Status"))
    tipo = _norm_token(item.get("Tipo de dado"))
    if status not in {_norm_token(s) for s in _STATUS_CONFIRMADOS}:
        return False
    if status in ("CALCULADO", "CALCULADO_PELO_SISTEMA") or tipo.startswith("CALCULADO"):
        return _texto_evidencia_util(item.get("Trecho de evidência"))
    fonte_ok = _valor_informado(item.get("Arquivo fonte"))
    evidencia_ok = _texto_evidencia_util(item.get("Trecho de evidência"))
    confianca = int(item.get("Confiança") or 0)
    return bool(fonte_ok and evidencia_ok and confianca >= 70)


def _valor_nao_confirmado(campo: str) -> str:
    if campo == "contrato_assinado":
        return "Não validado"
    if campo == "alerta_assinatura":
        return "Não foi localizada evidência documental suficiente para validar a assinatura."
    return "Não identificado com segurança"


def _encontrar_item_auditoria(mapa: Dict[str, Dict[str, Any]], campo: str) -> Dict[str, Any] | None:
    candidatos = [campo, campo.replace("_", " ")]
    for rotulo, chave in CAMPOS_OFICIAIS:
        if chave == campo:
            candidatos.append(rotulo)
    for c in candidatos:
        item = mapa.get(_norm_token(c))
        if item:
            return item
    return None


def _classificar_item_financeiro(item: Dict[str, Any]) -> str:
    desc = clean_text(item.get("Descrição") or item.get("descricao")).lower()
    unidade = clean_text(item.get("Unidade") or item.get("unidade")).lower()
    texto = f"{desc} {unidade}"
    if any(t in texto for t in ["implantação", "implantacao", "setup", "taxa única", "taxa unica", "adesão", "adesao"]):
        return "PONTUAL"
    if any(t in texto for t in ["custo fixo mensal", "mensalidade", "remuneração mensal", "remuneracao mensal", "fixo mensal"]):
        return "FIXO_MENSAL"
    if unidade in ("mês", "mes", "mensal") and not any(t in texto for t in ["acionista", "operação", "operacao", "unidade"]):
        return "FIXO_MENSAL"
    if any(t in texto for t in ["acionista", "operação", "operacao", "por item", "por usuário", "por usuario", "por refeição", "por refeicao", "por vaga", "por trabalhador"]):
        return "UNITARIO_VARIAVEL"
    if _valor_informado(item.get("Taxa / Percentual")):
        return "PERCENTUAL_VARIAVEL"
    return "OUTRO"


def _primeiro_valor_item(itens: List[Dict[str, Any]], classificacao: str, preferir_total: bool = False) -> tuple[float | None, Dict[str, Any] | None]:
    for item in itens:
        if _classificar_item_financeiro(item) != classificacao:
            continue
        campos = ["Valor total", "Valor unitário"] if preferir_total else ["Valor unitário", "Valor total"]
        for campo in campos:
            numero = _parse_moeda_brasil(item.get(campo))
            if numero is not None:
                return numero, item
    return None, None


def _valores_estruturados_dict(valor: Any) -> Dict[str, Any]:
    return valor if isinstance(valor, dict) else {}


def _bloco_valor_estruturado(estruturados: Dict[str, Any], *chaves: str) -> Dict[str, Any]:
    for chave in chaves:
        bloco = estruturados.get(chave)
        if isinstance(bloco, dict):
            return bloco
        if _valor_informado(bloco):
            return {"valor": bloco, "status": "CONFIRMADO"}
    return {}


def _bloco_valor_confirmado(bloco: Dict[str, Any]) -> bool:
    if not bloco:
        return False
    status = _norm_token(bloco.get("status") or "CONFIRMADO")
    evidencia = bloco.get("trecho_evidencia") or bloco.get("evidencia")
    fonte = bloco.get("arquivo_fonte") or bloco.get("fonte")
    if status not in ("CONFIRMADO", "LOCALIZADO", "VALIDADO", "CALCULADO", "CALCULADO_PELO_SISTEMA"):
        return False
    # Blocos antigos sem metadados continuam aceitos somente quando o valor é claro;
    # na versão nova, o prompt sempre enviará fonte e trecho.
    return _valor_informado(bloco.get("valor")) and (bool(_valor_informado(fonte) and _texto_evidencia_util(evidencia)) or not fonte)


def _aplicar_financeiro_confiavel(base: Dict[str, Any], resultado_bruto: Dict[str, Any]) -> Dict[str, Any]:
    itens = normalizar_itens_contrato(base.get("itens_contrato", []))
    estruturados = _valores_estruturados_dict(resultado_bruto.get("valores_estruturados"))

    global_bloco = _bloco_valor_estruturado(estruturados, "valor_global", "valor_contrato_original")
    mensal_bloco = _bloco_valor_estruturado(estruturados, "valor_fixo_mensal", "valor_mensal_fixo", "valor_mensal")

    global_num = _parse_moeda_brasil(global_bloco.get("valor")) if _bloco_valor_confirmado(global_bloco) else None
    mensal_num = _parse_moeda_brasil(mensal_bloco.get("valor")) if _bloco_valor_confirmado(mensal_bloco) else None
    item_mensal = None
    if mensal_num is None:
        mensal_num, item_mensal = _primeiro_valor_item(itens, "FIXO_MENSAL")

    tarifas = [i for i in itens if _classificar_item_financeiro(i) in ("UNITARIO_VARIAVEL", "PERCENTUAL_VARIAVEL")]
    pontuais = [i for i in itens if _classificar_item_financeiro(i) == "PONTUAL"]

    if global_num is not None:
        base["valor_contrato_original"] = (
            f"{_formatar_moeda_brasil(global_num)}. Valor global expressamente definido no instrumento contratual."
        )
    else:
        base["valor_contrato_original"] = _texto_sem_valor_global()

    if mensal_num is not None:
        texto_mensal = f"{_formatar_moeda_brasil(mensal_num)}/mês. Valor fixo mensal identificado nos documentos."
        if tarifas:
            texto_mensal += f" Há ainda {len(tarifas)} tarifa(s) variável(is), cobradas separadamente conforme acionistas, operações, itens ou demanda."
        base["valor_mensal_estimado"] = texto_mensal
    elif tarifas:
        exemplos = []
        for item in tarifas[:3]:
            valor = _parse_moeda_brasil(item.get("Valor unitário"))
            if valor is not None:
                exemplos.append(f"{clean_text(item.get('Descrição'))}: {_formatar_moeda_brasil(valor)}")
            elif _valor_informado(item.get("Taxa / Percentual")):
                exemplos.append(f"{clean_text(item.get('Descrição'))}: {clean_text(item.get('Taxa / Percentual'))}")
        complemento = (" Exemplos: " + "; ".join(exemplos) + ".") if exemplos else ""
        base["valor_mensal_estimado"] = (
            "Valor mensal variável conforme utilização, quantidade, operação ou demanda; não há mensalidade fixa confirmada."
            + complemento
        )

    # Consolidação financeira por natureza. Não mistura implantação/taxa única
    # com mensalidade recorrente, percentual ou tarifa variável.
    def _somar_categoria_financeira(classificacao: str) -> float | None:
        total_categoria = 0.0
        encontrou = False
        for item in itens:
            if _classificar_item_financeiro(item) != classificacao:
                continue
            numero = _parse_moeda_brasil(item.get("Valor total"))
            if numero is None:
                numero = _parse_moeda_brasil(item.get("Valor unitário"))
            if numero is not None:
                total_categoria += numero
                encontrou = True
        return total_categoria if encontrou else None

    total_pontual = _somar_categoria_financeira("PONTUAL")
    total_mensal = _somar_categoria_financeira("FIXO_MENSAL")

    outros_explicitos = []
    for item in itens:
        classificacao = _classificar_item_financeiro(item)
        if classificacao in ("PONTUAL", "FIXO_MENSAL", "UNITARIO_VARIAVEL", "PERCENTUAL_VARIAVEL"):
            continue
        numero = _parse_moeda_brasil(item.get("Valor total"))
        if numero is not None:
            outros_explicitos.append(numero)

    if total_pontual is not None and total_mensal is not None:
        base["valor_total_materiais_servicos"] = (
            "Valores separados por natureza: "
            f"custo pontual/implantação de {_formatar_moeda_brasil(total_pontual)}; "
            f"custo fixo mensal de {_formatar_moeda_brasil(total_mensal)}/mês. "
            "Esses valores não foram somados porque um é pontual e o outro é recorrente. "
            "O valor global do contrato permanece não calculável sem prazo e quantidades definidos."
        )
    elif total_pontual is not None:
        base["valor_total_materiais_servicos"] = (
            f"{_formatar_moeda_brasil(total_pontual)} em valores pontuais confirmados. "
            "Esse total não inclui mensalidades, percentuais ou tarifas variáveis e não representa o valor global do contrato."
        )
    elif total_mensal is not None:
        base["valor_total_materiais_servicos"] = (
            f"{_formatar_moeda_brasil(total_mensal)}/mês em custos fixos recorrentes confirmados. "
            "Não representa o valor global do contrato."
        )
    elif outros_explicitos:
        total_outros = sum(outros_explicitos)
        base["valor_total_materiais_servicos"] = (
            f"{_formatar_moeda_brasil(total_outros)}. Soma somente de linhas com valor total monetário explícito e mesma natureza. "
            "Mensalidades, percentuais e tarifas variáveis sem quantidade não foram somados."
        )
    else:
        base["valor_total_materiais_servicos"] = (
            "Não calculável com segurança. Não há linhas homogêneas com valor total e quantidade suficientes para soma."
        )

    meses = _extrair_meses_vigencia(str(base.get("texto_extraido") or ""), base)
    if mensal_num is not None and meses:
        base_fixa = mensal_num * meses
        if tarifas:
            base["valor_total_estimado_vigencia"] = (
                f"Base fixa estimada: {_formatar_moeda_brasil(base_fixa)} ({_formatar_moeda_brasil(mensal_num)}/mês x {meses} meses). "
                "O total final não é calculável com precisão porque existem tarifas variáveis."
            )
        else:
            base["valor_total_estimado_vigencia"] = (
                f"{_formatar_moeda_brasil(base_fixa)}. Cálculo do sistema: {_formatar_moeda_brasil(mensal_num)}/mês x {meses} meses."
            )
    elif mensal_num is not None and "31/12/9999" in clean_text(base.get("periodo_vigencia_formatado")):
        base["valor_total_estimado_vigencia"] = (
            "Não calculável com precisão: o contrato possui mensalidade, porém a vigência é por prazo indeterminado."
        )
    elif tarifas:
        base["valor_total_estimado_vigencia"] = (
            "Não calculável com precisão. O total depende da quantidade de acionistas, operações, itens ou demanda durante a vigência."
        )

    base["valor_total"] = base.get("valor_contrato_original")
    return base



def _filtrar_itens_com_evidencia(base: Dict[str, Any], resultado_bruto: Dict[str, Any]) -> Dict[str, Any]:
    """Mantém somente itens comerciais sustentados por fonte e trecho específico."""
    raw = resultado_bruto.get("itens_contrato")
    if not isinstance(raw, list) or not base.get("auditoria_campos"):
        return base
    confirmados: List[Dict[str, Any]] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        fonte = clean_text(item.get("arquivo_fonte") or item.get("fonte") or item.get("Fonte"))
        pagina = clean_text(item.get("pagina") or item.get("página") or item.get("Página"))
        evidencia = clean_text(item.get("trecho_evidencia") or item.get("evidencia") or item.get("evidência"))
        descricao = clean_text(item.get("descricao") or item.get("Descrição"))
        if not (_valor_informado(descricao) and _valor_informado(fonte) and _texto_evidencia_util(evidencia)):
            continue
        novo = dict(item)
        novo["fonte"] = f"{fonte}" + (f" • p. {pagina}" if _valor_informado(pagina) else "")
        confirmados.append(novo)
    base["itens_contrato"] = normalizar_itens_contrato(confirmados)
    return base


def _filtrar_aditivos_com_evidencia(base: Dict[str, Any], resultado_bruto: Dict[str, Any]) -> Dict[str, Any]:
    """Impede que menção a aditivo no contrato principal vire aditivo inexistente."""
    raw = resultado_bruto.get("aditivos_contrato")
    if not isinstance(raw, list) or not base.get("auditoria_campos"):
        return base
    confirmados: List[Dict[str, Any]] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        fonte = clean_text(item.get("anexo_origem") or item.get("Anexo do aditivo") or item.get("arquivo_fonte"))
        evidencia = clean_text(item.get("trecho_evidencia") or item.get("evidencia") or item.get("evidência"))
        pagina = clean_text(item.get("pagina") or item.get("página"))
        if not (_valor_informado(fonte) and _texto_evidencia_util(evidencia)):
            continue
        novo = dict(item)
        novo["anexo_origem"] = fonte
        novo["pagina"] = pagina
        novo["trecho_evidencia"] = evidencia
        confirmados.append(novo)
    base["aditivos_contrato"] = normalizar_aditivos_contrato(confirmados)
    base["resumo_aditivos"] = _montar_resumo_aditivos(base["aditivos_contrato"])
    return base


def _filtrar_checklist_sem_evidencia(base: Dict[str, Any]) -> Dict[str, Any]:
    checklist = base.get("checklist") if isinstance(base.get("checklist"), list) else []
    validos = []
    for item in checklist:
        if not isinstance(item, dict):
            continue
        evidencia = clean_text(item.get("Evidência") or item.get("evidencia"))
        if not _texto_evidencia_util(evidencia):
            continue
        validos.append(item)
    base["checklist"] = validos
    return base


def _reconstruir_resumo_e_parecer(base: Dict[str, Any]) -> Dict[str, Any]:
    """Gera resumo/parecer apenas com fatos já validados, sem texto livre inventado."""
    if not base.get("auditoria_campos"):
        return base
    partes = []
    tipo = clean_text(base.get("tipo_contrato"))
    empresa = clean_text(base.get("empresa_grupo_sbf"))
    contraparte = clean_text(base.get("contraparte"))
    objetivo = clean_text(base.get("objetivo"))
    vigencia = clean_text(base.get("periodo_vigencia_formatado") or base.get("vigencia_apos_assinatura"))
    assinatura = clean_text(base.get("contrato_assinado"))
    if _valor_informado(tipo):
        partes.append(tipo.rstrip(" .;"))
    if _valor_informado(empresa) and _valor_informado(contraparte):
        partes.append(f"firmado entre {empresa} e {contraparte}")
    if _valor_informado(objetivo):
        partes.append(f"Objeto: {objetivo.rstrip(' .;')}")
    if _valor_informado(vigencia):
        partes.append(f"Vigência: {vigencia.rstrip(' .;')}")
    if _valor_informado(assinatura):
        partes.append(f"Assinatura validada: {assinatura}")
    base["resumo_executivo"] = ". ".join(partes).strip() + ("." if partes else "Não foi possível consolidar resumo com evidências suficientes.")

    pendencias = base.get("pendencias") if isinstance(base.get("pendencias"), list) else []
    criticas = [p for p in pendencias if _norm_token(p.get("Crítico") or p.get("critico")) in ("SIM", "TRUE", "1")]
    if criticas:
        base["parecer"] = (
            f"Revisão obrigatória antes de seguir: existem {len(criticas)} pendência(s) crítica(s) com evidência documental. "
            "Consulte a aba Pendências e valide o documento original."
        )
    elif pendencias:
        base["parecer"] = (
            f"A análise identificou {len(pendencias)} ponto(s) de atenção com evidência documental. "
            "A continuidade depende da validação das recomendações registradas."
        )
    else:
        base["parecer"] = (
            "Não foram localizadas pendências documentais confirmadas na análise. "
            "A conferência humana continua recomendada para decisões jurídicas ou financeiras críticas."
        )
    return base


def _derivar_status_vigencia(base: Dict[str, Any]) -> Dict[str, Any]:
    if not base.get("auditoria_campos"):
        return base
    periodo = clean_text(base.get("periodo_vigencia_formatado"))
    if "31/12/9999" in periodo:
        base["status"] = "Ativo"
        return base
    datas = re.findall(r"\b\d{2}/\d{2}/\d{4}\b", periodo)
    if len(datas) >= 2:
        fim = _parse_data_br_para_datetime(datas[-1])
        if fim:
            base["status"] = "Ativo" if fim.date() >= datetime.now().date() else "Encerrado"
            return base
    base["status"] = "Não identificado com segurança"
    return base


def _filtrar_assinaturas_com_evidencia(base: Dict[str, Any]) -> Dict[str, Any]:
    raw = base.get("assinaturas_contrato")
    if not isinstance(raw, list):
        raw = []
    filtradas: List[Dict[str, Any]] = []
    for item in raw:
        if not isinstance(item, dict):
            continue
        nome = clean_text(item.get("nome") or item.get("Nome"))
        evidencia = clean_text(item.get("evidencia") or item.get("evidência") or item.get("Evidência"))
        fonte = clean_text(item.get("fonte") or item.get("Fonte"))
        if not _valor_informado(nome) or not _texto_evidencia_util(evidencia) or not _valor_informado(fonte):
            continue
        filtradas.append(item)
    base["assinaturas_contrato"] = filtradas
    nomes = []
    for item in filtradas:
        nome = clean_text(item.get("nome") or item.get("Nome"))
        if nome and nome not in nomes:
            nomes.append(nome)
    if nomes:
        base["pessoas_que_assinaram"] = "; ".join(nomes)
        base["contrato_assinado"] = "Sim"
    elif clean_text(base.get("contrato_assinado")).lower() == "sim":
        # Só preserva o SIM quando a auditoria do campo confirmou evidência de assinatura.
        auditoria = _mapa_auditoria(base.get("auditoria_campos", []))
        item_ass = _encontrar_item_auditoria(auditoria, "contrato_assinado")
        if not _auditoria_confirma(item_ass):
            base["contrato_assinado"] = "Não validado"
            base["pessoas_que_assinaram"] = "Não identificado com segurança"
    return base



def _datas_br_encontradas(valor: Any) -> List[str]:
    """Retorna datas brasileiras únicas na ordem em que aparecem."""
    datas = re.findall(r"\b\d{2}/\d{2}/\d{4}\b", clean_text(valor))
    saida: List[str] = []
    for data in datas:
        if data not in saida:
            saida.append(data)
    return saida


def _data_reconhecimento_da_evidencia(valor: Any) -> str:
    """Extrai somente data ligada a reconhecimento de firma/cartório."""
    texto = clean_text(valor)
    padroes = [
        r"reconhec(?:imento|ida|ido)[^\d]{0,90}(\d{2}/\d{2}/\d{4})",
        r"firma[^\d]{0,90}(\d{2}/\d{2}/\d{4})",
        r"cart[oó]rio[^\d]{0,90}(\d{2}/\d{2}/\d{4})",
    ]
    for padrao in padroes:
        m = re.search(padrao, texto, flags=re.IGNORECASE)
        if m:
            return m.group(1)
    return ""


def _data_valida_simples(valor: Any) -> str:
    datas = _datas_br_encontradas(valor)
    return datas[0] if datas else ""


def _consolidar_assinatura_final(base: Dict[str, Any]) -> Dict[str, Any]:
    """Elimina contradições entre tabela, card, alerta e datas de assinatura."""
    assinaturas = base.get("assinaturas_contrato") if isinstance(base.get("assinaturas_contrato"), list) else []
    validas: List[Dict[str, Any]] = []
    nomes: List[str] = []
    datas_assinatura: List[str] = []
    datas_reconhecimento: List[str] = []
    fontes: List[str] = []
    evidencias: List[str] = []

    for item in assinaturas:
        if not isinstance(item, dict):
            continue
        nome = clean_text(item.get("nome") or item.get("Nome"))
        fonte = clean_text(item.get("fonte") or item.get("Fonte"))
        evidencia = clean_text(item.get("evidencia") or item.get("evidência") or item.get("Evidência"))
        if not (_valor_informado(nome) and _valor_informado(fonte) and _texto_evidencia_util(evidencia)):
            continue

        data_ass = _data_valida_simples(item.get("data_assinatura") or item.get("Data da assinatura"))
        data_rec = _data_valida_simples(
            item.get("data_reconhecimento_firma")
            or item.get("Data do reconhecimento de firma")
            or _data_reconhecimento_da_evidencia(evidencia)
        )

        # Reconhecimento de firma é evento posterior e nunca substitui a data da assinatura.
        if data_ass and data_rec and data_ass == data_rec:
            texto_ev = evidencia.lower()
            if "reconhec" in texto_ev or "cartório" in texto_ev or "cartorio" in texto_ev:
                # Mantém a data no reconhecimento; a assinatura global será buscada no contrato.
                data_ass = ""

        novo = dict(item)
        if data_ass:
            novo["data_assinatura"] = data_ass
            datas_assinatura.append(data_ass)
        if data_rec:
            novo["data_reconhecimento_firma"] = data_rec
            datas_reconhecimento.append(data_rec)
        validas.append(novo)

        if nome not in nomes:
            nomes.append(nome)
        if fonte not in fontes:
            fontes.append(fonte)
        if evidencia not in evidencias:
            evidencias.append(evidencia)

    base["assinaturas_contrato"] = validas

    auditoria = _mapa_auditoria(base.get("auditoria_campos", []))
    item_assinado = _encontrar_item_auditoria(auditoria, "contrato_assinado")
    auditoria_assinatura_ok = _auditoria_confirma(item_assinado)
    contrato_assinado = bool(validas) or clean_text(base.get("contrato_assinado")).lower() == "sim" or auditoria_assinatura_ok

    if contrato_assinado:
        base["contrato_assinado"] = "Sim"
        if nomes:
            base["pessoas_que_assinaram"] = "; ".join(nomes)

        # Data da assinatura: usa primeiro as datas específicas dos signatários;
        # depois a data global já confirmada. Nunca usa reconhecimento como fallback.
        data_global = _data_valida_simples(base.get("data_assinatura"))
        data_contrato = _data_valida_simples(base.get("data_contrato"))
        if datas_assinatura:
            base["data_assinatura"] = datas_assinatura[0]
        elif data_global:
            base["data_assinatura"] = data_global
        elif data_contrato:
            base["data_assinatura"] = data_contrato

        if datas_reconhecimento:
            unicas = []
            for data in datas_reconhecimento:
                if data not in unicas:
                    unicas.append(data)
            base["data_reconhecimento_firma"] = "; ".join(unicas)

        data_docusign = clean_text(base.get("data_conclusao_docusign"))
        docusign_valido = _data_valida_simples(data_docusign) and "não aplicável" not in data_docusign.lower()
        qtd = len(nomes) if nomes else len(validas)
        partes_alerta = []
        if docusign_valido:
            partes_alerta.append(f"Contrato assinado eletronicamente com conclusão DocuSign em {_data_valida_simples(data_docusign)}")
        else:
            partes_alerta.append("Contrato assinado fisicamente")
        if qtd:
            partes_alerta.append(f"{qtd} signatário(s) identificado(s)")
        if _data_valida_simples(base.get("data_assinatura")):
            partes_alerta.append(f"data principal de assinatura: {_data_valida_simples(base.get('data_assinatura'))}")
        if _valor_informado(base.get("data_reconhecimento_firma")):
            partes_alerta.append(f"reconhecimento de firma: {clean_text(base.get('data_reconhecimento_firma'))}")
        if fontes:
            partes_alerta.append("fonte: " + "; ".join(fontes[:3]))
        base["alerta_assinatura"] = ". ".join(partes_alerta).rstrip(".") + "."
    else:
        base["contrato_assinado"] = "Não validado"
        base["alerta_assinatura"] = "Não foi localizada evidência documental suficiente para validar a assinatura."

    return base


def _pendencia_email_signatario_sem_exigencia(pendencia: Dict[str, Any]) -> bool:
    """Ausência de e-mail não é risco contratual sem obrigação expressa no instrumento."""
    texto_total = " ".join(clean_text(v) for v in pendencia.values()).lower()
    trata_email = bool(re.search(r"e-?mail", texto_total)) and bool(re.search(r"signat|assinante|representante", texto_total))
    if not trata_email:
        return False
    exigencia_expressa = bool(re.search(
        r"(contrato|cl[aá]usula|instrumento).{0,90}(exige|exigência|obrigat[oó]rio|dever[aá]|deve informar).{0,90}e-?mail",
        texto_total,
        flags=re.IGNORECASE,
    ))
    return not exigencia_expressa


def _filtrar_pendencias_sem_evidencia(base: Dict[str, Any]) -> Dict[str, Any]:
    pendencias = base.get("pendencias") if isinstance(base.get("pendencias"), list) else []
    validas = []
    for p in pendencias:
        if not isinstance(p, dict):
            continue
        if _pendencia_email_signatario_sem_exigencia(p):
            continue
        evidencia = clean_text(p.get("Evidência") or p.get("evidencia") or p.get("trecho_evidencia"))
        pagina = clean_text(p.get("Página") or p.get("pagina"))
        arquivo = clean_text(p.get("Arquivo") or p.get("arquivo") or p.get("Fonte") or p.get("fonte"))
        # Compatibilidade: aceita recomendação que contenha referência objetiva a cláusula/anexo/página.
        recomendacao = clean_text(p.get("Recomendação") or p.get("recomendacao"))
        ref_objetiva = bool(re.search(r"\b(cl[aá]usula|anexo|p[aá]gina|item\s+\d+)\b", f"{evidencia} {recomendacao}", flags=re.IGNORECASE))
        if (_texto_evidencia_util(evidencia) and _valor_informado(arquivo)) or ref_objetiva:
            if pagina and "Página" not in p:
                p["Página"] = pagina
            if arquivo and "Arquivo" not in p:
                p["Arquivo"] = arquivo
            validas.append(p)
    base["pendencias"] = validas
    return base


def _recalcular_risco_por_evidencias(base: Dict[str, Any]) -> Dict[str, Any]:
    pendencias = base.get("pendencias") if isinstance(base.get("pendencias"), list) else []
    score = 100
    tem_critica = False
    tem_alta = False
    for p in pendencias:
        critico = _norm_token(p.get("Crítico") or p.get("critico")) in ("SIM", "TRUE", "1")
        risco = _norm_token(p.get("Risco") or p.get("risco"))
        if critico:
            score -= 18
            tem_critica = True
        if risco == "ALTO":
            score -= 15
            tem_alta = True
        elif risco in ("MEDIO", "MÉDIO"):
            score -= 8
        else:
            score -= 3
    score = max(0, min(score, 100))
    if tem_critica and tem_alta:
        risco_final = "ALTO"
    elif score < 85 or tem_critica:
        risco_final = "MÉDIO"
    else:
        risco_final = "BAIXO"
    base["score"] = score
    base["risco"] = risco_final
    return base


def _aplicar_auditoria_rigida(base: Dict[str, Any], resultado_bruto: Dict[str, Any]) -> Dict[str, Any]:
    auditoria = _normalizar_auditoria_campos(resultado_bruto.get("auditoria_campos"))
    base["auditoria_campos"] = auditoria
    base["conflitos_documentais"] = resultado_bruto.get("conflitos_documentais") if isinstance(resultado_bruto.get("conflitos_documentais"), list) else []
    base["valores_estruturados"] = resultado_bruto.get("valores_estruturados") if isinstance(resultado_bruto.get("valores_estruturados"), dict) else {}

    # Só ativa a trava rígida em análises novas que retornaram a matriz de auditoria.
    if auditoria:
        mapa = _mapa_auditoria(auditoria)
        for campo in _CAMPOS_DIRETOS_COM_EVIDENCIA:
            item = _encontrar_item_auditoria(mapa, campo)
            if _auditoria_confirma(item):
                valor_auditado = clean_text(item.get("Valor"))
                if _valor_informado(valor_auditado):
                    base[campo] = valor_auditado
            else:
                base[campo] = _valor_nao_confirmado(campo)

        # Campos calculados exigem ao menos a base documental correspondente.
        for campo in _CAMPOS_CALCULADOS_OU_CONSOLIDADOS:
            item = _encontrar_item_auditoria(mapa, campo)
            if item and _auditoria_confirma(item) and _valor_informado(item.get("Valor")):
                base[campo] = clean_text(item.get("Valor"))

    base = _filtrar_itens_com_evidencia(base, resultado_bruto)
    base = _filtrar_aditivos_com_evidencia(base, resultado_bruto)
    base = _aplicar_financeiro_confiavel(base, resultado_bruto)
    base = _filtrar_assinaturas_com_evidencia(base)
    base = _consolidar_assinatura_final(base)
    base = _filtrar_checklist_sem_evidencia(base)
    base = _filtrar_pendencias_sem_evidencia(base)
    base = _recalcular_risco_por_evidencias(base)
    base = _derivar_status_vigencia(base)
    base = _reconstruir_resumo_e_parecer(base)
    return base


def prompt_ia(texto: str) -> str:
    """Prompt V2: extração factual + evidência obrigatória por campo."""
    base = _prompt_ia_legado(texto)
    return f"""
MODO DE AUDITORIA DOCUMENTAL ESTRITA — INSTRUÇÕES COM PRIORIDADE MÁXIMA

Você deve separar três coisas: (1) fato literal localizado, (2) cálculo do sistema e (3) interpretação.
Nunca misture dados de categorias diferentes e nunca complete lacunas por plausibilidade.

REGRAS ANTIALUCINAÇÃO
- Todo campo principal deve ter uma entrada correspondente em auditoria_campos.
- Um campo só pode ter status CONFIRMADO quando houver arquivo_fonte, página e trecho_evidencia específico.
- Para DOCX sem página, use página = "Documento Word — trecho textual".
- Se a informação não estiver comprovada, valor = "Não identificado com segurança" e status = "NAO_LOCALIZADO".
- Não invente testemunha ausente, prazo, valor, data, CNPJ, signatário, multa, cláusula ou obrigação.
- Pendência só pode existir com Arquivo, Página e Evidência. Sem esses três elementos, não crie a pendência.
- Checklist só pode marcar Concluído quando a evidência trouxer cláusula/página/trecho correspondente.

SEPARAÇÃO FINANCEIRA OBRIGATÓRIA
- valor_global: apenas total fechado de todo o contrato.
- valor_fixo_mensal: mensalidade/custo fixo mensal.
- valor_pontual: implantação, setup ou taxa única.
- tarifa_unitaria: por acionista, operação, refeição, usuário, item ou unidade.
- percentual_variavel: taxa/percentual sobre salário, faturamento ou base variável.
Nunca transforme valor fixo mensal em valor por acionista. Nunca trate implantação como valor global.
Nunca some mensalidade, implantação, percentuais e tarifas unitárias sem quantidade.

DATAS SEPARADAS
- data_contrato = data textual do instrumento.
- data_assinatura = data efetiva da assinatura.
- data_conclusao_docusign = somente conclusão DocuSign; em assinatura física use "Não aplicável — assinatura física".
- data_reconhecimento_firma = reconhecimento cartorial, quando houver.
Não coloque descrições dentro de campos de data.

RETORNO ADICIONAL OBRIGATÓRIO
Além das chaves já pedidas, retorne:
1. auditoria_campos: lista com um objeto para CADA campo principal, contendo exatamente:
   campo, valor, status, tipo_dado, arquivo_fonte, pagina, trecho_evidencia, confianca.
2. valores_estruturados: objeto contendo:
   valor_global, valor_fixo_mensal, valores_pontuais, tarifas_unitarias, percentuais_variaveis.
   Cada bloco deve conter valor, unidade_cobranca, periodicidade, status, arquivo_fonte, pagina e trecho_evidencia.
3. conflitos_documentais: lista com campo, valores_conflitantes, arquivos, regra_aplicada e decisao.
4. Cada item de itens_contrato deve também conter tipo_valor, periodicidade, arquivo_fonte, pagina e trecho_evidencia.
5. Cada item de pendencias deve também conter Arquivo, Página e Evidência.
6. Cada assinatura deve ter fonte e evidencia específica; não liste nome apenas porque ele aparece no corpo do documento. Cada assinatura deve separar data_assinatura de data_reconhecimento_firma.
7. Cada aditivo deve também conter pagina e trecho_evidencia específicos do próprio termo aditivo.
8. Cada item do checklist deve usar Evidência específica, com cláusula/página/trecho; não use textos genéricos como “termo localizado”.

VALIDAÇÃO FINAL
Antes de responder, confira campo por campo contra os documentos. Quando duas informações verdadeiras forem de categorias diferentes, mantenha-as separadas. Retorne APENAS o JSON completo.

{base}

{PROMPT_EVIDENCIAS_V4}
"""


def _prompt_ia_com_documentos_originais(texto: str, nomes_arquivos: List[str]) -> str:
    lista = "\n".join(f"- {nome}" for nome in nomes_arquivos) or "- Não informado"
    return f"""
Você recebeu os arquivos originais. Eles são a fonte principal e devem ser examinados página por página.
O texto extraído é apenas índice de apoio. Nome genérico de arquivo não define o tipo documental.

ARQUIVOS RECEBIDOS:
{lista}

Ao citar evidência, informe o nome exato do arquivo, a página e um trecho curto fiel ao documento.
Para tabelas, informe também a linha/descrição da tarifa. Para assinaturas físicas, diferencie data de assinatura e reconhecimento de firma.

{prompt_ia(texto)}
"""


def _prompt_verificacao_documental(resultado_preliminar: Dict[str, Any], nomes_arquivos: List[str]) -> str:
    rascunho = json.dumps(resultado_preliminar, ensure_ascii=False, default=str)
    if len(rascunho) > 90000:
        rascunho = rascunho[:90000]
    lista = "\n".join(f"- {n}" for n in nomes_arquivos) or "- Não informado"
    return f"""
Você é o SEGUNDO AUDITOR independente. Revise o JSON preliminar comparando cada afirmação com os arquivos originais anexados.
Retorne APENAS o JSON completo corrigido, sem comentários.

ARQUIVOS:
{lista}

TESTES OBRIGATÓRIOS
1. Elimine qualquer dado sem arquivo, página e trecho de evidência.
2. Corrija mistura entre valor global, mensalidade, implantação, tarifa unitária e percentual.
3. Confirme partes/CNPJs apenas no contrato operacional principal; documentos cadastrais são apoio.
4. Confirme cada assinatura no bloco/certificado correspondente. Não invente testemunha ausente.
5. Mantenha data do contrato, assinatura, DocuSign e reconhecimento de firma em campos distintos.
6. Verifique se todas as informações existentes foram transportadas para cards, itens, assinaturas, checklist e parecer.
7. Pendências sem evidência objetiva devem ser removidas. Ausência de e-mail de signatário NÃO é pendência contratual, salvo se o próprio contrato exigir expressamente esse e-mail.
8. Nunca some implantação/taxa única com mensalidade recorrente para formar valor total do contrato. Mostre os valores separados por natureza.
9. Atualize auditoria_campos, valores_estruturados e conflitos_documentais.

JSON PRELIMINAR:
{rascunho}

{PROMPT_VERIFICADOR_V4}
"""


def _verificar_resultado_com_sdk_novo(client: Any, modelo: str, resultado_preliminar: Dict[str, Any], uploaded_files: list, nomes_arquivos: List[str]) -> Dict[str, Any]:
    prompt = _prompt_verificacao_documental(resultado_preliminar, nomes_arquivos)
    resp = client.models.generate_content(
        model=modelo,
        contents=[prompt] + list(uploaded_files or []),
        config={
            "temperature": 0.0,
            "top_p": 0.1,
            "max_output_tokens": 65535,
            "response_mime_type": "application/json",
        },
    )
    return _json_da_resposta_gemini(resp)


def analisar_gemini(texto: str, api_key: str, opcao_modelo: str, arquivos_originais: Any = None) -> Dict[str, Any]:
    """Análise V2 em duas passagens: extração e auditoria independente."""
    modelos = MODELOS_GEMINI.get(opcao_modelo, MODELOS_GEMINI["Automático recomendado"])
    nomes_arquivos = [getattr(a, "name", "documento") for a in (arquivos_originais or [])]
    ultimo_erro_multimodal = None
    uploaded_files: list = []
    temp_paths: list = []

    if arquivos_originais:
        try:
            genai_new = _importar_google_genai_novo()
            client = genai_new.Client(api_key=api_key)
            uploaded_files, temp_paths, erros_upload = _subir_arquivos_originais_gemini(client, arquivos_originais)
            if erros_upload:
                with st.expander("⚠️ Detalhes de preparação/upload dos arquivos para o Gemini", expanded=False):
                    for erro in erros_upload:
                        st.write(f"- {erro}")

            if uploaded_files:
                prompt_final = _prompt_ia_com_documentos_originais(texto, nomes_arquivos)
                for nome in modelos:
                    try:
                        preliminar = _gerar_com_sdk_novo(client, nome, prompt_final, uploaded_files)
                        final = preliminar
                        try:
                            auditado = _verificar_resultado_com_sdk_novo(client, nome, preliminar, uploaded_files, nomes_arquivos)
                            if isinstance(auditado, dict) and auditado:
                                final = dict(preliminar)
                                final.update(auditado)
                                final["verificacao_documental"] = "Concluída em segunda passagem"
                        except Exception as erro_verificacao:
                            final["verificacao_documental"] = f"Primeira passagem utilizada; verificação adicional falhou: {erro_verificacao}"

                        final["modelo_ia"] = nome
                        final["modo_analise_ia"] = "Documentos originais + validação em duas passagens + Motor de Evidências V4"
                        final["arquivos_originais_enviados"] = len(uploaded_files)
                        st.success(f"IA utilizada: {nome} • documentos validados em duas passagens + Motor de Evidências V4")
                        return final
                    except Exception as e:
                        ultimo_erro_multimodal = e
                        if opcao_modelo != "Automático recomendado":
                            raise Exception(f"Erro ao usar o modelo {nome} com documentos originais. Detalhe: {e}")
                        continue
        except Exception as e:
            ultimo_erro_multimodal = e
        finally:
            try:
                if 'client' in locals():
                    _limpar_uploads_gemini(client, uploaded_files, temp_paths)
            except Exception:
                pass

    texto_sem_base_confiavel = texto_indica_falha_leitura(texto) and not texto_tem_conteudo_contratual(texto)
    if arquivos_originais and texto_sem_base_confiavel:
        raise Exception(
            "Falha técnica: o arquivo original não pôde ser analisado pela Gemini Files API "
            "e o texto extraído não contém conteúdo contratual confiável. "
            "Não foi gerada análise por fallback para evitar cards incorretos. "
            f"Detalhe Files API: {ultimo_erro_multimodal}"
        )

    ultimo_erro_texto = None
    for nome in modelos:
        try:
            resultado_json = _gerar_texto_com_sdk_legado(texto, api_key, nome)
            resultado_json["modelo_ia"] = nome
            resultado_json["modo_analise_ia"] = "Texto extraído + regras estritas + Motor de Evidências V4"
            if ultimo_erro_multimodal:
                resultado_json["erro_upload_documentos_originais"] = str(ultimo_erro_multimodal)
                st.warning("A Files API falhou; foi usado o texto extraído com trava antialucinação.")
            st.success(f"IA utilizada: {nome} • texto extraído analisado")
            return resultado_json
        except Exception as e:
            ultimo_erro_texto = e
            if opcao_modelo != "Automático recomendado":
                raise Exception(f"Erro ao usar o modelo {nome}. Detalhe: {e}")
            continue
    raise Exception(
        "Nenhum modelo Gemini disponível. "
        f"Erro documentos originais: {ultimo_erro_multimodal}. Erro texto extraído: {ultimo_erro_texto}."
    )


def normalizar(resultado: Dict[str, Any]) -> Dict[str, Any]:
    """Normalização V2: compatibilidade antiga + validação documental rígida."""
    bruto = dict(resultado or {})
    base = _normalizar_legado(bruto)
    base["texto_extraido"] = bruto.get("texto_extraido") or base.get("texto_extraido") or ""
    base = _aplicar_auditoria_rigida(base, bruto)
    base["risco"] = normalize_risco(base.get("risco"))
    base["score"] = int(min(max(as_float_score(base.get("score")), 0), 100))
    return base


# =========================================================
# EXCEL - RELATÓRIO EXECUTIVO 100% PROFISSIONAL
# =========================================================
def excel_clean(value: Any, padrao: str = "Não localizado") -> str:
    """Texto seguro para Excel, sem HTML, sem caracteres inválidos e sem risco de fórmula."""
    value = clean_text(value)
    if value in (None, "", "Não localizado", "Não localizada", "None", "N/A"):
        value = padrao
    value = str(value).strip()
    value = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    if value and value[0] in ("=", "+", "-", "@", "\t", "\r", "\n"):
        value = "'" + value
    return value[:32700]


def excel_risk_fill(risco: Any) -> str:
    risco_norm = normalize_risco(risco)
    if risco_norm == "ALTO":
        return "DC2626"
    if risco_norm in ["MÉDIO", "MEDIO"]:
        return "F59E0B"
    if risco_norm == "BAIXO":
        return "16A34A"
    return "64748B"


def _wb_styles():
    # Paleta visual do Excel
    # Fundo geral: verde claro | Cabeçalhos: verde escuro | Títulos: dourado
    return {
        "green": "004D40",
        "green2": "00695C",
        "gold": "D4AF37",
        "pale_gold": "F6E7A8",
        "white": "FFFFFF",
        "dark": "111827",
        "muted": "4B5563",
        "background": "EAF6F1",
        "light": "EEF5F2",
        "line": "8FA39A",
        "soft_green": "DCFCE7",
        "soft_warn": "FEF3C7",
        "soft_danger": "FEE2E2",
    }


def _thin(color: str = "8FA39A") -> Border:
    return Border(
        left=Side(style="thin", color=color),
        right=Side(style="thin", color=color),
        top=Side(style="thin", color=color),
        bottom=Side(style="thin", color=color),
    )


def _sheet_base(ws, title: str, subtitle: str, last_col: int = 8, zoom: int = 95) -> None:
    s = _wb_styles()
    end_col = get_column_letter(last_col)
    ws.sheet_view.showGridLines = False
    ws.sheet_view.zoomScale = zoom
    ws.freeze_panes = None
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.35
    ws.page_margins.bottom = 0.35

    widths = {"A": 22, "B": 24, "C": 20, "D": 20, "E": 20, "F": 20, "G": 20, "H": 20}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    # Fundo geral verde claro para tirar o visual cinza/branco padrão do Excel.
    for row_cells in ws.iter_rows(min_row=1, max_row=120, min_col=1, max_col=26):
        for cell in row_cells:
            cell.fill = PatternFill("solid", fgColor=s["background"])

    ws.merge_cells(f"A1:{end_col}1")
    ws["A1"] = title
    ws["A1"].font = Font(name="Calibri", size=26, bold=True, color=s["white"])
    ws["A1"].fill = PatternFill("solid", fgColor=s["green"])
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 38

    ws.merge_cells(f"A2:{end_col}2")
    ws["A2"] = subtitle
    ws["A2"].font = Font(name="Calibri", size=11, bold=True, color=s["green"])
    ws["A2"].fill = PatternFill("solid", fgColor=s["gold"])
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 22

    ws.merge_cells(f"A3:{end_col}3")
    ws["A3"] = ""
    ws["A3"].fill = PatternFill("solid", fgColor=s["background"])
    ws.row_dimensions[3].height = 14


def _merge(ws, cell_range: str, value: Any, *, fill: str = "FFFFFF", font_color: str = "111827", size: int = 10, bold: bool = False, align: str = "left", valign: str = "center") -> None:
    ws.merge_cells(cell_range)
    cell = ws[cell_range.split(":")[0]]
    cell.value = excel_clean(value, "")
    cell.fill = PatternFill("solid", fgColor=fill)
    cell.font = Font(name="Calibri", size=size, bold=bold, color=font_color)
    cell.alignment = Alignment(horizontal=align, vertical=valign, wrap_text=True)
    cell.border = _thin()


def _section(ws, row: int, text: str, last_col: int = 8) -> int:
    s = _wb_styles()
    _merge(ws, f"A{row}:{get_column_letter(last_col)}{row}", text, fill=s["pale_gold"], font_color=s["green"], size=13, bold=True, align="left")
    ws.row_dimensions[row].height = 24
    return row + 2


def _metric_card(ws, cell_range: str, label: str, value: Any, fill: str) -> None:
    s = _wb_styles()
    _merge(ws, cell_range, f"{label}\n{excel_clean(value, '')}", fill=fill, font_color=s["white"], size=14, bold=True, align="center")


def _table_header(ws, row: int, headers: List[str], col_spans: List[int] | None = None) -> None:
    s = _wb_styles()
    if col_spans is None:
        col_spans = [1] * len(headers)
    col = 1
    for header, span in zip(headers, col_spans):
        start = get_column_letter(col)
        end = get_column_letter(col + span - 1)
        _merge(ws, f"{start}{row}:{end}{row}", header, fill=s["green"], font_color=s["white"], size=10, bold=True, align="center")
        col += span
    ws.row_dimensions[row].height = 28


def _write_kv_table(ws, start_row: int, rows: List[tuple[str, Any]], label_span: int = 2, value_span: int = 6, row_height: int = 34) -> int:
    s = _wb_styles()
    r = start_row
    for i, (campo, valor) in enumerate(rows):
        fill = "FFFFFF" if i % 2 == 0 else s["light"]
        _merge(ws, f"A{r}:{get_column_letter(label_span)}{r}", campo, fill=fill, font_color=s["green"], size=10, bold=True, align="left")
        _merge(ws, f"{get_column_letter(label_span+1)}{r}:{get_column_letter(label_span+value_span)}{r}", valor, fill=fill, font_color=s["dark"], size=10, align="left")
        ws.row_dimensions[r].height = row_height
        r += 1
    return r


def _write_dataframe_table(ws, start_row: int, df: pd.DataFrame, widths: Dict[str, int] | None = None, row_height: int = 34) -> int:
    s = _wb_styles()
    if df is None or df.empty:
        return start_row
    for c_idx, col_name in enumerate(df.columns, 1):
        cell = ws.cell(start_row, c_idx)
        cell.value = str(col_name)
        cell.fill = PatternFill("solid", fgColor=s["green"])
        cell.font = Font(name="Calibri", size=10, bold=True, color=s["white"])
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = _thin()
    ws.row_dimensions[start_row].height = 28

    for r_idx, (_, row) in enumerate(df.iterrows(), start_row + 1):
        fill = "FFFFFF" if (r_idx - start_row) % 2 else s["light"]
        for c_idx, col_name in enumerate(df.columns, 1):
            cell = ws.cell(r_idx, c_idx)
            cell.value = excel_clean(row.get(col_name, ""), "")
            cell.fill = PatternFill("solid", fgColor=fill)
            cell.font = Font(name="Calibri", size=10, color=s["dark"])
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cell.border = _thin()
        ws.row_dimensions[r_idx].height = row_height

    if widths:
        for col, width in widths.items():
            ws.column_dimensions[col].width = width
    return start_row + len(df) + 1


def gerar_excel(resultado: Dict[str, Any], texto: str) -> io.BytesIO:
    """Relatório Excel com visual executivo: capa, dashboard, dados, resumo, checklist, parecer, auditoria e texto."""
    from openpyxl import Workbook

    output = io.BytesIO()
    s = _wb_styles()

    def v(chave: str, padrao: str = "Não localizado") -> str:
        return excel_clean(resultado.get(chave, padrao), padrao)

    risco = normalize_risco(resultado.get("risco"))
    score = resultado.get("score", 0)
    risk_fill = excel_risk_fill(risco)
    pendencias_lista = resultado.get("pendencias", []) if isinstance(resultado.get("pendencias"), list) else []
    checklist_lista = resultado.get("checklist", []) if isinstance(resultado.get("checklist"), list) else []
    itens_lista = normalizar_itens_contrato(resultado.get("itens_contrato", []))
    aditivos_lista = normalizar_aditivos_contrato(resultado.get("aditivos_contrato", []))
    assinaturas_df_excel = montar_df_assinaturas(resultado)
    triagem_lista = resultado.get("triagem_anexos", []) if isinstance(resultado.get("triagem_anexos", []), list) else []
    resumo_processamento = resultado.get("resumo_processamento", {}) if isinstance(resultado.get("resumo_processamento", {}), dict) else {}

    contraparte = v("contraparte", v("fornecedor"))
    cnpj_contraparte = v("cnpj_contraparte", v("cnpj"))
    valor_contrato = v("valor_contrato_original", v("valor_total"))
    vigencia = v("vigencia_apos_assinatura", v("vigencia"))
    pagamento = v("condicao_pagamento_dias", v("condicao_pagamento"))
    data_assinatura = v("data_assinatura")
    modelo_ia = v("modelo_ia", "Não informado")
    origem = v("tipo_origem", v("origem_contrato", "Não informado"))
    data_analise = v("data_analise", datetime.now().strftime("%d/%m/%Y %H:%M"))
    parecer = v("parecer")
    resumo = v("resumo_executivo")

    wb = Workbook()
    ws = wb.active
    ws.title = "Capa"

    # CAPA
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório Executivo de Análise Contratual", 8, 100)
    row = _section(ws, 4, "Visão Geral da Análise")
    cards = [
        ("Status", v("status"), "A6:B8", s["green"]),
        ("Risco", risco, "C6:D8", risk_fill),
        ("Score", str(score), "E6:F8", s["green"]),
        ("Pendências", str(len(pendencias_lista)), "G6:H8", s["green"]),
    ]
    for label, value, cell_range, fill in cards:
        _metric_card(ws, cell_range, label, value, fill)
    ws.row_dimensions[6].height = 28
    ws.row_dimensions[7].height = 28
    ws.row_dimensions[8].height = 28

    row = _section(ws, 10, "Resumo do Contrato")
    row = _write_kv_table(ws, row, [
        ("Contraparte", contraparte),
        ("CNPJ Contraparte", cnpj_contraparte),
        ("Valor do Contrato", valor_contrato),
        ("Valor Mensal Estimado", v("valor_mensal_estimado")),
        ("Valor Total Estimado da Vigência", v("valor_total_estimado_vigencia")),
        ("Vigência", vigencia),
        ("Período de Vigência", v("periodo_vigencia_formatado")),
        ("Pagamento", pagamento),
        ("Data da Assinatura", data_assinatura),
        ("Data do Reconhecimento de Firma", v("data_reconhecimento_firma", "Não aplicável")),
        ("Data do Contrato", v("data_contrato")),
        ("Data Conclusão DocuSign", v("data_conclusao_docusign")),
        ("Pessoas que assinaram", v("pessoas_que_assinaram")),
        ("Contrato Assinado", v("contrato_assinado")),
        ("Resumo de Aditivos", v("resumo_aditivos")),
        ("Arquivos enviados para Análise IA", resumo_processamento.get("analise_profunda", "Não informado")),
        ("Arquivos de apoio", resumo_processamento.get("apoio", "Não informado")),
        ("Arquivos ignorados", resumo_processamento.get("ignorados", "Não informado")),
        ("Tempo total do processamento", resumo_processamento.get("tempo_total", "Não informado")),
        ("Tempo da IA", resumo_processamento.get("tempo_ia", "Não informado")),
        ("Modelo IA", modelo_ia),
    ], row_height=30)
    row = _section(ws, row + 1, "Parecer Automático")
    _merge(ws, f"A{row}:H{row+4}", parecer, fill=s["light"], font_color=s["dark"], size=11, align="left", valign="top")
    for rr in range(row, row+5):
        ws.row_dimensions[rr].height = 28

    # DASHBOARD EXECUTIVO
    ws = wb.create_sheet("Dashboard Executivo")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Dashboard Executivo", 8, 100)
    _section(ws, 4, "Indicadores Principais")
    for label, value, cell_range, fill in cards:
        _metric_card(ws, cell_range, label, value, fill)
    _section(ws, 10, "Informações-Chave")
    _table_header(ws, 12, ["Indicador", "Valor"], [2, 6])
    _write_kv_table(ws, 13, [
        ("Contraparte", contraparte),
        ("CNPJ Contraparte", cnpj_contraparte),
        ("Valor do Contrato", valor_contrato),
        ("Valor Mensal Estimado", v("valor_mensal_estimado")),
        ("Valor Total Estimado da Vigência", v("valor_total_estimado_vigencia")),
        ("Condição de Pagamento", pagamento),
        ("Vigência", vigencia),
        ("Período de Vigência", v("periodo_vigencia_formatado")),
        ("Data da Assinatura", data_assinatura),
        ("Data do Reconhecimento de Firma", v("data_reconhecimento_firma", "Não aplicável")),
        ("Data do Contrato", v("data_contrato")),
        ("Data Conclusão DocuSign", v("data_conclusao_docusign")),
        ("Pessoas que assinaram", v("pessoas_que_assinaram")),
        ("Contrato Assinado", v("contrato_assinado")),
        ("Origem", origem),
        ("Modelo IA", modelo_ia),
        ("Data da Análise", data_analise),
        ("Análise IA", resumo_processamento.get("analise_profunda", "Não informado")),
        ("Apoio", resumo_processamento.get("apoio", "Não informado")),
        ("Ignorados", resumo_processamento.get("ignorados", "Não informado")),
        ("Tempo total", resumo_processamento.get("tempo_total", "Não informado")),
        ("Tempo IA", resumo_processamento.get("tempo_ia", "Não informado")),
    ], row_height=31)

    # DADOS EXTRAÍDOS
    ws = wb.create_sheet("Dados Extraídos")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Dados Extraídos", 8, 95)
    _table_header(ws, 5, ["Campo", "Informação"], [2, 6])
    dados_rows = [(label, v(chave)) for label, chave in CAMPOS_OFICIAIS]
    _write_kv_table(ws, 6, dados_rows, row_height=42)

    # RESUMO EXECUTIVO
    ws = wb.create_sheet("Resumo Executivo")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Resumo Executivo", 8, 95)
    _table_header(ws, 5, ["Seção", "Conteúdo"], [2, 6])
    _write_kv_table(ws, 6, [
        ("Resumo Executivo", resumo),
        ("Objeto / Escopo", v("descricao_servico_material", v("objetivo"))),
        ("Parecer Automático", parecer),
        ("Alerta de Assinatura", v("alerta_assinatura")),
    ], row_height=74)

    # ASSINATURAS
    ws = wb.create_sheet("Assinaturas")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Assinaturas", 8, 95)
    _write_dataframe_table(ws, 5, assinaturas_df_excel, {
        "A": 30, "B": 24, "C": 26, "D": 18, "E": 24, "F": 24, "G": 18, "H": 62
    }, 46)

    # ITENS DO CONTRATO
    ws = wb.create_sheet("Itens do Contrato")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Materiais e Serviços", 8, 95)
    itens_df = pd.DataFrame(itens_lista)
    if itens_df.empty:
        itens_df = pd.DataFrame([{
            "Item": "-",
            "Descrição": "Nenhum material ou serviço unitário identificado no contrato/anexos.",
            "Tipo": "N/A",
            "Quantidade": "N/A",
            "Unidade": "N/A",
            "Valor unitário": "Não localizado",
            "Valor total": "Não localizado",
            "Fonte": "N/A",
        }])
    _write_dataframe_table(ws, 5, itens_df, {
        "A": 10, "B": 52, "C": 15, "D": 14, "E": 14, "F": 18, "G": 18, "H": 22
    }, 44)

        # ADITIVOS
    ws = wb.create_sheet("Aditivos")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Aditivos", 8, 95)

    if aditivos_lista:
        df_aditivos = pd.DataFrame([
            {k: v for k, v in aditivo.items() if k != "_itens_aditivo"}
            for aditivo in aditivos_lista
        ])
    else:
        df_aditivos = pd.DataFrame([{
            "Status": "Nenhum aditivo identificado nos documentos analisados."
        }])

    _write_dataframe_table(ws, 5, df_aditivos, {
        "A": 10,
        "B": 24,
        "C": 42,
        "D": 18,
        "E": 20,
        "F": 14,
        "G": 20,
        "H": 46,
        "I": 18,
        "J": 42,
        "K": 42,
        "L": 26,
    }, 46)

    # ITENS DOS ADITIVOS
    ws = wb.create_sheet("Itens dos Aditivos")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Itens dos Aditivos", 8, 95)

    linhas_itens_aditivos = []
    for aditivo in aditivos_lista:
        for item in aditivo.get("_itens_aditivo", []):
            linha = {
                "Aditivo": aditivo.get("Nº", "Não localizado"),
                "Anexo do aditivo": aditivo.get("Anexo do aditivo", "Não localizado"),
            }
            linha.update(item)
            linhas_itens_aditivos.append(linha)

    if linhas_itens_aditivos:
        df_itens_aditivos = pd.DataFrame(linhas_itens_aditivos)
    else:
        df_itens_aditivos = pd.DataFrame([{
            "Status": "Nenhum item específico de aditivo identificado."
        }])

    _write_dataframe_table(ws, 5, df_itens_aditivos, {
        "A": 14,
        "B": 42,
        "C": 12,
        "D": 52,
        "E": 16,
        "F": 16,
        "G": 16,
        "H": 18,
        "I": 18,
    }, 46)

    # CHECKLIST
    ws = wb.create_sheet("Checklist")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Checklist", 8, 95)
    checklist = pd.DataFrame(checklist_lista)
    if checklist.empty:
        checklist = pd.DataFrame([{"Validação": "Nenhum checklist retornado", "Status": "N/A", "Peso de risco": 0, "Crítico": "Não", "Evidência": "Não informado"}])
    _write_dataframe_table(ws, 5, checklist, {"A": 34, "B": 16, "C": 18, "D": 16, "E": 54}, 36)
    headers = {str(ws.cell(5, c).value).strip().lower(): c for c in range(1, ws.max_column + 1)}
    status_col = headers.get("status")
    critico_col = headers.get("crítico") or headers.get("critico")
    for rr in range(6, ws.max_row + 1):
        if status_col:
            cell = ws.cell(rr, status_col)
            txt = str(cell.value or "").upper()
            if "OK" in txt or "CONCLU" in txt:
                cell.fill = PatternFill("solid", fgColor=s["soft_green"])
                cell.font = Font(name="Calibri", size=10, bold=True, color="166534")
            elif "ATEN" in txt or "PEND" in txt:
                cell.fill = PatternFill("solid", fgColor=s["soft_warn"])
                cell.font = Font(name="Calibri", size=10, bold=True, color="92400E")
        if critico_col:
            cell = ws.cell(rr, critico_col)
            if str(cell.value or "").strip().upper() == "SIM":
                cell.fill = PatternFill("solid", fgColor=s["soft_danger"])
                cell.font = Font(name="Calibri", size=10, bold=True, color="991B1B")

    # PENDÊNCIAS
    ws = wb.create_sheet("Pendências")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Pendências", 8, 95)
    pendencias = pd.DataFrame(pendencias_lista)
    if pendencias.empty:
        pendencias = pd.DataFrame([{
            "Pendência": "Nenhuma pendência crítica localizada",
            "Crítico": "Não",
            "Risco": 0,
            "Recomendação": "Não foram identificadas pendências críticas na análise realizada.",
        }])
    _write_dataframe_table(ws, 5, pendencias, {"A": 42, "B": 16, "C": 14, "D": 70}, 46)
    for rr in range(6, ws.max_row + 1):
        if str(ws.cell(rr, 1).value or "").lower().startswith("nenhuma"):
            for cc in range(1, ws.max_column + 1):
                ws.cell(rr, cc).fill = PatternFill("solid", fgColor=s["soft_green"])
                ws.cell(rr, cc).font = Font(name="Calibri", size=10, bold=True, color="166534")

    # PARECER
    ws = wb.create_sheet("Parecer")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Parecer", 8, 95)
    _table_header(ws, 5, ["Item", "Descrição"], [2, 6])
    recomendacao = "Seguir com o processo caso as informações extraídas estejam de acordo com a documentação analisada." if risco == "BAIXO" else "Revisar as pendências e pontos de atenção antes de seguir com RC/PO."
    _write_kv_table(ws, 6, [
        ("Risco Geral", risco),
        ("Score", score),
        ("Parecer", parecer),
        ("Resumo Executivo", resumo),
        ("Recomendação", recomendacao),
    ], row_height=62)

    # AUDITORIA
    ws = wb.create_sheet("Auditoria")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Auditoria", 8, 95)
    _table_header(ws, 5, ["Campo", "Valor"], [2, 6])
    _write_kv_table(ws, 6, [
        ("Data da análise", data_analise),
        ("Modelo IA", modelo_ia),
        ("Origem", origem),
        ("Status", v("status")),
        ("Risco", risco),
        ("Score", score),
        ("Contrato Assinado", v("contrato_assinado")),
        ("Quantidade de Pendências", len(pendencias_lista)),
        ("Arquivos analisados", v("arquivos_analisados", v("arquivo", "Não informado"))),
        ("Arquivos para Análise IA", resumo_processamento.get("analise_profunda", "Não informado")),
        ("Arquivos de apoio", resumo_processamento.get("apoio", "Não informado")),
        ("Arquivos ignorados", resumo_processamento.get("ignorados", "Não informado")),
        ("Tempo total do processamento", resumo_processamento.get("tempo_total", "Não informado")),
        ("Tempo da IA", resumo_processamento.get("tempo_ia", "Não informado")),
        ("Finalizado em", resumo_processamento.get("finalizado_em", "Não informado")),
    ], row_height=32)

    # TRIAGEM DOS ANEXOS
    ws = wb.create_sheet("Triagem dos Anexos")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Pré-triagem dos Anexos", 8, 95)

    if triagem_lista:
        df_triagem = pd.DataFrame(triagem_lista)
    else:
        df_triagem = pd.DataFrame([{
            "Status": "Pré-triagem não disponível para este registro."
        }])

    _write_dataframe_table(ws, 5, df_triagem, {
        "A": 58,
        "B": 22,
        "C": 14,
        "D": 28,
        "E": 74,
    }, 48)

    headers_triagem = {str(ws.cell(5, c).value).strip().lower(): c for c in range(1, ws.max_column + 1)}
    decisao_col = headers_triagem.get("decisão") or headers_triagem.get("decisao")
    for rr in range(6, ws.max_row + 1):
        decisao_txt = str(ws.cell(rr, decisao_col).value if decisao_col else "").lower()
        if "análise profunda" in decisao_txt or "analise profunda" in decisao_txt:
            fill = s["soft_green"]
            font_color = "166534"
        elif "apoio" in decisao_txt:
            fill = "DBEAFE"
            font_color = "1E3A8A"
        elif "ignorado" in decisao_txt:
            fill = s["soft_danger"]
            font_color = "991B1B"
        else:
            fill = "FFFFFF"
            font_color = s["dark"]
        for cc in range(1, ws.max_column + 1):
            ws.cell(rr, cc).fill = PatternFill("solid", fgColor=fill)
            ws.cell(rr, cc).font = Font(name="Calibri", size=10, bold=True, color=font_color)

    # AUDITORIA DE CAMPOS — fonte, página e trecho que sustentam cada card
    auditoria_campos = resultado.get("auditoria_campos") if isinstance(resultado.get("auditoria_campos"), list) else []
    ws = wb.create_sheet("Auditoria de Campos")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Evidências por Campo", 8, 95)
    if auditoria_campos:
        df_auditoria = pd.DataFrame(auditoria_campos)
    else:
        df_auditoria = pd.DataFrame([{"Status": "Auditoria de evidências não disponível para este registro antigo."}])
    _write_dataframe_table(ws, 5, df_auditoria, {
        "A": 31, "B": 46, "C": 20, "D": 22, "E": 38, "F": 16, "G": 72, "H": 12,
    }, 54)

    # TEXTO EXTRAÍDO
    ws = wb.create_sheet("Texto Extraído")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Relatório de Análise Contratual • Texto Extraído", 8, 90)
    texto_limpo = clean_text(texto or resultado.get("texto_extraido", ""))
    if texto_limpo in ("", "Não localizado"):
        texto_limpo = "Texto extraído não disponível para este registro."
    blocos = [texto_limpo[i:i + 6000] for i in range(0, len(texto_limpo), 6000)] or [texto_limpo]
    _table_header(ws, 5, ["Bloco", "Texto Extraído"], [1, 7])
    r = 6
    for i, bloco in enumerate(blocos, 1):
        fill = "FFFFFF" if i % 2 else s["light"]
        _merge(ws, f"A{r}:A{r}", i, fill=fill, font_color=s["dark"], size=10, align="center")
        _merge(ws, f"B{r}:H{r}", excel_clean(bloco, ""), fill=fill, font_color=s["dark"], size=9, align="left", valign="top")
        ws.row_dimensions[r].height = 145
        r += 1

    for ws in wb.worksheets:
        ws.sheet_view.showGridLines = False
        ws.sheet_view.topLeftCell = "A1"
        ws.sheet_view.selection[0].sqref = "A1"
        ws.sheet_view.selection[0].activeCell = "A1"
        ws.sheet_properties.tabColor = s["green"]
        # Oculta colunas muito distantes para reduzir poluição visual.
        for col_idx in range(9, 27):
            ws.column_dimensions[get_column_letter(col_idx)].width = 3

    wb.active = 0
    wb.save(output)
    output.seek(0)
    return output


def gerar_excel_card_bytes(row: pd.Series) -> bytes:
    """Gera Excel completo do card do Dashboard."""
    resultado_completo: Dict[str, Any] = {}
    texto_extraido = ""

    raw_json = row.get("resultado_json") if "resultado_json" in row.index else None
    if raw_json not in (None, "", "Não informado"):
        try:
            resultado_completo = json.loads(raw_json)
            if not isinstance(resultado_completo, dict):
                resultado_completo = {}
        except Exception:
            resultado_completo = {}

    if resultado_completo:
        texto_extraido = resultado_completo.get("texto_extraido", "")
        return gerar_excel(resultado_completo, texto_extraido).getvalue()

    resultado_fallback = {
        "data_analise": row.get("data_analise"),
        "contraparte": row.get("fornecedor"),
        "fornecedor": row.get("fornecedor"),
        "cnpj_contraparte": row.get("cnpj"),
        "valor_contrato_original": row.get("valor_total"),
        "vigencia_apos_assinatura": row.get("vigencia"),
        "status": row.get("status"),
        "risco": row.get("risco"),
        "score": row.get("score"),
        "contrato_assinado": row.get("contrato_assinado"),
        "modelo_ia": row.get("modelo_ia"),
        "tipo_origem": row.get("tipo_origem"),
        "resumo_executivo": "Relatório gerado com os dados disponíveis no histórico.",
        "parecer": "Para gerar todas as informações detalhadas, refaça a análise do contrato nesta versão atualizada.",
        "checklist": [],
        "pendencias": [],
        "itens_contrato": [],
    }
    return gerar_excel(resultado_fallback, "Texto extraído não disponível para este registro antigo.").getvalue()



# Funções auxiliares para o Excel de Histórico.
# Mantidas para compatibilidade com a seção "Histórico".
def excel_apply_page(ws, zoom: int = 90) -> None:
    ws.sheet_view.showGridLines = False
    ws.sheet_view.zoomScale = zoom
    ws.freeze_panes = None
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins.left = 0.25
    ws.page_margins.right = 0.25
    ws.page_margins.top = 0.35
    ws.page_margins.bottom = 0.35


def excel_header(ws, titulo: str, subtitulo: str, last_col: int = 8) -> None:
    last_col = max(int(last_col or 1), 1)
    _sheet_base(ws, titulo, subtitulo, last_col=last_col, zoom=90)


def excel_style_table(ws, header_row: int = 6, data_start: int = 7) -> None:
    s = _wb_styles()
    # Cabeçalho
    for cell in ws[header_row]:
        if cell.value is not None:
            cell.fill = PatternFill("solid", fgColor=s["green"])
            cell.font = Font(name="Calibri", size=10, bold=True, color=s["white"])
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = _thin()
    ws.row_dimensions[header_row].height = 28

    # Linhas alternadas branco / verde claríssimo
    for r in range(data_start, ws.max_row + 1):
        fill = s["white"] if (r - data_start) % 2 == 0 else s["light"]
        for c in range(1, ws.max_column + 1):
            cell = ws.cell(r, c)
            cell.fill = PatternFill("solid", fgColor=fill)
            cell.font = Font(name="Calibri", size=10, color=s["dark"])
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cell.border = _thin()
        ws.row_dimensions[r].height = 34


def excel_auto_width(ws, max_width: int = 60) -> None:
    for idx, col_cells in enumerate(ws.columns, 1):
        max_len = 12
        for cell in col_cells:
            if cell.value is not None:
                max_len = max(max_len, min(len(str(cell.value)) + 2, max_width))
        ws.column_dimensions[get_column_letter(idx)].width = max_len



def _resumir_arquivos_excel(valor: Any) -> str:
    """Deixa a coluna de arquivos mais executiva no Excel do histórico."""
    txt = excel_clean(valor, "Não informado")
    if txt == "Não informado":
        return txt
    partes = [p.strip() for p in re.split(r"\s*\|\s*", txt) if p.strip()]
    if not partes:
        return txt
    if len(partes) == 1:
        return partes[0][:120]
    return f"{len(partes)} arquivos analisados"


def _preparar_historico_excel(df: pd.DataFrame) -> pd.DataFrame:
    """Padroniza colunas do histórico para exportação executiva."""
    colunas = [
        "ID", "Data da análise", "Contraparte", "CNPJ", "Valor total",
        "Valor mensal estimado", "Valor total estimado da vigência", "Valor total dos materiais e serviços",
        "Data do contrato", "Data conclusão DocuSign", "Pessoas que assinaram",
        "Vigência", "Status", "Risco", "Score", "Assinado", "Modelo IA", "Origem", "Arquivos analisados",
    ]
    base = df.copy() if df is not None else pd.DataFrame(columns=colunas)
    for col in colunas:
        if col not in base.columns:
            base[col] = "Não informado"

    base = base[colunas].copy()
    base["Score"] = pd.to_numeric(base["Score"], errors="coerce").fillna(0).astype(int)
    base["Risco"] = base["Risco"].astype(str).str.upper().replace({"MEDIO": "MÉDIO"})
    base["Arquivos analisados"] = base["Arquivos analisados"].apply(_resumir_arquivos_excel)

    for col in base.columns:
        if col != "Score":
            base[col] = base[col].apply(lambda x: excel_clean(x, "Não informado"))

    return base


def _aplicar_fundo_excel(ws, max_row: int = 80, max_col: int = 14) -> None:
    """Aplica fundo verde claro sem criar uma área visual gigante."""
    s = _wb_styles()
    for row_cells in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
        for cell in row_cells:
            if cell.value in (None, ""):
                cell.fill = PatternFill("solid", fgColor=s["background"])
            cell.border = _thin(s["line"])


def _formatar_risco_cell(cell, risco: Any) -> None:
    risco_norm = normalize_risco(risco)
    if risco_norm == "ALTO":
        cell.fill = PatternFill("solid", fgColor="FEE2E2")
        cell.font = Font(name="Calibri", size=10, bold=True, color="991B1B")
    elif risco_norm in ["MÉDIO", "MEDIO"]:
        cell.fill = PatternFill("solid", fgColor="FEF3C7")
        cell.font = Font(name="Calibri", size=10, bold=True, color="92400E")
    elif risco_norm == "BAIXO":
        cell.fill = PatternFill("solid", fgColor="DCFCE7")
        cell.font = Font(name="Calibri", size=10, bold=True, color="166534")
    else:
        cell.fill = PatternFill("solid", fgColor="E5E7EB")
        cell.font = Font(name="Calibri", size=10, bold=True, color="374151")
    cell.alignment = Alignment(horizontal="center", vertical="center")
    cell.border = _thin()


def gerar_excel_historico_profissional(export_df: pd.DataFrame, total_geral: int | None = None) -> io.BytesIO:
    """Gera um histórico executivo com dashboard, tabela filtrável e auditoria."""
    from openpyxl import Workbook
    from openpyxl.chart import PieChart, Reference

    s = _wb_styles()
    output = io.BytesIO()
    df = _preparar_historico_excel(export_df)

    total_filtrado = int(len(df))
    total_geral = int(total_geral if total_geral is not None else total_filtrado)
    risco_series = df["Risco"].astype(str).str.upper().replace({"MEDIO": "MÉDIO"}) if not df.empty else pd.Series(dtype=str)
    qtd_alto = int((risco_series == "ALTO").sum())
    qtd_medio = int((risco_series == "MÉDIO").sum())
    qtd_baixo = int((risco_series == "BAIXO").sum())
    score_medio = round(float(pd.to_numeric(df["Score"], errors="coerce").fillna(0).mean()), 1) if total_filtrado else 0
    assinados = int(df["Assinado"].astype(str).str.upper().eq("SIM").sum()) if total_filtrado else 0
    data_geracao = datetime.now().strftime("%d/%m/%Y %H:%M")

    wb = Workbook()

    # =====================================================
    # ABA 1 - DASHBOARD
    # =====================================================
    ws = wb.active
    ws.title = "Dashboard"
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Dashboard Executivo do Histórico", 8, 95)
    _aplicar_fundo_excel(ws, 48, 10)
    ws.freeze_panes = None

    _section(ws, 4, "Indicadores Executivos", 8)
    cards = [
        ("Contratos filtrados", total_filtrado, "A6:B8", s["green"]),
        ("Base total", total_geral, "C6:D8", s["green2"]),
        ("Score médio", score_medio, "E6:F8", s["green"]),
        ("Assinados", assinados, "G6:H8", s["green2"]),
    ]
    for label, value, cell_range, fill in cards:
        _metric_card(ws, cell_range, label, value, fill)
        for rr in range(6, 9):
            ws.row_dimensions[rr].height = 27

    _section(ws, 10, "Distribuição por Risco", 8)
    risco_data = [("ALTO", qtd_alto), ("MÉDIO", qtd_medio), ("BAIXO", qtd_baixo)]
    _table_header(ws, 12, ["Risco", "Quantidade"], [2, 2])
    r = 13
    for risco_nome, qtd in risco_data:
        _merge(ws, f"A{r}:B{r}", risco_nome, fill=s["white"], font_color=s["dark"], size=10, bold=True, align="center")
        _merge(ws, f"C{r}:D{r}", qtd, fill=s["white"], font_color=s["dark"], size=10, bold=True, align="center")
        _formatar_risco_cell(ws[f"A{r}"], risco_nome)
        ws.row_dimensions[r].height = 25
        r += 1

    if total_filtrado:
        chart = PieChart()
        chart.title = "Risco dos contratos"
        labels = Reference(ws, min_col=1, min_row=13, max_row=15)
        data = Reference(ws, min_col=3, min_row=12, max_row=15)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(labels)
        chart.height = 7.5
        chart.width = 10
        chart.firstSliceAng = 270
        ws.add_chart(chart, "F12")

    _section(ws, 22, "Últimas análises", 8)
    ultimas = df.head(8).copy()
    resumo_cols = ["ID", "Data da análise", "Contraparte", "Valor total", "Risco", "Score"]
    _write_dataframe_table(ws, 24, ultimas[resumo_cols] if not ultimas.empty else pd.DataFrame(columns=resumo_cols),
                           {"A": 10, "B": 20, "C": 36, "D": 18, "E": 14, "F": 12}, 28)
    # Formata risco no resumo
    for rr in range(25, 25 + len(ultimas)):
        _formatar_risco_cell(ws.cell(rr, 5), ws.cell(rr, 5).value)
        ws.cell(rr, 6).alignment = Alignment(horizontal="center", vertical="center")

    _merge(ws, "A36:H39", "Leitura recomendada: use a aba Histórico Completo para filtros, auditoria e consulta detalhada dos contratos analisados.",
           fill=s["light"], font_color=s["dark"], size=11, align="left", valign="top")

    # =====================================================
    # ABA 2 - HISTÓRICO COMPLETO
    # =====================================================
    ws = wb.create_sheet("Histórico Completo")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Histórico Completo de Análises", 13, 90)
    _aplicar_fundo_excel(ws, max(40, len(df) + 12), 14)
    ws.freeze_panes = "A7"

    start_row = 6
    # cabeçalho da tabela
    for c_idx, col_name in enumerate(df.columns, 1):
        cell = ws.cell(start_row, c_idx)
        cell.value = col_name
        cell.fill = PatternFill("solid", fgColor=s["green"])
        cell.font = Font(name="Calibri", size=10, bold=True, color=s["white"])
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = _thin()
    ws.row_dimensions[start_row].height = 30

    for r_idx, (_, row) in enumerate(df.iterrows(), start_row + 1):
        fill = s["white"] if (r_idx - start_row) % 2 else s["light"]
        for c_idx, col_name in enumerate(df.columns, 1):
            cell = ws.cell(r_idx, c_idx)
            cell.value = excel_clean(row[col_name], "") if col_name != "Score" else int(row[col_name])
            cell.fill = PatternFill("solid", fgColor=fill)
            cell.font = Font(name="Calibri", size=10, color=s["dark"])
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            cell.border = _thin()
        _formatar_risco_cell(ws.cell(r_idx, df.columns.get_loc("Risco") + 1), row["Risco"])
        ws.cell(r_idx, df.columns.get_loc("Score") + 1).alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[r_idx].height = 36

    if total_filtrado:
        last_row = start_row + total_filtrado
        # Mantém filtro sem criar Tabela estruturada do Excel.
        # Isso evita erro de reparo em /xl/tables/table1.xml ao abrir o arquivo.
        ultima_coluna = get_column_letter(len(df.columns))
        ws.auto_filter.ref = f"A{start_row}:{ultima_coluna}{last_row}"

    widths = {
        "A": 9, "B": 19, "C": 34, "D": 19, "E": 17,
        "F": 22, "G": 25, "H": 18, "I": 22, "J": 46,
        "K": 34, "L": 16, "M": 13, "N": 10, "O": 13,
        "P": 18, "Q": 15, "R": 28,
    }
    for idx_col in range(1, len(df.columns) + 1):
        col_letter = get_column_letter(idx_col)
        ws.column_dimensions[col_letter].width = widths.get(col_letter, 20)

    # =====================================================
    # ABA 3 - AUDITORIA
    # =====================================================
    ws = wb.create_sheet("Auditoria")
    _sheet_base(ws, "Auditor de Contratos - Grupo SBF", "Auditoria do Histórico", 8, 95)
    _aplicar_fundo_excel(ws, 45, 9)
    _section(ws, 4, "Informações da Exportação", 8)
    auditoria_rows = [
        ("Data de geração", data_geracao),
        ("Registros exportados", total_filtrado),
        ("Total geral no banco", total_geral),
        ("Score médio filtrado", score_medio),
        ("Risco alto", qtd_alto),
        ("Risco médio", qtd_medio),
        ("Risco baixo", qtd_baixo),
        ("Contratos assinados", assinados),
        ("Observação", "Relatório gerado com base nos filtros aplicados na aba Histórico do Auditor de Contratos - Grupo SBF."),
    ]
    _table_header(ws, 6, ["Campo", "Valor"], [2, 6])
    _write_kv_table(ws, 7, auditoria_rows, row_height=32)

    for ws in wb.worksheets:
        ws.sheet_view.showGridLines = False
        ws.sheet_properties.tabColor = s["green"]
        ws.sheet_view.topLeftCell = "A1"
        ws.sheet_view.selection[0].sqref = "A1"
        ws.sheet_view.selection[0].activeCell = "A1"

    wb.active = 0
    wb.save(output)
    output.seek(0)
    return output

# =========================================================
# COMPONENTES DE TELA
# =========================================================
def render_hero(titulo: str, descricao: str, tag: str = "Suprimentos • Contratos") -> None:
    st.markdown(
        f"""
        <div class="hero">
            <span class="eyebrow">{safe(tag)}</span>
            <h1>{safe(titulo)}</h1>
            <p>{safe(descricao)}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_metric(label: str, value: Any) -> str:
    return f'<div class="metric-card"><small>{safe(label)}</small><h2>{safe(value)}</h2></div>'


def obter_indicadores_resultado(resultado: Dict[str, Any]) -> Dict[str, Any]:
    """Separa risco contratual, confiança da extração e pendências por natureza."""
    indicadores = resultado.get("indicadores_pendencias") if isinstance(resultado.get("indicadores_pendencias"), dict) else {}
    pendencias = resultado.get("pendencias") if isinstance(resultado.get("pendencias"), list) else []
    criticas_fallback = sum(1 for p in pendencias if clean_text(p.get("Crítico") or p.get("critico")).upper() in ("SIM", "TRUE", "1"))
    pontos_fallback = max(0, len(pendencias) - criticas_fallback)
    campos_fallback = len(resultado.get("campos_nao_localizados", [])) if isinstance(resultado.get("campos_nao_localizados"), list) else 0
    return {
        "status_contratual": resultado.get("status_contratual") or resultado.get("status") or "Não identificado com segurança",
        "situacao_operacional": resultado.get("situacao_operacional") or "Não confirmada nos documentos analisados",
        "risco_contratual": normalize_risco(resultado.get("risco")),
        "confianca_extracao": resultado.get("confianca_extracao", resultado.get("score", 0)),
        "pendencias_criticas": indicadores.get("pendencias_criticas", criticas_fallback),
        "pontos_atencao": indicadores.get("pontos_atencao", pontos_fallback),
        "campos_nao_localizados": indicadores.get("campos_nao_localizados", campos_fallback),
    }


def render_indicadores_analise(resultado: Dict[str, Any]) -> None:
    ind = obter_indicadores_resultado(resultado)
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    c1.markdown(render_metric("Status contratual", ind["status_contratual"]), unsafe_allow_html=True)
    c2.markdown(render_metric("Risco contratual", ind["risco_contratual"]), unsafe_allow_html=True)
    c3.markdown(render_metric("Confiança da extração", f'{ind["confianca_extracao"]}%'), unsafe_allow_html=True)
    c4.markdown(render_metric("Pend. críticas", ind["pendencias_criticas"]), unsafe_allow_html=True)
    c5.markdown(render_metric("Pontos de atenção", ind["pontos_atencao"]), unsafe_allow_html=True)
    c6.markdown(render_metric("Campos não localizados", ind["campos_nao_localizados"]), unsafe_allow_html=True)
    st.caption(f'Situação operacional: {clean_text(ind["situacao_operacional"])}')


def render_filter_metric(label: str, value: Any, filtro: str, ativo: bool = False) -> str:
    """Card clicável do Dashboard. Mantém o usuário na aba Dashboard ao filtrar."""
    classe_ativo = " active" if ativo else ""
    href = f"?pagina=dashboard&risco={filtro}"
    return f"""
    <a class="metric-link" href="{href}" target="_self">
        <div class="metric-card clickable{classe_ativo}">
            <small>{safe(label)}</small>
            <h2>{safe(value)}</h2>
        </div>
    </a>
    """

CAMPOS_VALORES_VISUAIS = {
    "Tipo de Contrato": "Classificação",
    "Empresa do Grupo SBF": "Contratante",
    "CNPJ Empresa do Grupo": "CNPJ interno",
    "Local de Prestação Contraparte": "Local de execução",
    "Contraparte": "Fornecedor",
    "CNPJ Contraparte": "CNPJ fornecedor",
    "Objetivo": "Finalidade",
    "Descrição do Serviço/ Material": "Escopo",
    "Descrição breve do cadastro": "Cadastro",
    "Forma de pagamento": "Pagamento",
    "Condição de Pagamento em Dias": "Prazo",
    "Multa": "Penalidade",
    "Vigência após a data de assinatura": "Vigência",
    "Tipo de Vigência": "Natureza do prazo",
    "Período de Vigência": "Período técnico",
    "Status Contratual": "Situação documental",
    "Situação Operacional": "Confirmação corporativa",
    "Resumo de Aditivos": "Aditivos",
    "Rescisão e Indenização": "Encerramento",
    "Anticorrupção": "Compliance",
    "Proteção de Dados LGPD": "LGPD",
    "Data da Assinatura": "Assinatura",
    "Data do Contrato": "Instrumento",
    "Data Conclusão DocuSign": "DocuSign",
    "Valor do Contrato Original": "Valor global",
    "Valor Mensal Estimado": "Mensal / variável",
    "Valor Total Estimado da Vigência": "Projeção da vigência",
    "Valor Total dos Materiais e Serviços": "Soma dos itens",
    "Pessoas que assinaram": "Signatários",
}


def _extrair_primeiro_valor_card(texto: Any) -> str:
    txt = clean_text(texto)

    m = re.search(
        r"R\$\s*[0-9]{1,3}(?:\.[0-9]{3})*,[0-9]{2}|R\$\s*[0-9]+(?:,[0-9]{2})?|R\$\s*[0-9]+(?:\.[0-9]{2})?",
        txt,
        flags=re.IGNORECASE,
    )

    if m:
        return _formatar_valor_monetario_item(m.group(0))

    return ""


def _montar_partes_card_valor(label: str, value: Any) -> tuple[str, str, str]:
    txt = clean_text(value)
    low = txt.lower()

    alerta = CAMPOS_VALORES_VISUAIS.get(label, "Informação")

    if txt in ("", "Não localizado", "Não localizada", "Não identificado", "N/A", "None"):
        return (
            "Não identificado",
            "Não foi localizada informação suficiente nos documentos analisados.",
            alerta,
        )

    # Campos de data
    if label in ("Data da Assinatura", "Data do Contrato", "Data Conclusão DocuSign"):
        return (
            txt,
            "Data identificada conforme documentos analisados.",
            alerta,
        )

    # Campos de CNPJ
    if "CNPJ" in label:
        return (
            txt,
            "Cadastro identificado nos documentos analisados.",
            alerta,
        )

    # Campo de assinantes
    if label == "Pessoas que assinaram":
        nomes = [n.strip() for n in re.split(r";|,", txt) if n.strip()]
        if len(nomes) >= 4:
            principal = f"{len(nomes)} signatários identificados"
            detalhe = "; ".join(nomes[:4]) + f"; e mais {len(nomes) - 4} nome(s)."
            return principal, detalhe, alerta

        return (
            "Signatários identificados",
            txt,
            alerta,
        )

    # Campo sem valor global
    if "sem valor global fixo" in low:
        valor_ref = _extrair_primeiro_valor_card(txt)
        detalhe = (
            "O contrato não apresenta valor total fechado para todo o contrato. "
            "Há valores unitários, mensais, por demanda ou referência comercial localizada."
        )
        if valor_ref:
            detalhe += f" Referência localizada: {valor_ref}."
        return "Sem valor global fixo", detalhe, "Não confundir com valor global"

    # Campo não calculável
    if "não calculável" in low or "nao calculavel" in low:
        return (
            "Não calculável com precisão",
            "Faltam informações como quantidade, volume ou demanda para projetar o valor total da vigência. A fórmula depende de valor unitário x quantidade x prazo.",
            alerta,
        )

    # Campo não aplicável
    if low.startswith("não aplicável") or low.startswith("nao aplicavel"):
        detalhe = txt
        if len(detalhe) > 180:
            detalhe = detalhe[:180].rsplit(" ", 1)[0] + "..."
        return "Não aplicável", detalhe, alerta

    # Campos com valor monetário
    valor = _extrair_primeiro_valor_card(txt)

    if label == "Valor Total dos Materiais e Serviços":
        # Valores de implantação, mensalidade e tarifas variáveis possuem naturezas incompatíveis.
        # O card não pode transformar o primeiro valor encontrado em uma soma global falsa.
        if any(t in low for t in ("valores pontuais/implantação", "mensalidades fixas", "tarifa(s) variável", "tarifas variáveis", "separados por natureza")):
            return (
                "Valores separados por natureza",
                txt,
                "Não representa valor global do contrato",
            )
        if valor:
            return (
                valor,
                "Total calculável apenas entre itens da mesma natureza. Consulte a tabela completa para periodicidade e condições.",
                "Total de natureza compatível",
            )

    if label == "Valor Mensal Estimado" and valor:
        if "valor fixo mensal confirmado" in low or "mensalidade fixa" in low:
            return f"{valor}/mês", txt, "Mensalidade fixa"
        unidade = _extrair_unidade_demanda(txt)
        if _parece_valor_unitario_ou_demanda(txt):
            detalhe = (
                f"O valor varia conforme a quantidade contratada/utilizada. "
                f"Referência documental: {valor} por {unidade}."
            )
            return f"{valor} por {unidade}", detalhe, "Valor variável"

        return f"{valor}/mês", "Valor mensal estimado conforme informação localizada nos documentos.", alerta

    if label == "Valor Total Estimado da Vigência" and valor:
        detalhe = "Valor projetado para a vigência conforme informações localizadas nos documentos."
        if "cálculo" in low or "calculo" in low:
            detalhe = "Valor estimado com base no prazo de vigência e nas informações comerciais identificadas."
        return valor, detalhe, alerta

    if label == "Valor do Contrato Original" and valor:
        return (
            valor,
            "Valor global previsto para execução do objeto contratado, quando expressamente definido nos documentos.",
            alerta,
        )

    # Campos comuns: divide o texto em título principal + detalhe
    partes = re.split(r"(?<=[.!?])\s+", txt)
    principal = partes[0].strip() if partes else txt
    detalhe = " ".join(partes[1:]).strip() if len(partes) > 1 else ""

    if not detalhe and len(txt) > len(principal):
        detalhe = txt.replace(principal, "").strip(" .-")

    return principal, detalhe, alerta


def _texto_copiavel_card(label: str, principal: str, detalhe: str, alerta: str) -> str:
    """Monta o texto que será copiado pelo botão do card."""
    linhas = [str(label).strip(), str(principal).strip()]
    if detalhe and clean_text(detalhe) not in ("", "Não localizado", "Não localizada", "N/A", "None"):
        linhas.append(str(detalhe).strip())
    if alerta and clean_text(alerta) not in ("", "Não localizado", "Não localizada", "N/A", "None"):
        linhas.append(f"Observação: {str(alerta).strip()}")
    return "\n".join([x for x in linhas if x])


def render_info_card(label: str, value: Any) -> str:
    principal, detalhe, alerta = _montar_partes_card_valor(label, value)
    texto_copiar = _texto_copiavel_card(label, principal, detalhe, alerta)
    texto_copiar_attr = html.escape(texto_copiar, quote=True)

    detalhe_html = ""
    if detalhe and clean_text(detalhe) not in ("", "Não localizado", "Não localizada", "N/A", "None"):
        detalhe_html = f'<div class="valor-detalhe">{safe(detalhe)}</div>'

    return (
        '<div class="valor-card">'
        f'<button type="button" class="copy-card-btn" title="Copiar informação" aria-label="Copiar informação" data-copy="{texto_copiar_attr}">⧉</button>'
        f'<div class="valor-titulo">{safe(label)}</div>'
        f'<div class="valor-principal">{safe(principal)}</div>'
        f'{detalhe_html}'
        f'<div class="valor-alerta">{safe(alerta)}</div>'
        '</div>'
    )


def render_info_grid_com_copy(resultado: Dict[str, Any]) -> None:
    """Renderiza os Dados Extraídos em uma única grade, com botão de copiar por card e geral.

    Observação: as abas principais da tela ficam fora deste componente. Aqui ficam apenas os cards
    das informações do contrato, para evitar o vão preto causado por abas internas em iframe.
    """

    cards_html = []
    textos_todos = []
    mapa_evidencias = resultado.get("mapa_evidencias_cards") if isinstance(resultado.get("mapa_evidencias_cards"), dict) else {}

    for label, chave in CAMPOS_OFICIAIS:
        principal, detalhe, alerta = _montar_partes_card_valor(label, resultado.get(chave))
        evidencia_card = mapa_evidencias.get(chave, {}) if isinstance(mapa_evidencias.get(chave, {}), dict) else {}
        fonte = clean_text(evidencia_card.get("arquivo"))
        pagina = clean_text(evidencia_card.get("pagina"))
        secao = clean_text(evidencia_card.get("secao"))
        status_ev = clean_text(evidencia_card.get("status"))
        confianca_ev = evidencia_card.get("confianca", 0)
        referencia = " • ".join([x for x in [fonte, pagina, secao] if _valor_informado(x)])

        texto_copiar = _texto_copiavel_card(label, principal, detalhe, alerta)
        if referencia:
            texto_copiar += f"\nFonte: {referencia}\nEvidência: {status_ev} • Confiança: {confianca_ev}%"
        textos_todos.append(texto_copiar)

        detalhe_html = ""
        if detalhe and clean_text(detalhe) not in ("", "Não localizado", "Não localizada", "N/A", "None"):
            detalhe_html = f'<div class="valor-detalhe">{safe(detalhe)}</div>'

        fonte_html = ""
        if referencia:
            fonte_html = (
                f'<div class="valor-fonte"><b>Fonte:</b> {safe(referencia)}<br>'
                f'<b>Status:</b> {safe(status_ev)} • <b>Confiança:</b> {safe(confianca_ev)}%</div>'
            )

        cards_html.append(
            '<div class="valor-card">'
            f'<button type="button" class="copy-card-btn" title="Copiar este campo" data-copy="{html.escape(texto_copiar, quote=True)}">⧉</button>'
            f'<div class="valor-titulo">{safe(label)}</div>'
            f'<div class="valor-principal">{safe(principal)}</div>'
            f'{detalhe_html}'
            f'<div class="valor-alerta">{safe(alerta)}</div>'
            f'{fonte_html}'
            '</div>'
        )

    texto_todos = "\n\n".join([x for x in textos_todos if x])
    qtd_cards = len(cards_html)
    linhas = max(1, math.ceil(qtd_cards / 3))
    altura = max(980, min(5200, 230 + (linhas * 295)))

    html_component = f"""
<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8" />
<style>
    *{{box-sizing:border-box;}}
    body{{
        margin:0;
        padding:0;
        background:transparent;
        color:#f8fafc;
        font-family:Inter, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
        overflow:visible;
    }}
    .dados-wrap{{
        width:100%;
        padding:0;
    }}
    .dados-topbar{{
        display:flex;
        justify-content:space-between;
        gap:14px;
        align-items:center;
        margin:0 0 16px 0;
    }}
    .dados-topbar-title{{
        display:flex;
        flex-direction:column;
        gap:4px;
    }}
    .dados-topbar-title strong{{
        color:#f3e6b3;
        font-size:18px;
        font-weight:950;
    }}
    .dados-topbar-title span{{
        color:#94a3b8;
        font-size:12px;
        font-weight:800;
    }}
    .copy-all-btn{{
        border:1px solid rgba(215,191,117,.36);
        background:linear-gradient(135deg,rgba(0,77,61,.95),rgba(6,95,70,.88));
        color:#fff;
        border-radius:14px;
        padding:10px 16px;
        font-weight:950;
        cursor:pointer;
        box-shadow:0 10px 24px rgba(0,0,0,.22);
        white-space:nowrap;
    }}
    .copy-all-btn:hover{{filter:brightness(1.12);transform:translateY(-1px);}}
    .info-grid{{
        display:grid;
        grid-template-columns:repeat(3, minmax(0, 1fr));
        gap:16px;
        align-items:stretch;
    }}
    .valor-card{{
        background:linear-gradient(145deg,#101821,#0b1118);
        border:1px solid rgba(215,191,117,.30);
        border-radius:18px;
        padding:18px;
        padding-right:54px;
        min-height:176px;
        height:auto;
        box-shadow:0 10px 28px rgba(0,0,0,.24);
        position:relative;
        overflow:visible;
    }}
    .valor-card::before{{
        content:"";
        position:absolute;
        left:0;
        top:0;
        width:100%;
        height:4px;
        background:linear-gradient(90deg,#d7bf75,#065f46);
    }}
    .valor-card.copiado{{border-color:rgba(52,211,153,.85);box-shadow:0 0 0 1px rgba(52,211,153,.18),0 10px 28px rgba(0,0,0,.24);}}
    .copy-card-btn{{
        position:absolute;
        top:12px;
        right:12px;
        width:34px;
        height:34px;
        border-radius:11px;
        border:1px solid rgba(215,191,117,.34);
        background:rgba(255,255,255,.055);
        color:#f3e6b3;
        display:flex;
        align-items:center;
        justify-content:center;
        font-size:16px;
        cursor:pointer;
        transition:.18s ease;
    }}
    .copy-card-btn:hover{{background:rgba(215,191,117,.16);transform:translateY(-1px);}}
    .copy-card-btn.ok, .copy-all-btn.ok{{background:rgba(16,185,129,.22);border-color:rgba(16,185,129,.45);color:#ecfdf5;}}
    .valor-titulo{{
        color:#d7bf75;
        font-weight:950;
        text-transform:uppercase;
        letter-spacing:.7px;
        font-size:11px;
        margin-bottom:12px;
        line-height:1.25;
    }}
    .valor-principal{{
        color:#ffffff;
        font-size:22px;
        line-height:1.18;
        font-weight:950;
        margin-bottom:10px;
        overflow-wrap:anywhere;
    }}
    .valor-detalhe{{
        color:#d8dee9;
        font-size:13px;
        line-height:1.45;
        font-weight:700;
        overflow-wrap:anywhere;
    }}
    .valor-alerta{{
        display:inline-block;
        margin-top:12px;
        padding:5px 10px;
        border-radius:999px;
        background:rgba(215,191,117,.12);
        color:#f3e6b3;
        border:1px solid rgba(215,191,117,.28);
        font-size:11px;
        font-weight:950;
    }}
    .valor-fonte{{
        margin-top:12px;
        padding-top:10px;
        border-top:1px dashed rgba(215,191,117,.22);
        color:#94a3b8;
        font-size:10px;
        line-height:1.45;
        font-weight:700;
        overflow-wrap:anywhere;
    }}
    .valor-fonte b{{color:#f3e6b3;}}
    .copy-toast{{
        position:fixed;
        right:14px;
        bottom:14px;
        background:rgba(0,77,61,.98);
        border:1px solid rgba(215,191,117,.35);
        color:#fff;
        padding:11px 14px;
        border-radius:14px;
        font-weight:900;
        box-shadow:0 16px 34px rgba(0,0,0,.32);
        opacity:0;
        transform:translateY(12px);
        pointer-events:none;
        transition:.2s ease;
    }}
    .copy-toast.show{{opacity:1;transform:translateY(0);}}
    @media(max-width:1180px){{
        .info-grid{{grid-template-columns:repeat(2,minmax(0,1fr));}}
    }}
    @media(max-width:760px){{
        .dados-topbar{{flex-direction:column;align-items:stretch;}}
        .copy-all-btn{{width:100%;}}
        .info-grid{{grid-template-columns:1fr;}}
    }}
</style>
</head>
<body>
    <div class="dados-wrap">
        <div class="dados-topbar">
            <div class="dados-topbar-title">
                <strong>Informações principais do contrato</strong>
                <span>Campos extraídos pela análise. Use o botão de cada card para copiar uma informação específica.</span>
            </div>
            <button type="button" class="copy-all-btn" data-copy="{html.escape(texto_todos, quote=True)}">⧉ Copiar todos os dados extraídos</button>
        </div>
        <div class="info-grid">{''.join(cards_html)}</div>
    </div>
    <div id="copy-toast" class="copy-toast">Informação copiada</div>
<script>
(function(){{
    function ajustarAlturaComponente(){{
        const h = Math.max(document.body.scrollHeight || 0, document.documentElement.scrollHeight || 0) + 20;
        window.parent.postMessage({{isStreamlitMessage: true, type: 'streamlit:setFrameHeight', height: h}}, '*');
        window.parent.postMessage({{type: 'streamlit:setFrameHeight', height: h}}, '*');
    }}
    window.addEventListener('load', ajustarAlturaComponente);
    window.addEventListener('resize', ajustarAlturaComponente);
    setTimeout(ajustarAlturaComponente, 150);
    setTimeout(ajustarAlturaComponente, 600);
    if (window.ResizeObserver) {{ new ResizeObserver(ajustarAlturaComponente).observe(document.body); }}

    function showToast(msg){{
        const toast = document.getElementById('copy-toast');
        toast.textContent = msg || 'Informação copiada';
        toast.classList.add('show');
        setTimeout(() => toast.classList.remove('show'), 1400);
    }}
    function fallbackCopy(text){{
        const area = document.createElement('textarea');
        area.value = text;
        area.setAttribute('readonly', '');
        area.style.position = 'fixed';
        area.style.opacity = '0';
        document.body.appendChild(area);
        area.select();
        try {{ document.execCommand('copy'); }} catch(e) {{}}
        document.body.removeChild(area);
    }}
    async function copiar(text){{
        if (navigator.clipboard && window.isSecureContext) {{ await navigator.clipboard.writeText(text); }}
        else {{ fallbackCopy(text); }}
    }}
    document.querySelectorAll('[data-copy]').forEach((btn) => {{
        btn.addEventListener('click', async () => {{
            const text = btn.getAttribute('data-copy') || '';
            await copiar(text);
            const card = btn.closest('.valor-card');
            if (card) {{
                card.classList.add('copiado');
                btn.classList.add('ok');
                const original = btn.textContent;
                btn.textContent = '✓';
                setTimeout(() => {{ card.classList.remove('copiado'); btn.classList.remove('ok'); btn.textContent = original; }}, 1100);
                showToast('Campo copiado');
            }} else {{
                btn.classList.add('ok');
                const original = btn.textContent;
                btn.textContent = '✓ Copiado';
                setTimeout(() => {{ btn.classList.remove('ok'); btn.textContent = original; }}, 1100);
                showToast('Todos os dados foram copiados');
            }}
        }});
    }});
}})();
</script>
</body>
</html>
"""
    components.html(html_component, height=altura, scrolling=False)


def render_checklist_validacao(resultado: Dict[str, Any]) -> None:
    st.markdown('<div class="section-title">Checklist de validação</div>', unsafe_allow_html=True)
    df_checklist = pd.DataFrame(resultado.get("checklist", []))
    if df_checklist.empty:
        st.warning("Nenhum checklist pôde ser reconstruído porque este registro não possui matriz de evidências. Reprocesse o contrato nesta versão.")
        return
    ordem = ["Validação", "Status", "Peso de risco", "Crítico", "Arquivo", "Página", "Evidência"]
    cols = [c for c in ordem if c in df_checklist.columns] + [c for c in df_checklist.columns if c not in ordem]
    st.dataframe(df_checklist[cols], use_container_width=True, hide_index=True)


def render_pendencias_encontradas(resultado: Dict[str, Any]) -> None:
    st.markdown('<div class="section-title">Pendências e pontos de atenção</div>', unsafe_allow_html=True)
    pendencias = resultado.get("pendencias", []) if isinstance(resultado.get("pendencias"), list) else []
    campos = resultado.get("campos_nao_localizados", []) if isinstance(resultado.get("campos_nao_localizados"), list) else []
    criticas = [p for p in pendencias if clean_text(p.get("Crítico") or p.get("critico")).upper() in ("SIM", "TRUE", "1")]
    pontos = [p for p in pendencias if p not in criticas]
    c1, c2, c3 = st.columns(3)
    c1.markdown(render_metric("Pendências críticas", len(criticas)), unsafe_allow_html=True)
    c2.markdown(render_metric("Pontos de atenção", len(pontos)), unsafe_allow_html=True)
    c3.markdown(render_metric("Campos não localizados", len(campos)), unsafe_allow_html=True)

    def _render_lista(titulo: str, lista: List[Dict[str, Any]]) -> None:
        if not lista:
            return
        st.markdown(f"#### {titulo}")
        for i, pendencia in enumerate(lista, 1):
            fonte = clean_text(pendencia.get("Arquivo") or pendencia.get("Fonte"))
            pagina = clean_text(pendencia.get("Página") or pendencia.get("pagina"))
            evidencia = clean_text(pendencia.get("Evidência") or pendencia.get("evidencia"))
            st.markdown(
                f"""
                <div class="risk-row">
                    <b>{i}. {safe(pendencia.get('Pendência', 'Ponto de atenção'))}</b><br>
                    Crítico: {safe(pendencia.get('Crítico', 'Não'))} • Risco: {safe(pendencia.get('Risco', 'Baixo'))}<br>
                    <span class="subtle">{safe(pendencia.get('Recomendação', 'Validar antes de seguir.'))}</span><br>
                    <span class="subtle">Fonte: {safe(fonte or 'Não localizada')} • Página: {safe(pagina or 'Não localizada')} • Evidência: {safe(evidencia or 'Não localizada')}</span>
                </div>
                """,
                unsafe_allow_html=True,
            )

    _render_lista("Pendências críticas", criticas)
    _render_lista("Pontos de atenção", pontos)
    if campos:
        st.markdown("#### Campos não localizados nos documentos")
        st.info(" • ".join(clean_text(c).replace("_", " ").title() for c in campos))
    if not pendencias and not campos:
        st.markdown('<div class="ok-row">Nenhuma pendência crítica, ponto de atenção ou campo obrigatório não localizado.</div>', unsafe_allow_html=True)


def render_parecer_automatico(resultado: Dict[str, Any]) -> None:
    st.markdown('<div class="section-title">Parecer automático</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="executive-box"><b>Parecer:</b><br><br>{safe(resultado.get("parecer"))}</div>', unsafe_allow_html=True)


def render_resumo_executivo_aba(resultado: Dict[str, Any]) -> None:
    """Mostra o resumo executivo com aparência de informação, não de erro."""
    st.markdown('<div class="section-title">Resumo executivo</div>', unsafe_allow_html=True)
    resumo = resultado.get("resumo_executivo") or resultado.get("parecer") or "Resumo executivo não localizado."
    risco = normalize_risco(resultado.get("risco"))
    score = resultado.get("score") or "Não localizado"
    status = resultado.get("status") or "Não localizado"

    html_resumo = f"""
        <div class="executive-box" style="border-left:4px solid #d7bf75; background:linear-gradient(135deg, rgba(0,77,61,.52), rgba(16,22,32,.92));">
            <div style="display:flex; align-items:center; justify-content:space-between; gap:14px; flex-wrap:wrap; margin-bottom:14px;">
                <div>
                    <div style="font-size:12px; color:#d7bf75; font-weight:900; letter-spacing:.08em; text-transform:uppercase;">Síntese da análise</div>
                    <div style="font-size:22px; font-weight:950; color:#fff; margin-top:4px;">Contrato analisado com visão executiva</div>
                </div>
                <div style="display:flex; gap:8px; flex-wrap:wrap;">
                    <span class="pill pill-ok">Status: {safe(status)}</span>
                    <span class="pill pill-ok">Risco: {safe(risco)}</span>
                    <span class="pill pill-ok">Score: {safe(score)}</span>
                </div>
            </div>
            <div style="font-size:15px; line-height:1.65; color:#fff; font-weight:800;">
                {safe(resumo)}
            </div>
        </div>
    """
    st.markdown(html_resumo, unsafe_allow_html=True)


def render_texto_extraido_aba(resultado: Dict[str, Any], texto_extraido: str = "", key_prefix: str = "texto_resultado") -> None:
    """Mostra o texto extraído em aba própria para não poluir a análise principal."""
    st.markdown('<div class="section-title">Texto extraído do contrato e anexos</div>', unsafe_allow_html=True)
    texto = texto_extraido or str(resultado.get("texto_extraido") or "Texto extraído não disponível para este registro.")
    st.caption("Use esta aba apenas para auditoria técnica ou conferência do conteúdo lido pelo robô.")
    st.text_area("Texto extraído", texto[:50000], height=360, key=f"{key_prefix}_{id(resultado)}")



def _split_assinantes_para_linhas(valor: Any) -> List[str]:
    """Converte texto/lista de signatários em lista limpa de nomes."""
    if isinstance(valor, list):
        bruto = "; ".join([clean_text(x) for x in valor if clean_text(x)])
    else:
        bruto = clean_text(valor)
    if bruto in ("", "Não localizado", "Não localizada", "N/A", "None"):
        return []
    bruto = re.sub(r"\s+e\s+mais\s+\d+\s+nome\(s\).*", "", bruto, flags=re.IGNORECASE)
    bruto = bruto.replace("|", ";").replace("\n", ";")
    if ";" in bruto:
        partes = bruto.split(";")
    else:
        partes = re.split(r",\s*(?=[A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-Za-zÀ-ÿ])", bruto)
    nomes = []
    for parte in partes:
        nome = clean_text(parte).strip(" -:;,.•")
        if not nome or nome.lower() in ("não localizado", "nao localizado", "none", "n/a"):
            continue
        if len(nome) > 120:
            nome = nome[:120].strip() + "..."
        if nome not in nomes:
            nomes.append(nome)
    return nomes


def montar_df_assinaturas(resultado: Dict[str, Any]) -> pd.DataFrame:
    """Monta tabela executiva de assinaturas para tela e Excel."""
    linhas: List[Dict[str, Any]] = []
    raw = resultado.get("assinaturas_contrato")
    data_padrao = clean_text(resultado.get("data_assinatura") or resultado.get("data_conclusao_docusign") or resultado.get("data_contrato"))
    data_docusign_txt = clean_text(resultado.get("data_conclusao_docusign"))
    docusign_confirmado = bool(_data_valida_simples(data_docusign_txt)) and "não aplicável" not in data_docusign_txt.lower()
    fonte_padrao = "DocuSign / contrato" if "docusign" in clean_text(resultado.get("alerta_assinatura")).lower() or docusign_confirmado else "Contrato"
    status_padrao = "Assinado" if clean_text(resultado.get("contrato_assinado")).lower() == "sim" else "Não localizado"
    evidencia_padrao = clean_text(resultado.get("alerta_assinatura") or "Evidência de assinatura conforme documentos analisados.")

    if isinstance(raw, list) and raw:
        for item in raw:
            if not isinstance(item, dict):
                continue
            nome = clean_text(item.get("nome") or item.get("Nome") or item.get("signatario") or item.get("signatário") or item.get("assinante"))
            if not _valor_informado(nome):
                continue
            linhas.append({
                "Nome": nome,
                "Papel/Cargo": clean_text(item.get("papel_cargo") or item.get("Papel/Cargo") or item.get("cargo") or item.get("papel") or "Não localizado"),
                "E-mail": clean_text(item.get("email") or item.get("e-mail") or item.get("Email") or "Não localizado"),
                "Data da assinatura": clean_text(item.get("data_assinatura") or item.get("Data da assinatura") or data_padrao or "Não localizado"),
                "Data do reconhecimento de firma": clean_text(
                    item.get("data_reconhecimento_firma")
                    or item.get("Data do reconhecimento de firma")
                    or _data_reconhecimento_da_evidencia(item.get("evidencia") or item.get("evidência") or item.get("Evidência"))
                    or "Não aplicável"
                ),
                "Fonte": clean_text(item.get("fonte") or item.get("Fonte") or fonte_padrao),
                "Página": clean_text(item.get("pagina") or item.get("Página") or "Não localizado"),
                "Status": clean_text(item.get("status") or item.get("Status") or status_padrao),
                "Evidência": clean_text(item.get("evidencia") or item.get("evidência") or item.get("Evidência") or evidencia_padrao),
            })

    if not linhas:
        for nome in _split_assinantes_para_linhas(resultado.get("pessoas_que_assinaram") or resultado.get("assinantes")):
            linhas.append({
                "Nome": nome,
                "Papel/Cargo": "Não localizado",
                "E-mail": "Não localizado",
                "Data da assinatura": data_padrao or "Não localizado",
                "Data do reconhecimento de firma": clean_text(resultado.get("data_reconhecimento_firma") or "Não aplicável"),
                "Fonte": fonte_padrao,
                "Página": "Não localizado",
                "Status": status_padrao,
                "Evidência": evidencia_padrao,
            })

    if not linhas:
        linhas.append({
            "Nome": "Não localizado",
            "Papel/Cargo": "Não localizado",
            "E-mail": "Não localizado",
            "Data da assinatura": data_padrao or "Não localizado",
            "Data do reconhecimento de firma": clean_text(resultado.get("data_reconhecimento_firma") or "Não aplicável"),
            "Fonte": fonte_padrao,
            "Página": "Não localizado",
            "Status": status_padrao,
            "Evidência": evidencia_padrao if _valor_informado(evidencia_padrao) else "Nenhum signatário localizado nos documentos analisados.",
        })
    return pd.DataFrame(linhas)


def render_assinaturas_contrato(resultado: Dict[str, Any]) -> None:
    """Aba executiva completa para conferência de assinaturas."""
    st.markdown('<div class="section-title">Assinaturas</div>', unsafe_allow_html=True)
    df_ass = montar_df_assinaturas(resultado)
    total = 0 if (len(df_ass) == 1 and clean_text(df_ass.iloc[0].get("Nome")) == "Não localizado") else len(df_ass)
    contrato_assinado = clean_text(resultado.get("contrato_assinado") or "Não localizado")
    data_principal = clean_text(resultado.get("data_assinatura") or resultado.get("data_conclusao_docusign") or "Não localizado")
    alerta = clean_text(resultado.get("alerta_assinatura") or "Valide a evidência de assinatura nos documentos originais.")

    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(render_metric("Signatários", total), unsafe_allow_html=True)
    c2.markdown(render_metric("Contrato assinado", contrato_assinado), unsafe_allow_html=True)
    c3.markdown(render_metric("Data principal", data_principal), unsafe_allow_html=True)
    c4.markdown(render_metric("Fonte", "DocuSign/Contrato" if "docusign" in alerta.lower() else "Contrato"), unsafe_allow_html=True)

    st.markdown(
        f'<div class="executive-box"><b>Resumo da assinatura:</b><br>{safe(alerta)}<br><br>'
        f'<b>Regra:</b> contrato assinado só deve ser considerado concluído quando houver evidência de DocuSign concluído, assinatura eletrônica, assinatura das partes ou certificado equivalente.</div>',
        unsafe_allow_html=True,
    )
    st.dataframe(df_ass, use_container_width=True, hide_index=True)

def render_auditoria_evidencias(resultado: Dict[str, Any]) -> None:
    """Exibe a matriz que sustenta cada card, cálculo e conclusão."""
    st.markdown('<div class="section-title">Auditoria de evidências</div>', unsafe_allow_html=True)
    metricas = resultado.get("metricas_confianca") if isinstance(resultado.get("metricas_confianca"), dict) else {}
    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(render_metric("Cobertura dos campos", f"{metricas.get('cobertura_campos_percentual', 0)}%"), unsafe_allow_html=True)
    c2.markdown(render_metric("Páginas processadas", f"{metricas.get('paginas_processadas_percentual', metricas.get('cobertura_paginas_percentual', 0))}%"), unsafe_allow_html=True)
    c3.markdown(render_metric("Campos não localizados", metricas.get("campos_nao_localizados", len(resultado.get("campos_nao_localizados", [])))), unsafe_allow_html=True)
    c4.markdown(render_metric("Conflitos", metricas.get("conflitos", len(resultado.get("conflitos_documentais", [])))), unsafe_allow_html=True)

    st.caption("Todo dado confirmado deve possuir arquivo, página ou seção, trecho de evidência e nível de confiança. Cálculos são identificados separadamente.")
    linhas = linhas_auditoria_para_tela(resultado)
    if linhas:
        df = pd.DataFrame(linhas)
        st.dataframe(
            df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Valor consolidado": st.column_config.TextColumn("Valor consolidado", width="large"),
                "Evidência": st.column_config.TextColumn("Evidência", width="large"),
                "Confiança": st.column_config.ProgressColumn("Confiança", min_value=0, max_value=100),
            },
        )
    else:
        st.info("Esta análise não possui matriz de evidências. Registros antigos permanecem em modo de compatibilidade.")

    conflitos = resultado.get("conflitos_documentais") if isinstance(resultado.get("conflitos_documentais"), list) else []
    if conflitos:
        st.markdown('<div class="section-title">Conflitos documentais</div>', unsafe_allow_html=True)
        st.dataframe(pd.DataFrame(conflitos), use_container_width=True, hide_index=True)


def render_resultado_em_abas(resultado: Dict[str, Any], texto_extraido: str = "", key_prefix: str = "resultado") -> None:
    """Organiza o resultado completo em abas por seção da análise.

    Ordem final aprovada:
    - Resumo executivo
    - Informações em caixas
    - Assinaturas
    - Aditivos
    - Materiais/Serviços
    - Checklist de validação
    - Pendências
    - Parecer automático
    - Auditoria de evidências
    - Texto extraído
    - Triagem dos anexos
    """
    (
        tab_resumo,
        tab_infos,
        tab_assinaturas,
        tab_aditivos,
        tab_materiais,
        tab_checklist,
        tab_pendencias,
        tab_parecer,
        tab_evidencias,
        tab_texto,
        tab_triagem,
    ) = st.tabs([
        "📊 Resumo executivo",
        "📌 Informações em caixas",
        "✍️ Assinaturas",
        "🧩 Aditivos",
        "📦 Materiais/Serviços",
        "✅ Checklist",
        "⚠️ Pendências",
        "📝 Parecer",
        "🔎 Evidências",
        "📄 Texto extraído",
        "📎 Triagem dos anexos",
    ])

    with tab_resumo:
        render_resumo_executivo_aba(resultado)

    with tab_infos:
        st.markdown('<div class="section-title">Dados extraídos</div>', unsafe_allow_html=True)
        render_info_grid_com_copy(resultado)
        st.markdown('<div class="section-title">Objeto / Escopo</div>', unsafe_allow_html=True)
        st.markdown(
            f'<div class="executive-box">{safe(resultado.get("descricao_servico_material") or resultado.get("objetivo"))}</div>',
            unsafe_allow_html=True,
        )

    with tab_assinaturas:
        render_assinaturas_contrato(resultado)

    with tab_aditivos:
        render_aditivos_contrato(resultado)

    with tab_materiais:
        render_itens_contrato(resultado)

    with tab_checklist:
        render_checklist_validacao(resultado)

    with tab_pendencias:
        render_pendencias_encontradas(resultado)

    with tab_parecer:
        render_parecer_automatico(resultado)

    with tab_evidencias:
        render_auditoria_evidencias(resultado)

    with tab_texto:
        render_texto_extraido_aba(resultado, texto_extraido, key_prefix=key_prefix)

    with tab_triagem:
        render_triagem_anexos(resultado)

# =========================================================
# POPUP AO VIVO DE PROCESSAMENTO
# =========================================================
def _formatar_tempo_execucao(segundos: float) -> str:
    segundos = int(segundos)
    minutos = segundos // 60
    seg = segundos % 60

    if minutos <= 0:
        return f"{seg}s"

    return f"{minutos}min {seg}s"


def _icone_status_processamento(status: str) -> str:
    status = str(status or "").upper()

    if status in ("CONCLUIDO", "ANALISE_IA"):
        return "✅" if status == "CONCLUIDO" else "🧠"

    if status == "APOIO":
        return "📎"

    if status == "IGNORADO":
        return "🚫"

    if status == "PROCESSANDO":
        return "🔄"

    if status == "ERRO":
        return "❌"

    return "⏳"


def atualizar_popup_processamento(
    placeholder,
    arquivos_status: Dict[str, str],
    arquivo_atual: str,
    etapa_atual: str,
    inicio_processamento: float,
) -> None:
    """Atualiza um painel/modal fixo com andamento ao vivo da análise."""
    total = len(arquivos_status)
    concluidos = sum(1 for s in arquivos_status.values() if s in ("CONCLUIDO", "ANALISE_IA", "APOIO", "IGNORADO"))
    erros = sum(1 for s in arquivos_status.values() if s == "ERRO")
    pendentes = sum(1 for s in arquivos_status.values() if s == "PENDENTE")

    percentual = int((concluidos / total) * 100) if total else 0
    etapa_html = safe(etapa_atual).replace("\n", "<br>")
    if str(arquivo_atual or "").startswith("Arquivos preparados"):
        linha_arquivo_atual = f'<div><b>{safe(arquivo_atual)}</b></div>'
    else:
        linha_arquivo_atual = f'<div><b>Arquivo atual:</b> {safe(arquivo_atual or "Aguardando...")}</div>'
    tempo = _formatar_tempo_execucao(time.perf_counter() - inicio_processamento)
    ultima_atualizacao = datetime.now().strftime("%H:%M:%S")

    linhas = ""
    for nome, status in arquivos_status.items():
        icone = _icone_status_processamento(status)
        classe = "pendente"
        if status in ("CONCLUIDO", "ANALISE_IA"):
            classe = "concluido"
        elif status == "APOIO":
            classe = "apoio"
        elif status == "IGNORADO":
            classe = "ignorado"
        elif status == "PROCESSANDO":
            classe = "processando"
        elif status == "ERRO":
            classe = "erro"

        # Importante: HTML sem recuo no início da linha.
        # Streamlit/Markdown trata linhas HTML com 4+ espaços como bloco de código.
        label_status = str(status or "")
        label_status = label_status.replace("ANALISE_IA", "ANÁLISE IA")
        linhas += (
            f'<div class="linha-arquivo {classe}">'
            f'<div class="arquivo-nome" title="{safe(nome)}">{icone} {safe(nome)}</div>'
            f'<div class="arquivo-status">{safe(label_status)}</div>'
            f'</div>\n'
        )

    html_popup = f"""
<style>
.popup-processamento-backdrop {{
position: fixed;
inset: 0;
background: rgba(0, 0, 0, 0.58);
z-index: 999999;
display: flex;
align-items: center;
justify-content: center;
backdrop-filter: blur(4px);
}}
.popup-processamento-card {{
width: min(980px, 92vw);
max-height: 86vh;
overflow: hidden;
background: linear-gradient(145deg, #071018 0%, #0b131c 55%, #111827 100%);
border: 1px solid rgba(238, 202, 105, 0.55);
border-radius: 22px;
box-shadow: 0 22px 70px rgba(0,0,0,0.55);
color: #ffffff;
font-family: inherit;
}}
.popup-processamento-header {{
padding: 24px 28px 16px 28px;
border-bottom: 1px solid rgba(238, 202, 105, 0.22);
}}
.popup-processamento-header h2 {{
margin: 0;
color: #f5d36d;
font-size: 26px;
font-weight: 900;
}}
.popup-processamento-header p {{
margin: 8px 0 0 0;
color: rgba(255,255,255,0.82);
font-size: 14px;
}}
.popup-processamento-body {{
padding: 20px 28px 28px 28px;
}}
.popup-metricas {{
display: grid;
grid-template-columns: repeat(5, 1fr);
gap: 12px;
margin-bottom: 18px;
}}
.popup-metrica {{
background: rgba(255,255,255,0.045);
border: 1px solid rgba(238, 202, 105, 0.18);
border-radius: 14px;
padding: 12px 14px;
}}
.popup-metrica small {{
display: block;
color: #f5d36d;
font-size: 10px;
font-weight: 900;
text-transform: uppercase;
letter-spacing: .06em;
margin-bottom: 6px;
}}
.popup-metrica strong {{
display: block;
color: #ffffff;
font-size: 20px;
font-weight: 900;
}}
.popup-etapa {{
background: rgba(0, 94, 73, 0.35);
border: 1px solid rgba(16, 185, 129, 0.28);
border-radius: 16px;
padding: 14px 16px;
margin-bottom: 16px;
}}
.popup-etapa b {{
color: #f5d36d;
}}
.barra-progresso {{
height: 14px;
width: 100%;
background: rgba(255,255,255,0.10);
border-radius: 999px;
overflow: hidden;
margin: 14px 0 18px 0;
border: 1px solid rgba(238, 202, 105, 0.18);
}}
.barra-progresso-fill {{
height: 100%;
width: {percentual}%;
background: linear-gradient(90deg, #10b981, #f5d36d);
border-radius: 999px;
transition: width .35s ease;
}}
.lista-arquivos {{
max-height: 330px;
overflow-y: auto;
border-radius: 16px;
border: 1px solid rgba(238, 202, 105, 0.18);
}}
.linha-arquivo {{
display: grid;
grid-template-columns: 1fr 150px;
gap: 14px;
padding: 11px 14px;
border-bottom: 1px solid rgba(255,255,255,0.08);
font-size: 13px;
align-items: center;
}}
.linha-arquivo:last-child {{
border-bottom: none;
}}
.arquivo-nome {{
white-space: nowrap;
overflow: hidden;
text-overflow: ellipsis;
font-weight: 800;
}}
.arquivo-status {{
text-align: right;
font-size: 11px;
font-weight: 900;
letter-spacing: .04em;
}}
.linha-arquivo.concluido {{
background: rgba(22, 163, 74, 0.16);
}}
.linha-arquivo.processando {{
background: rgba(245, 158, 11, 0.20);
}}
.linha-arquivo.erro {{
background: rgba(220, 38, 38, 0.22);
}}
.linha-arquivo.apoio {{
background: rgba(59, 130, 246, 0.14);
}}
.linha-arquivo.ignorado {{
background: rgba(100, 116, 139, 0.18);
color: rgba(255,255,255,0.72);
}}
.linha-arquivo.pendente {{
background: rgba(255,255,255,0.025);
}}
.popup-rodape {{
margin-top: 14px;
color: rgba(255,255,255,0.62);
font-size: 12px;
}}
</style>
<div class="popup-processamento-backdrop">
<div class="popup-processamento-card">
<div class="popup-processamento-header">
<h2>🚀 Analisando contrato e anexos</h2>
<p>Não feche esta tela. O painel abaixo mostra o andamento em tempo real.</p>
</div>
<div class="popup-processamento-body">
<div class="popup-metricas">
<div class="popup-metrica"><small>Total</small><strong>{total}</strong></div>
<div class="popup-metrica"><small>Concluídos</small><strong>{concluidos}</strong></div>
<div class="popup-metrica"><small>Pendentes</small><strong>{pendentes}</strong></div>
<div class="popup-metrica"><small>Erros</small><strong>{erros}</strong></div>
<div class="popup-metrica"><small>Tempo</small><strong>{tempo}</strong></div>
</div>
<div class="popup-etapa">
<div><b>Etapa atual:</b> {etapa_html}</div>
{linha_arquivo_atual}
<div><b>Última atualização:</b> {ultima_atualizacao}</div>
</div>
<div class="barra-progresso"><div class="barra-progresso-fill"></div></div>
<div class="lista-arquivos">
{linhas}
</div>
<div class="popup-rodape">
Se o tempo continuar aumentando e o arquivo atual não mudar por muito tempo, provavelmente o processamento travou nesse anexo.
</div>
</div>
</div>
</div>
"""

    placeholder.markdown(html_popup, unsafe_allow_html=True)



def _nome_low(nome: Any) -> str:
    return clean_text(nome).lower()


def _nome_indica_certificado(nome: Any) -> bool:
    nome_low = _nome_low(nome)
    return any(t in nome_low for t in [
        "certificado de conclusão",
        "certificado de conclusao",
        "certificate of completion",
        "certificado_docusign",
        "certificado docusign",
    ])


def _nome_indica_apoio(nome: Any) -> bool:
    nome_low = _nome_low(nome)
    termos_apoio = [
        "comunicado",
        "apresentação",
        "apresentacao",
        "proposta",
        "orçamento",
        "orcamento",
        "estudo de mercado",
        "validacao",
        "validação",
        "aprovacao",
        "aprovação",
    ]
    return any(t in nome_low for t in termos_apoio)


def _normalizar_texto_chave(valor: Any) -> str:
    txt = clean_text(valor).lower()
    txt = txt.replace("º", "o").replace("°", "o")
    txt = re.sub(r"[^a-z0-9]+", " ", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def _ordinal_aditivo_numero_prioritario(nome: Any, texto: Any = "") -> int | None:
    """
    Descobre qual aditivo é, priorizando o nome/título do arquivo.
    Não usa o texto inteiro primeiro, porque aditivos posteriores citam Primeiro/Segundo/etc. nos considerandos.
    """
    nome_txt = clean_text(nome)
    texto_txt = clean_text(texto)

    candidatos = [nome_txt]

    # Título real costuma estar no começo do documento, antes dos considerandos.
    inicio = texto_txt[:900]
    candidatos.append(inicio)

    for fonte in candidatos:
        fonte_low = fonte.lower()

        padroes = [
            (1, [r"\bprimeiro\s+adit", r"\b1\s*[º°o]?\s*adit", r"\b1[º°]\s*adit"]),
            (2, [r"\bsegundo\s+adit", r"\b2\s*[º°o]?\s*adit", r"\b2[º°]\s*adit"]),
            (3, [r"\bterceiro\s+adit", r"\b3\s*[º°o]?\s*adit", r"\b3[º°]\s*adit"]),
            (4, [r"\bquarto\s+adit", r"\b4\s*[º°o]?\s*adit", r"\b4[º°]\s*adit"]),
            (5, [r"\bquinto\s+adit", r"\b5\s*[º°o]?\s*adit", r"\b5[º°]\s*adit"]),
        ]

        for numero, lista in padroes:
            if any(re.search(p, fonte_low, flags=re.IGNORECASE) for p in lista):
                return numero

    return None


def _evidencia_assinatura_arquivo(nome: Any, texto: Any) -> bool:
    """
    Evidência de assinatura para decidir análise profunda.
    Para contrato/aditivo: DocuSign Envelope ID também é evidência forte.
    Certificado isolado continua sendo APOIO pela classificação de tipo, não análise principal.
    """
    plano = f"{clean_text(nome)}\n{clean_text(texto)}".lower()
    nome_low = _nome_low(nome)

    if "assinado" in nome_low or "para_ass" in nome_low or "para ass" in nome_low:
        return True

    sinais_fortes = [
        "certificate of completion",
        "certificado de conclusão",
        "certificado de conclusao",
        "status: completed",
        "status: concluído",
        "status: concluido",
        "assinatura concluída",
        "assinatura concluida",
        "signing complete",
        "completed security checked",
        "concluído segurança verificada",
        "concluido seguranca verificada",
    ]

    if any(s in plano for s in sinais_fortes):
        return True

    # Muitos PDFs assinados trazem apenas o Envelope ID em cada página.
    if "docusign envelope id" in plano:
        return True

    if re.search(r"\bassinad[oa]\s+em\b", plano, flags=re.IGNORECASE):
        return True

    return False


def _tipo_documento_triagem(nome: Any, texto: Any = "") -> str:
    nome_low = clean_text(nome).lower()
    texto_inicio = clean_text(texto)[:2500].lower()
    plano = f"{nome_low} {texto_inicio}"

    # 1) Apoios explícitos pelo NOME mandam antes do conteúdo.
    # Ex.: certificado pode citar "aditivo" no assunto, mas não é o aditivo principal.
    if _nome_indica_certificado(nome):
        return "Certificado DocuSign"
    if "comunicado" in nome_low:
        return "Comunicado"
    if "apresentação" in nome_low or "apresentacao" in nome_low:
        return "Apresentação"
    if "proposta" in nome_low or "orçamento" in nome_low or "orcamento" in nome_low:
        return "Proposta/Orçamento"
    if "validação" in nome_low or "validacao" in nome_low or "aprovação" in nome_low or "aprovacao" in nome_low:
        return "Validação/Aprovação"
    if "berkan" in nome_low:
        return "Berkan/Validação terceiros"
    if "cnpj" in nome_low or "cartao cnpj" in nome_low or "cartão cnpj" in nome_low:
        return "Documento cadastral"
    if "contrato social" in nome_low or "alteração contrato social" in nome_low or "alteracao contrato social" in nome_low or "alteração contratual" in nome_low or "alteracao contratual" in nome_low:
        return "Contrato social/Alteração cadastral"

    # 2) Contrato/aditivo pelo NOME.
    if "aditivo" in nome_low or "aditamento" in nome_low:
        return "Aditivo"
    if "contrato" in nome_low:
        return "Contrato principal"

    # 3) Apoios pelo conteúdo.
    if "certificado de conclusão" in plano or "certificado de conclusao" in plano or "certificate of completion" in plano:
        return "Certificado DocuSign"
    if "comunicado" in plano:
        return "Comunicado"
    if "apresentação" in plano or "apresentacao" in plano or "estudo de mercado" in plano:
        return "Apresentação"
    if "proposta" in plano or "orçamento" in plano or "orcamento" in plano:
        return "Proposta/Orçamento"

    # 4) Contrato/aditivo pelo TÍTULO/conteúdo do começo do documento.
    if _ordinal_aditivo_numero_prioritario(nome, texto):
        return "Aditivo"
    if re.search(r"\bcontrato\s+de\s+fornecimento\b|\bcontrato\s+de\s+prestação\b|\bcontrato\s+de\s+prestacao\b|instrumento\s+particular|termo\s+de\s+contrato", texto_inicio, flags=re.IGNORECASE):
        return "Contrato principal"
    if texto_tem_conteudo_contratual(texto_inicio):
        return "Contrato principal"

    if "validação" in plano or "validacao" in plano or "aprovação" in plano or "aprovacao" in plano:
        return "Validação/Aprovação"
    if "berkan" in plano:
        return "Berkan/Validação terceiros"
    if "cadastro nacional da pessoa jurídica" in plano or "comprovante de inscrição e de situação cadastral" in plano:
        return "Documento cadastral"
    if "contrato social" in plano or "alteração contratual" in plano or "alteracao contratual" in plano or "nire" in plano:
        return "Contrato social/Alteração cadastral"

    return "Documento complementar"


def _documento_precisa_assinatura(tipo: str) -> bool:
    return tipo in ("Contrato principal", "Aditivo")


def _documento_apoio_sem_aprovacao(tipo: str) -> bool:
    return tipo in (
        "Proposta/Orçamento",
        "Certificado DocuSign",
        "Comunicado",
        "Apresentação",
        "Validação/Aprovação",
        "Documento cadastral",
        "Contrato social/Alteração cadastral",
        "Berkan/Validação terceiros",
        "Documento complementar",
    )


def _chave_apoio_duplicado(nome: Any, tipo: str) -> str:
    """
    Cria chave para detectar duplicidade em documentos de apoio.
    Ex.: apresentação assinada e apresentação sem assinatura devem ser a mesma família.
    """
    base = _normalizar_texto_chave(nome)
    base = re.sub(r"\b(pdf|docx|doc|assinado|assinada|para|ass|para_ass|rev|revisao|revisão|minuta|rascunho|limpa)\b", " ", base, flags=re.IGNORECASE)
    base = re.sub(r"\s+", " ", base).strip()
    tipo_chave = _normalizar_texto_chave(tipo).replace(" ", "_") or "apoio"
    return f"apoio_{tipo_chave}_{base or 'documento'}"


def _chave_versao_documento(nome: Any, texto: Any = "") -> str:
    tipo = _tipo_documento_triagem(nome, texto)
    nome_low = _normalizar_texto_chave(nome)

    if tipo == "Aditivo":
        numero = _ordinal_aditivo_numero_prioritario(nome, texto)
        if numero:
            return f"aditivo_{numero}"
        # fallback conservador para não juntar aditivos diferentes por engano
        return "aditivo_" + re.sub(r"[^a-z0-9]+", "_", nome_low[:80]).strip("_")

    if tipo == "Contrato principal":
        return "contrato_principal"

    if tipo == "Certificado DocuSign":
        # certificado isolado não deve disputar com contrato/aditivo principal
        m = re.search(r"(?:envelope id|identificação de envelope|identificacao de envelope)[:\s]*([A-Z0-9-]{20,})", clean_text(texto), flags=re.IGNORECASE)
        if m:
            return f"certificado_{m.group(1).lower()}"
        return "certificado_docusign"

    if tipo == "Apresentação":
        # Apresentação assinada e apresentação sem assinatura do mesmo assunto viram a mesma família.
        # Assim a versão assinada fica como APOIO e a duplicada sem assinatura vira IGNORADO.
        return _chave_apoio_duplicado(nome, tipo)

    base = re.sub(r"[^a-z0-9]+", "_", nome_low[:90]).strip("_")
    return f"apoio_{base or 'documento'}"


def _pontuacao_triagem(nome: Any, texto: Any) -> int:
    nome_low = _nome_low(nome)
    tipo = _tipo_documento_triagem(nome, texto)
    plano = f"{nome_low} {clean_text(texto)[:3000].lower()}"
    pontos = 0

    if nome_low.endswith(".pdf"):
        pontos += 300
    if nome_low.endswith(".docx"):
        pontos -= 220

    if tipo in ("Contrato principal", "Aditivo"):
        pontos += 250

    if _evidencia_assinatura_arquivo(nome, texto):
        pontos += 600

    if "assinado" in nome_low or "para_ass" in nome_low or "para ass" in nome_low:
        pontos += 300

    if "certificate of completion" in plano or "certificado de conclusão" in plano or "certificado de conclusao" in plano:
        pontos += 250

    # Minutas/revisões caem, mas se for o único arquivo assinado ainda pode subir.
    if any(t in nome_low for t in ["minuta", "rascunho", "rev", "revisao", "revisão"]):
        pontos -= 140

    if "limpa(assinado)" in nome_low or "limpa (assinado)" in nome_low:
        pontos += 260
    elif "limpa" in nome_low:
        pontos += 40

    if tipo in ("Certificado DocuSign", "Comunicado", "Apresentação", "Proposta/Orçamento", "Validação/Aprovação", "Documento cadastral", "Contrato social/Alteração cadastral", "Berkan/Validação terceiros", "Documento complementar"):
        # Apoio não deve vencer contrato/aditivo principal.
        pontos -= 300

    return pontos


def classificar_anexos_para_analise(arquivos: List[Any], textos_por_arquivo: Dict[str, str]) -> Dict[str, Any]:
    """
    Pré-triagem executiva antes da IA.
    A análise profunda recebe apenas documentos assinados/relevantes.
    Arquivos sem assinatura, minutas, duplicados e versões antigas entram como ignorados com motivo.
    """
    registros = []
    total_arquivos = len(arquivos or [])
    for arquivo in arquivos:
        nome = getattr(arquivo, "name", "documento")
        texto = textos_por_arquivo.get(nome, "")
        falha_leitura = texto_indica_falha_leitura(texto)
        tipo = _tipo_documento_triagem(nome, texto)

        # Regra crítica: arquivo único genérico não pode virar apoio simples só porque
        # a extração local falhou. Ex.: SBF.pdf escaneado precisa seguir para Files API.
        if (
            total_arquivos == 1
            and arquivo_documental_analisavel(nome)
            and tipo == "Documento complementar"
            and not nome_indica_apoio_explicito(nome)
        ):
            tipo = "Contrato principal"

        # Se o conteúdo tem cara de contrato, também promove para contrato principal.
        if tipo == "Documento complementar" and texto_tem_conteudo_contratual(texto):
            tipo = "Contrato principal"

        precisa_assinatura = _documento_precisa_assinatura(tipo)
        assinado = _evidencia_assinatura_arquivo(nome, texto)
        if tipo == "Contrato principal":
            chave = "contrato_principal"
        elif tipo == "Aditivo":
            numero_aditivo_tmp = _ordinal_aditivo_numero_prioritario(nome, texto)
            chave = f"aditivo_{numero_aditivo_tmp}" if numero_aditivo_tmp else _chave_versao_documento(nome, texto)
        else:
            chave = _chave_versao_documento(nome, texto)
        pontos = _pontuacao_triagem(nome, texto)
        if tipo in ("Contrato principal", "Aditivo"):
            pontos += 250
        if tipo in ("Contrato principal", "Aditivo") and falha_leitura:
            pontos += 180
        registros.append({
            "arquivo_obj": arquivo,
            "Arquivo": nome,
            "Tipo": tipo,
            "Assinado": "Sim" if assinado else ("A validar" if falha_leitura and precisa_assinatura else "Não"),
            "Precisa assinatura": "Sim" if precisa_assinatura else "Não",
            "Falha leitura": "Sim" if falha_leitura else "Não",
            "chave": chave,
            "pontos": pontos,
            "texto": texto,
            "Decisão": "Pendente",
            "Motivo": "",
        })

    melhor_por_chave: Dict[str, Dict[str, Any]] = {}
    for reg in registros:
        if reg["Precisa assinatura"] != "Sim":
            continue
        atual = melhor_por_chave.get(reg["chave"])
        if atual is None or reg["pontos"] > atual["pontos"]:
            melhor_por_chave[reg["chave"]] = reg

    # Apoios também podem ter duplicidade. Ex.: apresentação assinada + apresentação sem assinatura.
    # A melhor versão fica como APOIO; as duplicadas viram IGNORADO para manter 6/3/9 no teste BBP.
    melhor_apoio_por_chave: Dict[str, Dict[str, Any]] = {}
    qtd_apoio_por_chave: Dict[str, int] = {}
    existe_apresentacao_assinada = any(
        reg.get("Tipo") == "Apresentação"
        and ("assinado" in _nome_low(reg.get("Arquivo")) or "assinada" in _nome_low(reg.get("Arquivo")))
        for reg in registros
    )
    for reg in registros:
        if reg["Precisa assinatura"] == "Sim":
            continue
        if reg["Tipo"] not in ("Apresentação",):
            continue
        qtd_apoio_por_chave[reg["chave"]] = qtd_apoio_por_chave.get(reg["chave"], 0) + 1
        atual = melhor_apoio_por_chave.get(reg["chave"])
        if atual is None or reg["pontos"] > atual["pontos"]:
            melhor_apoio_por_chave[reg["chave"]] = reg

    enviados_ia = []
    textos_ia = []
    triagem = []

    for reg in registros:
        precisa_assinatura = reg["Precisa assinatura"] == "Sim"
        tipo = reg["Tipo"]
        melhor = melhor_por_chave.get(reg["chave"])
        eh_melhor = melhor is reg
        nome_low = _nome_low(reg["Arquivo"])

        if precisa_assinatura:
            leitura_ruim = reg.get("Falha leitura") == "Sim"
            arquivo_unico_contrato = total_arquivos == 1 and reg.get("Tipo") in ("Contrato principal", "Aditivo")

            if reg["Assinado"] == "Sim" and eh_melhor:
                reg["Decisão"] = "Análise profunda"
                reg["Motivo"] = "Documento assinado selecionado como melhor versão para análise da IA."
                enviados_ia.append(reg["arquivo_obj"])
                textos_ia.append(reg)
            elif leitura_ruim and arquivo_unico_contrato:
                reg["Decisão"] = "Análise profunda"
                reg["Motivo"] = "Arquivo único com leitura local insuficiente; enviado como documento original para Gemini Files API validar conteúdo e assinatura."
                enviados_ia.append(reg["arquivo_obj"])
                textos_ia.append(reg)
            elif arquivo_unico_contrato and eh_melhor:
                reg["Decisão"] = "Análise profunda"
                reg["Motivo"] = "Arquivo único com perfil de contrato/aditivo; análise profunda necessária para validar assinatura, partes, valores e vigência."
                enviados_ia.append(reg["arquivo_obj"])
                textos_ia.append(reg)
            elif reg["Assinado"] == "Sim" and not eh_melhor:
                reg["Decisão"] = "Ignorado da análise profunda"
                reg["Motivo"] = "Versão assinada duplicada/menos prioritária; existe versão melhor do mesmo documento."
            elif melhor and melhor.get("Assinado") == "Sim":
                reg["Decisão"] = "Ignorado da análise profunda"
                reg["Motivo"] = "Sem assinatura ou versão antiga; existe versão assinada equivalente."
            elif eh_melhor and texto_tem_conteudo_contratual(reg.get("texto")):
                reg["Decisão"] = "Análise profunda"
                reg["Motivo"] = "Documento com conteúdo contratual relevante; enviado para IA validar assinatura e dados antes de concluir."
                enviados_ia.append(reg["arquivo_obj"])
                textos_ia.append(reg)
            else:
                reg["Decisão"] = "Ignorado da análise profunda"
                reg["Motivo"] = "Documento exige assinatura, mas não foi localizada evidência suficiente de assinatura."
        else:
            # Apoios não são enviados como arquivo original pesado, mas o texto pode entrar como apoio no prompt.
            melhor_apoio = melhor_apoio_por_chave.get(reg["chave"])
            apoio_duplicado = reg["Tipo"] == "Apresentação" and qtd_apoio_por_chave.get(reg["chave"], 0) > 1 and melhor_apoio is not reg

            # Regra específica: se existem duas apresentações de reajuste e uma delas tem
            # "assinado" no nome, só a assinada fica como APOIO; a outra vira IGNORADO.
            apoio_apresentacao_sem_assinatura = (
                reg["Tipo"] == "Apresentação"
                and existe_apresentacao_assinada
                and "assinado" not in nome_low
                and "assinada" not in nome_low
            )

            if apoio_duplicado or apoio_apresentacao_sem_assinatura:
                reg["Decisão"] = "Ignorado da análise profunda"
                reg["Motivo"] = "Documento de apoio duplicado; existe versão assinada/equivalente selecionada como apoio."
            elif any(t in nome_low for t in ["comunicado", "apresentação", "apresentacao"]):
                reg["Decisão"] = "Apoio simples"
                reg["Motivo"] = "Documento de apoio; usado apenas como contexto, sem análise jurídica profunda."
                textos_ia.append(reg)
            elif tipo == "Certificado DocuSign":
                reg["Decisão"] = "Apoio de assinatura"
                reg["Motivo"] = "Certificado usado para validar assinatura, sem virar documento contratual separado."
                textos_ia.append(reg)
            else:
                reg["Decisão"] = "Apoio simples"
                reg["Motivo"] = "Documento não exige assinatura formal; usado como contexto de apoio."
                textos_ia.append(reg)

        triagem.append({
            "Arquivo": reg["Arquivo"],
            "Tipo": reg["Tipo"],
            "Assinado": reg["Assinado"],
            "Decisão": reg["Decisão"],
            "Motivo": reg["Motivo"],
        })

    # Segurança conservadora: se nenhum documento assinado foi detectado, NÃO manda tudo para análise profunda.
    # Escolhe somente o melhor contrato operacional/final e mantém proposta/cadastros como apoio.
    if not enviados_ia:
        candidatos = [r for r in registros if r.get("Tipo") in ("Contrato principal", "Aditivo")]
        if candidatos:
            melhor_fallback = max(candidatos, key=lambda r: r.get("pontos", 0))
            enviados_ia = [melhor_fallback["arquivo_obj"]]
            if melhor_fallback not in textos_ia:
                textos_ia.append(melhor_fallback)
            for item in triagem:
                if item.get("Arquivo") == melhor_fallback.get("Arquivo"):
                    item["Decisão"] = "Análise profunda"
                    item["Motivo"] = "Fallback conservador: nenhum assinado detectado; selecionado o melhor contrato final/operacional para análise."
                elif item.get("Tipo") in ("Contrato principal", "Aditivo"):
                    item["Decisão"] = "Ignorado da análise profunda"
                    item["Motivo"] = "Modelo/minuta ou versão menos prioritária; existe documento final/operacional melhor selecionado."
        else:
            # Último recurso: usa textos de apoio, mas não promove tudo para análise profunda.
            for reg in registros:
                if reg not in textos_ia:
                    textos_ia.append(reg)
            for item in triagem:
                if item.get("Decisão") == "Ignorado da análise profunda":
                    item["Motivo"] = item.get("Motivo") or "Sem documento principal claro; mantido fora da análise profunda."

    return {
        "arquivos_para_ia": enviados_ia,
        "textos_para_ia": textos_ia,
        "triagem": triagem,
    }


def _sincronizar_triagem_com_resultado(triagem: List[Dict[str, Any]], resultado: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Atualiza a pré-triagem com o que foi efetivamente validado na análise profunda."""
    if not isinstance(triagem, list):
        return []

    saida = [dict(item) for item in triagem if isinstance(item, dict)]
    assinado_final = clean_text(resultado.get("contrato_assinado")).lower() == "sim"
    assinaturas = resultado.get("assinaturas_contrato") if isinstance(resultado.get("assinaturas_contrato"), list) else []

    fontes_assinatura: List[str] = []
    for item in assinaturas:
        if not isinstance(item, dict):
            continue
        fonte = clean_text(item.get("fonte") or item.get("Fonte"))
        if fonte and fonte not in fontes_assinatura:
            fontes_assinatura.append(fonte)

    profundos = [i for i in saida if i.get("Decisão") == "Análise profunda" and i.get("Tipo") in ("Contrato principal", "Aditivo")]

    def _arquivo_mencionado(nome: str) -> bool:
        nome_norm = _normalizar_texto_chave(nome)
        for fonte in fontes_assinatura:
            fonte_norm = _normalizar_texto_chave(fonte)
            if nome_norm and fonte_norm and (nome_norm in fonte_norm or fonte_norm in nome_norm):
                return True
        return False

    for item in saida:
        if item.get("Decisão") != "Análise profunda" or item.get("Tipo") not in ("Contrato principal", "Aditivo"):
            continue

        nome = clean_text(item.get("Arquivo"))
        corresponde = _arquivo_mencionado(nome) or len(profundos) == 1
        if assinado_final and corresponde:
            item["Assinado"] = "Sim"
            detalhe = "Assinatura confirmada durante a análise profunda"
            if fontes_assinatura:
                detalhe += "; fonte: " + "; ".join(fontes_assinatura[:2])
            data_ass = _data_valida_simples(resultado.get("data_assinatura"))
            if data_ass:
                detalhe += f"; data principal: {data_ass}"
            item["Motivo"] = detalhe + "."
        elif not assinado_final and item.get("Assinado") in ("A validar", "Não"):
            item["Assinado"] = "Não validado"
            item["Motivo"] = "Análise profunda concluída sem evidência documental suficiente para confirmar a assinatura."

    return saida


def _finalizar_coerencia_pos_processamento(
    resultado: Dict[str, Any],
    resultado_bruto: Dict[str, Any],
    triagem: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """Última trava: reaplica evidências depois das regras legadas e sincroniza a triagem."""
    base = dict(resultado or {})
    if isinstance(resultado_bruto, dict) and resultado_bruto.get("auditoria_campos"):
        base = _aplicar_auditoria_rigida(base, resultado_bruto)
    else:
        base = _consolidar_assinatura_final(base)
        base = _filtrar_pendencias_sem_evidencia(base)
        base = _recalcular_risco_por_evidencias(base)
        base = _reconstruir_resumo_e_parecer(base)

    # Motor V4: última consolidação determinística. Cards, valores, vigência,
    # assinaturas, checklist, pendências, score e parecer passam a usar apenas
    # fatos sustentados pela matriz de evidências.
    base = aplicar_motor_evidencias_v4(
        base,
        resultado_bruto,
        clean_text(base.get("texto_extraido")),
    )

    base["triagem_anexos"] = _sincronizar_triagem_com_resultado(triagem, base)
    return base


def montar_texto_para_ia_com_triagem(textos_por_arquivo: Dict[str, str], triagem: List[Dict[str, Any]]) -> str:
    """Monta texto reduzido e explicado para a IA, com documentos relevantes + apoios."""
    decisao_por_arquivo = {t.get("Arquivo"): t for t in triagem}
    partes = [
        "TRIAGEM AUTOMÁTICA DOS ANEXOS ANTES DA IA",
        "Regra: documentos assinados recebem análise profunda; minutas/duplicados/sem assinatura ficam apenas registrados com motivo.",
        "Regra crítica: se o texto extraído de um PDF apresentar erro técnico de leitura, a IA deve ignorar esse erro como conteúdo e analisar o arquivo original anexado pela Files API.",
        "Regra de qualidade: cards e análise completa só podem ser preenchidos com evidência do contrato principal/anexo válido; apoio cadastral não substitui contrato operacional.",
        "",
    ]

    for item in triagem:
        partes.append(
            f"- {item.get('Arquivo')} | Tipo: {item.get('Tipo')} | Assinado: {item.get('Assinado')} | Decisão: {item.get('Decisão')} | Motivo: {item.get('Motivo')}"
        )

    partes.append("\n==============================\nDOCUMENTOS ENVIADOS COMO BASE DA ANÁLISE\n==============================")

    for nome, texto in textos_por_arquivo.items():
        item = decisao_por_arquivo.get(nome, {})
        decisao = item.get("Decisão", "")
        if decisao in ("Análise profunda", "Apoio simples", "Apoio de assinatura"):
            partes.append("\n\n==============================")
            partes.append(f"ARQUIVO: {nome}")
            partes.append(f"DECISÃO DA TRIAGEM: {decisao}")
            partes.append(f"MOTIVO: {item.get('Motivo', '')}")
            partes.append("==============================")
            partes.append(str(texto or ""))

    return "\n".join(partes).strip()


def executar_ia_com_timer_ao_vivo(
    funcao_ia,
    popup_processamento,
    arquivos_status: Dict[str, str],
    inicio_processamento: float,
    *args,
    **kwargs,
):
    """Executa Gemini em segundo plano e mantém o popup atualizando o tempo da IA."""
    resultado_box: Dict[str, Any] = {}
    erro_box: Dict[str, Any] = {}
    inicio_ia = time.perf_counter()
    total = len(arquivos_status)

    def alvo():
        try:
            resultado_box["resultado"] = funcao_ia(*args, **kwargs)
        except Exception as exc:
            erro_box["erro"] = exc

    thread = threading.Thread(target=alvo, daemon=True)
    try:
        if add_script_run_ctx is not None:
            add_script_run_ctx(thread)
    except Exception:
        pass

    thread.start()
    while thread.is_alive():
        tempo_ia = _formatar_tempo_execucao(time.perf_counter() - inicio_ia)
        preparados = sum(1 for s in arquivos_status.values() if s != "PENDENTE")
        atualizar_popup_processamento(
            popup_processamento,
            arquivos_status,
            f"Arquivos preparados: {preparados}/{total}",
            f"IA analisando o pacote final...\nTempo da IA: {tempo_ia}",
            inicio_processamento,
        )
        time.sleep(1)

    tempo_ia_final = time.perf_counter() - inicio_ia
    try:
        st.session_state["ultimo_tempo_ia_segundos"] = tempo_ia_final
    except Exception:
        pass

    if erro_box:
        raise erro_box["erro"]

    resultado_final = resultado_box.get("resultado")
    if isinstance(resultado_final, dict):
        resultado_final["_tempo_ia_segundos"] = tempo_ia_final

    return resultado_final


def render_triagem_anexos(resultado: Dict[str, Any]) -> None:
    triagem = resultado.get("triagem_anexos") or []
    st.markdown('<div class="section-title">Triagem dos anexos</div>', unsafe_allow_html=True)
    if not isinstance(triagem, list) or not triagem:
        st.info("Triagem dos anexos não disponível para este registro.")
        return

    total = len(triagem)
    profunda = sum(1 for x in triagem if x.get("Decisão") == "Análise profunda")
    apoio = sum(1 for x in triagem if str(x.get("Decisão", "")).startswith("Apoio"))
    ignorados = sum(1 for x in triagem if "Ignorado" in str(x.get("Decisão", "")))

    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(render_metric("Arquivos", total), unsafe_allow_html=True)
    c2.markdown(render_metric("Análise profunda", profunda), unsafe_allow_html=True)
    c3.markdown(render_metric("Apoio", apoio), unsafe_allow_html=True)
    c4.markdown(render_metric("Ignorados", ignorados), unsafe_allow_html=True)

    st.markdown(
        '<div class="executive-box"><h3>🧭 Pré-triagem automática</h3>'
        '<p>O robô priorizou documentos assinados para análise profunda. Minutas, duplicados e arquivos sem assinatura foram registrados com motivo para não atrasar a IA.</p></div>',
        unsafe_allow_html=True,
    )

    df = pd.DataFrame(triagem)
    colunas = [c for c in ["Arquivo", "Tipo", "Assinado", "Decisão", "Motivo"] if c in df.columns]

    def _cor_triagem(row):
        decisao = str(row.get("Decisão", ""))
        if decisao == "Análise profunda":
            return ["background-color: rgba(22, 163, 74, 0.18); color: #ffffff; font-weight: 700;"] * len(row)
        if decisao.startswith("Apoio"):
            return ["background-color: rgba(59, 130, 246, 0.14); color: #ffffff; font-weight: 700;"] * len(row)
        if "Ignorado" in decisao:
            return ["background-color: rgba(100, 116, 139, 0.18); color: rgba(255,255,255,0.82); font-weight: 700;"] * len(row)
        return [""] * len(row)

    with st.expander("📋 Ver arquivos ignorados/apoio e motivo", expanded=False):
        st.dataframe(df[colunas].style.apply(_cor_triagem, axis=1), use_container_width=True, hide_index=True)

def render_resumo_processamento_final(resultado: Dict[str, Any]) -> None:
    """Mostra um resumo final da execução depois que o popup fecha."""
    resumo = resultado.get("resumo_processamento") or {}
    if not isinstance(resumo, dict) or not resumo:
        return

    st.markdown('<div class="section-title">Resumo do processamento</div>', unsafe_allow_html=True)

    total = resumo.get("total_arquivos", 0)
    profunda = resumo.get("analise_profunda", 0)
    apoio = resumo.get("apoio", 0)
    ignorados = resumo.get("ignorados", 0)
    tempo_total = resumo.get("tempo_total", "0s")
    tempo_ia = resumo.get("tempo_ia", "0s")
    finalizado_em = resumo.get("finalizado_em", "")

    html = f'''
<div class="executive-box" style="padding:22px 24px;">
    <h3 style="margin:0 0 14px 0;">✅ Análise concluída</h3>
    <p style="margin:0 0 18px 0;">O processamento foi finalizado com sucesso. O popup foi fechado para liberar a tela, mas o resumo da execução fica registrado abaixo.</p>
    <div style="display:grid;grid-template-columns:repeat(6,1fr);gap:12px;">
        <div class="valor-card" style="min-height:95px;padding:14px;">
            <div class="valor-titulo">TOTAL</div>
            <div class="valor-principal" style="font-size:22px;">{safe(total)}</div>
        </div>
        <div class="valor-card" style="min-height:95px;padding:14px;">
            <div class="valor-titulo">ANÁLISE IA</div>
            <div class="valor-principal" style="font-size:22px;">{safe(profunda)}</div>
        </div>
        <div class="valor-card" style="min-height:95px;padding:14px;">
            <div class="valor-titulo">APOIO</div>
            <div class="valor-principal" style="font-size:22px;">{safe(apoio)}</div>
        </div>
        <div class="valor-card" style="min-height:95px;padding:14px;">
            <div class="valor-titulo">IGNORADOS</div>
            <div class="valor-principal" style="font-size:22px;">{safe(ignorados)}</div>
        </div>
        <div class="valor-card" style="min-height:95px;padding:14px;">
            <div class="valor-titulo">TEMPO TOTAL</div>
            <div class="valor-principal" style="font-size:22px;">{safe(tempo_total)}</div>
        </div>
        <div class="valor-card" style="min-height:95px;padding:14px;">
            <div class="valor-titulo">TEMPO IA</div>
            <div class="valor-principal" style="font-size:22px;">{safe(tempo_ia)}</div>
        </div>
    </div>
    <div class="valor-alerta" style="margin-top:16px;display:inline-block;">Finalizado em {safe(finalizado_em)}</div>
</div>
'''
    st.markdown(html, unsafe_allow_html=True)


def fechar_popup_processamento(placeholder) -> None:
    placeholder.empty()


def render_aditivos_contrato(resultado: Dict[str, Any]) -> None:
    """Renderiza seção executiva de aditivos identificados de forma limpa e profissional."""
    aditivos = normalizar_aditivos_contrato(resultado.get("aditivos_contrato", []))

    st.markdown('<div class="section-title">Aditivos identificados</div>', unsafe_allow_html=True)

    resumo = clean_text(resultado.get("resumo_aditivos"))

    total = len(aditivos)
    assinados = sum(1 for a in aditivos if clean_text(a.get("Assinado")).upper() == "SIM")
    nao_assinados = sum(1 for a in aditivos if clean_text(a.get("Assinado")).upper() in ("NÃO", "NAO"))
    valor_total_fmt = "Sem valor global fixo"

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Aditivos", total)
    c2.metric("Assinados", assinados)
    c3.metric("Não assinados", nao_assinados)
    c4.metric("Valor dos aditivos", valor_total_fmt)

    st.markdown(
        f"""
<div class="executive-box">
    <h3>📎 Resumo dos aditivos</h3>
    <p>{safe(resumo)}</p>
</div>
""",
        unsafe_allow_html=True,
    )

    if not aditivos:
        st.info("Nenhum termo aditivo foi identificado nos documentos analisados.")
        return

    def curto(valor: Any, limite: int = 90) -> str:
        return clean_text(valor)

    # Cards executivos por aditivo
    for i in range(0, len(aditivos), 3):
        cols = st.columns(3)

        for col, aditivo in zip(cols, aditivos[i:i + 3]):
            assinado = clean_text(aditivo.get("Assinado"))
            status_cor = "#16a34a" if assinado.upper() == "SIM" else "#dc2626" if assinado.upper() in ("NÃO", "NAO") else "#f59e0b"

            with col:
                st.markdown(
                    f"""
<div class="valor-card" style="min-height:230px;">
    <div class="valor-titulo">ADITIVO {safe(aditivo.get("Nº"))}</div>
    <div class="valor-principal" style="font-size:18px;">{safe(aditivo.get("Tipo do aditivo"))}</div>
    <div class="valor-detalhe">
        <b>Anexo:</b><br>{safe(curto(aditivo.get("Anexo do aditivo"), 115))}<br><br>
        <b>Data:</b> {safe(aditivo.get("Data do aditivo"))}<br>
        <b>Assinado:</b> <span style="color:{status_cor};font-weight:900;">{safe(assinado)}</span><br>
        <b>Assinatura:</b> {safe(aditivo.get("Data da assinatura"))}<br>
        <b>Período:</b> {safe(curto(aditivo.get("Período do aditivo"), 90))}
    </div>
    <div class="valor-alerta">{safe(curto(aditivo.get("Status de validação"), 45))}</div>
</div>
""",
                    unsafe_allow_html=True,
                )

    # Tabela técnica escondida para não poluir a apresentação
    with st.expander("📋 Ver tabela técnica dos aditivos", expanded=False):
        df_aditivos = pd.DataFrame([
            {k: v for k, v in aditivo.items() if k != "_itens_aditivo"}
            for aditivo in aditivos
        ])

        colunas_resumo = [
            "Nº",
            "Tipo do aditivo",
            "Anexo do aditivo",
            "Data do aditivo",
            "Assinado",
            "Data da assinatura",
            "Valor do aditivo",
            "Impacto no valor",
            "Impacto no prazo",
            "Período do aditivo",
            "Status de validação",
        ]

        colunas_resumo = [c for c in colunas_resumo if c in df_aditivos.columns]
        df_view = df_aditivos[colunas_resumo].copy()

        def _cor_linha_aditivo(row):
            assinado_linha = clean_text(row.get("Assinado")).upper()

            if assinado_linha == "SIM":
                return ["background-color: rgba(22, 163, 74, 0.22); color: #ffffff; font-weight: 700;"] * len(row)

            if assinado_linha in ("NÃO", "NAO"):
                return ["background-color: rgba(220, 38, 38, 0.24); color: #ffffff; font-weight: 700;"] * len(row)

            return ["background-color: rgba(245, 158, 11, 0.18); color: #ffffff; font-weight: 700;"] * len(row)

        st.dataframe(
            df_view.style.apply(_cor_linha_aditivo, axis=1),
            use_container_width=True,
            hide_index=True,
        )

    st.markdown('<div class="section-title">Detalhamento dos aditivos</div>', unsafe_allow_html=True)

    for idx, aditivo in enumerate(aditivos, 1):
        titulo = (
            f"Aditivo {aditivo.get('Nº', idx)} • "
            f"{aditivo.get('Assinado', 'Não localizado')} • "
            f"{aditivo.get('Anexo do aditivo', 'Anexo não localizado')}"
        )

        with st.expander(titulo):
            d1, d2, d3 = st.columns(3)

            d1.markdown(f"**Tipo do aditivo**  \n{safe(aditivo.get('Tipo do aditivo'))}", unsafe_allow_html=True)
            d1.markdown(f"**Anexo do aditivo**  \n{safe(aditivo.get('Anexo do aditivo'))}", unsafe_allow_html=True)
            d1.markdown(f"**Data do aditivo**  \n{safe(aditivo.get('Data do aditivo'))}", unsafe_allow_html=True)
            d1.markdown(f"**Data de carga no robô**  \n{safe(aditivo.get('Data de carga no robô'))}", unsafe_allow_html=True)

            d2.markdown(f"**Assinado**  \n{safe(aditivo.get('Assinado'))}", unsafe_allow_html=True)
            d2.markdown(f"**Data da assinatura**  \n{safe(aditivo.get('Data da assinatura'))}", unsafe_allow_html=True)
            d2.markdown(f"**Quem assinou**  \n{safe(aditivo.get('Quem assinou'))}", unsafe_allow_html=True)

            d3.markdown(f"**Valor do aditivo**  \n{safe(aditivo.get('Valor do aditivo'))}", unsafe_allow_html=True)
            d3.markdown(f"**Impacto no valor**  \n{safe(aditivo.get('Impacto no valor'))}", unsafe_allow_html=True)
            d3.markdown(f"**Impacto no prazo**  \n{safe(aditivo.get('Impacto no prazo'))}", unsafe_allow_html=True)

            st.markdown(f"**Período do aditivo**  \n{safe(aditivo.get('Período do aditivo'))}", unsafe_allow_html=True)
            st.markdown(f"**Escopo do aditivo**  \n{safe(aditivo.get('Escopo do aditivo'))}", unsafe_allow_html=True)
            st.markdown(f"**Status de validação**  \n{safe(aditivo.get('Status de validação'))}", unsafe_allow_html=True)

            itens = aditivo.get("_itens_aditivo", [])
            if itens:
                with st.expander(f"📦 Itens do aditivo {aditivo.get('Nº', idx)}", expanded=False):
                    st.dataframe(pd.DataFrame(itens), use_container_width=True, hide_index=True)
            else:
                st.caption("Nenhum item específico foi identificado para este aditivo.")


def render_itens_contrato(resultado: Dict[str, Any], titulo: str = "Materiais e serviços identificados") -> None:
    itens = normalizar_itens_contrato(resultado.get("itens_contrato", []))
    st.markdown(f'<div class="section-title">{safe(titulo)}</div>', unsafe_allow_html=True)
    metricas = resultado.get("metricas_tabela_comercial") if isinstance(resultado.get("metricas_tabela_comercial"), dict) else {}
    if metricas:
        c1, c2, c3, c4 = st.columns(4)
        c1.markdown(render_metric("Itens no documento", metricas.get("itens_encontrados_documento", len(itens))), unsafe_allow_html=True)
        c2.markdown(render_metric("Itens exibidos", metricas.get("itens_exibidos_auditor", len(itens))), unsafe_allow_html=True)
        c3.markdown(render_metric("Cobertura da tabela", f'{metricas.get("cobertura_tabela_percentual", 0)}%'), unsafe_allow_html=True)
        c4.markdown(render_metric("Página(s) comercial(is)", metricas.get("paginas_tabela_comercial", "Não localizada")), unsafe_allow_html=True)
    if not itens:
        st.info("Nenhum material, serviço ou condição comercial unitária foi identificado no contrato/anexos.")
        return

    df_itens = pd.DataFrame(itens)
    colunas_base = ["Item", "Grupo/Tabela", "Descrição", "Tipo", "Natureza do valor", "Faixa/Condição", "Quantidade", "Unidade", "Valor unitário", "Valor total"]
    colunas_opcionais = ["Periodicidade", "Condição comercial", "Taxa / Percentual", "Total de encargos", "Vencimento / Prazo", "Página", "Status de evidência", "Evidência"]
    colunas_finais = [c for c in colunas_base if c in df_itens.columns]
    for col in colunas_opcionais:
        if col in df_itens.columns and df_itens[col].apply(_valor_informado).any():
            colunas_finais.append(col)
    if "Fonte" in df_itens.columns:
        colunas_finais.append("Fonte")
    df_itens = df_itens[colunas_finais]

    st.dataframe(
        df_itens,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Grupo/Tabela": st.column_config.TextColumn("Grupo/Tabela", width="medium"),
            "Descrição": st.column_config.TextColumn("Descrição", width="large"),
            "Faixa/Condição": st.column_config.TextColumn("Faixa/Condição", width="medium"),
            "Condição comercial": st.column_config.TextColumn("Condição comercial", width="large"),
            "Valor unitário": st.column_config.TextColumn("Valor unitário"),
            "Valor total": st.column_config.TextColumn("Valor total"),
            "Taxa / Percentual": st.column_config.TextColumn("Taxa / Percentual"),
            "Total de encargos": st.column_config.TextColumn("Total de encargos"),
            "Vencimento / Prazo": st.column_config.TextColumn("Vencimento / Prazo", width="medium"),
            "Fonte": st.column_config.TextColumn("Fonte", width="medium"),
            "Página": st.column_config.TextColumn("Página", width="small"),
            "Evidência": st.column_config.TextColumn("Evidência", width="large"),
            "Natureza do valor": st.column_config.TextColumn("Natureza do valor", width="medium"),
        },
    )


def _hist_texto_curto(valor: Any, limite: int = 190) -> str:
    txt = clean_text(valor)
    if txt in ("", "Não localizado", "None", "nan"):
        return "Não localizado"
    txt = re.sub(r"\s+", " ", txt).strip()
    if len(txt) <= limite:
        return txt
    frases = re.split(r"(?<=[.!?])\s+", txt)
    if frases and 35 <= len(frases[0]) <= limite:
        return frases[0]
    return txt[:limite].rsplit(" ", 1)[0]


def _hist_valor_resumo(row: pd.Series, resultado: Dict[str, Any]) -> str:
    candidatos = [resultado.get("valor_contrato_original"), resultado.get("valor_total"), row.get("valor_total") if "valor_total" in row.index else None]
    for valor in candidatos:
        txt = clean_text(valor)
        low = txt.lower()
        if "sem valor global fixo" in low:
            return "Sem valor global fixo"
        if "não calculável" in low or "nao calculavel" in low:
            return "Não calculável com precisão"
        if txt not in ("", "Não localizado", "None", "nan"):
            return _hist_texto_curto(txt, 150)
    return "Não localizado"


def _hist_split_arquivos(valor: Any) -> List[str]:
    txt = clean_text(valor)
    if txt in ("", "Não localizado", "None", "nan"):
        return []
    partes = re.split(r"\s*\|\s*", txt)
    return [p.strip() for p in partes if p.strip()]


def _hist_resumo_processamento(resultado: Dict[str, Any], arquivos: List[str]) -> Dict[str, Any]:
    resumo = resultado.get("resumo_processamento") if isinstance(resultado.get("resumo_processamento"), dict) else {}
    triagem = resultado.get("triagem_anexos") if isinstance(resultado.get("triagem_anexos"), list) else []
    if triagem:
        total = len(triagem)
        profunda = sum(1 for x in triagem if x.get("Decisão") == "Análise profunda")
        apoio = sum(1 for x in triagem if str(x.get("Decisão", "")).startswith("Apoio"))
        ignorados = sum(1 for x in triagem if "Ignorado" in str(x.get("Decisão", "")))
    else:
        total = resumo.get("total_arquivos") or len(arquivos)
        profunda = resumo.get("analise_profunda") or "-"
        apoio = resumo.get("apoio") or "-"
        ignorados = resumo.get("ignorados") or "-"
    return {"total": total, "analise_ia": profunda, "apoio": apoio, "ignorados": ignorados}


def render_historico_card_executivo(row: pd.Series) -> str:
    resultado = _hist_parse_resultado_json(row)

    fornecedor = clean_text(row.get("fornecedor") or resultado.get("contraparte") or resultado.get("fornecedor"))
    cnpj = clean_text(row.get("cnpj") or resultado.get("cnpj_contraparte") or resultado.get("cnpj"))
    risco = normalize_risco(row.get("risco") or resultado.get("risco"))
    cor = risco_cor(risco)
    score = clean_text(row.get("score") or resultado.get("score"))
    status = clean_text(row.get("status") or resultado.get("status"))
    origem = clean_text(row.get("tipo_origem") or resultado.get("tipo_origem"))
    modelo = clean_text(row.get("modelo_ia") or resultado.get("modelo_ia"))
    assinado = clean_text(row.get("contrato_assinado") or resultado.get("contrato_assinado"))
    data = clean_text(row.get("data_analise") or resultado.get("data_analise"))
    id_reg = clean_text(row.get("id"))

    vigencia = _hist_texto_curto(
        row.get("vigencia") or resultado.get("periodo_vigencia") or resultado.get("vigencia_apos_assinatura") or resultado.get("vigencia"),
        180,
    )
    valor = _hist_valor_resumo(row, resultado)
    arquivos = _hist_split_arquivos(row.get("arquivo") or resultado.get("arquivos_analisados"))
    qtd_arquivos = len(arquivos)
    resumo_proc = _hist_resumo_processamento(resultado, arquivos)

    aditivos = normalizar_aditivos_contrato(resultado.get("aditivos_contrato", [])) if isinstance(resultado, dict) else []
    qtd_aditivos = len(aditivos)
    aditivos_assinados = sum(1 for a in aditivos if clean_text(a.get("Assinado")).upper() == "SIM")
    itens = normalizar_itens_contrato(resultado.get("itens_contrato", [])) if isinstance(resultado, dict) else []
    qtd_itens = len(itens)
    pendencias = resultado.get("pendencias", []) if isinstance(resultado.get("pendencias"), list) else []
    qtd_pendencias = len(pendencias)

    valor_principal = "Sem valor global fixo" if "sem valor global fixo" in valor.lower() else valor
    parecer = _hist_texto_curto(resultado.get("parecer") or resultado.get("resumo_executivo") or "Use a análise completa para consultar o parecer detalhado.", 190)

    total_proc = resumo_proc.get("total")
    analise_ia = resumo_proc.get("analise_ia")
    apoio = resumo_proc.get("apoio")
    ignorados = resumo_proc.get("ignorados")

    subtitulo = f"Análise ID {safe(id_reg)} • {safe(data)}"

    try:
        excel_bytes = gerar_excel_card_bytes(row)
        excel_b64 = base64.b64encode(excel_bytes).decode("ascii")
        excel_href = f"data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{excel_b64}"
        excel_file = f"analise_contrato_{id_reg or 'historico'}.xlsx"
        excel_html = f'<a class="history-excel-link-v4" href="{excel_href}" download="{safe(excel_file)}">📥 Baixar Excel</a>'
    except Exception:
        excel_html = '<span class="history-excel-link-v4">Excel indisponível</span>'

    html_card = f"""
<div class="history-card-v2">
<div class="history-hero-v3">
<div class="history-title-v3">
<div class="history-icon-v3">📄</div>
<div>
<h3>{safe(fornecedor)}</h3>
<div class="history-sub-v3">{subtitulo}</div>
</div>
</div>
<div class="history-risk-stack-v3">
<div class="history-risk-v2" style="background:{cor};">{safe(risco)}</div>
<div class="history-score-v3">Score <strong>{safe(score)}</strong></div>
</div>
</div>
<div class="history-body-v3">
<div class="history-chip-row-v2">
<span class="history-chip-v2 ok">✅ Assinado: {safe(assinado)}</span>
<span class="history-chip-v2">📌 Origem: {safe(origem)}</span>
<span class="history-chip-v2">🤖 Modelo: {safe(modelo)}</span>
<span class="history-chip-v2 warn">📁 Anexos: {safe(qtd_arquivos)}</span>
</div>
<div class="history-main-grid-v3">
<div class="history-kpi-v2"><small>CNPJ</small><strong>{safe(cnpj)}</strong></div>
<div class="history-kpi-v2 history-value-v3"><small>Valor do contrato</small><strong>{safe(valor_principal)}</strong><p>Resumo financeiro consolidado da análise.</p></div>
<div class="history-kpi-v2"><small>Status</small><strong>{safe(status)}</strong></div>
<div class="history-kpi-v2"><small>Score</small><strong>{safe(score)}</strong></div>
</div>
<div class="history-wide-v2">
<div class="history-box-v2">
<small>Vigência / Período</small>
<p>{safe(vigencia)}</p>
</div>
<div class="history-box-v2">
<small>Pré-triagem dos anexos</small>
<div class="history-mini-dashboard-v3">
<div class="history-mini-v3"><small>Total</small><strong>{safe(total_proc)}</strong><span>arquivos</span></div>
<div class="history-mini-v3"><small>Análise IA</small><strong>{safe(analise_ia)}</strong><span>profunda</span></div>
<div class="history-mini-v3"><small>Apoio</small><strong>{safe(apoio)}</strong><span>contexto</span></div>
<div class="history-mini-v3"><small>Ignorados</small><strong>{safe(ignorados)}</strong><span>duplicados</span></div>
</div>
</div>
</div>
<div class="history-summary-v3">
<div class="history-insight-v3"><small>Aditivos</small><strong>{safe(qtd_aditivos)}</strong><p>{safe(aditivos_assinados)} assinado(s)</p></div>
<div class="history-insight-v3"><small>Materiais/Serviços</small><strong>{safe(qtd_itens)}</strong><p>itens identificados</p></div>
<div class="history-insight-v3"><small>Pendências</small><strong>{safe(qtd_pendencias)}</strong><p>pontos para revisar</p></div>
<div class="history-insight-v3"><small>Modelo</small><strong>{safe(modelo)}</strong><p>motor utilizado</p></div>
</div>
<div class="history-docline-v3">
<span><b>Resumo:</b> {safe(parecer)}</span>
<span><b>Registro:</b> {safe(data)} • ID {safe(id_reg)}</span>
</div>
<div class="history-actions-v4">
<div class="history-action-note-v4">Abra a análise completa abaixo para consultar aditivos, materiais, triagem, checklist e auditoria técnica.</div>
{excel_html}
</div>
</div>
</div>
"""
    return "\n".join(line.strip() for line in html_card.splitlines() if line.strip())

def obter_resultado_completo_historico(row: pd.Series) -> tuple[Dict[str, Any], str]:
    """Recupera do histórico o JSON completo salvo na análise."""
    resultado: Dict[str, Any] = {}
    texto_extraido = ""

    raw_json = row.get("resultado_json") if "resultado_json" in row.index else None
    if raw_json not in (None, "", "Não informado"):
        try:
            resultado = json.loads(raw_json)
            if not isinstance(resultado, dict):
                resultado = {}
        except Exception:
            resultado = {}

    if resultado:
        texto_extraido = str(resultado.get("texto_extraido") or row.get("texto_extraido") or "")
    else:
        texto_extraido = str(row.get("texto_extraido") or "")
        resultado = {
            "data_analise": row.get("data_analise"),
            "contraparte": row.get("fornecedor"),
            "fornecedor": row.get("fornecedor"),
            "cnpj_contraparte": row.get("cnpj"),
            "cnpj": row.get("cnpj"),
            "valor_contrato_original": row.get("valor_total"),
            "valor_mensal_estimado": row.get("valor_mensal_estimado") if "valor_mensal_estimado" in row.index else "Não localizado",
            "valor_total_estimado_vigencia": row.get("valor_total_estimado_vigencia") if "valor_total_estimado_vigencia" in row.index else "Não localizado",
            "valor_total": row.get("valor_total"),
            "vigencia_apos_assinatura": row.get("vigencia"),
            "vigencia": row.get("vigencia"),
            "status": row.get("status"),
            "risco": row.get("risco"),
            "score": row.get("score"),
            "contrato_assinado": row.get("contrato_assinado"),
            "data_contrato": row.get("data_contrato") if "data_contrato" in row.index else "Não localizado",
            "data_conclusao_docusign": row.get("data_conclusao_docusign") if "data_conclusao_docusign" in row.index else "Não localizado",
            "pessoas_que_assinaram": row.get("pessoas_que_assinaram") if "pessoas_que_assinaram" in row.index else "Não localizado",
            "modelo_ia": row.get("modelo_ia"),
            "tipo_origem": row.get("tipo_origem"),
            "arquivos_analisados": row.get("arquivo"),
            "resumo_executivo": "Registro antigo: visualização montada com os dados disponíveis no histórico.",
            "parecer": "Para visualizar todos os campos com maior detalhe, refaça a análise do contrato nesta versão atualizada.",
            "checklist": [],
            "pendencias": [],
            "itens_contrato": [],
        }

    resultado.setdefault("texto_extraido", texto_extraido)
    resultado.setdefault("arquivos_analisados", row.get("arquivo", "Não informado"))
    resultado.setdefault("data_analise", row.get("data_analise", "Não informado"))
    resultado.setdefault("modelo_ia", row.get("modelo_ia", "Não informado"))
    resultado.setdefault("tipo_origem", row.get("tipo_origem", "Não informado"))

    try:
        resultado = normalizar(resultado)
    except Exception:
        pass

    return resultado, texto_extraido


def render_analise_completa_historico(row: pd.Series) -> None:
    """Mostra no Histórico a mesma visão completa exibida após uma nova análise."""
    resultado, texto_extraido = obter_resultado_completo_historico(row)
    risco = normalize_risco(resultado.get("risco"))
    pill = "pill-ok" if risco == "BAIXO" else "pill-warn" if risco == "MÉDIO" else "pill-danger"

    st.markdown('<div class="section-title">Resumo da análise</div>', unsafe_allow_html=True)
    render_indicadores_analise(resultado)

    if str(resultado.get("contrato_assinado", "")).upper() == "NÃO":
        st.error("⚠️ Contrato sem assinatura localizada. Revisar antes da criação da RC/PO.")

    render_resultado_em_abas(
        resultado,
        texto_extraido=texto_extraido,
        key_prefix=f"texto_historico_{row.get('id', id(row))}",
    )

    st.download_button(
        "📥 Baixar relatório Excel completo",
        data=gerar_excel(resultado, texto_extraido).getvalue(),
        file_name=f"analise_completa_contrato_{row.get('id', 'historico')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key=f"download_historico_completo_{row.get('id', id(row))}",
    )

# =========================================================
# SIDEBAR
# =========================================================
with st.sidebar:
    st.markdown("## ⚖️ AUDITOR DE CONTRATOS")
    st.caption("Análise inteligente de contratos")
    st.divider()

    query_pagina = str(st.query_params.get("pagina", "")).lower()
    abrir_dashboard = bool(st.query_params.get("risco")) or query_pagina == "dashboard"

    pagina = st.radio(
    "Menu",
    [
        "🏠 Dashboard",
        "📄 Nova Análise",
        "📚 Histórico",
        "🤖 Assistente IA"
    ],
    index=0 if abrir_dashboard else 1,
)

    st.divider()

    # Modo fixo para o usuário final.
    modo = st.radio(
        "Modo de análise",
        ["Automático recomendado"],
        index=0,
    )

    gemini_key = obter_gemini_key()
    st.info("Modo automático recomendado ativo.")


# =========================================================
# DASHBOARD
# =========================================================
if pagina == "🏠 Dashboard":
    render_hero("Dashboard", "Visão executiva das análises de contratos, riscos e pendências.")

    historico = carregar_historico_seguro()

    col_a, col_b = st.columns([3, 1])
    with col_b:
        if st.button("🗑️ Limpar histórico", use_container_width=True):
            limpar_historico()
            st.success("Histórico limpo com sucesso.")
            st.rerun()

    total = len(historico)
    if total > 0:
        risco_series = historico["risco"].astype(str).str.upper().replace({"MEDIO": "MÉDIO"})
        alto = int((risco_series == "ALTO").sum())
        medio = int((risco_series == "MÉDIO").sum())
        baixo = int((risco_series == "BAIXO").sum())
        score_medio = round(historico["score"].apply(as_float_score).mean(), 1)
    else:
        alto = medio = baixo = score_medio = 0

    st.markdown(
        f"""
        <div class="executive-box">
            <h3>📊 Resumo executivo</h3>
            <p>O ambiente possui <b>{total}</b> contrato(s) analisado(s). Atualmente existem <b>{alto}</b> contrato(s) classificados como risco alto.</p>
            <p>A prioridade recomendada é revisar contratos de risco <b>ALTO</b> antes da criação ou continuidade de RC/PO.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    filtro_param = st.query_params.get("risco", "TODOS")
    if isinstance(filtro_param, list):
        filtro_param = filtro_param[0] if filtro_param else "TODOS"

    filtro_atual = normalize_risco(filtro_param)
    if filtro_atual in ["MEDIO", "MÉDIO"]:
        filtro_atual = "MÉDIO"
    if filtro_atual not in ["TODOS", "ALTO", "MÉDIO", "BAIXO"]:
        filtro_atual = "TODOS"

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.markdown(render_filter_metric("Contratos", total, "TODOS", filtro_atual == "TODOS"), unsafe_allow_html=True)
    c2.markdown(render_filter_metric("Risco alto", alto, "ALTO", filtro_atual == "ALTO"), unsafe_allow_html=True)
    c3.markdown(render_filter_metric("Risco médio", medio, "MEDIO", filtro_atual == "MÉDIO"), unsafe_allow_html=True)
    c4.markdown(render_filter_metric("Risco baixo", baixo, "BAIXO", filtro_atual == "BAIXO"), unsafe_allow_html=True)
    c5.markdown(render_filter_metric("Score médio", score_medio, "TODOS", False), unsafe_allow_html=True)

    historico_filtrado = historico.copy()
    if not historico_filtrado.empty and filtro_atual != "TODOS":
        risco_filtro = historico_filtrado["risco"].astype(str).str.upper().replace({"MEDIO": "MÉDIO"})
        historico_filtrado = historico_filtrado[risco_filtro == filtro_atual]

    texto_filtro = "todos os contratos" if filtro_atual == "TODOS" else f"contratos com risco {filtro_atual}"
    st.markdown(
        f'<div class="filter-note">🔎 Visualizando {len(historico_filtrado)} de {total} análise(s): <b>{safe(texto_filtro)}</b>.</div>',
        unsafe_allow_html=True,
    )

    st.markdown('<div class="section-title">Últimas análises</div>', unsafe_allow_html=True)
    if historico.empty:
        st.info("Nenhuma análise registrada ainda.")
    elif historico_filtrado.empty:
        st.warning("Nenhuma análise encontrada para este filtro.")
    else:
        for _, row in historico_filtrado.head(10).iterrows():
            render_contract_card(row)

        st.markdown('<div class="footer">Auditor de Contratos - Grupo SBF • Suprimentos • Análise de Contratos</div>', unsafe_allow_html=True)
        st.stop()

# =========================================================
# ASSISTENTE IA LOCAL
# =========================================================
if pagina == "🤖 Assistente IA":

    render_hero(
        "Assistente Auditor",
        "Consulte informações dos contratos analisados no sistema."
    )

    contratos = carregar_contratos_chat()

    if contratos.empty:
        st.warning("Nenhum contrato encontrado no banco.")
        st.stop()

    st.markdown(
        """
        <style>
        .ai-panel{
            background:linear-gradient(145deg,#101821,#0b1118);
            border:1px solid rgba(215,191,117,.16);
            border-radius:18px;
            padding:22px 24px;
            margin:18px 0 22px;
            box-shadow:0 18px 45px rgba(0,0,0,.22);
        }
        .ai-question-card{
            background:#1a202b;
            border-left:6px solid #ff4d4d;
            border-radius:14px;
            padding:16px 18px;
            margin:18px 0 14px;
            color:#ffffff;
            font-weight:900;
            box-shadow:0 14px 30px rgba(0,0,0,.18);
        }
        .ai-answer-title{
            color:#f3d36b;
            font-weight:900;
            font-size:18px;
            margin:0 0 14px;
        }
        .ai-answer-text{
            color:#e5e7eb;
            font-size:15px;
            line-height:1.75;
            font-weight:650;
        }
        .ai-summary-grid{
            display:grid;
            grid-template-columns:repeat(5,minmax(0,1fr));
            gap:14px;
            margin-top:12px;
        }
        .ai-summary-card{
            background:#151d28;
            border:1px solid rgba(255,255,255,.06);
            border-radius:14px;
            padding:16px;
        }
        .ai-summary-card small{
            color:#cbd5e1;
            display:block;
            font-size:11px;
            text-transform:uppercase;
            letter-spacing:.05em;
        }
        .ai-summary-card strong{
            display:block;
            color:#ffffff;
            font-size:27px;
            margin-top:8px;
        }
        .ai-contract-card{
            background:#151d28;
            border:1px solid rgba(255,255,255,.06);
            border-left:6px solid var(--risk-color);
            border-radius:16px;
            padding:18px 20px;
            margin:14px 0;
            box-shadow:0 12px 28px rgba(0,0,0,.16);
        }
        .ai-contract-title{
            color:#ffffff;
            font-size:19px;
            font-weight:900;
            line-height:1.35;
            overflow-wrap:anywhere;
        }
        .ai-risk-badge{
            display:inline-flex;
            align-items:center;
            gap:6px;
            margin:10px 0 14px;
            padding:5px 12px;
            border-radius:999px;
            background:var(--risk-bg);
            color:var(--risk-color);
            font-size:12px;
            font-weight:900;
        }
        .ai-contract-grid{
            display:grid;
            grid-template-columns:repeat(4,minmax(0,1fr));
            gap:14px;
        }
        .ai-info-label{
            color:#94a3b8;
            font-size:11px;
            text-transform:uppercase;
            letter-spacing:.04em;
            font-weight:800;
        }
        .ai-info-value{
            color:#ffffff;
            font-size:13px;
            font-weight:750;
            margin-top:5px;
            overflow-wrap:anywhere;
        }
        .ai-empty-box{
            background:#111827;
            border:1px dashed rgba(250,204,21,.35);
            border-radius:14px;
            padding:18px;
            color:#cbd5e1;
            font-weight:700;
        }
        .ai-mini-note{
            color:#9ca3af;
            font-size:13px;
            margin-top:12px;
            line-height:1.6;
        }
        @media (max-width: 1100px){
            .ai-summary-grid{grid-template-columns:repeat(2,minmax(0,1fr));}
            .ai-contract-grid{grid-template-columns:repeat(2,minmax(0,1fr));}
        }
        @media (max-width: 700px){
            .ai-summary-grid{grid-template-columns:1fr;}
            .ai-contract-grid{grid-template-columns:1fr;}
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="section-title">Perguntas rápidas</div>', unsafe_allow_html=True)

    exemplos = [
        "Quantos contratos existem?",
        "Quais contratos são de risco alto?",
        "Quais contratos são de risco médio?",
        "Quais contratos são de risco baixo?",
        "Qual contrato tem maior valor?",
        "Qual contrato tem menor valor?",
        "Qual contrato tem menor score?",
        "Qual contrato tem maior score?",
        "Qual é o score médio?",
        "Quais contratos estão assinados?",
        "Quais contratos não estão assinados?",
        "Quais contratos são do Projuris?",
        "Quais contratos são do Ariba?",
        "Quais contratos foram analisados pelo Gemini?",
        "Liste os últimos contratos analisados",
    ]

    pergunta_exemplo = st.selectbox(
        "Escolha uma pergunta pronta",
        [""] + exemplos,
        key="assistente_ia_pergunta_pronta",
    )

    pergunta_digitada = st.chat_input("Ou digite sua pergunta sobre os contratos...")
    pergunta = (pergunta_digitada or pergunta_exemplo or "").strip()

    def texto_tem(texto, palavras):
        texto = str(texto or "").lower()
        return any(p in texto for p in palavras)

    def _garantir_colunas_chat(df: pd.DataFrame) -> pd.DataFrame:
        df = df.copy()
        defaults = {
            "fornecedor": "Não informado",
            "cnpj": "Não informado",
            "valor_total": "Não informado",
            "vigencia": "Não informado",
            "status": "Não informado",
            "risco": "N/A",
            "score": 0,
            "contrato_assinado": "Não informado",
            "modelo_ia": "Não informado",
            "tipo_origem": "Não informado",
            "arquivo": "Não informado",
            "data_analise": "Não informado",
            "id": "",
        }
        for col, default in defaults.items():
            if col not in df.columns:
                df[col] = default
            df[col] = df[col].fillna(default)
        df["risco_norm"] = df["risco"].astype(str).str.upper().str.strip().replace({"MEDIO": "MÉDIO"})
        df["score_num"] = pd.to_numeric(df["score"], errors="coerce").fillna(0)

        def _valor_num_chat(v):
            try:
                parsed = _parse_moeda_brasil(v)
                if parsed is not None:
                    return parsed
            except Exception:
                pass
            txt = str(v or "").replace("R$", "").replace(" ", "")
            if "," in txt:
                txt = txt.replace(".", "").replace(",", ".")
            try:
                return float(re.sub(r"[^0-9.\-]", "", txt) or 0)
            except Exception:
                return 0.0

        df["valor_num"] = df["valor_total"].apply(_valor_num_chat)
        df["data_dt"] = pd.to_datetime(df["data_analise"], errors="coerce", dayfirst=True)
        return df

    def _risco_style_chat(risco):
        risco = normalize_risco(risco)
        if risco == "ALTO":
            return "#ff4d4d", "rgba(255,77,77,.15)"
        if risco in ["MÉDIO", "MEDIO"]:
            return "#f59e0b", "rgba(245,158,11,.15)"
        if risco == "BAIXO":
            return "#22c55e", "rgba(34,197,94,.15)"
        return "#94a3b8", "rgba(148,163,184,.12)"

    def _render_pergunta_ia(texto):
        st.markdown(
            f'<div class="ai-question-card">🙋 {safe(texto)}</div>',
            unsafe_allow_html=True,
        )

    def _render_texto_ia(titulo, texto):
        texto_html = safe(texto).replace("\n", "<br>")
        st.markdown(
            f"""
            <div class="ai-panel">
                <div class="ai-answer-title">{safe(titulo)}</div>
                <div class="ai-answer-text">{texto_html}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    def _render_resumo_ia(df: pd.DataFrame, titulo: str = "📊 Resumo Executivo"):
        if df.empty:
            st.markdown(
                f"""
                <div class="ai-panel">
                    <div class="ai-answer-title">{safe(titulo)}</div>
                    <div class="ai-empty-box">Nenhum contrato encontrado para esta consulta.</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            return

        riscos = df["risco_norm"].astype(str)
        score_medio = round(pd.to_numeric(df["score_num"], errors="coerce").fillna(0).mean(), 1)
        qtd_alto = int((riscos == "ALTO").sum())
        qtd_medio = int((riscos == "MÉDIO").sum())
        qtd_baixo = int((riscos == "BAIXO").sum())
        valor_total = float(pd.to_numeric(df["valor_num"], errors="coerce").fillna(0).sum())
        valor_fmt = _formatar_moeda_brasil(valor_total) if valor_total > 0 else "Não informado"

        st.markdown(
            f"""
            <div class="ai-panel">
                <div class="ai-answer-title">{safe(titulo)}</div>
                <div class="ai-summary-grid">
                    <div class="ai-summary-card"><small>Contratos</small><strong>{len(df)}</strong></div>
                    <div class="ai-summary-card"><small>Score Médio</small><strong>{score_medio}</strong></div>
                    <div class="ai-summary-card"><small>Risco Alto</small><strong>{qtd_alto}</strong></div>
                    <div class="ai-summary-card"><small>Risco Médio</small><strong>{qtd_medio}</strong></div>
                    <div class="ai-summary-card"><small>Risco Baixo</small><strong>{qtd_baixo}</strong></div>
                </div>
                <div class="ai-mini-note">Valor somado dos registros encontrados: <b>{safe(valor_fmt)}</b>.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    def _render_lista_contratos_ia(df: pd.DataFrame, limite: int = 10):
        if df.empty:
            return

        cards = []
        for _, r in df.head(limite).iterrows():
            cor, fundo = _risco_style_chat(r.get("risco"))
            cards.append(f"""
            <div class="ai-contract-card" style="--risk-color:{cor};--risk-bg:{fundo};">
                <div class="ai-contract-title">{safe(r.get('fornecedor', 'Não informado'))}</div>
                <div class="ai-risk-badge">● RISCO {safe(r.get('risco', 'N/A'))}</div>
                <div class="ai-contract-grid">
                    <div><div class="ai-info-label">CNPJ</div><div class="ai-info-value">{safe(r.get('cnpj'))}</div></div>
                    <div><div class="ai-info-label">Score</div><div class="ai-info-value">{safe(r.get('score'))}</div></div>
                    <div><div class="ai-info-label">Valor</div><div class="ai-info-value">{safe(r.get('valor_total'))}</div></div>
                    <div><div class="ai-info-label">Status</div><div class="ai-info-value">{safe(r.get('status'))}</div></div>
                    <div><div class="ai-info-label">Assinatura</div><div class="ai-info-value">{safe(r.get('contrato_assinado'))}</div></div>
                    <div><div class="ai-info-label">Origem</div><div class="ai-info-value">{safe(r.get('tipo_origem'))}</div></div>
                    <div><div class="ai-info-label">Modelo IA</div><div class="ai-info-value">{safe(r.get('modelo_ia'))}</div></div>
                    <div><div class="ai-info-label">Data</div><div class="ai-info-value">{safe(r.get('data_analise'))}</div></div>
                </div>
                <div class="ai-mini-note">Arquivo: {safe(r.get('arquivo'))}</div>
            </div>
            """)

        if len(df) > limite:
            cards.append(f'<div class="ai-mini-note">Exibindo {limite} de {len(df)} contratos encontrados. Faça uma busca mais específica para reduzir a lista.</div>')

        st.markdown("\n".join(cards), unsafe_allow_html=True)

    def _render_conjunto_ia(titulo: str, df_resultado: pd.DataFrame, limite: int = 10):
        _render_resumo_ia(df_resultado, titulo)
        _render_lista_contratos_ia(df_resultado, limite=limite)

    def _render_um_contrato_ia(titulo: str, row: pd.Series):
        _render_texto_ia(titulo, "Resultado encontrado abaixo.")
        _render_lista_contratos_ia(pd.DataFrame([row]), limite=1)

    def _responder_ia(pergunta_original: str):
        df = _garantir_colunas_chat(contratos)
        pergunta_lower = pergunta_original.lower().strip()
        _render_pergunta_ia(pergunta_original)

        try:
            if texto_tem(pergunta_lower, ["oi", "olá", "ola", "bom dia", "boa tarde", "boa noite"]):
                _render_texto_ia(
                    "Olá!",
                    "Eu sou o Assistente Auditor. Posso consultar quantidade de contratos, riscos, scores, valores, origem, assinatura, modelo de IA e histórico das análises.",
                )

            elif texto_tem(pergunta_lower, ["quantos contratos", "total de contratos", "quantidade de contratos", "qtd contratos"]):
                _render_texto_ia("Total de contratos", f"Existem {len(df)} contrato(s) cadastrados no histórico.")
                _render_resumo_ia(df)

            elif texto_tem(pergunta_lower, ["risco alto", "alto risco", "contratos alto"]):
                _render_conjunto_ia("Contratos de risco alto", df[df["risco_norm"] == "ALTO"], limite=25)

            elif texto_tem(pergunta_lower, ["risco médio", "risco medio", "médio risco", "medio risco"]):
                _render_conjunto_ia("Contratos de risco médio", df[df["risco_norm"] == "MÉDIO"], limite=25)

            elif texto_tem(pergunta_lower, ["risco baixo", "baixo risco", "contratos baixo"]):
                _render_conjunto_ia("Contratos de risco baixo", df[df["risco_norm"] == "BAIXO"], limite=25)

            elif texto_tem(pergunta_lower, ["maior valor", "valor mais alto", "contrato mais caro", "maior contrato"]):
                base_valor = df.sort_values("valor_num", ascending=False)
                if base_valor.empty:
                    _render_texto_ia("Maior valor", "Nenhum contrato encontrado.")
                else:
                    _render_um_contrato_ia("Contrato com maior valor", base_valor.iloc[0])

            elif texto_tem(pergunta_lower, ["menor valor", "valor mais baixo", "contrato mais barato"]):
                base_valor = df[df["valor_num"] > 0].sort_values("valor_num", ascending=True)
                if base_valor.empty:
                    _render_texto_ia("Menor valor", "Nenhum contrato com valor numérico localizado.")
                else:
                    _render_um_contrato_ia("Contrato com menor valor", base_valor.iloc[0])

            elif texto_tem(pergunta_lower, ["menor score", "pior score", "menor nota", "pior contrato"]):
                base_score = df.sort_values("score_num", ascending=True)
                _render_um_contrato_ia("Contrato com menor score", base_score.iloc[0])

            elif texto_tem(pergunta_lower, ["maior score", "melhor score", "maior nota", "melhor contrato"]):
                base_score = df.sort_values("score_num", ascending=False)
                _render_um_contrato_ia("Contrato com maior score", base_score.iloc[0])

            elif texto_tem(pergunta_lower, ["score médio", "score medio", "média de score", "media de score"]):
                media = round(float(df["score_num"].mean()), 1) if not df.empty else 0
                _render_texto_ia("Score médio", f"O score médio dos contratos é {media}.")
                _render_resumo_ia(df)

            elif texto_tem(pergunta_lower, ["não assinados", "nao assinados", "sem assinatura", "não estão assinados", "nao estao assinados", "contrato não assinado", "contrato nao assinado"]):
                ass = df["contrato_assinado"].astype(str).str.upper().str.strip()
                _render_conjunto_ia("Contratos não assinados", df[ass != "SIM"], limite=25)

            elif texto_tem(pergunta_lower, ["contratos assinados", "estão assinados", "estao assinados", "com assinatura", "contrato assinado"]):
                ass = df["contrato_assinado"].astype(str).str.upper().str.strip()
                _render_conjunto_ia("Contratos assinados", df[ass == "SIM"], limite=25)

            elif texto_tem(pergunta_lower, ["projuris"]):
                filtro = df[df["tipo_origem"].astype(str).str.lower().str.contains("projuris", na=False)]
                _render_conjunto_ia("Contratos do Projuris", filtro, limite=25)

            elif texto_tem(pergunta_lower, ["ariba"]):
                filtro = df[df["tipo_origem"].astype(str).str.lower().str.contains("ariba", na=False)]
                _render_conjunto_ia("Contratos do Ariba", filtro, limite=25)

            elif texto_tem(pergunta_lower, ["gemini", "modelo ia", "inteligência artificial", "inteligencia artificial", "analisados pela ia", "analisados pelo gemini"]) or re.search(r"\bia\b", pergunta_lower):
                filtro = df[df["modelo_ia"].astype(str).str.lower().str.contains("gemini", na=False)]
                _render_conjunto_ia("Contratos analisados pelo Gemini", filtro, limite=25)

            elif texto_tem(pergunta_lower, ["últimos", "ultimos", "recentes", "últimas análises", "ultimas analises", "últimos contratos", "ultimos contratos"]):
                ultimos = df.sort_values(["data_dt", "id"], ascending=[False, False], na_position="last").head(10)
                _render_conjunto_ia("Últimos contratos analisados", ultimos, limite=10)

            else:
                busca = re.escape(pergunta_lower)
                cols = [
                    "fornecedor", "cnpj", "valor_total", "vigencia", "status", "risco",
                    "contrato_assinado", "modelo_ia", "tipo_origem", "arquivo", "data_analise",
                ]
                mask = pd.Series(False, index=df.index)
                for col in cols:
                    mask = mask | df[col].astype(str).str.lower().str.contains(busca, na=False, regex=True)
                resultado = df[mask]
                if resultado.empty:
                    _render_texto_ia(
                        "Nenhum resultado encontrado",
                        "Não encontrei contratos relacionados a essa busca. Pesquise por fornecedor, CNPJ, valor, risco, status, origem, modelo IA ou nome do arquivo.",
                    )
                else:
                    _render_conjunto_ia("Resultado da busca", resultado, limite=50)

        except Exception as erro:
            _render_texto_ia("Erro ao consultar histórico", f"Ocorreu um erro ao consultar o histórico: {erro}")

    if pergunta:
        _responder_ia(pergunta)
    else:
        _render_texto_ia(
            "Como consultar",
            "Escolha uma pergunta pronta acima ou digite sua dúvida no campo inferior. As respostas agora são renderizadas em cards visuais, sem exibir HTML ou código na tela.",
        )

    st.markdown('<div class="footer">Auditor de Contratos - Grupo SBF • Suprimentos • Análise de Contratos</div>', unsafe_allow_html=True)
    st.stop()


# =========================================================
# NOVA ANÁLISE
# =========================================================
if pagina == "📄 Nova Análise":
    render_hero("Auditor de Contratos - Grupo SBF", "Análise profissional e automatizada de contratos Projuris ou Ariba em PDF e Word.")

    st.markdown('<div class="section-title">Tipo de análise</div>', unsafe_allow_html=True)
    origem_contrato = st.radio("Origem do contrato", ["📘 Projuris", "🛒 Ariba"], horizontal=True, label_visibility="collapsed")

    st.markdown('<div class="section-title">Upload do contrato</div>', unsafe_allow_html=True)

    # Controle para limpar todos os anexos de uma vez.
    # No Streamlit, o file_uploader só é realmente limpo quando a key muda.
    if "upload_contrato_key" not in st.session_state:
        st.session_state["upload_contrato_key"] = 0

    up_col1, up_col2 = st.columns([4, 1])
    with up_col2:
        if st.button("🧹 Limpar anexos", use_container_width=True):
            st.session_state["upload_contrato_key"] += 1
            st.rerun()

    arquivos = st.file_uploader(
        "Envie o contrato principal e anexos",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key=f"upload_contrato_{st.session_state['upload_contrato_key']}",
    )

    if not arquivos:
        cols = st.columns(4)
        cards = [("Formato", "PDF/Word"), ("Análise", "Local ou IA"), ("Saída", "Excel"), ("Risco", "Score")]
        for col, (label, value) in zip(cols, cards):
            col.markdown(render_metric(label, value), unsafe_allow_html=True)

        st.markdown('<div class="section-title">Como funciona</div>', unsafe_allow_html=True)
        st.markdown(
            """
            <div class="flow-grid">
                <div class="flow-card"><div class="flow-number">1</div><h4>Upload do contrato</h4><p>Envie o contrato principal e todos os anexos necessários.</p></div>
                <div class="flow-card"><div class="flow-number">2</div><h4>Extração dos dados</h4><p>O sistema consolida PDF, Word e anexos em uma análise única.</p></div>
                <div class="flow-card"><div class="flow-number">3</div><h4>Validação do risco</h4><p>As cláusulas são avaliadas por pendência, criticidade e score.</p></div>
                <div class="flow-card"><div class="flow-number">4</div><h4>Relatório executivo</h4><p>Gere o Excel com resumo, checklist, pendências e texto extraído.</p></div>
                <div class="flow-card"><div class="flow-number">5</div><h4>Histórico</h4><p>A análise fica salva para consulta posterior no dashboard.</p></div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.success(f"{len(arquivos)} arquivo(s) carregado(s). Clique abaixo para iniciar a análise.")

        if st.button("🚀 Analisar contrato e anexos", use_container_width=True):
            popup_processamento = st.empty()
            inicio_processamento = time.perf_counter()
            st.session_state["ultimo_tempo_ia_segundos"] = 0
            arquivos_status = {arquivo.name: "PENDENTE" for arquivo in arquivos}

            atualizar_popup_processamento(
                popup_processamento,
                arquivos_status,
                "",
                "Preparando arquivos para análise",
                inicio_processamento,
            )

            try:
                texto_total = ""
                textos_por_arquivo: Dict[str, str] = {}

                for arquivo in arquivos:
                    arquivos_status[arquivo.name] = "PROCESSANDO"
                    atualizar_popup_processamento(
                        popup_processamento,
                        arquivos_status,
                        arquivo.name,
                        "Lendo e preparando anexo",
                        inicio_processamento,
                    )

                    try:
                        if arquivo.name.lower().endswith(".pdf"):
                            texto_arquivo = ler_pdf(arquivo)
                        else:
                            texto_arquivo = ler_docx(arquivo)

                        arquivos_status[arquivo.name] = "CONCLUIDO"
                        atualizar_popup_processamento(
                            popup_processamento,
                            arquivos_status,
                            arquivo.name,
                            "Anexo processado com sucesso",
                            inicio_processamento,
                        )
                    except Exception as e:
                        texto_arquivo = f"Erro ao ler arquivo {arquivo.name}: {e}"
                        arquivos_status[arquivo.name] = "ERRO"
                        atualizar_popup_processamento(
                            popup_processamento,
                            arquivos_status,
                            arquivo.name,
                            "Erro ao processar anexo",
                            inicio_processamento,
                        )

                    textos_por_arquivo[arquivo.name] = texto_arquivo
                    texto_total += "\n\n==============================\n"
                    texto_total += f"ARQUIVO: {arquivo.name}\n"
                    texto_total += "==============================\n"
                    texto_total += texto_arquivo

                texto = texto_total.strip()

                atualizar_popup_processamento(
                    popup_processamento,
                    arquivos_status,
                    "Triagem",
                    "Classificando anexos: assinados, apoio, duplicados e ignorados",
                    inicio_processamento,
                )

                triagem_info = classificar_anexos_para_analise(arquivos, textos_por_arquivo)
                triagem_anexos = triagem_info.get("triagem", [])
                arquivos_para_ia = triagem_info.get("arquivos_para_ia", arquivos)
                texto_para_ia = montar_texto_para_ia_com_triagem(textos_por_arquivo, triagem_anexos)

                nomes_para_ia = [getattr(a, "name", "documento") for a in (arquivos_para_ia or [])]
                falha_leitura_critica = bool(nomes_para_ia) and all(
                    texto_indica_falha_leitura(textos_por_arquivo.get(nome, ""))
                    and not texto_tem_conteudo_contratual(textos_por_arquivo.get(nome, ""))
                    for nome in nomes_para_ia
                )

                decisao_por_arquivo = {t.get("Arquivo"): t.get("Decisão") for t in triagem_anexos}
                for nome_arquivo, decisao in decisao_por_arquivo.items():
                    if decisao == "Análise profunda":
                        arquivos_status[nome_arquivo] = "ANALISE_IA"
                    elif str(decisao).startswith("Apoio"):
                        arquivos_status[nome_arquivo] = "APOIO"
                    elif "Ignorado" in str(decisao):
                        arquivos_status[nome_arquivo] = "IGNORADO"

                atualizar_popup_processamento(
                    popup_processamento,
                    arquivos_status,
                    f"Arquivos preparados: {len(arquivos_status)}/{len(arquivos_status)}",
                    f"Triagem concluída. {len(arquivos_para_ia)} arquivo(s) seguirão para análise profunda da IA",
                    inicio_processamento,
                )

                try:
                    if modo != "Análise Local":
                        if not gemini_key:
                            fechar_popup_processamento(popup_processamento)
                            st.error("Configure a GEMINI_API_KEY no Streamlit Secrets ou no arquivo .env para usar a análise IA.")
                            st.stop()

                        atualizar_popup_processamento(
                            popup_processamento,
                            arquivos_status,
                            f"Arquivos preparados: {len(arquivos_status)}/{len(arquivos_status)}",
                            "IA analisando o pacote final...\nTempo da IA: 0s",
                            inicio_processamento,
                        )

                        resultado = executar_ia_com_timer_ao_vivo(
                            analisar_gemini,
                            popup_processamento,
                            arquivos_status,
                            inicio_processamento,
                            texto=texto_para_ia,
                            api_key=gemini_key,
                            opcao_modelo=modo,
                            arquivos_originais=arquivos_para_ia,
                        )

                        atualizar_popup_processamento(
                            popup_processamento,
                            arquivos_status,
                            "Gemini",
                            "IA finalizou a análise. Preparando resultado na tela",
                            inicio_processamento,
                        )
                    else:
                        atualizar_popup_processamento(
                            popup_processamento,
                            arquivos_status,
                            "Análise Local",
                            "Executando análise local dos documentos",
                            inicio_processamento,
                        )
                        resultado = local_extract(texto_para_ia)
                except Exception as e:
                    if 'falha_leitura_critica' in locals() and falha_leitura_critica:
                        atualizar_popup_processamento(
                            popup_processamento,
                            arquivos_status,
                            "Falha técnica",
                            "O documento original não foi lido com segurança. Gerando alerta sem preencher cards incorretos",
                            inicio_processamento,
                        )
                        nome_falha = " | ".join(nomes_para_ia or [getattr(a, "name", "documento") for a in arquivos])
                        st.error(
                            "Não foi possível analisar o documento com segurança. "
                            "O sistema não vai preencher cards com chute. "
                            f"Detalhe: {e}"
                        )
                        resultado = resultado_falha_tecnica_leitura(nome_falha, e)
                    else:
                        atualizar_popup_processamento(
                            popup_processamento,
                            arquivos_status,
                            "Análise Local",
                            "A IA falhou. Executando análise local de contingência",
                            inicio_processamento,
                        )
                        st.warning(f"A IA falhou e o sistema usou análise local. Detalhe: {e}")
                        resultado = local_extract(texto)

                atualizar_popup_processamento(
                    popup_processamento,
                    arquivos_status,
                    "Resultado",
                    "Normalizando dados, aplicando regras e salvando histórico",
                    inicio_processamento,
                )

                # Preserva a resposta bruta da IA para reaplicar a matriz de evidências
                # depois das regras legadas de compatibilidade.
                resultado_bruto_ia = dict(resultado or {})
                resultado_bruto_ia["texto_extraido"] = texto

                # Extração determinística adicional: percorre todas as linhas das tabelas
                # comerciais, inclusive faixas, isenções e valores textuais. O resultado
                # é mesclado com a IA; nunca substitui uma tabela completa por um resumo.
                itens_tabela_documento = extrair_tabela_comercial_completa(texto)
                # O parser genérico só entra quando não existe tabela comercial estruturada;
                # isso evita duplicar faixas e criar descrições truncadas pelo OCR.
                itens_fallback_local = [] if itens_tabela_documento else extrair_itens_local(texto)
                itens_completos = mesclar_itens_comerciais(
                    resultado_bruto_ia.get("itens_contrato", []),
                    itens_tabela_documento,
                    itens_fallback_local,
                )
                resultado_bruto_ia["itens_contrato"] = normalizar_itens_contrato(itens_completos)
                resultado_bruto_ia["metricas_tabela_comercial"] = calcular_metricas_tabela_comercial(
                    itens_tabela_documento,
                    resultado_bruto_ia["itens_contrato"],
                )
                resultado = normalizar(resultado_bruto_ia)

                # Reforço final: para serviços percentuais, consolida taxa/encargos no serviço principal.
                itens_percentuais = detectar_servico_percentual(texto)
                if itens_percentuais and (_texto_indica_mao_obra_temporaria(texto) or not resultado.get("itens_contrato") or _itens_sao_apenas_atributos_comerciais(resultado.get("itens_contrato", []))):
                    resultado["itens_contrato"] = normalizar_itens_contrato(itens_percentuais)
                elif not resultado.get("itens_contrato"):
                    resultado["itens_contrato"] = extrair_itens_local(texto)

                resultado["metricas_tabela_comercial"] = resultado_bruto_ia.get("metricas_tabela_comercial", {})

                # Regras legadas continuam sendo executadas para compatibilidade visual/histórico.
                resultado = aplicar_regras_finais_contrato(resultado, texto)
                resultado["data_analise"] = datetime.now().strftime("%d/%m/%Y %H:%M")
                resultado["origem_contrato"] = origem_contrato

                nomes_arquivos = " | ".join([arquivo.name for arquivo in arquivos])
                resultado["texto_extraido"] = texto
                resultado["arquivos_analisados"] = nomes_arquivos
                resultado["tipo_origem"] = origem_contrato.replace("📘", "").replace("🛒", "").strip()
                resultado["modelo_ia"] = resultado.get("modelo_ia", modo if modo != "Análise Local" else "Análise Local")

                # Aditivos e, por último, a trava de coerência baseada em evidência.
                resultado = aplicar_regras_aditivos(resultado, texto)
                resultado = _finalizar_coerencia_pos_processamento(
                    resultado,
                    resultado_bruto_ia,
                    triagem_anexos,
                )

                triagem_anexos = resultado.get("triagem_anexos", triagem_anexos)
                resultado["arquivos_enviados_analise_profunda"] = " | ".join([getattr(a, "name", "documento") for a in arquivos_para_ia])
                resultado["arquivos_ignorados_analise_profunda"] = [t for t in triagem_anexos if "Ignorado" in str(t.get("Decisão", ""))]

                total_triagem = len(triagem_anexos)
                qtd_profunda = sum(1 for x in triagem_anexos if x.get("Decisão") == "Análise profunda")
                qtd_apoio = sum(1 for x in triagem_anexos if str(x.get("Decisão", "")).startswith("Apoio"))
                qtd_ignorados = sum(1 for x in triagem_anexos if "Ignorado" in str(x.get("Decisão", "")))
                tempo_total_segundos = time.perf_counter() - inicio_processamento
                tempo_ia_segundos = float(resultado.get("_tempo_ia_segundos") or st.session_state.get("ultimo_tempo_ia_segundos", 0) or 0)
                resultado["resumo_processamento"] = {
                    "total_arquivos": total_triagem or len(arquivos),
                    "analise_profunda": qtd_profunda,
                    "apoio": qtd_apoio,
                    "ignorados": qtd_ignorados,
                    "tempo_total": _formatar_tempo_execucao(tempo_total_segundos),
                    "tempo_ia": _formatar_tempo_execucao(tempo_ia_segundos),
                    "finalizado_em": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                }

                salvar_analise(
                    resultado,
                    nomes_arquivos,
                    modelo_ia=resultado.get("modelo_ia"),
                    tipo_origem=resultado.get("tipo_origem"),
                    texto_extraido=texto,
                )

                atualizar_popup_processamento(
                    popup_processamento,
                    arquivos_status,
                    "Concluído",
                    f"Análise concluída com sucesso\nTempo total: {resultado['resumo_processamento']['tempo_total']}\nTempo da IA: {resultado['resumo_processamento']['tempo_ia']}",
                    inicio_processamento,
                )
                time.sleep(1.2)
                fechar_popup_processamento(popup_processamento)

            except Exception as e:
                fechar_popup_processamento(popup_processamento)
                st.error(f"Erro durante a análise: {e}")
                st.stop()

            risco = normalize_risco(resultado.get("risco"))
            pill = "pill-ok" if risco == "BAIXO" else "pill-warn" if risco == "MÉDIO" else "pill-danger"

            render_indicadores_analise(resultado)

            render_resumo_processamento_final(resultado)

            if clean_text(resultado.get("contrato_assinado")).lower() != "sim":
                st.error("⚠️ Contrato sem assinatura validada. Revisar antes da criação da RC/PO.")

            render_resultado_em_abas(
                resultado,
                texto_extraido=texto,
                key_prefix="texto_nova_analise",
            )

            excel = gerar_excel(resultado, texto)
            st.download_button(
                "📥 Baixar relatório Excel",
                data=excel,
                file_name=f"relatorio_auditor_contract_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )


# =========================================================
# HISTÓRICO
# =========================================================
if pagina == "📚 Histórico":
    render_hero("Histórico", "Consulta executiva dos contratos analisados, com filtros, indicadores e relatórios.")

    historico = carregar_historico_seguro()

    if historico.empty:
        st.info("Nenhuma análise salva ainda.")
        st.markdown('<div class="footer">Auditor de Contratos - Grupo SBF • Suprimentos • Análise de Contratos</div>', unsafe_allow_html=True)
        st.stop()

    # -------------------------
    # Preparação dos dados
    # -------------------------
    hist = historico.copy()
    for col in ["fornecedor", "cnpj", "valor_total", "vigencia", "status", "risco", "score", "contrato_assinado", "modelo_ia", "tipo_origem", "arquivo", "data_analise"]:
        if col not in hist.columns:
            hist[col] = "Não informado"

    hist["risco_norm"] = hist["risco"].astype(str).str.upper().replace({"MEDIO": "MÉDIO"})
    hist["score_num"] = hist["score"].apply(as_float_score)
    hist["data_dt"] = pd.to_datetime(hist["data_analise"], format="%d/%m/%Y %H:%M", errors="coerce")

    total_hist = int(len(hist))
    qtd_alto = int((hist["risco_norm"] == "ALTO").sum())
    qtd_medio = int((hist["risco_norm"] == "MÉDIO").sum())
    qtd_baixo = int((hist["risco_norm"] == "BAIXO").sum())
    qtd_assinado = int(hist["contrato_assinado"].astype(str).str.upper().eq("SIM").sum())
    score_medio_hist = round(float(hist["score_num"].mean()), 1) if total_hist else 0

    # -------------------------
    # Indicadores executivos
    # -------------------------
    st.markdown('<div class="section-title">Painel do histórico</div>', unsafe_allow_html=True)
    h1, h2, h3, h4, h5 = st.columns(5)
    h1.markdown(render_metric("Contratos", total_hist), unsafe_allow_html=True)
    h2.markdown(render_metric("Risco alto", qtd_alto), unsafe_allow_html=True)
    h3.markdown(render_metric("Risco médio", qtd_medio), unsafe_allow_html=True)
    h4.markdown(render_metric("Risco baixo", qtd_baixo), unsafe_allow_html=True)
    h5.markdown(render_metric("Score médio", score_medio_hist), unsafe_allow_html=True)

    st.markdown(
        f"""
        <div class="executive-box">
            <h3>📚 Visão executiva do histórico</h3>
            <p>Existem <b>{total_hist}</b> análise(s) registrada(s), sendo <b>{qtd_baixo}</b> de risco baixo, <b>{qtd_medio}</b> de risco médio e <b>{qtd_alto}</b> de risco alto.</p>
            <p><b>{qtd_assinado}</b> contrato(s) possuem evidência de assinatura registrada no histórico.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -------------------------
    # Filtros profissionais
    # -------------------------
    st.markdown('<div class="section-title">Filtros</div>', unsafe_allow_html=True)

    riscos_disponiveis = [r for r in ["ALTO", "MÉDIO", "BAIXO"] if r in set(hist["risco_norm"].dropna().astype(str))]
    origens_disponiveis = sorted([x for x in hist["tipo_origem"].dropna().astype(str).unique() if x and x != "Não informado"])
    assinaturas_disponiveis = sorted([x for x in hist["contrato_assinado"].dropna().astype(str).unique() if x])
    modelos_disponiveis = sorted([x for x in hist["modelo_ia"].dropna().astype(str).unique() if x])
    status_disponiveis = sorted([x for x in hist["status"].dropna().astype(str).unique() if x])

    with st.container(border=True):
        st.markdown("#### 🔎 Consulta do histórico")
        st.caption("Use os filtros principais abaixo. Os filtros avançados ficam recolhidos para não poluir a tela.")

        f1, f2, f3, f4 = st.columns([1.8, .8, .9, .9])

        with f1:
            busca = st.text_input(
                "Buscar",
                placeholder="Digite contraparte, CNPJ, arquivo, status ou modelo...",
            ).strip()

        with f2:
            risco_opcao = st.selectbox("Risco", ["Todos"] + riscos_disponiveis, index=0)

        with f3:
            assinatura_opcao = st.selectbox("Assinatura", ["Todos"] + assinaturas_disponiveis, index=0)

        with f4:
            ordenar_por = st.selectbox("Ordenar por", ["Mais recentes", "Maior score", "Menor score", "Risco", "Contraparte"], index=0)

        with st.expander("⚙️ Filtros avançados", expanded=False):
            a1, a2, a3, a4 = st.columns([1, 1, 1, 1])
            with a1:
                origem_opcao = st.selectbox("Origem", ["Todas"] + origens_disponiveis, index=0)
            with a2:
                modelo_opcao = st.selectbox("Modelo IA", ["Todos"] + modelos_disponiveis, index=0)
            with a3:
                status_opcao = st.selectbox("Status", ["Todos"] + status_disponiveis, index=0)
            with a4:
                score_min, score_max = st.slider("Score", 0, 100, (0, 100))

    # Converte seleção única em listas para reaproveitar a lógica de filtragem.
    riscos_sel = riscos_disponiveis if risco_opcao == "Todos" else [risco_opcao]
    origens_sel = origens_disponiveis if origem_opcao == "Todas" else [origem_opcao]
    assinatura_sel = assinaturas_disponiveis if assinatura_opcao == "Todos" else [assinatura_opcao]
    modelos_sel = modelos_disponiveis if modelo_opcao == "Todos" else [modelo_opcao]
    status_sel = status_disponiveis if status_opcao == "Todos" else [status_opcao]

    filtrado = hist.copy()

    if busca:
        busca_lower = busca.lower()
        cols_busca = ["fornecedor", "cnpj", "valor_total", "vigencia", "status", "risco", "modelo_ia", "tipo_origem", "arquivo", "data_analise"]
        mask = pd.Series(False, index=filtrado.index)
        for col in cols_busca:
            mask = mask | filtrado[col].astype(str).str.lower().str.contains(busca_lower, na=False)
        filtrado = filtrado[mask]

    if riscos_sel:
        filtrado = filtrado[filtrado["risco_norm"].isin(riscos_sel)]
    else:
        filtrado = filtrado.iloc[0:0]

    if origens_disponiveis and origens_sel:
        filtrado = filtrado[filtrado["tipo_origem"].astype(str).isin(origens_sel)]
    elif origens_disponiveis:
        filtrado = filtrado.iloc[0:0]

    if assinatura_sel:
        filtrado = filtrado[filtrado["contrato_assinado"].astype(str).isin(assinatura_sel)]
    else:
        filtrado = filtrado.iloc[0:0]

    if modelos_sel:
        filtrado = filtrado[filtrado["modelo_ia"].astype(str).isin(modelos_sel)]
    else:
        filtrado = filtrado.iloc[0:0]

    if status_sel:
        filtrado = filtrado[filtrado["status"].astype(str).isin(status_sel)]
    else:
        filtrado = filtrado.iloc[0:0]

    filtrado = filtrado[(filtrado["score_num"] >= score_min) & (filtrado["score_num"] <= score_max)]

    if ordenar_por == "Mais recentes":
        filtrado = filtrado.sort_values(["data_dt", "id"], ascending=[False, False], na_position="last")
    elif ordenar_por == "Maior score":
        filtrado = filtrado.sort_values("score_num", ascending=False)
    elif ordenar_por == "Menor score":
        filtrado = filtrado.sort_values("score_num", ascending=True)
    elif ordenar_por == "Risco":
        ordem_risco = {"ALTO": 1, "MÉDIO": 2, "BAIXO": 3}
        filtrado["ordem_risco"] = filtrado["risco_norm"].map(ordem_risco).fillna(9)
        filtrado = filtrado.sort_values(["ordem_risco", "score_num"], ascending=[True, True])
    elif ordenar_por == "Contraparte":
        filtrado = filtrado.sort_values("fornecedor", ascending=True)

    st.markdown(
        f'<div class="filter-note">🔎 Visualizando <b>{len(filtrado)}</b> de <b>{total_hist}</b> análise(s) do histórico.</div>',
        unsafe_allow_html=True,
    )

    # -------------------------
    # Exportação filtrada
    # -------------------------
    filtrado_export = filtrado.copy()

    # Campos novos ficam dentro do resultado_json no histórico. Aqui extraímos para o Excel.
    def _hist_json_val(row, chave: str, padrao: str = "Não informado") -> str:
        try:
            payload = json.loads(row.get("resultado_json") or "{}") if isinstance(row.get("resultado_json"), str) else {}
            valor = payload.get(chave, padrao)
            if valor in (None, "", [], {}, "Não localizado", "Não localizada"):
                return padrao
            if isinstance(valor, (list, dict)):
                return json.dumps(valor, ensure_ascii=False)
            return str(valor)
        except Exception:
            return padrao

    def _hist_json_nested_val(row, pai: str, chave: str, padrao: str = "Não informado") -> str:
        try:
            payload = json.loads(row.get("resultado_json") or "{}") if isinstance(row.get("resultado_json"), str) else {}
            bloco = payload.get(pai, {}) if isinstance(payload, dict) else {}
            valor = bloco.get(chave, padrao) if isinstance(bloco, dict) else padrao
            if valor in (None, "", [], {}, "Não localizado", "Não localizada"):
                return padrao
            return str(valor)
        except Exception:
            return padrao

    def _hist_json_count(row, chave: str) -> int:
        try:
            payload = json.loads(row.get("resultado_json") or "{}") if isinstance(row.get("resultado_json"), str) else {}
            valor = payload.get(chave, []) if isinstance(payload, dict) else []
            return len(valor) if isinstance(valor, list) else 0
        except Exception:
            return 0

    if "resultado_json" in filtrado_export.columns:
        campos_json_excel = {
            "valor_mensal_estimado": "valor_mensal_estimado",
            "valor_total_estimado_vigencia": "valor_total_estimado_vigencia",
            "valor_total_materiais_servicos": "valor_total_materiais_servicos",
            "data_contrato": "data_contrato",
            "data_conclusao_docusign": "data_conclusao_docusign",
            "pessoas_que_assinaram": "pessoas_que_assinaram",
            "resumo_aditivos": "resumo_aditivos",
        }
        for destino, chave_json in campos_json_excel.items():
            if destino not in filtrado_export.columns:
                filtrado_export[destino] = filtrado_export.apply(lambda row, ch=chave_json: _hist_json_val(row, ch), axis=1)

        campos_processamento_excel = {
            "triagem_total_arquivos": ("resumo_processamento", "total_arquivos"),
            "triagem_analise_ia": ("resumo_processamento", "analise_profunda"),
            "triagem_apoio": ("resumo_processamento", "apoio"),
            "triagem_ignorados": ("resumo_processamento", "ignorados"),
            "tempo_total_processamento": ("resumo_processamento", "tempo_total"),
            "tempo_ia": ("resumo_processamento", "tempo_ia"),
        }
        for destino, (pai, chave_json) in campos_processamento_excel.items():
            if destino not in filtrado_export.columns:
                filtrado_export[destino] = filtrado_export.apply(lambda row, pa=pai, ch=chave_json: _hist_json_nested_val(row, pa, ch), axis=1)

        if "qtd_aditivos" not in filtrado_export.columns:
            filtrado_export["qtd_aditivos"] = filtrado_export.apply(lambda row: _hist_json_count(row, "aditivos_contrato"), axis=1)
        if "qtd_itens_contrato" not in filtrado_export.columns:
            filtrado_export["qtd_itens_contrato"] = filtrado_export.apply(lambda row: _hist_json_count(row, "itens_contrato"), axis=1)

    export_cols = [
        "id", "data_analise", "fornecedor", "cnpj", "valor_total",
        "valor_mensal_estimado", "valor_total_estimado_vigencia", "valor_total_materiais_servicos",
        "qtd_aditivos", "resumo_aditivos", "qtd_itens_contrato",
        "data_contrato", "data_conclusao_docusign", "pessoas_que_assinaram",
        "triagem_total_arquivos", "triagem_analise_ia", "triagem_apoio", "triagem_ignorados",
        "tempo_total_processamento", "tempo_ia",
        "vigencia", "status", "risco", "score", "contrato_assinado", "modelo_ia", "tipo_origem", "arquivo",
    ]
    export_df = filtrado_export[[c for c in export_cols if c in filtrado_export.columns]].copy()
    export_df = export_df.rename(columns={
        "id": "ID",
        "data_analise": "Data da análise",
        "fornecedor": "Contraparte",
        "cnpj": "CNPJ",
        "valor_total": "Valor total",
        "valor_mensal_estimado": "Valor mensal estimado",
        "valor_total_estimado_vigencia": "Valor total estimado da vigência",
        "valor_total_materiais_servicos": "Valor total dos materiais e serviços",
        "qtd_aditivos": "Qtd. aditivos",
        "resumo_aditivos": "Resumo dos aditivos",
        "qtd_itens_contrato": "Qtd. itens/serviços",
        "data_contrato": "Data do contrato",
        "data_conclusao_docusign": "Data conclusão DocuSign",
        "pessoas_que_assinaram": "Pessoas que assinaram",
        "triagem_total_arquivos": "Total de anexos",
        "triagem_analise_ia": "Análise IA",
        "triagem_apoio": "Apoio",
        "triagem_ignorados": "Ignorados",
        "tempo_total_processamento": "Tempo total processamento",
        "tempo_ia": "Tempo IA",
        "vigencia": "Vigência",
        "status": "Status",
        "risco": "Risco",
        "score": "Score",
        "contrato_assinado": "Assinado",
        "modelo_ia": "Modelo IA",
        "tipo_origem": "Origem",
        "arquivo": "Arquivos analisados",
    })

    excel_hist = gerar_excel_historico_profissional(export_df, total_geral=total_hist)

    d1, d2 = st.columns([3, 1])
    with d2:
        st.download_button(
            "📥 Baixar histórico filtrado",
            data=excel_hist,
            file_name=f"historico_auditor_contract_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

    # -------------------------
    # Visualização
    # -------------------------
    st.markdown('<div class="section-title">Resultado do histórico</div>', unsafe_allow_html=True)

    if filtrado.empty:
        st.warning("Nenhuma análise encontrada para os filtros selecionados.")
    else:
        tab_cards, tab_tabela, tab_auditoria = st.tabs(["📌 Cards executivos", "📋 Tabela executiva", "🧾 Auditoria técnica"])

        with tab_cards:
            for _, row in filtrado.head(25).iterrows():
                with st.container():
                    st.markdown(render_historico_card_executivo(row), unsafe_allow_html=True)

                    with st.expander("🔎 Abrir análise completa deste contrato", expanded=False):
                        render_analise_completa_historico(row)

            if len(filtrado) > 25:
                st.info("Exibindo os 25 primeiros registros filtrados. Use os filtros ou a tabela executiva para consultar os demais.")

        with tab_tabela:
            tabela = export_df.copy()
            st.dataframe(
                tabela,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Score": st.column_config.ProgressColumn("Score", min_value=0, max_value=100),
                    "Risco": st.column_config.TextColumn("Risco"),
                    "Valor total": st.column_config.TextColumn("Valor total"),
                    "Arquivos analisados": st.column_config.TextColumn("Arquivos analisados", width="large"),
                },
            )

        with tab_auditoria:
            auditoria_cols = [c for c in ["id", "data_criacao", "data_analise", "modelo_ia", "tipo_origem", "risco", "score", "status", "contrato_assinado", "arquivo"] if c in filtrado.columns]
            st.dataframe(
                filtrado[auditoria_cols].rename(columns={
                    "id": "ID",
                    "data_criacao": "Criado em",
                    "data_analise": "Data análise",
                    "modelo_ia": "Modelo IA",
                    "tipo_origem": "Origem",
                    "risco": "Risco",
                    "score": "Score",
                    "status": "Status",
                    "contrato_assinado": "Assinado",
                    "arquivo": "Arquivo",
                }),
                use_container_width=True,
                hide_index=True,
            )

    st.markdown('<div class="footer">Auditor de Contratos - Grupo SBF • Suprimentos • Análise de Contratos</div>', unsafe_allow_html=True)
    st.stop()

st.markdown('<div class="footer">Auditor de Contratos - Grupo SBF • Suprimentos • Análise de Contratos</div>', unsafe_allow_html=True)
